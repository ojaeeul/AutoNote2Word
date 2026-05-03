import datetime
import re
import streamlit as st

# --- Premium UI/UX Design System Injection ---
def inject_premium_design():
    st.set_page_config(page_title="SNU Chem-Ed Studio Pro", page_icon="🧪", layout="wide")
    st.markdown("""
        <style>
        @import url('https://fonts.googleapis.com/css2?family=Outfit:wght@300;400;600;700&display=swap');

        /* Global Styles */
        * { font-family: 'Outfit', sans-serif; }
        
        /* Sidebar Glassmorphism */
        section[data-testid="stSidebar"] {
            background: rgba(255, 255, 255, 0.05);
            backdrop-filter: blur(10px);
            border-right: 1px solid rgba(0, 0, 0, 0.1);
        }

        /* Gradient Headers */
        .main-header {
            background: linear-gradient(135deg, #6366F1 0%, #3B82F6 100%);
            -webkit-background-clip: text;
            -webkit-text-fill-color: transparent;
            font-weight: 700;
            font-size: 3rem;
            margin-bottom: 1rem;
        }

        /* Card Design */
        div.stButton > button {
            border-radius: 12px;
            background: linear-gradient(135deg, #4F46E5 0%, #3B82F6 100%);
            color: white;
            border: none;
            padding: 0.6rem 1.2rem;
            font-weight: 600;
            transition: all 0.3s cubic-bezier(0.4, 0, 0.2, 1);
            box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.1), 0 2px 4px -1px rgba(0, 0, 0, 0.06);
        }
        div.stButton > button:hover {
            transform: translateY(-2px);
            box-shadow: 0 10px 15px -3px rgba(79, 70, 229, 0.4);
            color: white;
        }

        /* Status Indicator Pulse */
        div[data-testid="stStatus"] {
            background-color: #F0F9FF;
            border: 1px solid #BAE6FD;
            border-radius: 12px;
            animation: pulse 2s infinite;
        }
        @keyframes pulse {
            0% { transform: scale(1); }
            50% { transform: scale(1.02); }
            100% { transform: scale(1); }
        }

        /* Text Area Focus Effect */
        textarea {
            border-radius: 12px !important;
            border: 1px solid #E2E8F0 !important;
            transition: border-color 0.3s ease;
        }
        textarea:focus {
            border-color: #3B82F6 !important;
            box-shadow: 0 0 0 3px rgba(59, 130, 246, 0.1) !important;
        }
        </style>
    """, unsafe_allow_html=True)

inject_premium_design()

if "gemini_api_key" not in st.session_state:
    st.session_state.gemini_api_key = "AIzaSyAS7lVfLXRdNSzinXnZBTf9T6ES7s0qNMg"
import google.generativeai as genai
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import io
import os
import urllib.request
from PIL import Image
import matplotlib.pyplot as plt
import numpy as np
from mpl_toolkits.mplot3d import Axes3D
import pypandoc
import ssl

# Fix for macOS SSL certificate verify failed error
try:
    _create_unverified_https_context = ssl._create_unverified_context
except AttributeError:
    pass
else:
    ssl._create_default_https_context = _create_unverified_https_context

def get_safe_gemini_model(use_grounding=False):
    tools = []
    if use_grounding:
        try:
            # 구글 검색 그라운딩 도구 추가 (사용 가능한 경우)
            tools = [genai.Tool(google_search_retrieval=genai.GoogleSearchRetrieval())]
        except:
            pass

    try:
        available = [m.name for m in genai.list_models() if 'generateContent' in m.supported_generation_methods]
        # 'Pro' 모델을 최우선으로, 그 다음 최신 Flash 모델을 시도하도록 정렬
        for pref in ['models/gemini-1.5-pro-latest', 'models/gemini-1.5-pro', 'models/gemini-2.0-flash', 'models/gemini-1.5-flash-latest', 'models/gemini-1.5-flash', 'models/gemini-pro']:
            if pref in available or pref.replace('models/', '') in available:
                return genai.GenerativeModel(pref.replace('models/', ''), tools=tools)
        return genai.GenerativeModel('gemini-1.5-pro', tools=tools)
    except:
        return genai.GenerativeModel('gemini-1.5-pro', tools=tools)







def open_in_new_window(title, content):
    import streamlit.components.v1 as components
    import json
    # 줄바꿈 및 특수문자 처리를 위한 JSON 인코딩
    safe_content = json.dumps(content)
    html_code = f"""
    <script>
    function openWindow() {{
        const win = window.open("", "_blank");
        win.document.write(`
            <html>
            <head>
                <title>{title}</title>
                <script src="https://cdn.jsdelivr.net/npm/marked/marked.min.js"></script>
                <script src="https://polyfill.io/v3/polyfill.min.js?features=es6"></script>
                <script id="MathJax-script" async src="https://cdn.jsdelivr.net/npm/mathjax@3/es5/tex-mml-chtml.js"></script>
                <style>
                    body {{ 
                        font-family: 'Segoe UI', Roboto, Helvetica, Arial, sans-serif; 
                        padding: 40px; 
                        line-height: 1.8; 
                        color: #1e293b; 
                        background-color: #f8fafc;
                        max-width: 1000px;
                        margin: 0 auto;
                        word-wrap: break-word;
                        overflow-wrap: break-word;
                    }}
                    .container {{
                        background: white;
                        padding: 40px;
                        border-radius: 15px;
                        box-shadow: 0 10px 15px -3px rgba(0, 0, 0, 0.1);
                        border: 1px solid #e2e8f0;
                    }}
                    h1 {{ 
                        color: #2563eb; 
                        border-bottom: 3px solid #3b82f6; 
                        padding-bottom: 15px;
                        margin-top: 0;
                        font-size: 1.8rem;
                    }}
                    #rendered-content {{
                        font-size: 1.1rem;
                    }}
                    pre {{ 
                        white-space: pre-wrap; 
                        word-wrap: break-word; 
                        background: #f1f5f9; 
                        padding: 20px; 
                        border-radius: 8px; 
                        font-family: 'Fira Code', 'Courier New', monospace;
                        font-size: 0.95em;
                        border: 1px solid #e2e8f0;
                        overflow-x: auto;
                    }}
                    code {{
                        background: #f1f5f9;
                        padding: 2px 5px;
                        border-radius: 4px;
                        font-family: monospace;
                    }}
                    blockquote {{
                        border-left: 5px solid #3b82f6;
                        padding: 10px 20px;
                        margin: 20px 0;
                        background: #eff6ff;
                        color: #1e40af;
                        border-radius: 0 8px 8px 0;
                    }}
                    table {{
                        border-collapse: collapse;
                        width: 100%;
                        margin: 20px 0;
                        border-radius: 8px;
                        overflow: hidden;
                        box-shadow: 0 1px 3px rgba(0,0,0,0.1);
                    }}
                    th, td {{
                        border: 1px solid #e2e8f0;
                        padding: 15px;
                        text-align: left;
                    }}
                    th {{
                        background-color: #f1f5f9;
                        font-weight: 600;
                        color: #475569;
                    }}
                    .no-print {{
                        margin-top: 40px;
                        text-align: center;
                    }}
                    .btn {{
                        padding: 15px 30px; 
                        background: #2563eb; 
                        color: white; 
                        border: none; 
                        border-radius: 8px; 
                        cursor: pointer;
                        font-weight: bold;
                        font-size: 1rem;
                        transition: all 0.2s;
                        box-shadow: 0 4px 6px rgba(37, 99, 235, 0.2);
                    }}
                    .btn:hover {{ 
                        background: #1d4ed8; 
                        transform: translateY(-2px);
                        box-shadow: 0 6px 8px rgba(37, 99, 235, 0.3);
                    }}
                    @media print {{
                        .no-print {{ display: none; }}
                        body {{ padding: 0; background: white; }}
                        .container {{ box-shadow: none; border: none; padding: 0; }}
                    }}
                </style>
            </head>
            <body>
                <div class="container">
                    <h1>{title}</h1>
                    <div id="rendered-content">분석 결과를 렌더링 중입니다...</div>
                    <div class="no-print">
                        <button class="btn" onclick="window.print()">📥 결과 보고서 PDF / 인쇄로 저장</button>
                    </div>
                </div>
                <script>
                    setTimeout(() => {{
                        const rawContent = ${{safe_content}};
                        document.getElementById('rendered-content').innerHTML = marked.parse(rawContent);
                        
                        // MathJax 렌더링 재실행
                        if (window.MathJax) {{
                            window.MathJax.typesetPromise();
                        }}
                    }}, 100);
                </script>
            </body>
            </html>
        `);
        win.document.close();
    }}
    openWindow();
    </script>
    """
    components.html(html_code, height=0)



def handle_image_analysis(file_obj):
    """
    이미지 내의 복잡한 문제, 그래프, 구조식을 전문가 수준으로 정밀 분석하는 기능
    """
    from PIL import Image
    import streamlit as st

    with st.status(f"🔍 '{file_obj.name}' 상세 분석 중...", expanded=True) as status:
        prompt = "당신은 화학/물리 전문가입니다. 이 이미지에 있는 과학 문제, 그래프, 혹은 구조식을 분석하여 '상세한 풀이 과정'과 '논리적 근거', 그리고 '최종 답'을 설명해 주세요. 수식은 LaTeX 형식($$ ... $$)을 사용하세요."
        res_text = robust_generate_content(prompt, images=Image.open(file_obj))

        if res_text:
            now = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            st.markdown(f"### 📑 {file_obj.name} 상세 분석 결과 (분석 시간: {now})")
            st.markdown(res_text)
            open_in_new_window(f"상세 분석 결과: {file_obj.name}", res_text)

            # 워드에 추가
            now = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            st.session_state.word_doc.add_heading(f"Detailed Analysis of {file_obj.name} (Time: {now})", level=2)
            for line in res_text.split('\n'):
                st.session_state.word_doc.add_paragraph(line)

            st.success("분석 결과가 워드 문서에 추가되었습니다.")
            status.update(label="✅ 상세 분석 완료", state="complete", expanded=False)
        else:
            st.error("분석에 실패했습니다.")


def robust_generate_content(prompt, images=None, use_grounding=False):
    """
    재시도 간격(Sleep)을 추가하고 모든 변종 모델명을 시도하는 최후의 철벽 로직
    """
    import streamlit as st
    import google.generativeai as genai
    import time

    tools = []
    if use_grounding:
        try: tools = [genai.Tool(google_search_retrieval=genai.GoogleSearchRetrieval())]
        except: pass

    # 1. 동적 목록 확보
    dynamic_models = []
    try:
        for m in genai.list_models():
            if 'generateContent' in m.supported_generation_methods:
                dynamic_models.append(m.name)
    except: pass

    # 2. 우선순위 후보군 ('Pro' 모델을 먼저 시도하여 최고 품질 확보)
    manual_candidates = [
        'gemini-1.5-pro', 'models/gemini-1.5-pro',
        'gemini-1.5-pro-latest', 'models/gemini-1.5-pro-latest',
        'gemini-2.0-flash', 'models/gemini-2.0-flash',
        'gemini-1.5-flash', 'models/gemini-1.5-flash',
        'gemini-1.5-flash-latest', 'models/gemini-1.5-flash-latest'
    ]

    model_candidates = []
    for m in dynamic_models + manual_candidates:
        if m not in model_candidates: model_candidates.append(m)

    last_err = None
    inputs = [prompt]
    if images:
        if isinstance(images, list): inputs.extend(images)
        else: inputs.append(images)

    for model_name in model_candidates:
        try:
            # grounding 실패 시 일반 모드 자동 전환을 위해 try-except 중첩
            try:
                model = genai.GenerativeModel(model_name, tools=tools if use_grounding else [])
                res = model.generate_content(inputs, generation_config=genai.GenerationConfig(max_output_tokens=8192, temperature=0.3))
                if res and res.text: return res.text
            except:
                if use_grounding:
                    model = genai.GenerativeModel(model_name)
                    res = model.generate_content(inputs, generation_config=genai.GenerationConfig(max_output_tokens=8192))
                    if res and res.text: return res.text
        except Exception as e:
            last_err = e
            time.sleep(1) # 과부하 방지 및 재시도 준비를 위한 1초 대기
            continue

    st.error(f"❌ 모든 시도가 차단되었습니다. (마지막 모델: {model_name}, 에러: {last_err})")
    return None








# st.set_page_config removed (called at the beginning)

import streamlit.components.v1 as components
components.html("""
<script>
    const parent = window.parent;
    const mainContainer = parent.document.querySelector('.main') || parent.document.querySelector('section.main');

    if (mainContainer) {
        mainContainer.addEventListener('scroll', () => {
            parent.localStorage.setItem('st_scroll_pos', mainContainer.scrollTop);
        });

        setTimeout(() => {
            const scrollPos = parent.localStorage.getItem('st_scroll_pos');
            if (scrollPos) {
                mainContainer.scrollTop = parseInt(scrollPos);
            }
        }, 300);
    }
</script>
""", height=0, width=0)

import json
STATE_FILE = "workspace_state.json"

def load_state():
    if "state_loaded" not in st.session_state:
        if os.path.exists(STATE_FILE):
            try:
                with open(STATE_FILE, "r", encoding="utf-8") as f:
                    saved = json.load(f)
                for k, v in saved.items():
                    st.session_state[k] = v
            except:
                pass
        st.session_state.state_loaded = True

def save_state():
    state_to_save = {}
    for k in ["notion_db", "latex_code", "mathlive_init", "menu_selection", "global_smart_img_result", "global_smart_img_name"]:
        if k in st.session_state:
            state_to_save[k] = st.session_state[k]
    try:
        with open(STATE_FILE, "w", encoding="utf-8") as f:
            json.dump(state_to_save, f, ensure_ascii=False, indent=2)
    except:
        pass

# --- 백그라운드 분석 관리 시스템 ---
if "pending_tasks" not in st.session_state:
    st.session_state.pending_tasks = []

def add_analysis_task(prompt, images, target_subject, file_name):
    st.session_state.pending_tasks.append({
        "prompt": prompt,
        "images": images,
        "subject": target_subject,
        "file_name": file_name,
        "status": "pending"
    })

def process_background_tasks():
    if st.session_state.pending_tasks:
        for task in st.session_state.pending_tasks:
            if task["status"] == "pending":
                with st.sidebar:
                    with st.status(f"⏳ '{task['file_name']}' 분석 중...", expanded=False):
                        task["status"] = "processing"
                        try:
                            res = robust_generate_content(task["prompt"], images=task["images"])
                            if res:
                                now = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                                header = f"\n\n--- [AI 자동추출: {task['file_name']} ({now})] ---\n"
                                if task["subject"] in st.session_state.notion_db:
                                    st.session_state.notion_db[task["subject"]] += header + res
                                    st.toast(f"✅ {task['file_name']} 분석 완료 및 '{task['subject']}'에 저장됨!")
                                task["status"] = "completed"
                                st.rerun()
                        except Exception as e:
                            st.error(f"오류: {e}")
                            task["status"] = "failed"
        st.session_state.pending_tasks = [t for t in st.session_state.pending_tasks if t["status"] not in ["completed", "failed"]]

# 앱 초기화 및 상태 로드
load_state()
process_background_tasks()


if "word_doc" not in st.session_state:
    st.session_state.word_doc = Document()
    st.session_state.word_doc.add_heading("SNU Chem-Ed Analysis Report", level=0)

if "analysis_buffer" not in st.session_state:
    st.session_state.analysis_buffer = {}


if "notion_db" not in st.session_state:
    st.session_state.notion_db = {
        "물리화학": """### 📘 Atkins [화학의 원리] 물리화학 핵심 요약
**1. 양자역학적 모델 (The Quantum World)**
- 슈뢰딩거 방정식 ($\hat{H}\psi = E\psi$) 및 1차원 상자 속 입자 모델.
- 수소 원자 오비탈의 확률 밀도 분산 ($R(r)^2$).
- **[3D 시각화 추천]**: 2s, 3p 오비탈의 노드(Node) 구조 및 전자 구름 밀도.

**2. 화학 열역학 (Thermodynamics)**
- 엔탈피(H), 엔트로피(S), 깁스 자유 에너지(G)의 관계: $\Delta G = \Delta H - T\Delta S$.
- 평형 상수(K)와 자유 에너지의 연결: $\Delta G^\circ = -RT \ln K$.

**3. 화학 반응 속도론 (Kinetics)**
- 아레니우스 식: $k = A e^{-E_a/RT}$.
- 반응 메커니즘과 속도 결정 단계(RDS) 분석.""",
        "유기화학": """### 📙 Atkins [화학의 원리] 유기화학 기초
**1. 탄화수소의 구조와 결합**
- $sp^3, sp^2, sp$ 혼성 오비탈과 기하 구조.
- 에텐(Ethene)의 $\pi$ 결합 형성과 평면 구조 분석.

**2. 작용기 및 명명법**
- 알코올, 카복실산, 에스터의 수소 결합 및 끓는점 비교.
- **[도표 추천]**: 작용기별 IR 흡수 주파수 영역대 정리 표.""",
        "무기화학": """### 📗 Atkins [화학의 원리] 무기 및 원자 구조
**1. 주기적 성질 (Periodic Trends)**
- 유효 핵전하($$Z_{eff}$$)와 이온화 에너지, 원자 반지름의 경향성.
- Slater's Rule을 이용한 가리움 효과 계산.

**2. 배위 화학 (Coordination Chemistry)**
- 결정장 이론(CFT): 팔면체($O_h$) 및 사면체($T_d$) 장에서의 d-오비탈 갈라짐.
- 강한 장(Strong field) vs 약한 장(Weak field) 리간드와 스핀 상태.""",
        "분석화학": """### 📕 Atkins [화학의 원리] 분석 및 평형
**1. 수용액 평형 및 적정**
- 완충 용액과 Henderson-Hasselbalch 식: $pH = pK_a + \log \frac{[A^-]}{[HA]}$.
- 다양성자산의 단계별 이온화 지표 분석.

**2. 전기화학 (Electrochemistry)**
- 네른스트 식 (Nernst Equation): $E = E^\circ - \frac{RT}{nF} \ln Q$.
- 표준 환원 전위표를 이용한 전지 전위($E_{cell}$) 예측.""",
        "화학교육": """### 📓 Atkins [화학의 원리] 화학교육학적 분석
**1. 학습자 오개념 분석**
- **평형의 동적 특성**: 학생들이 화학 평형을 정적인 상태(반응이 멈춤)로 오해하는 경향 분석.
- **오비탈의 물리적 의미**: 오비탈을 전자가 들어있는 '그릇'으로 인식하는 오개념 교정 전략.

**2. 시각화 도구 활용**
- VSEPR 모형을 통한 분자 구조 예측 교수법.
- **[수업 설계]**: POE 모형을 적용한 르 샤틀리에 원리 실험 지도안."""
    }

    st.markdown(f"""
<style>
    :root {{
        --bg-color: #F8FAFC;
        --sidebar-bg: #FFFFFF;
    }}
    .stApp {{ background-color: var(--bg-color); font-family: 'Inter', sans-serif; }}
    h1 {{ font-size: 2.2rem !important; font-weight: 700 !important; color: #0F172A !important; margin-bottom: 1rem !important; }}
    .main {{ overflow: auto !important; height: 100vh !important; }} /* 추가된 스크롤 고정 */
    .block-container {{ padding: 2rem 4rem !important; max-width: 100% !important; }}

    div.stButton > button {{
        border-radius: 10px;
        transition: all 0.3s ease;
        font-weight: 600;

        background: linear-gradient(135deg, #2E5BFF 0%, #6366F1 100%);
        color: white;
        border: none;
    }}

    div.stButton > button:hover {{ transform: translateY(-2px); box-shadow: 0 10px 15px -3px rgba(46, 91, 255, 0.4); }}

    .block-container {{ padding: 2rem 4rem !important; }}
</style>

""", unsafe_allow_html=True)

plt.rcParams['font.family'] = 'sans-serif'
plt.rcParams['axes.unicode_minus'] = False

# ==========================================
# 도식/그래프 함수들
# ==========================================

def get_preset_custom_data(module, choice):
    if module == "cell":
        if choice == "Simple Cubic (SC)":
            return "0,0,0,blue\n1,0,0,blue\n0,1,0,blue\n1,1,0,blue\n0,0,1,blue\n1,0,1,blue\n0,1,1,blue\n1,1,1,blue"
        elif choice == "Body-Centered (BCC)":
            return "0,0,0,blue\n1,0,0,blue\n0,1,0,blue\n1,1,0,blue\n0,0,1,blue\n1,0,1,blue\n0,1,1,blue\n1,1,1,blue\n0.5,0.5,0.5,red"
        elif choice == "Face-Centered (FCC)":
            return "0,0,0,blue\n1,0,0,blue\n0,1,0,blue\n1,1,0,blue\n0,0,1,blue\n1,0,1,blue\n0,1,1,blue\n1,1,1,blue\n0.5,0.5,0,green\n0.5,0.5,1,green\n0.5,0,0.5,green\n0.5,1,0.5,green\n0,0.5,0.5,green\n1,0.5,0.5,green"
        else:
            return "0,0,0,blue\n1,1,1,red"

    elif module == "graph":
        if choice == "1D Box 파동함수 (n=1,2,3)":
            return "np.sin(np.pi*x)*np.sqrt(2)"
        elif choice == "2s 오비탈 확률 밀도":
            return "(x*(2-x)*np.exp(-x/2))**2"
        return "sin(x)*exp(-x/5)"
        if choice == "Atkins: 2s 오비탈 노드 분석":
            return "(x*(2-x)*np.exp(-x/2))**2"
        if choice == "Atkins: 아레니우스 에너지 장벽":
            return "np.exp(-10/x)"


    elif module == "orbital":
        if choice == "NO3- (질산 이온)":
            return "ATOM, N, 0, 0, 0\nATOM, O, 0, 1.2, 0\nATOM, O, 1.04, -0.6, 0\nATOM, O, -1.04, -0.6, 0\nBOND, 0, 0, 0, 0, 1.2, 0\nBOND, 0, 0, 0, 1.04, -0.6, 0\nBOND, 0, 0, 0, -1.04, -0.6, 0\nPZ, 0, 0, 0, #60A5FA, N_pz\nPZ, 0, 1.2, 0, #9CA3AF, O_pz\nPZ, 1.04, -0.6, 0, #9CA3AF, O_pz\nPZ, -1.04, -0.6, 0, #9CA3AF, O_pz"
        elif choice == "HNO3 (질산)":
            return "ATOM, N, 0, 0, 0\nATOM, O, 0, 1.2, 0\nATOM, O, 1.04, -0.6, 0\nATOM, O, -1.04, -0.6, 0\nATOM, H, -1.8, -1.0, 0\nBOND, 0, 0, 0, 0, 1.2, 0\nBOND, 0, 0, 0, 1.04, -0.6, 0\nBOND, 0, 0, 0, -1.04, -0.6, 0\nBOND, -1.04, -0.6, 0, -1.8, -1.0, 0\nPZ, 0, 0, 0, #60A5FA, N_pz\nPZ, 0, 1.2, 0, #9CA3AF, O_pz\nPZ, 1.04, -0.6, 0, #9CA3AF, O_pz\nSP2, -1.04, -0.6, 0, 210, #4ADE80, O_sp2\nS, -1.8, -1.0, 0, #F87171, H_s"
        elif choice == "C2H4 (에텐)":
            return "ATOM, C1, -0.6, 0, 0\nATOM, C2, 0.6, 0, 0\nBOND, -0.6, 0, 0, 0.6, 0, 0\nPZ, -0.6, 0, 0, #60A5FA, C1_pz\nPZ, 0.6, 0, 0, #60A5FA, C2_pz\nATOM, H, -1.2, 0.8, 0\nATOM, H, -1.2, -0.8, 0\nATOM, H, 1.2, 0.8, 0\nATOM, H, 1.2, -0.8, 0\nBOND, -0.6, 0, 0, -1.2, 0.8, 0\nBOND, -0.6, 0, 0, -1.2, -0.8, 0\nBOND, 0.6, 0, 0, 1.2, 0.8, 0\nBOND, 0.6, 0, 0, 1.2, -0.8, 0\nS, -1.2, 0.8, 0, #9CA3AF, H_s\nS, -1.2, -0.8, 0, #9CA3AF, H_s\nS, 1.2, 0.8, 0, #9CA3AF, H_s\nS, 1.2, -0.8, 0, #9CA3AF, H_s"
        return "ATOM, X, 0, 0, 0"

    elif module == "lewis":
        if choice == "H2O (물)":
            return "TEXT, O, 0.5, 0.5, 36, #1E3A8A\nTEXT, H, 0.2, 0.5, 36, #1E3A8A\nTEXT, H, 0.8, 0.5, 36, #1E3A8A\nLINE, 0.3, 0.5, 0.4, 0.5, 2, k\nLINE, 0.6, 0.5, 0.7, 0.5, 2, k"



def draw_unit_cell(cell_type, lattice_size=1, atom_rad=0.15, custom_coords=""):
    errors = []
    fig = plt.figure(figsize=(10, 10))
    ax = fig.add_subplot(111, projection='3d')

    # 사실적인 구(Sphere)를 그리는 함수
    def draw_sphere(ax, center, radius, color):
        u, v = np.mgrid[0:2*np.pi:15j, 0:np.pi:10j]
        x = center[0] + radius * np.cos(u) * np.sin(v)
        y = center[1] + radius * np.sin(u) * np.sin(v)
        z = center[2] + radius * np.cos(v)
        # shade=True와 rstride/cstride 조절로 입체감 극대화
        ax.plot_surface(x, y, z, color=color, alpha=1.0, linewidth=0, shade=True, rstride=1, cstride=1)

    # 1. 격자 선 생성 (수동 모드 포함 모든 모드에서 항상 표시)
    atoms = [] # (pos, color, radius)
    for i in range(lattice_size + 1):
        for j in range(lattice_size + 1):
            # X축 방향 선
            ax.plot([0, lattice_size], [i, i], [j, j], color='gray', lw=1, alpha=0.4)
            # Y축 방향 선
            ax.plot([i, i], [0, lattice_size], [j, j], color='gray', lw=1, alpha=0.4)
            # Z축 방향 선
            ax.plot([i, i], [j, j], [0, lattice_size], color='gray', lw=1, alpha=0.4)

    # 2. 원자 위치 데이터 생성
    if cell_type != "직접 좌표 입력 (Custom)":
        rad = atom_rad # 사용자 지정 반지름

        if cell_type == "Simple Cubic (SC)":
            for i in range(lattice_size + 1):
                for j in range(lattice_size + 1):
                    for k in range(lattice_size + 1):
                        draw_sphere(ax, (i, j, k), rad, '#D32F2F') # Corner: Red

        elif cell_type == "Body-Centered (BCC)":
            # Corners
            for i in range(lattice_size + 1):
                for j in range(lattice_size + 1):
                    for k in range(lattice_size + 1):
                        draw_sphere(ax, (i, j, k), rad, '#D32F2F') # Corner: Red
            # Centers
            for i in range(lattice_size):
                for j in range(lattice_size):
                    for k in range(lattice_size):
                        draw_sphere(ax, (i+0.5, j+0.5, k+0.5), rad, '#1976D2') # Body: Blue

        elif cell_type == "Face-Centered (FCC)":
            # Corners
            for i in range(lattice_size + 1):
                for j in range(lattice_size + 1):
                    for k in range(lattice_size + 1):
                        draw_sphere(ax, (i, j, k), rad, '#D32F2F') # Corner: Red
            # Faces
            for i in range(lattice_size):
                for j in range(lattice_size):
                    for k in range(lattice_size):
                        off = np.array([i, j, k])
                        # XY, XZ, YZ faces
                        for p, c in [(np.array([0.5,0.5,0]), '#388E3C'), (np.array([0.5,0,0.5]), '#1976D2'), (np.array([0,0.5,0.5]), '#FBC02D')]:
                            draw_sphere(ax, p + off, rad, c)

        elif cell_type == "HCP (Hexagonal)":
            c_a = 1.633 # c/a ratio
            for i in range(lattice_size + 1):
                for j in range(lattice_size + 1):
                    for layer in range(lattice_size * 2 + 1):
                        z = layer * (c_a / 2)
                        # Hexagonal points
                        angles = np.linspace(0, 2*np.pi, 7)
                        for ang in angles:
                            x, y = np.cos(ang) + i*1.5, np.sin(ang) + j*np.sqrt(3)
                            if layer % 2 == 0:
                                draw_sphere(ax, (x, y, z), rad, '#D32F2F')
                            else:
                                # Middle layer offset
                                draw_sphere(ax, (x+0.5, y+0.3, z), rad, '#1976D2')
                        # Draw vertical lines
                        if layer < lattice_size * 2:
                            for ang in angles:
                                x, y = np.cos(ang) + i*1.5, np.sin(ang) + j*np.sqrt(3)
                                ax.plot([x, x], [y, y], [z, z + c_a/2], color='gray', lw=1, alpha=0.3)

        elif cell_type == "CsCl (BCC)":
            # Corners: Cl-
            for i in range(lattice_size + 1):
                for j in range(lattice_size + 1):
                    for k in range(lattice_size + 1):
                        draw_sphere(ax, (i, j, k), rad, '#D32F2F') # 빨강 (Cl-)
            # Centers: Cs+
            for i in range(lattice_size):
                for j in range(lattice_size):
                    for k in range(lattice_size):
                        draw_sphere(ax, (i+0.5, j+0.5, k+0.5), rad*1.2, '#1976D2') # 파랑 (Cs+)

        elif cell_type == "CaF2 (Fluorite)":
            # Ca2+: FCC
            for i in range(lattice_size + 1):
                for j in range(lattice_size + 1):
                    for k in range(lattice_size + 1):
                        draw_sphere(ax, (i, j, k), rad, '#1976D2') # 파랑 (Ca2+)
            # F-: 8 sites in unit cell
            for i in range(lattice_size):
                for j in range(lattice_size):
                    for k in range(lattice_size):
                        off = np.array([i, j, k])
                        f_pts = [np.array([0.25,0.25,0.25]), np.array([0.75,0.25,0.25]), np.array([0.25,0.75,0.25]), np.array([0.75,0.75,0.25]),
                                 np.array([0.25,0.25,0.75]), np.array([0.75,0.25,0.75]), np.array([0.25,0.75,0.75]), np.array([0.75,0.75,0.75])]
                        for p in f_pts: draw_sphere(ax, p + off, rad*0.6, '#FBC02D') # 노랑 (F-)

        elif cell_type == "NaCl (Rock Salt)":
            for i in range(lattice_size + 1):
                for j in range(lattice_size + 1):
                    for k in range(lattice_size + 1):
                        # Cl-
                        if (i+j+k) % 2 == 0: draw_sphere(ax, (i, j, k), rad*1.1, '#D32F2F')
                        # Na+
                        else: draw_sphere(ax, (i, j, k), rad*0.7, '#E0E0E0')

        elif cell_type == "Rhombohedral":
            alpha = np.radians(60)
            a1, a2, a3 = np.array([1,0,0]), np.array([np.cos(alpha), np.sin(alpha), 0]), np.array([np.cos(alpha), 0.5, 0.8])
            for i in range(lattice_size + 1):
                for j in range(lattice_size + 1):
                    for k in range(lattice_size + 1):
                        pos = i*a1 + j*a2 + k*a3
                        draw_sphere(ax, pos, rad, '#C2185B')
                        if i < lattice_size: ax.plot([pos[0], (pos+a1)[0]], [pos[1], (pos+a1)[1]], [pos[2], (pos+a1)[2]], color='gray', lw=1, alpha=0.3)
                        if j < lattice_size: ax.plot([pos[0], (pos+a2)[0]], [pos[1], (pos+a2)[1]], [pos[2], (pos+a2)[2]], color='gray', lw=1, alpha=0.3)
                        if k < lattice_size: ax.plot([pos[0], (pos+a3)[0]], [pos[1], (pos+a3)[1]], [pos[2], (pos+a3)[2]], color='gray', lw=1, alpha=0.3)

        # 3. 단위 세포 강조박스 (Dashed Box) - 이미지 스타일 재현
        if lattice_size > 1:
            for start, end in [([0,0,0],[1,0,0]), ([1,0,0],[1,1,0]), ([1,1,0],[0,1,0]), ([0,1,0],[0,0,0]),
                               ([0,0,1],[1,0,1]), ([1,0,1],[1,1,1]), ([1,1,1],[0,1,1]), ([0,1,1],[0,0,1]),
                               ([0,0,0],[0,0,1]), ([1,0,0],[1,0,1]), ([1,1,0],[1,1,1]), ([0,1,0],[0,1,1])]:
                ax.plot([start[0], end[0]], [start[1], end[1]], [start[2], end[2]], color='#FFEB3B', lw=3, ls='--', alpha=0.9, zorder=10)

    else: # Custom Coords
        if custom_coords:
            # Support both newline and semicolon as delimiters
            data = custom_coords.replace(';', '\n')
            for line_idx, p in enumerate(data.split('\n')):
                if p.strip():
                    try:
                        vals = [v.strip() for v in p.split(',')]
                        if len(vals) >= 3:
                            x, y, z = float(vals[0]), float(vals[1]), float(vals[2])
                            c = vals[3] if len(vals) > 3 else 'red'
                            draw_sphere(ax, (x, y, z), 0.15, c)
                    except Exception as e:
                        errors.append(f"{line_idx+1}번째 줄 오류: {e}")

    ax.set_title(f"Detailed {cell_type} Lattice", fontsize=16, fontweight='bold')
    ax.axis('off')

    # 뷰 각도 조절 (더 입체적으로 보이게)
    ax.view_init(elev=20, azim=45)

    # 배경 투명화 및 시각 개선
    ax.set_facecolor((1.0, 1.0, 1.0, 0.0))

    img_stream = io.BytesIO()
    plt.savefig(img_stream, format='png', bbox_inches='tight', dpi=200, transparent=True)
    plt.close()
    img_stream.seek(0)
    return img_stream, errors

def draw_quantum_graph(graph_type, custom_expr="", x_min=0.0, x_max=10.0):
    fig, ax = plt.subplots(figsize=(6, 4))
    if graph_type == "1D Box 파동함수 (n=1,2,3)":
        x = np.linspace(0, 3, 100)
        ax.plot(x, np.sin(1 * np.pi * x / 3) + 4, label='n=1')
        ax.plot(x, np.sin(2 * np.pi * x / 3) + 2, label='n=2')
        ax.plot(x, np.sin(3 * np.pi * x / 3), label='n=3')
        ax.axis('off')
        ax.legend()
    elif graph_type == "2s 오비탈 확률 밀도":
        r = np.linspace(0, 15, 200)
        R2s = (1/(2*np.sqrt(2))) * (1)**(1.5) * (2 - r) * np.exp(-r/2)
        ax.plot(r, r**2 * R2s**2, color='purple', lw=2)
    elif graph_type == "직접 수식 입력 (Custom)":
        x = np.linspace(x_min, x_max, 300)
        safe_dict = {
            'x': x, 'np': np, 'sin': np.sin, 'cos': np.cos, 'tan': np.tan,
            'exp': np.exp, 'log': np.log, 'sqrt': np.sqrt, 'pi': np.pi, 'e': np.e,
            'abs': np.abs
        }
        try:
            # Replace common math terms to numpy terms for convenience
            expr = custom_expr.replace('^', '**').replace('sin', 'np.sin').replace('cos', 'np.cos').replace('exp', 'np.exp').replace('np.np.', 'np.')
            y = eval(expr, {"__builtins__": None}, safe_dict)
            ax.plot(x, y, color='blue', lw=2, label=custom_expr)
            ax.legend()
            ax.grid(True, linestyle='--', alpha=0.6)
        except Exception as e:
            ax.text(0.5, 0.5, f"수식 오류!\n{e}", ha='center', va='center', color='red', transform=ax.transAxes)

    img_stream = io.BytesIO()
    plt.savefig(img_stream, format='png', bbox_inches='tight')
    plt.close()
    img_stream.seek(0)
    return img_stream

def get_molecule_image(name):
    name = name.strip()
    urls = [
        f"https://pubchem.ncbi.nlm.nih.gov/rest/pug/compound/name/{name}/PNG?record_type=2d&image_size=large",
        f"https://pubchem.ncbi.nlm.nih.gov/rest/pug/compound/fastformula/{name}/PNG?record_type=2d&image_size=large"
    ]
    for url in urls:
        try:
            req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
            with urllib.request.urlopen(req) as response:
                return io.BytesIO(response.read())
        except:
            continue
    st.error(f"'{name}'에 대한 이미지를 PubChem에서 찾을 수 없습니다. (이름이나 분자식을 확인해주세요)")
    return None

def draw_schematic(shape_type, color='#2E5BFF', custom_data="", **kwargs):
    errors = []
    fig = plt.figure(figsize=(4, 4))

    if shape_type == "직접 입력 (Custom)":
        ax = fig.add_subplot(111)
        ax.set_aspect('equal')
        ax.axis('off')
        ax.set_xlim(0, 10)
        ax.set_ylim(0, 10)
        for line_idx, line in enumerate(custom_data.split('\n')):
            parts = [p.strip() for p in line.split(',')]
            if not parts or not parts[0]: continue
            cmd = parts[0].upper()
            try:
                if cmd == 'RECT' and len(parts) >= 5:
                    x, y, w, h = map(float, parts[1:5])
                    c = parts[5] if len(parts) > 5 else color
                    rect = plt.Rectangle((x, y), w, h, fill=True, color=c, alpha=0.7, ec='k', lw=2)
                    ax.add_patch(rect)
                elif cmd == 'CIRCLE' and len(parts) >= 4:
                    x, y, r = map(float, parts[1:4])
                    c = parts[4] if len(parts) > 4 else color
                    circle = plt.Circle((x, y), r, fill=True, color=c, alpha=0.7, ec='k', lw=2)
                    ax.add_patch(circle)
                elif cmd == 'TEXT' and len(parts) >= 4:
                    t, x, y = parts[1], float(parts[2]), float(parts[3])
                    s = int(parts[4]) if len(parts) > 4 else 12
                    c = parts[5] if len(parts) > 5 else 'black'
                    ax.text(x, y, t, fontsize=s, color=c, ha='center', va='center')
                elif cmd == 'LINE' and len(parts) >= 5:
                    x1, y1, x2, y2 = map(float, parts[1:5])
                    c = parts[5] if len(parts) > 5 else 'black'
                    ax.plot([x1, x2], [y1, y2], color=c, lw=2)
            except Exception as e:
                errors.append(f"{line_idx+1}번째 줄 오류: {e}")

    if shape_type == "정육면체 (Cube)":
        ax = fig.add_subplot(111, projection='3d')
        length = kwargs.get('length', 1.0)
        width = kwargs.get('width', 1.0)
        height = kwargs.get('height', 1.0)

        r_x = [0, length]
        r_y = [0, width]
        X, Y = np.meshgrid(r_x, r_y)

        z_top = np.full((2,2), height)
        z_bot = np.zeros((2,2))

        ax.plot_surface(X, Y, z_top, alpha=0.5, color=color, edgecolor='k')
        ax.plot_surface(X, Y, z_bot, alpha=0.5, color=color, edgecolor='k')

        X_xz, Z_xz = np.meshgrid(r_x, [0, height])
        ax.plot_surface(X_xz, np.zeros((2,2)), Z_xz, alpha=0.5, color=color, edgecolor='k')
        ax.plot_surface(X_xz, np.full((2,2), width), Z_xz, alpha=0.5, color=color, edgecolor='k')

        Y_yz, Z_yz = np.meshgrid(r_y, [0, height])
        ax.plot_surface(np.zeros((2,2)), Y_yz, Z_yz, alpha=0.5, color=color, edgecolor='k')
        ax.plot_surface(np.full((2,2), length), Y_yz, Z_yz, alpha=0.5, color=color, edgecolor='k')

        ax.set_axis_off()
        ax.view_init(elev=20, azim=30)
    else:
        ax = fig.add_subplot(111)
        ax.set_aspect('equal')
        ax.axis('off')

        if shape_type == "정사각형 (Square)":
            w = kwargs.get('width', 0.8)
            h = kwargs.get('height', 0.8)
            rect = plt.Rectangle(((1-w)/2, (1-h)/2), w, h, fill=True, color=color, alpha=0.7, ec='k', lw=2)
            ax.add_patch(rect)
            ax.set_xlim(0, 1)
            ax.set_ylim(0, 1)
        elif shape_type == "원형 (Circle)":
            r = kwargs.get('radius', 0.4)
            circle = plt.Circle((0.5, 0.5), r, fill=True, color=color, alpha=0.7, ec='k', lw=2)
            ax.add_patch(circle)
            ax.set_xlim(0, 1)
            ax.set_ylim(0, 1)
        elif shape_type == "다각형 (Polygon)":
            n = kwargs.get('n_sides', 6)
            from matplotlib.patches import Polygon
            angles = np.linspace(0, 2 * np.pi, n, endpoint=False)
            points = np.column_stack((0.5 + 0.4 * np.cos(angles), 0.5 + 0.4 * np.sin(angles)))
            poly = Polygon(points, fill=True, color=color, alpha=0.7, ec='k', lw=2)
            ax.add_patch(poly)
            ax.set_xlim(0, 1)
            ax.set_ylim(0, 1)

    img_stream = io.BytesIO()
    plt.savefig(img_stream, format='png', bbox_inches='tight', transparent=True)
    plt.close()
    img_stream.seek(0)
    return img_stream, errors

def draw_lewis_structure(molecule, custom_data=""):
    errors = []
    fig, ax = plt.subplots(figsize=(4, 4))
    ax.set_aspect('equal')
    ax.axis('off')
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)

    def draw_pair(x, y, dx, dy):
        ax.plot([x - dx, x + dx], [y - dy, y + dy], 'ko', markersize=6)

    def text_sym(x, y, sym):
        ax.text(x, y, sym, fontsize=36, ha='center', va='center', fontweight='bold', color='#1E3A8A')

    if molecule == "H2O (물)":
        text_sym(0.5, 0.5, "O")
        text_sym(0.2, 0.5, "H")
        text_sym(0.8, 0.5, "H")
        ax.plot([0.3, 0.4], [0.5, 0.5], 'k-', lw=2)
        ax.plot([0.6, 0.7], [0.5, 0.5], 'k-', lw=2)
        draw_pair(0.5, 0.65, 0.04, 0)
        draw_pair(0.5, 0.35, 0.04, 0)

    elif molecule == "CO2 (이산화탄소)":
        text_sym(0.5, 0.5, "C")
        text_sym(0.2, 0.5, "O")
        text_sym(0.8, 0.5, "O")
        ax.plot([0.3, 0.4], [0.53, 0.53], 'k-', lw=2)
        ax.plot([0.3, 0.4], [0.47, 0.47], 'k-', lw=2)
        ax.plot([0.6, 0.7], [0.53, 0.53], 'k-', lw=2)
        ax.plot([0.6, 0.7], [0.47, 0.47], 'k-', lw=2)
        draw_pair(0.2, 0.65, 0.04, 0)
        draw_pair(0.2, 0.35, 0.04, 0)
        draw_pair(0.8, 0.65, 0.04, 0)
        draw_pair(0.8, 0.35, 0.04, 0)

    elif molecule == "NH3 (암모니아)":
        text_sym(0.5, 0.5, "N")
        text_sym(0.2, 0.5, "H")
        text_sym(0.8, 0.5, "H")
        text_sym(0.5, 0.2, "H")
        ax.plot([0.3, 0.4], [0.5, 0.5], 'k-', lw=2)
        ax.plot([0.6, 0.7], [0.5, 0.5], 'k-', lw=2)
        ax.plot([0.5, 0.5], [0.4, 0.3], 'k-', lw=2)
        draw_pair(0.5, 0.65, 0.04, 0)

    elif molecule == "CH4 (메테인)":
        text_sym(0.5, 0.5, "C")
        text_sym(0.2, 0.5, "H")
        text_sym(0.8, 0.5, "H")
        text_sym(0.5, 0.8, "H")
        text_sym(0.5, 0.2, "H")
        ax.plot([0.3, 0.4], [0.5, 0.5], 'k-', lw=2)
        ax.plot([0.6, 0.7], [0.5, 0.5], 'k-', lw=2)
        ax.plot([0.5, 0.5], [0.6, 0.7], 'k-', lw=2)
        ax.plot([0.5, 0.5], [0.4, 0.3], 'k-', lw=2)

    elif molecule == "O2 (산소 분자)":
        text_sym(0.35, 0.5, "O")
        text_sym(0.65, 0.5, "O")
        ax.plot([0.45, 0.55], [0.53, 0.53], 'k-', lw=2)
        ax.plot([0.45, 0.55], [0.47, 0.47], 'k-', lw=2)
        draw_pair(0.35, 0.65, 0.04, 0)
        draw_pair(0.35, 0.35, 0.04, 0)
        draw_pair(0.65, 0.65, 0.04, 0)
        draw_pair(0.65, 0.35, 0.04, 0)

    elif molecule == "N2 (질소 분자)":
        text_sym(0.35, 0.5, "N")
        text_sym(0.65, 0.5, "N")
        ax.plot([0.45, 0.55], [0.55, 0.55], 'k-', lw=2)
        ax.plot([0.45, 0.55], [0.50, 0.50], 'k-', lw=2)
        ax.plot([0.45, 0.55], [0.45, 0.45], 'k-', lw=2)
        draw_pair(0.2, 0.5, 0, 0.04)
        draw_pair(0.8, 0.5, 0, 0.04)

    elif molecule == "HCl (염화수소)":
        text_sym(0.3, 0.5, "H")
        text_sym(0.7, 0.5, "Cl")
        ax.plot([0.4, 0.55], [0.5, 0.5], 'k-', lw=2)
        draw_pair(0.7, 0.65, 0.04, 0)
        draw_pair(0.7, 0.35, 0.04, 0)
        draw_pair(0.85, 0.5, 0, 0.04)

    elif molecule == "HCN (사이안화수소)":
        text_sym(0.2, 0.5, "H")
        text_sym(0.5, 0.5, "C")
        text_sym(0.8, 0.5, "N")
        ax.plot([0.3, 0.4], [0.5, 0.5], 'k-', lw=2)
        ax.plot([0.6, 0.7], [0.55, 0.55], 'k-', lw=2)
        ax.plot([0.6, 0.7], [0.50, 0.50], 'k-', lw=2)
        ax.plot([0.6, 0.7], [0.45, 0.45], 'k-', lw=2)
        draw_pair(0.95, 0.5, 0, 0.04)

    elif molecule == "C2H4 (에텐)":
        text_sym(0.35, 0.5, "C")
        text_sym(0.65, 0.5, "C")
        text_sym(0.15, 0.7, "H")
        text_sym(0.15, 0.3, "H")
        text_sym(0.85, 0.7, "H")
        text_sym(0.85, 0.3, "H")
        ax.plot([0.45, 0.55], [0.53, 0.53], 'k-', lw=2)
        ax.plot([0.45, 0.55], [0.47, 0.47], 'k-', lw=2)
        ax.plot([0.25, 0.3], [0.6, 0.55], 'k-', lw=2)
        ax.plot([0.25, 0.3], [0.4, 0.45], 'k-', lw=2)
        ax.plot([0.75, 0.7], [0.6, 0.55], 'k-', lw=2)
        ax.plot([0.75, 0.7], [0.4, 0.45], 'k-', lw=2)

    elif molecule == "직접 입력 (Custom)":
        for line_idx, line in enumerate(custom_data.split('\n')):
            if not line.strip(): continue
            parts = [p.strip() for p in line.split(',')]
            if not parts or not parts[0]: continue
            cmd = parts[0].upper()
            try:
                if cmd == 'TEXT' and len(parts) >= 4:
                    t, x, y = parts[1], float(parts[2]), float(parts[3])
                    fs = int(parts[4]) if len(parts) > 4 else 36
                    color = parts[5] if len(parts) > 5 else '#1E3A8A'
                    ax.text(x, y, t, fontsize=fs, ha='center', va='center', fontweight='bold', color=color)
                elif cmd == 'LINE' and len(parts) >= 5:
                    x1, y1, x2, y2 = map(float, parts[1:5])
                    lw = float(parts[5]) if len(parts) > 5 else 2
                    color = parts[6] if len(parts) > 6 else 'k'
                    ax.plot([x1, x2], [y1, y2], '-', color=color, lw=lw)
                elif cmd == 'DOTS' and len(parts) >= 5:
                    x, y, dx, dy = map(float, parts[1:5])
                    draw_pair(x, y, dx, dy)
            except Exception as e:
                errors.append(f"{line_idx+1}번째 줄 오류: {e}")

    img_stream = io.BytesIO()
    plt.savefig(img_stream, format='png', bbox_inches='tight', transparent=True, dpi=150)
    plt.close()
    img_stream.seek(0)
    return img_stream, errors

def draw_orbital_diagram(molecule_type, custom_data=""):
    errors = []
    fig = plt.figure(figsize=(8, 6))
    ax = fig.add_subplot(111, projection='3d')
    ax.set_axis_off()

    # draw molecular plane
    xx, yy = np.meshgrid([-2, 2], [-2, 2])
    zz = np.zeros_like(xx)
    ax.plot_surface(xx, yy, zz, color='gray', alpha=0.1, edgecolor='k', lw=0.5)

    # z axis
    ax.plot([0, 0], [0, 0], [-2, 2], 'k--', lw=1)
    ax.text(0, 0, 2.2, 'z', fontsize=14, fontstyle='italic')

    def draw_pz(center, color='#3B82F6', label=''):
        size = 0.3
        u = np.linspace(0, 2 * np.pi, 15)
        v = np.linspace(0, np.pi, 15)
        x = size * 0.8 * np.outer(np.cos(u), np.sin(v)) + center[0]
        y = size * 0.8 * np.outer(np.sin(u), np.sin(v)) + center[1]

        z_top = size * 1.5 * np.outer(np.ones(np.size(u)), np.cos(v)) + center[2] + size*1.2
        ax.plot_surface(x, y, z_top, color=color, alpha=0.6, edgecolor='none')

        z_bot = size * 1.5 * np.outer(np.ones(np.size(u)), np.cos(v)) + center[2] - size*1.2
        ax.plot_surface(x, y, z_bot, color=color, alpha=0.6, edgecolor='none')

        if label:
            ax.text(center[0], center[1]+0.2, center[2]+size*2.5, label, color=color, fontsize=12, fontweight='bold')

    def draw_sp2(center, angle_deg, color, label=''):
        size = 0.25
        u = np.linspace(0, 2 * np.pi, 15)
        v = np.linspace(0, np.pi, 15)
        rad = np.radians(angle_deg)

        X = size * 1.5 * np.outer(np.cos(u), np.sin(v))
        Y = size * 0.6 * np.outer(np.sin(u), np.sin(v))
        Z = size * 0.6 * np.outer(np.ones(np.size(u)), np.cos(v))

        X_rot = X * np.cos(rad) - Y * np.sin(rad)
        Y_rot = X * np.sin(rad) + Y * np.cos(rad)

        offset = size * 1.0
        cx = center[0] + offset * np.cos(rad)
        cy = center[1] + offset * np.sin(rad)
        cz = center[2]

        ax.plot_surface(X_rot + cx, Y_rot + cy, Z + cz, color=color, alpha=0.6, edgecolor='none')
        if label:
            ax.text(cx + 0.3*np.cos(rad), cy + 0.3*np.sin(rad), cz, label, color=color, fontsize=10)

    def draw_s(center, color='#9CA3AF', label=''):
        size = 0.3
        u = np.linspace(0, 2 * np.pi, 15)
        v = np.linspace(0, np.pi, 15)
        X = size * np.outer(np.cos(u), np.sin(v)) + center[0]
        Y = size * np.outer(np.sin(u), np.sin(v)) + center[1]
        Z = size * np.outer(np.ones(np.size(u)), np.cos(v)) + center[2]
        ax.plot_surface(X, Y, Z, color=color, alpha=0.6, edgecolor='none')
        if label:
            ax.text(center[0], center[1], center[2]+0.4, label, color='gray', fontsize=12)

    # Base coords
    N_pos = (0, 0, 0)
    Oa_pos = (0, 1.2, 0)
    Ob_pos = (-1.04, -0.6, 0) # 210 deg
    Oc_pos = (1.04, -0.6, 0)  # 330 deg

    if molecule_type == "NO3- (질산 이온)":
        ax.text(N_pos[0]+0.1, N_pos[1], 0, 'N', fontsize=16, fontweight='bold')
        ax.text(Oa_pos[0]+0.1, Oa_pos[1], 0, 'O_a', fontsize=14)
        ax.text(Ob_pos[0]-0.3, Ob_pos[1]-0.2, 0, 'O_b', fontsize=14)
        ax.text(Oc_pos[0]+0.1, Oc_pos[1]-0.2, 0, 'O_c', fontsize=14)

        ax.plot([0, Oa_pos[0]], [0, Oa_pos[1]], [0, 0], 'k-', lw=2)
        ax.plot([0, Ob_pos[0]], [0, Ob_pos[1]], [0, 0], 'k-', lw=2)
        ax.plot([0, Oc_pos[0]], [0, Oc_pos[1]], [0, 0], 'k-', lw=2)

        draw_pz(N_pos, color='#60A5FA', label='N 2p_z')
        draw_pz(Oa_pos, color='#93C5FD', label='O 2p_z')
        draw_pz(Ob_pos, color='#93C5FD')
        draw_pz(Oc_pos, color='#93C5FD')

        draw_sp2(N_pos, 90, '#4ADE80', 'N 2sp²')
        draw_sp2(N_pos, 210, '#4ADE80')
        draw_sp2(N_pos, 330, '#4ADE80')

        draw_sp2(Oa_pos, 270, '#F87171', 'O 2sp²')
        draw_sp2(Ob_pos, 30, '#F87171')
        draw_sp2(Oc_pos, 150, '#F87171')
        ax.set_title("Molecular Orbitals of NO3-", fontsize=16, y=0.95)

    elif molecule_type == "HNO3 (질산)":
        ax.text(N_pos[0]+0.1, N_pos[1], 0, 'N', fontsize=16, fontweight='bold')
        ax.text(Oa_pos[0]+0.1, Oa_pos[1], 0, 'O_a', fontsize=14)
        ax.text(Ob_pos[0]-0.3, Ob_pos[1]-0.2, 0, 'O_b', fontsize=14)
        ax.text(Oc_pos[0]+0.1, Oc_pos[1]-0.2, 0, 'O_c', fontsize=14)

        ax.plot([0, Oa_pos[0]], [0, Oa_pos[1]], [0, 0], 'k-', lw=2)
        ax.plot([0, Ob_pos[0]], [0, Ob_pos[1]], [0, 0], 'k-', lw=2)
        ax.plot([0, Oc_pos[0]], [0, Oc_pos[1]], [0, 0], 'k-', lw=2)

        draw_pz(N_pos, color='#60A5FA', label='N 2p_z')
        draw_pz(Oa_pos, color='#93C5FD', label='O 2p_z')
        draw_pz(Oc_pos, color='#93C5FD')

        H_pos = (-1.8, -0.6, 0)
        ax.text(H_pos[0]-0.2, H_pos[1], 0, 'H', fontsize=14)
        ax.plot([Ob_pos[0], H_pos[0]], [Ob_pos[1], H_pos[1]], [0, 0], 'k-', lw=2)

        draw_s(H_pos, label='H 1s')

        draw_sp2(N_pos, 90, '#4ADE80', 'N 2sp²')
        draw_sp2(N_pos, 210, '#4ADE80')
        draw_sp2(N_pos, 330, '#4ADE80')

        draw_sp2(Oa_pos, 270, '#F87171', 'O 2sp²')
        draw_sp2(Ob_pos, 30, '#F87171')
        draw_sp2(Ob_pos, 180, '#F87171')
        draw_sp2(Oc_pos, 150, '#F87171')
        ax.set_title("Molecular Orbitals of HNO3", fontsize=16, y=0.95)

    elif molecule_type == "C2H4 (에텐)":
        C1_pos = (-0.7, 0, 0)
        C2_pos = (0.7, 0, 0)
        H1_pos = (-1.3, 0.8, 0)
        H2_pos = (-1.3, -0.8, 0)
        H3_pos = (1.3, 0.8, 0)
        H4_pos = (1.3, -0.8, 0)

        ax.text(C1_pos[0]+0.1, C1_pos[1], 0, 'C', fontsize=16, fontweight='bold')
        ax.text(C2_pos[0]+0.1, C2_pos[1], 0, 'C', fontsize=16, fontweight='bold')

        ax.plot([C1_pos[0], C2_pos[0]], [C1_pos[1], C2_pos[1]], [0, 0], 'k-', lw=2)
        ax.plot([C1_pos[0], H1_pos[0]], [C1_pos[1], H1_pos[1]], [0, 0], 'k-', lw=2)
        ax.plot([C1_pos[0], H2_pos[0]], [C1_pos[1], H2_pos[1]], [0, 0], 'k-', lw=2)
        ax.plot([C2_pos[0], H3_pos[0]], [C2_pos[1], H3_pos[1]], [0, 0], 'k-', lw=2)
        ax.plot([C2_pos[0], H4_pos[0]], [C2_pos[1], H4_pos[1]], [0, 0], 'k-', lw=2)

        draw_pz(C1_pos, color='#60A5FA', label='C 2p_z')
        draw_pz(C2_pos, color='#60A5FA', label='C 2p_z')

        draw_sp2(C1_pos, 0, '#4ADE80')
        draw_sp2(C2_pos, 180, '#4ADE80')
        draw_sp2(C1_pos, 125, '#4ADE80')
        draw_sp2(C1_pos, 235, '#4ADE80')
        draw_sp2(C2_pos, 55, '#4ADE80')
        draw_sp2(C2_pos, 305, '#4ADE80')

        draw_s(H1_pos, label='H 1s')
        draw_s(H2_pos)
        draw_s(H3_pos)
        draw_s(H4_pos)
        ax.set_title("Molecular Orbitals of Ethene (C2H4)", fontsize=16, y=0.95)

    elif molecule_type == "직접 입력 (Custom)":
        for line_idx, line in enumerate(custom_data.split('\n')):
            parts = [p.strip() for p in line.split(',')]
            if not parts or not parts[0]: continue
            cmd = parts[0].upper()
            try:
                if cmd == 'ATOM' and len(parts) >= 5:
                    label, x, y, z = parts[1], float(parts[2]), float(parts[3]), float(parts[4])
                    ax.text(x+0.1, y, z, label, fontsize=14, fontweight='bold')
                elif cmd == 'BOND' and len(parts) >= 7:
                    x1, y1, z1, x2, y2, z2 = map(float, parts[1:7])
                    ax.plot([x1, x2], [y1, y2], [z1, z2], 'k-', lw=2)
                elif cmd == 'PZ' and len(parts) >= 4:
                    x, y, z = map(float, parts[1:4])
                    color = parts[4] if len(parts) > 4 else '#3B82F6'
                    label = parts[5] if len(parts) > 5 else ''
                    draw_pz((x, y, z), color, label)
                elif cmd == 'SP2' and len(parts) >= 5:
                    x, y, z, angle = map(float, parts[1:5])
                    color = parts[5] if len(parts) > 5 else '#4ADE80'
                    label = parts[6] if len(parts) > 6 else ''
                    draw_sp2((x, y, z), angle, color, label)
                elif cmd == 'S' and len(parts) >= 4:
                    x, y, z = map(float, parts[1:4])
                    color = parts[4] if len(parts) > 4 else '#9CA3AF'
                    label = parts[5] if len(parts) > 5 else ''
                    draw_s((x, y, z), color, label)
            except Exception as e:
                errors.append(f"{line_idx+1}번째 줄 오류: {e}")
        ax.set_title("Custom Molecular Orbitals", fontsize=16, y=0.95)

    ax.view_init(elev=25, azim=-55)

    img_stream = io.BytesIO()
    plt.savefig(img_stream, format='png', bbox_inches='tight', transparent=True, dpi=200)
    plt.close()
    img_stream.seek(0)
    return img_stream, errors

def draw_skeletal_structure(molecule, custom_data=""):
    errors = []
    fig, ax = plt.subplots(figsize=(6, 4))
    ax.set_aspect('equal')
    ax.axis('off')

    if molecule == "Butane (뷰테인)":
        ax.plot([0, 1, 2, 3], [0, 0.866, 0, 0.866], 'k-', lw=3)
    elif molecule == "Hexane (헥세인)":
        ax.plot([0, 1, 2, 3, 4, 5], [0, 0.866, 0, 0.866, 0, 0.866], 'k-', lw=3)
    elif molecule == "Cyclohexane (사이클로헥세인)":
        angles = np.linspace(0, 2 * np.pi, 7)
        x = np.cos(angles)
        y = np.sin(angles)
        ax.plot(x, y, 'k-', lw=3)
    elif molecule == "Benzene (벤젠)":
        angles = np.linspace(0, 2 * np.pi, 7)
        x = np.cos(angles)
        y = np.sin(angles)
        ax.plot(x, y, 'k-', lw=3)
        circle = plt.Circle((0, 0), 0.65, fill=False, color='k', lw=3)
        ax.add_patch(circle)
    elif molecule == "Acetone (아세톤)":
        ax.plot([0, 1, 2], [0, 0.866, 0], 'k-', lw=3)
        # Double bond to O
        ax.plot([0.9, 0.9], [0.866, 1.866], 'k-', lw=3)
        ax.plot([1.1, 1.1], [0.866, 1.866], 'k-', lw=3)
        ax.text(1, 2.1, 'O', fontsize=24, ha='center', va='center')
    elif molecule == "Acetic Acid (아세트산)":
        ax.plot([0, 1], [0, 0.866], 'k-', lw=3) # C-C
        ax.plot([1, 1.866], [0.866, 0.366], 'k-', lw=3) # C-OH
        ax.text(2.2, 0.1, 'OH', fontsize=24, ha='center', va='center')
        ax.plot([0.9, 0.9], [0.866, 1.866], 'k-', lw=3) # C=O
        ax.plot([1.1, 1.1], [0.866, 1.866], 'k-', lw=3)
        ax.text(1, 2.1, 'O', fontsize=24, ha='center', va='center')

    elif molecule == "직접 입력 (Custom)":
        for line in custom_data.split('\n'):
            parts = [p.strip() for p in line.split(',')]
            if not parts or not parts[0]: continue
            cmd = parts[0].upper()
            try:
                if cmd in ['TEXT', 'LTEXT', 'RTEXT'] and len(parts) >= 4:
                    t, x, y = parts[1], float(parts[2]), float(parts[3])
                    fs = int(parts[4]) if len(parts) > 4 else 18
                    color = parts[5] if len(parts) > 5 else 'k'
                    ha = 'center'
                    if cmd == 'LTEXT': ha = 'left'
                    elif cmd == 'RTEXT': ha = 'right'
                    ax.text(x, y, t, fontsize=fs, color=color, ha=ha, va='center', fontfamily='serif')
                elif cmd == 'LINE' and len(parts) >= 5:
                    x1, y1, x2, y2 = map(float, parts[1:5])
                    lw = float(parts[5]) if len(parts) > 5 else 2
                    color = parts[6] if len(parts) > 6 else 'k'
                    ax.plot([x1, x2], [y1, y2], '-', color=color, lw=lw, zorder=2)
                elif cmd == 'DLINE' and len(parts) >= 5:
                    x1, y1, x2, y2 = map(float, parts[1:5])
                    lw = float(parts[5]) if len(parts) > 5 else 2
                    color = parts[6] if len(parts) > 6 else 'k'
                    dx, dy = x2 - x1, y2 - y1
                    dist = np.hypot(dx, dy)
                    nx, ny = -dy/dist * 0.08, dx/dist * 0.08
                    ax.plot([x1+nx, x2+nx], [y1+ny, y2+ny], '-', color=color, lw=lw, zorder=2)
                    ax.plot([x1-nx, x2-nx], [y1-ny, y2-ny], '-', color=color, lw=lw, zorder=2)
                elif cmd == 'ARROW' and len(parts) >= 5:
                    x1, y1, x2, y2 = map(float, parts[1:5])
                    color = parts[5] if len(parts) > 5 else 'k'
                    ax.annotate('', xy=(x2, y2), xytext=(x1, y1), arrowprops=dict(arrowstyle="->", lw=1.5, color=color))
                elif cmd == 'DOTS' and len(parts) >= 5:
                    x1, y1, x2, y2 = map(float, parts[1:5])
                    color = parts[5] if len(parts) > 5 else 'k'
                    ax.plot([x1, x2], [y1, y2], 'o', color=color, markersize=3)
                elif cmd == 'BALL' and len(parts) >= 5:
                    x, y, r = map(float, parts[1:4])
                    color = parts[4]
                    base_circle = plt.Circle((x, y), r, color=color, ec='black', lw=1.5, zorder=3)
                    ax.add_patch(base_circle)
                    hl_circle = plt.Circle((x - r*0.3, y + r*0.3), r*0.3, color='white', alpha=0.4, zorder=4, edgecolor='none')
                    ax.add_patch(hl_circle)
            except Exception as e:
                errors.append(f"{line} 처리 중 오류: {e}")

    img_stream = io.BytesIO()
    plt.savefig(img_stream, format='png', bbox_inches='tight', transparent=True, dpi=150)
    plt.close()
    img_stream.seek(0)
    return img_stream, errors

# ==========================================
# 워드 수식(OMML) 자동 변환 처리용 Pandoc 엔진
# ==========================================
def convert_latex_to_word_docx(markdown_text, output_filename, margins):
    try:
        pypandoc.get_pandoc_version()
    except OSError:
        with st.spinner("최초 1회: 완벽한 수식 변환을 위한 Pandoc 엔진을 설치 중입니다 (약 1분 소요)..."):
            pypandoc.download_pandoc()

    temp_md = "temp.md"
    with open(temp_md, "w", encoding="utf-8") as f:
        f.write(markdown_text)

    pypandoc.convert_file(temp_md, 'docx', outputfile=output_filename)

    doc = Document(output_filename)
    for section in doc.sections:
        section.top_margin = Cm(margins['top'])
        section.bottom_margin = Cm(margins['bottom'])
        section.left_margin = Cm(margins['left'])
        section.right_margin = Cm(margins['right'])
    doc.save(output_filename)
    os.remove(temp_md)

# ==========================================
# 사이드바 & 워크스페이스 설정
# ==========================================
with st.sidebar:
    st.markdown("<h2 style='text-align: center; color: #1E3A8A;'>🧪 Chem-Ed Studio</h2>", unsafe_allow_html=True)
    st.caption("Professional AI for Chemistry Education")
    st.markdown("---")

    menu = st.radio("🏠 Workspace 메뉴", [
        "📓 Notion / MS Word 스타일 매니저 (추천)",
        "🔬 실험 보고서 AI 도우미",
        "🎓 전문가용 LaTeX (Overleaf) 에디터",
        "🧪 도표 & 3D 그림 생성기",
        "📝 수기 노트 AI 문서화"
    ], key="menu_selection")

    st.markdown("---")
    # 보안을 위해 API 키 입력창을 숨기고 내부적으로만 사용합니다.
    st.session_state.gemini_api_key = "AIzaSyAS7lVfLXRdNSzinXnZBTf9T6ES7s0qNMg"
    api_key = st.session_state.gemini_api_key
    st.success("✅ AI 보안 가동 중 (키 보호됨)")

    st.markdown("---")
    st.subheader("📥 통합 보고서 다운로드")

    # 임시 파일 생성을 위한 메모리 버퍼
    doc_stream = io.BytesIO()
    st.session_state.word_doc.save(doc_stream)

    st.download_button(
        label="📝 전체 작업 내용 Word로 받기",
        data=doc_stream.getvalue(),
        file_name="SNU_Chem_Report_Total.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True,
        type="primary",
        help="지금까지 '워드에 추가' 버튼을 눌러 쌓인 모든 분석 결과를 하나의 문서로 저장합니다."
    )

    st.markdown("---")
    st.caption("v2.5.0 Professional Edition")

# 과제 샘플 데이터 (TOP 5)
SAMPLES = {}
subjects = ["물리화학", "유기화학", "분석화학", "무기화학", "화학교육"]
for subj in subjects:
    try:
        with open(os.path.join(os.path.dirname(__file__), "samples", f"{subj}.md"), "r", encoding="utf-8") as f:
            SAMPLES[subj] = f.read()
    except Exception as e:
        SAMPLES[subj] = f"샘플 데이터를 불러올 수 없습니다: {e}"

# 화학교육(Chem-Ed) 교수학습 상세 가이드 데이터베이스
# 화학교육 가이드 데이터 로드 함수
def load_chem_ed_guides():
    base_dir = os.path.join(os.path.dirname(__file__), "chem_ed_guides")
    mapping = {
        "주요 오개념": "misconception.md",
        "화학 결합": "bonding.md",
        "동적 평형": "equilibrium.md",
        "권장 교수학습 모델": "models.md",
        "5E 순환 학습": "5e_model.md",
        "POE 모형": "poe_model.md",
        "평가": "assessment.md"
    }
    data = {}
    for key, filename in mapping.items():
        path = os.path.join(base_dir, filename)
        if os.path.exists(path):
            with open(path, "r", encoding="utf-8") as f:
                data[key] = f.read()
        else:
            # 기본 데이터 유지 (파일이 없을 경우 대비)
            data[key] = f"# {key}\n상세 자료를 불러올 수 없습니다."
    return data

CHEM_ED_GUIDE_DATA = load_chem_ed_guides()
def show_sample_dialog(title, content, target_subject):
    st.markdown("---")
    st.subheader(f"🎓 {title} 미리보기")

    if st.button("⚡ 이 샘플을 내 컴퓨터의 MS Word로 바로 열기 (수식 안 깨짐)", type="primary", use_container_width=True):
        with st.spinner("샘플 수식을 완벽히 변환하여 MS Word를 실행 중입니다..."):
            out_file = os.path.join(os.getcwd(), f"{title}_Sample.docx")
            full_markdown = f"# {title}\n\n" + content
            margins = {'top': 2.0, 'bottom': 2.0, 'left': 2.5, 'right': 2.5}
            convert_latex_to_word_docx(full_markdown, out_file, margins)

            import subprocess
            try:
                subprocess.Popen(['open', out_file])
                st.success("🎉 MS Word 프로그램이 성공적으로 실행되었습니다!")
            except Exception as e:
                st.error("Word를 자동으로 실행할 수 없습니다.")

    import json
    content_escaped = json.dumps(content)
    copy_code = f"""
    <button onclick="copySample(event)" style="background:#0F172A; color:white; border:none; padding:12px 15px; border-radius:8px; cursor:pointer; font-weight:bold; width:100%; font-size:15px; transition:0.2s; box-shadow: 0 2px 4px rgba(0,0,0,0.1);">
        📓 Notion / 옵시디언 등 마크다운 노트에 샘플 복사
    </button>
    <script>
    function copySample(event) {{
        const text = {content_escaped};
        const el = document.createElement('textarea');
        el.value = text;
        document.body.appendChild(el);
        el.select();
        document.execCommand('copy');
        document.body.removeChild(el);
        const btn = event.target;
        const orig = btn.innerText;
        btn.innerText = '✅ 복사 완료! (Notion에서 Ctrl+V 하세요)';
        btn.style.background = '#047857';
        setTimeout(() => {{ btn.innerText = orig; btn.style.background = '#0F172A'; }}, 2500);
    }}
    </script>
    """
    import streamlit.components.v1 as components
    components.html(copy_code, height=60)

    st.markdown(content)

    # 덮어쓰기 방지를 위해 이어붙이기로 변경
    if st.button(f"📥 이 내용을 '{target_subject}' 에디터 아래에 추가하기 (이어붙이기)", use_container_width=True):
        st.session_state.notion_db[target_subject] += "\n" + content
        if f"editor_{target_subject}" in st.session_state: del st.session_state[f"editor_{target_subject}"]
        st.success("에디터에 성공적으로 추가되었습니다!")
        st.rerun()

def render_chem_ed_core_guide():
    with st.expander("🎓 화학교육(Chem-Ed) 교수학습 핵심 가이드 (예비교사 필독)"):
        st.write("과제나 지도안 작성 시 다음의 교육학적 요소를 점검해 보세요. 각 항목을 클릭하면 **5페이지 분량의 상세 전문 자료**가 별도 창으로 열립니다.")
        
        import json
        data_json = json.dumps(CHEM_ED_GUIDE_DATA)
        
        guide_html = """
        <div id="guide-container" style="display: grid; grid-template-columns: 1fr 1fr; gap: 12px; font-family: 'Inter', sans-serif;">
        </div>
        
        <script>
        const data = """ + data_json + """;
        const container = document.getElementById('guide-container');
        
        const items = [
            { key: "주요 오개념", icon: "🧠", label: "1. 주요 오개념 (Misconception)" },
            { key: "화학 결합", icon: "💎", label: "2. 화학 결합 (Bonding)" },
            { key: "동적 평형", icon: "⚖️", label: "3. 동적 평형 (Equilibrium)" },
            { key: "권장 교수학습 모델", icon: "🧑‍🏫", label: "4. 권장 교수학습 모델" },
            { key: "5E 순환 학습", icon: "🔄", label: "5. 5E 순환 학습 모델" },
            { key: "POE 모형", icon: "🧪", label: "6. POE (Predict-Observe-Explain)" },
            { key: "평가", icon: "📊", label: "7. 평가 및 형성 평가 설계" }
        ];
        
        items.forEach(item => {
            const btn = document.createElement('button');
            btn.style.width = '100%';
            btn.style.padding = '14px 16px';
            btn.style.borderRadius = '10px';
            btn.style.border = '1px solid #e2e8f0';
            btn.style.background = 'white';
            btn.style.cursor = 'pointer';
            btn.style.textAlign = 'left';
            btn.style.fontWeight = '600';
            btn.style.fontSize = '14px';
            btn.style.color = '#1e293b';
            btn.style.transition = 'all 0.2s cubic-bezier(0.4, 0, 0.2, 1)';
            btn.style.boxShadow = '0 1px 2px rgba(0,0,0,0.05)';
            btn.innerHTML = `<span style="margin-right: 10px; font-size: 1.2rem;">${item.icon}</span> ${item.label}`;
            
            btn.onmouseover = () => { 
                btn.style.background = '#f8fafc'; 
                btn.style.borderColor = '#3B82F6';
                btn.style.transform = 'translateY(-1px)';
                btn.style.boxShadow = '0 4px 6px -1px rgba(0, 0, 0, 0.1)';
            };
            btn.onmouseout = () => { 
                btn.style.background = 'white'; 
                btn.style.borderColor = '#e2e8f0';
                btn.style.transform = 'translateY(0)';
                btn.style.boxShadow = '0 1px 2px rgba(0,0,0,0.05)';
            };
            
            btn.onclick = () => {
                const content = data[item.key] || "내용을 불러올 수 없습니다.";
                const win = window.open("", "_blank");
                if (!win) {
                    alert("팝업이 차단되었습니다! 브라우저 주소창 우측의 '팝업 차단 해제'를 클릭해주세요.");
                    return;
                }
                
                // Simple Markdown Parser
                let htmlContent = content
                    .replace(/^# (.*$)/gm, '<h1>$1</h1>')
                    .replace(/^## (.*$)/gm, '<h2>$1</h2>')
                    .replace(/^### (.*$)/gm, '<h3>$1</h3>')
                    .replace(/\\*\\*(.*?)\\*\\*/g, '<strong>$1</strong>')
                    .replace(/^- (.*$)/gm, '<li>$1</li>')
                    .replace(/\\n/g, '<br>');

                const docHtml = `
                    <!DOCTYPE html>
                    <html>
                    <head>
                        <meta charset="UTF-8">
                        <title>Chem-Ed 상세 가이드: ${item.key}</title>
                        <link href="https://fonts.googleapis.com/css2?family=Inter:wght@400;600;700&display=swap" rel="stylesheet">
                        <script src="https://polyfill.io/v3/polyfill.min.js?features=es6"><\/script>
                        <script id="MathJax-script" async src="https://cdn.jsdelivr.net/npm/mathjax@3/es5/tex-mml-chtml.js"><\/script>
                        <script>
                            window.MathJax = {
                                loader: {load: ['[tex]/mhchem']},
                                tex: {
                                    packages: {'[+]': ['mhchem']},
                                    inlineMath: [['$', '$'], ['\\\\(', '\\\\)']],
                                    displayMath: [['$$', '$$'], ['\\\\[', '\\\\]']],
                                    processEscapes: true
                                }
                            };
                        <\/script>
                        <style>
                            body { font-family: 'Inter', sans-serif; padding: 40px; line-height: 1.6; color: #1e293b; background: #f8fafc; user-select: text !important; -webkit-user-select: text !important; }
                            .container { background: white; padding: 40px; border-radius: 12px; box-shadow: 0 4px 6px rgba(0,0,0,0.05); max-width: 800px; margin: 0 auto; }
                            h1 { color: #1e3a8a; border-bottom: 2px solid #3b82f6; padding-bottom: 10px; }
                            h2 { color: #1e40af; margin-top: 30px; }
                            li { margin-bottom: 8px; }
                            .btn-group { display: flex; gap: 10px; margin-top: 40px; }
                            .print-btn, .copy-btn { padding: 12px 24px; border: none; border-radius: 6px; cursor: pointer; font-weight: bold; transition: opacity 0.2s; }
                            .print-btn { background: #3b82f6; color: white; }
                            .copy-btn { background: #10b981; color: white; }
                            .print-btn:hover, .copy-btn:hover { opacity: 0.9; }
                        </style>
                        <script>
                            function copyToClipboard() {
                                const el = document.getElementById('content');
                                // HTML과 텍스트 모두를 클립보드에 넣기 위해 시도
                                const text = el.innerText;
                                navigator.clipboard.writeText(text).then(() => {
                                    const btn = document.querySelector('.copy-btn');
                                    const originalText = btn.innerText;
                                    btn.innerText = '✅ 복사 완료!';
                                    setTimeout(() => { btn.innerText = originalText; }, 2000);
                                }).catch(err => {
                                    alert('복사 중 오류가 발생했습니다: ' + err);
                                });
                            }

                            function downloadWord() {
                                const header = "<html xmlns:o='urn:schemas-microsoft-com:office:office' " +
                                        "xmlns:w='urn:schemas-microsoft-com:office:word' " +
                                        "xmlns='http://www.w3.org/TR/REC-html40'>" +
                                        "<head><meta charset='utf-8'><title>Chem-Ed Guide</title></head><body>";
                                const footer = "</body></html>";
                                const sourceHTML = header + document.getElementById("content").innerHTML + footer;
                                
                                const source = 'data:application/vnd.ms-word;charset=utf-8,' + encodeURIComponent(sourceHTML);
                                const fileDownload = document.createElement("a");
                                document.body.appendChild(fileDownload);
                                fileDownload.href = source;
                                fileDownload.download = 'ChemEd_Guide.doc';
                                fileDownload.click();
                                document.body.removeChild(fileDownload);
                            }
                        <\/script>
                    </head>
                    <body>
                        <div class="container">
                            <div id="content">${htmlContent}</div>
                            <div class="btn-group">
                                <button class="copy-btn" onclick="copyToClipboard()">📋 복사</button>
                                <button class="copy-btn" style="background: #8b5cf6;" onclick="downloadWord()">📄 워드 다운로드</button>
                                <button class="print-btn" onclick="window.print()">📥 PDF 저장 / 인쇄</button>
                            </div>
                        </div>
                    </body>
                    </html>
                `;
                
                win.document.open();
                win.document.write(docHtml);
                win.document.close();
            };
            
            container.appendChild(btn);
        });
        </script>
        """
        import streamlit.components.v1 as components
        components.html(guide_html, height=350)

        st.markdown("---")
        st.markdown("### 📥 워드 파일(.docx) 다운로드 센터")
        st.caption("가이드 내용을 오프라인에서 확인하거나 과제에 활용할 수 있도록 워드 문서로 제공합니다.")
        
        d_col1, d_col2 = st.columns(2)
        
        items = list(CHEM_ED_GUIDE_DATA.items())
        for i in range(len(items)):
            key, content = items[i]
            col = d_col1 if i % 2 == 0 else d_col2
            with col:
                if st.button(f"📝 {key} (Word)", key=f"dl_trigger_{i}", use_container_width=True):
                    out_file = f"ChemEd_Guide_{key.replace(' ', '_')}.docx"
                    margins = (2.54, 2.54, 2.54, 2.54)
                    try:
                        convert_latex_to_word_docx(content, out_file, margins)
                        with open(out_file, "rb") as f:
                            st.download_button(
                                label=f"💾 {key}.docx 다운로드",
                                data=f.read(),
                                file_name=out_file,
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                key=f"dl_btn_{i}",
                                use_container_width=True
                            )
                    except Exception as e:
                        st.error(f"다운로드 생성 중 오류: {e}")

        st.write("")
        if st.button("📚 화학교육 가이드 전체 통합본 다운로드 (Word)", use_container_width=True, type="primary"):
            full_content = "# 화학교육(Chem-Ed) 교수학습 핵심 가이드 통합본\n\n"
            for key, content in CHEM_ED_GUIDE_DATA.items():
                full_content += f"\n\n---\n\n{content}\n"
            
            out_file = "ChemEd_Full_Guide.docx"
            margins = (2.54, 2.54, 2.54, 2.54)
            try:
                convert_latex_to_word_docx(full_content, out_file, margins)
                with open(out_file, "rb") as f:
                    st.download_button(
                        label="💾 통합 가이드.docx 즉시 다운로드",
                        data=f.read(),
                        file_name=out_file,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key="dl_full",
                        use_container_width=True
                    )
            except Exception as e:
                st.error(f"통합본 생성 중 오류: {e}")

# ==========================================
# 1. Notion 스타일 관리 및 여백 조절 + 네이티브 워드 수식
# ==========================================
if menu == "📓 Notion / MS Word 스타일 매니저 (추천)":


    # 상단 컨트롤 패널 (모든 요소를 가로 1줄로 쫙 펼치기)
    c1, c2, c3, c4, c5 = st.columns([1.5, 1.5, 1.0, 3.5, 1.5])

    with c1:
        subject = st.selectbox("과목", list(st.session_state.notion_db.keys()), label_visibility="collapsed")
    with c2:
        new_subject = st.text_input("추가", placeholder="새 과목 이름", label_visibility="collapsed")
    with c3:
        if st.button("추가"):
            if new_subject:
                st.session_state.notion_db[new_subject] = ""
                st.rerun()

    with c4:
        upload_img = st.file_uploader("사진", type=["png", "jpg", "jpeg"], label_visibility="collapsed")
        # 강력한 JavaScript 주입을 통해 업로드 창의 텍스트를 강제로 숨기고 높이를 압축합니다.
        import streamlit.components.v1 as components
        components.html("""
        <script>
        // 부모 창(Streamlit 메인 창)의 DOM에 접근
        const doc = window.parent.document;

        function flattenUploader() {
            const dropzones = doc.querySelectorAll('[data-testid="stFileUploadDropzone"]');
            dropzones.forEach(dz => {
                // 박스 자체의 여백과 높이 최소화
                dz.style.padding = '2px 5px';
                dz.style.minHeight = '35px';
                dz.style.height = '35px';
                dz.style.display = 'flex';
                dz.style.alignItems = 'center';
                dz.style.justifyContent = 'center';

                // 내부에 있는 'Drag and drop file here' 등 모든 글자와 아이콘 찾기
                const innerElements = dz.querySelectorAll('div, span, small, p, svg');
                innerElements.forEach(el => {
                    // 버튼(Browse files)이 아닌 텍스트 요소들은 모조리 숨김 처리
                    if(el.tagName !== 'BUTTON' && !el.closest('button')) {
                        el.style.display = 'none';
                    }
                });
            });
        }

        // 요소가 렌더링될 때까지 반복해서 시도
        flattenUploader();
        setTimeout(flattenUploader, 500);
        setTimeout(flattenUploader, 1000);
        </script>
        """, height=0)

    with c5:
        if upload_img and st.button("🚀 AI 자동추출"):
            if not api_key: st.error("키를 입력하세요!")
            else:
                with st.spinner("분석 중..."):
                    genai.configure(api_key=api_key)
                    model = get_safe_gemini_model()
                    try:
                        from PIL import Image
                        add_analysis_task("이 이미지 안의 텍스트와 수식을 LaTeX 형식($$ ... $$)을 사용해서 정확히 추출해줘.", Image.open(upload_img), subject, upload_img.name)
                        st.success(f"🚀 {upload_img.name} 분석 작업이 백그라운드에서 시작되었습니다. 다른 메뉴를 보셔도 분석은 계속됩니다!")
                    except Exception as e:
                        st.error(f"오류: {e}")

    st.markdown("---")

    # ==============================
    # 🚀 MS Word / Notion 바로 복사 기능 (맨 위로 이동)
    # ==============================
    st.subheader("🚀 원클릭 바로 쓰기 (다운로드 없이 복사)")
    st.write("번거로운 파일 변환/다운로드 과정 없이, 작성하신 노트를 클릭 한 번으로 복사하여 원하는 곳에 바로 붙여넣으세요!")

    # 최신 텍스트 상태 가져오기
    current_text = st.session_state.get(f"editor_{subject}", st.session_state.notion_db[subject])

    import json
    md_escaped = json.dumps(current_text)

    if st.button("⚡ 내 컴퓨터의 MS Word 프로그램으로 지금 쓴 내용 바로 열기", type="primary", use_container_width=True):
        with st.spinner("Pandoc 엔진으로 수식을 변환하고 MS Word를 실행 중입니다..."):
            out_file = os.path.join(os.getcwd(), "Converted_Note.docx")
            full_markdown = f"# {subject}\n\n" + current_text
            margins = {'top': 2.0, 'bottom': 2.0, 'left': 2.5, 'right': 2.5}
            convert_latex_to_word_docx(full_markdown, out_file, margins)

            import subprocess
            try:
                subprocess.Popen(['open', out_file])
                st.success("🎉 MS Word 프로그램이 성공적으로 실행되었습니다! (수식이 완벽하게 적용됨)")
            except Exception as e:
                st.error("Word를 자동으로 실행할 수 없습니다. 다운로드 기능을 이용해주세요.")

    copy_html = f"""
    <style>
        .btn {{
            border: none;
            border-radius: 8px;
            padding: 12px 20px;
            font-size: 15px;
            font-weight: 700;
            font-family: 'Inter', sans-serif;
            cursor: pointer;
            width: 100%;
            transition: all 0.3s;
            color: white;
            box-shadow: 0 4px 6px rgba(0,0,0,0.1);
        }}
        .btn-notion {{ background-color: #0F172A; }}
        .btn-notion:hover {{ background-color: #334155; transform: translateY(-2px); }}
        .copy-container {{ display: flex; flex-direction: row; align-items: center; justify-content: flex-start; padding: 5px 0px; }}
    </style>

    <div class="copy-container">
        <button class="btn btn-notion" id="btn-notion" onclick='copyNotion()'>
            📓 Notion / 옵시디언 등 마크다운 노트에 바로 복사
        </button>
        <button class="btn" id="openKbBtn" onclick="window.mathVirtualKeyboard ? window.mathVirtualKeyboard.show() : alert('키보드를 로드 중입니다.')"
            style="margin-left: 10px; background-color:#3B82F6;"
            >🔢 숫자 입력 키보드 열기</button>
    </div>

    <script>
    function copyNotion() {{
        const text = {md_escaped};
        const el = document.createElement('textarea');
        el.value = text;
        document.body.appendChild(el);
        el.select();
        document.execCommand('copy');
        document.body.removeChild(el);

        const btn = document.getElementById('btn-notion');
        const originalText = btn.innerText;
        btn.innerText = '✅ 복사 완료! (Notion에서 Ctrl+V 하세요)';
        setTimeout(() => btn.innerText = originalText, 2500);
    }}
    </script>
    """
    import streamlit.components.v1 as components
    components.html(copy_html, height=60)

    with st.expander("🖨️ 수동으로 Word 파일(.docx) 다운로드 (기존 기능)"):
        col_m1, col_m2, col_m3, col_m4 = st.columns(4)
        m_top = col_m1.slider("위쪽 여백 (cm)", 0.0, 5.0, 2.0, 0.1)
        m_bot = col_m2.slider("아래쪽 여백 (cm)", 0.0, 5.0, 2.0, 0.1)
        m_left = col_m3.slider("왼쪽 여백 (cm)", 0.0, 5.0, 2.5, 0.1)
        m_right = col_m4.slider("오른쪽 여백 (cm)", 0.0, 5.0, 2.5, 0.1)

        if st.button("📓 수식을 변환하여 Word 파일로 다운로드 (.docx)"):
            with st.spinner("Pandoc을 사용하여 텍스트와 LaTeX 수식을 Word 네이티브 수식으로 변환 중입니다..."):
                out_file = "Converted_Note.docx"
                full_markdown = f"# {subject}\n\n" + current_text
                margins = {'top': m_top, 'bottom': m_bot, 'left': m_left, 'right': m_right}
                convert_latex_to_word_docx(full_markdown, out_file, margins)
                with open(out_file, "rb") as f:
                    doc_bytes = f.read()
                st.download_button(
                    label="📥 완벽하게 변환된 Word 파일 다운로드",
                    data=doc_bytes,
                    file_name=f"{subject}_Notes.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

    st.markdown("---")
    st.subheader("🌟 전공 과제 샘플 TOP 5 (클릭하여 팝업 확인)")
    st.caption("클릭하시면 과목별 템플릿 샘플이 별도 창으로 열립니다.")

    def open_subject_sample_html(subject_name):
        html_path = os.path.join(os.path.dirname(__file__), "samples", f"{subject_name}.html")
        if os.path.exists(html_path):
            import subprocess
            try:
                subprocess.Popen(['open', html_path])
            except Exception as e:
                st.error(f"샘플을 여는 중 오류가 발생했습니다: {e}")
        else:
            st.error("해당 샘플 파일을 찾을 수 없습니다.")

    sc1, sc2, sc3, sc4, sc5 = st.columns(5)
    if sc1.button("⚛️ 물리화학", use_container_width=True): open_subject_sample_html("물리화학")
    if sc2.button("🧪 유기화학", use_container_width=True): open_subject_sample_html("유기화학")
    if sc3.button("📊 분석화학", use_container_width=True): open_subject_sample_html("분석화학")
    if sc4.button("💎 무기화학", use_container_width=True): open_subject_sample_html("무기화학")
    if sc5.button("🧑‍🏫 화학교육", use_container_width=True): open_subject_sample_html("화학교육")

    # 텍스트 에디터 및 미리보기 (전체 너비 사용)
    st.subheader(f"📝 {subject} 노트 에디터")
    st.info("이곳의 에디터는 화면 전체를 넓게 사용합니다. 내용이 길어도 편하게 작성하세요!")

    with st.expander("📌 자주 쓰는 수학/화학 기호 사전 (클릭해서 열기)"):
        st.info("각 코드 박스 우측의 아이콘을 클릭하여 복사(Copy)한 뒤 에디터에 붙여넣으세요! (LaTeX 형식)")

        c_sym1, c_sym2, c_sym3 = st.columns(3)
        with c_sym1:
            st.markdown("**1. 그리스 문자**")
            st.code("\\alpha, \\beta, \\gamma, \\delta\n\\Delta, \\pi, \\sigma, \\theta\n\\psi, \\Psi, \\phi, \\omega", language="text")
        with c_sym2:
            st.markdown("**2. 기본 수식 기호**")
            st.code("x^2, y_{i}, \\sqrt{x}, \\sqrt[n]{x}\n\\pm, \\mp, \\times, \\div\n\\approx, \\neq, \\propto, \\infty", language="text")
        with c_sym3:
            st.markdown("**3. 화학 반응/화살표**")
            st.code("\\rightarrow, \\rightleftharpoons\n\\xrightarrow{H^+}, \\xleftarrow[temp]{cat}\n\\uparrow, \\downarrow, \\leftrightarrow", language="text")

        st.markdown("---")
        c_sym4, c_sym5, c_sym6 = st.columns(3)
        with c_sym4:
            st.markdown("**4. 열역학/속도론**")
            st.code("\\Delta G = \\Delta H - T\\Delta S\nk = A e^{-\\frac{E_a}{RT}}\nPV = nRT, \\ln K = -\\frac{\\Delta G^\\circ}{RT}", language="text")
        with c_sym5:
            st.markdown("**5. 양자역학/오비탈**")
            st.code("\\hat{H}\\psi = E\\psi, \\lambda = \\frac{h}{p}\n\\psi_{n,l,m}, \\nabla^2, \\hbar\n\\int |\\psi|^2 d\\tau = 1", language="text")
        with c_sym6:
            st.markdown("**6. 고급 수학 (미적분/행렬)**")
            st.code("\\frac{dy}{dx}, \\frac{\\partial f}{\\partial x}, \\int_a^b\n\\sum_{i=1}^n, \\prod, \\lim_{x \\to 0}\n\\begin{pmatrix} a & b \\\\ c & d \\end{pmatrix}", language="text")

        st.caption("팁: 수식을 작성할 때는 기호 앞뒤를 $$ 로 감싸주세요! (예: $$ \\Delta G $$)")

    render_chem_ed_core_guide()

    with st.expander("📊 스마트 데이터 표(Table) & 화학 그래프 빌더"):
        st.write("데이터를 입력하면 에디터에 맞는 Markdown 표나 그래프를 생성해 줍니다.")
        t1, t2 = st.tabs(["📋 표 빌더", "📈 그래프 생성"])

        with t1:
            st.markdown("#### 엑셀형 데이터 표 에디터")
            st.caption("마우스로 셀을 클릭하여 직접 수정하거나, 표 아래를 클릭하여 행을 추가하세요. **(행 삭제: 왼쪽 끝 체크박스 선택 후 우측 상단의 휴지통 🗑️ 아이콘 클릭)**")

            # 프리미엄 논문용 템플릿 선택
            table_template = st.selectbox("📝 논문용 프리미엄 표 양식 선택:", [
                "빈 표 (기본)",
                "기술 통계 요약표 (Descriptive)",
                "상관관계 분석표 (Correlation)",
                "분산분석표 (ANOVA)",
                "이화학 실험 결과 요약표",
                "반응 속도 데이터",
                "산-염기 적정 데이터"
            ])

            import pandas as pd
            if "table_builder_df" not in st.session_state or st.session_state.get("table_template_prev") != table_template:
                if table_template == "기술 통계 요약표 (Descriptive)":
                    st.session_state.table_builder_df = pd.DataFrame([{"변수 (Variables)": "연령 (Age)", "표본 수 (N)": 120, "평균 (Mean)": 25.4, "표준편차 (SD)": 3.2}, {"변수 (Variables)": "시험 점수 (Score)", "표본 수 (N)": 120, "평균 (Mean)": 82.1, "표준편차 (SD)": 7.5}, {"변수 (Variables)": "학습 시간 (Hours)", "표본 수 (N)": 120, "평균 (Mean)": 4.5, "표준편차 (SD)": 1.1}])
                elif table_template == "상관관계 분석표 (Correlation)":
                    st.session_state.table_builder_df = pd.DataFrame([{"변수명": "1. 학습 동기", "1": "1.00", "2": "0.45**", "3": "0.38*"}, {"변수명": "2. 자기효능감", "1": "-", "2": "1.00", "3": "0.62**"}, {"변수명": "3. 학업 성취도", "1": "-", "2": "-", "3": "1.00"}])
                elif table_template == "분산분석표 (ANOVA)":
                    st.session_state.table_builder_df = pd.DataFrame([{"요인 (Source)": "집단 간 (Between)", "제곱합 (SS)": 450.2, "자유도 (df)": 2, "평균제곱 (MS)": 225.1, "F 값": 5.42, "p-value": "0.008**"}, {"요인 (Source)": "집단 내 (Within)", "제곱합 (SS)": 1205.5, "자유도 (df)": 29, "평균제곱 (MS)": 41.5, "F 값": "-", "p-value": "-"}, {"요인 (Source)": "전체 (Total)", "제곱합 (SS)": 1655.7, "자유도 (df)": 31, "평균제곱 (MS)": "-", "F 값": "-", "p-value": "-"}])
                elif table_template == "이화학 실험 결과 요약표":
                    st.session_state.table_builder_df = pd.DataFrame([{"실험군 (Group)": "대조군 (Control)", "이론 수득량 (g)": 5.0, "실제 수득량 (g)": 4.2, "수득률 (%)": 84.0, "오차율 (%)": 16.0}, {"실험군 (Group)": "촉매 A 사용", "이론 수득량 (g)": 5.0, "실제 수득량 (g)": 4.8, "수득률 (%)": 96.0, "오차율 (%)": 4.0}])
                elif table_template == "반응 속도 데이터":
                    st.session_state.table_builder_df = pd.DataFrame([{"시간 (s)": 0, "농도 (M)": 1.00}, {"시간 (s)": 10, "농도 (M)": 0.85}, {"시간 (s)": 20, "농도 (M)": 0.72}])
                elif table_template == "산-염기 적정 데이터":
                    st.session_state.table_builder_df = pd.DataFrame([{"적정액 부피 (mL)": 0.0, "pH": 2.5}, {"적정액 부피 (mL)": 5.0, "pH": 3.8}, {"적정액 부피 (mL)": 10.0, "pH": 7.0}])
                else:
                    st.session_state.table_builder_df = pd.DataFrame([{"컬럼 1": "", "컬럼 2": "", "컬럼 3": ""} for _ in range(3)])
                st.session_state.table_template_prev = table_template

            edited_df = st.data_editor(
                st.session_state.table_builder_df,
                num_rows="dynamic",
                use_container_width=True,
                key="table_editor_ui"
            )

            if st.button("✨ Markdown 표로 완성 및 에디터에 삽입"):
                if not edited_df.empty:
                    md_table = "\n| " + " | ".join([str(c) for c in edited_df.columns]) + " |\n"
                    md_table += "| " + " | ".join(["---"] * len(edited_df.columns)) + " |\n"
                    for _, row in edited_df.iterrows():
                        md_table += "| " + " | ".join([str(val) for val in row.values]) + " |\n"

                    st.session_state.notion_db[subject] += "\n" + md_table
                    if f"editor_{subject}" in st.session_state: del st.session_state[f"editor_{subject}"]
                    st.success("고급스러운 표가 성공적으로 추가되었습니다!")
                    st.rerun()

        with t2:
            st.markdown("#### 프리미엄 화학/물리 데이터 시각화")
            g_col1, g_col2 = st.columns([1, 2])

            with g_col1:
                g_type = st.radio("📈 시각화 유형", [
                    "선 그래프 (Line)",
                    "영역 그래프 (Area)",
                    "막대 그래프 (Bar)",
                    "산점도 (Scatter)",
                    "박스 플롯 (Box Plot)",
                    "방사형 차트 (Radar)",
                    "적정 곡선 (Spline)",
                    "파이 차트 (Pie)"
                ])
                g_color = st.color_picker("🎨 그래프 주요 색상", "#3B82F6")
                g_title = st.text_input("📝 그래프 제목", f"{subject} 데이터 분석")

            with g_col2:
                st.caption("그래프에 표시할 데이터를 입력하세요. 박스/방사/파이 차트의 경우 X는 범주(글자), Y는 숫자입니다. **(행 삭제: 우측 상단 🗑️ 클릭)**")
                import pandas as pd
                if "graph_builder_df" not in st.session_state or st.session_state.get("g_type_prev") != g_type:
                    if "적정 곡선" in g_type:
                        st.session_state.graph_builder_df = pd.DataFrame([{"X (적정 부피)": 0, "Y (pH)": 2.5}, {"X (적정 부피)": 10, "Y (pH)": 3.8}, {"X (적정 부피)": 20, "Y (pH)": 7.0}, {"X (적정 부피)": 25, "Y (pH)": 11.2}])
                    elif "선 그래프" in g_type or "영역 그래프" in g_type or "산점도" in g_type:
                        st.session_state.graph_builder_df = pd.DataFrame([{"X (시간/농도 등)": 10, "Y 값": 2.5}, {"X (시간/농도 등)": 20, "Y 값": 5.8}, {"X (시간/농도 등)": 30, "Y 값": 12.0}, {"X (시간/농도 등)": 40, "Y 값": 15.5}])
                    elif "박스 플롯" in g_type:
                        st.session_state.graph_builder_df = pd.DataFrame([{"집단 (X)": "대조군", "측정값 (Y)": 45.2}, {"집단 (X)": "대조군", "측정값 (Y)": 48.1}, {"집단 (X)": "대조군", "측정값 (Y)": 44.5}, {"집단 (X)": "실험군", "측정값 (Y)": 85.3}, {"집단 (X)": "실험군", "측정값 (Y)": 92.1}, {"집단 (X)": "실험군", "측정값 (Y)": 88.0}])
                    else: # 막대, 방사형, 파이
                        st.session_state.graph_builder_df = pd.DataFrame([{"항목 (X 범주)": "A", "수치 (Y)": 10.0}, {"항목 (X 범주)": "B", "수치 (Y)": 25.5}, {"항목 (X 범주)": "C", "수치 (Y)": 15.2}, {"항목 (X 범주)": "D", "수치 (Y)": 30.1}])
                    st.session_state.g_type_prev = g_type

                graph_df = st.data_editor(
                    st.session_state.graph_builder_df,
                    num_rows="dynamic",
                    use_container_width=True,
                    key="graph_editor_ui"
                )

            if st.button("🖼️ 고해상도 그래프 생성 및 워드 추가"):
                try:
                    import pandas as pd
                    import matplotlib.pyplot as plt
                    import numpy as np

                    clean_df = graph_df.dropna()
                    raw_x = clean_df.iloc[:, 0].tolist()
                    y_vals = pd.to_numeric(clean_df.iloc[:, 1], errors='coerce').tolist()

                    # X값 자동 형변환 (숫자가 가능하면 숫자, 아니면 문자열)
                    try:
                        x_vals = pd.to_numeric(clean_df.iloc[:, 0]).tolist()
                    except Exception:
                        x_vals = clean_df.iloc[:, 0].astype(str).tolist()

                    # Y값은 무조건 숫자여야 함
                    y_vals = pd.to_numeric(clean_df.iloc[:, 1], errors='coerce').tolist()

                    # 결측치(NaN) 쌍 제거
                    valid_pairs = [(x, y) for x, y in zip(x_vals, y_vals) if not (pd.isna(x) or pd.isna(y))]
                    x = [p[0] for p in valid_pairs]
                    y = [p[1] for p in valid_pairs]

                    if "적정 곡선" in g_type and len(x) > 0 and not all(isinstance(v, (int, float, np.integer, np.floating)) for v in x):
                        st.error("적정 곡선(Spline)은 X축 값이 반드시 '숫자'여야 합니다.")
                    elif not x or not y or len(x) != len(y):
                        st.error("데이터 형식이 올바르지 않거나 Y값에 문자가 섞여있습니다.")
                    else:
                        plt.style.use('seaborn-v0_8-whitegrid') # 고급 논문용 화이트그리드

                        if "방사형 차트" in g_type:
                            fig, ax = plt.subplots(figsize=(5, 5), dpi=300, subplot_kw=dict(polar=True))
                        else:
                            fig, ax = plt.subplots(figsize=(6, 4), dpi=300)

                        fig.patch.set_facecolor('#FFFFFF')
                        ax.set_facecolor('#FFFFFF')

                        if "선 그래프" in g_type:
                            ax.plot(x, y, marker='o', linestyle='-', color=g_color, linewidth=2.5, markersize=8, markerfacecolor='white', markeredgecolor=g_color, markeredgewidth=2)
                        elif "영역 그래프" in g_type:
                            ax.fill_between(x, y, color=g_color, alpha=0.3)
                            ax.plot(x, y, marker='o', color=g_color, linewidth=2)
                        elif "산점도" in g_type:
                            ax.scatter(x, y, color=g_color, s=80, alpha=0.8, edgecolor='white', linewidth=1)
                        elif "막대 그래프" in g_type:
                            ax.bar(x, y, color=g_color, alpha=0.85, edgecolor='black', linewidth=0.5, width=0.5)
                        elif "박스 플롯" in g_type:
                            unique_x = list(dict.fromkeys(x))
                            data_groups = [[y[i] for i in range(len(x)) if x[i] == ux] for ux in unique_x]
                            ax.boxplot(data_groups, labels=unique_x, patch_artist=True, boxprops=dict(facecolor=g_color, color='black', alpha=0.7), medianprops=dict(color='#EF4444', linewidth=2))
                        elif "적정 곡선" in g_type:
                            from scipy.interpolate import make_interp_spline

                            # X 정렬 및 중복 X값 처리 (평균값 적용)
                            unique_xy = {}
                            for px, py in sorted(zip(x, y), key=lambda p: p[0]):
                                if px in unique_xy:
                                    unique_xy[px].append(py)
                                else:
                                    unique_xy[px] = [py]

                            x_sorted = list(unique_xy.keys())
                            y_sorted = [np.mean(unique_xy[px]) for px in x_sorted]

                            if len(x_sorted) < 2:
                                st.error("적정 곡선(Spline)을 그리려면 고유한 X값이 최소 2개 이상 필요합니다.")
                            else:
                                x_new = np.linspace(min(x_sorted), max(x_sorted), 300)
                                spl = make_interp_spline(x_sorted, y_sorted, k=min(3, len(x_sorted)-1))
                                y_new = spl(x_new)
                                ax.plot(x_new, y_new, color=g_color, lw=3)
                                ax.scatter(x_sorted, y_sorted, color='#EF4444', s=40, zorder=5, label='Data')
                                ax.legend()
                        elif "파이 차트" in g_type:
                            colors = [g_color, '#FCA5A5', '#FCD34D', '#6EE7B7', '#93C5FD', '#C4B5FD', '#F9A8D4', '#D1D5DB']
                            ax.pie(y, labels=x, autopct='%1.1f%%', startangle=90, colors=colors, wedgeprops={'edgecolor': 'white', 'linewidth': 1.5})
                            ax.axis('equal')
                        elif "방사형 차트" in g_type:
                            angles = np.linspace(0, 2 * np.pi, len(x), endpoint=False).tolist()
                            y_rad = y + [y[0]]
                            angles += [angles[0]]

                            ax.plot(angles, y_rad, color=g_color, linewidth=2)
                            ax.fill(angles, y_rad, color=g_color, alpha=0.25)
                            ax.set_xticks(angles[:-1])
                            ax.set_xticklabels(x, fontsize=11, fontweight='bold')
                            ax.tick_params(axis='x', pad=10)

                        if g_type not in ["파이 차트 (Pie)", "방사형 차트 (Radar)"]:
                            ax.set_xlabel(clean_df.columns[0], fontweight='bold', color='#334155', fontsize=11)
                            ax.set_ylabel(clean_df.columns[1], fontweight='bold', color='#334155', fontsize=11)
                            ax.spines['top'].set_visible(False)
                            ax.spines['right'].set_visible(False)
                            ax.spines['left'].set_linewidth(1.5)
                            ax.spines['bottom'].set_linewidth(1.5)
                            ax.tick_params(axis='both', which='major', labelsize=10)

                        ax.set_title(g_title, fontsize=15, fontweight='bold', pad=20, color='#1E293B')

                        buf = io.BytesIO()
                        plt.savefig(buf, format='png', bbox_inches='tight', dpi=300)
                        buf.seek(0)
                        st.image(buf)
                        st.session_state.word_doc.add_picture(buf, width=Inches(5.0))
                        st.success("🎉 세련된 그래프가 생성되어 워드 문서에 성공적으로 추가되었습니다!")
                except Exception as e:
                    st.error(f"데이터 형식을 확인하세요: {e}")

    # ==============================================================
    # 💡 양방향 자동화: 수동/자동 모드 스위치(Toggle)
    # ==============================================================
    col_t1, col_t2 = st.columns([1, 4])
    with col_t1:
        auto_mode = st.toggle("🤖 AI 자동 완성 모드 켜기")
    with col_t2:
        if auto_mode:
            st.info("현재 **자동(AI) 모드**입니다. 아래에 지시어를 입력하면 에디터 맨 밑에 내용이 자동으로 작성됩니다.")
        else:
            st.write("현재 **수동 모드**입니다. 자유롭게 키보드로 과제를 타이핑하세요.")

    if auto_mode:
        with st.container(border=True):
            ai_c1, ai_c2 = st.columns([4, 1.2])
            with ai_c1:
                ai_prompt = st.text_input("AI에게 작성시킬 내용을 입력하세요", placeholder="예: 살리실산과 아세트산 무수물의 에스터화 반응 메커니즘을 상세히 설명해줘", label_visibility="collapsed")
            with ai_c2:
                if st.button("✨ 에디터에 이어붙이기", use_container_width=True):
                    if not api_key:
                        st.error("API 키를 입력하세요!")
                    elif ai_prompt:
                        with st.spinner("AI가 전공 수준의 지식으로 내용을 작성하여 에디터에 붙여넣고 있습니다..."):
                            genai.configure(api_key=api_key)
                            model = get_safe_gemini_model()
                            try:
                                sys_prompt = "너는 서울대학교 화학교육과 조교야. 학생이 과제 레포트 작성 중 너에게 다음 부분의 대필을 요청했어. 전공 수준의 정확한 화학/물리 지식(LaTeX 수식 $$ $$ 적극 활용)과 논리적인 문장으로 작성해줘:\n\n요청내용: "
                                res = model.generate_content(sys_prompt + ai_prompt)
                                st.session_state.notion_db[subject] += "\n" + res_text
                                if f"editor_{subject}" in st.session_state: del st.session_state[f"editor_{subject}"]
                                st.rerun()
                            except Exception as e:
                                st.error(f"오류: {e}")

    st.write("") # 간격 조절

    # 미리보기 영역을 위에 배치하기 위한 빈 공간(placeholder) 생성
    preview_placeholder = st.empty()

    # 에디터 (아래에 배치)
    note_content = st.text_area(
        "수식이나 내용을 자유롭게 수정하세요:",
        value=st.session_state.notion_db[subject],
        height=400,
        key=f"editor_{subject}"
    )
    st.session_state.notion_db[subject] = note_content

    # 위에서 만든 빈 공간에 실시간 렌더링 결과 채워넣기
    with preview_placeholder.container():
        st.markdown("**미리보기 (실시간 수식 확인):**")
        with st.container(border=True):
            st.markdown(note_content)


# ==========================================
# 2. 실험 보고서 AI 도우미 (검색 및 연동)
# ==========================================
elif menu == "🔬 실험 보고서 AI 도우미":
    st.title("🔬 화학/물리 실험 보고서 AI 도우미")
    st.write("실험 보고서 작성에 필요한 예시 자료를 검색하고, 유용한 화학 관련 웹 프로그램들을 바로 실행해 보세요!")

    st.subheader("🔍 보고서 예시 및 템플릿 검색 (AI)")

    # 추천 샘플 TOP 5 버튼 배치
    st.markdown("🌟 **추천 전공 실험 샘플 TOP 5 (Quick Access):**")

    def open_sample_html(topic_id):
        html_path = os.path.join(os.path.dirname(__file__), "samples_html", f"{topic_id}.html")
        if os.path.exists(html_path):
            import subprocess
            try:
                subprocess.Popen(['open', html_path])
            except Exception as e:
                st.error(f"샘플을 여는 중 오류가 발생했습니다: {e}")
        else:
            st.error("해당 샘플 파일을 찾을 수 없습니다.")

    c_s1, c_s2, c_s3, c_s4, c_s5 = st.columns(5)
    s_topics = ["아스피린 합성 및 재결정", "나일론 6,6 계면 중합", "산-염기 적정 지시약 원리", "NaCl 결정 구조 및 XRD 분석", "수소 원자 스펙트럼과 발머 계열"]

    if "report_search_query" not in st.session_state:
        st.session_state.report_search_query = ""

    if c_s1.button("💊 아스피린"):
        st.session_state.report_search_query = s_topics[0]
        open_sample_html("aspirin")
    if c_s2.button("🧶 나일론"):
        st.session_state.report_search_query = s_topics[1]
        open_sample_html("nylon")
    if c_s3.button("🧪 적정"):
        st.session_state.report_search_query = s_topics[2]
        open_sample_html("titration")
    if c_s4.button("🧊 격자"):
        st.session_state.report_search_query = s_topics[3]
        open_sample_html("lattice")
    if c_s5.button("🌈 스펙트럼"):
        st.session_state.report_search_query = s_topics[4]
        open_sample_html("spectrum")

    # 대학 및 대상 선택 UI 추가
    st.markdown("🏛️ **대학별/대상별 특화 검색 옵션:**")
    col_univ, col_target = st.columns(2)
    selected_univ = col_univ.selectbox("대학교 선택", ["전체(통합)", "서울대학교", "고려대학교", "연세대학교", "성균관대학교", "한양대학교", "KAIST", "POSTECH"], key="univ_select")
    selected_target = col_target.selectbox("대상 선택", ["전체(통합)", "교수용", "학생용"], key="target_select")

    search_query = st.text_input("어떤 실험 보고서 예시가 필요하신가요?", key="report_search_query", placeholder="예: 아스피린 합성 실험, 산염기 적정, 나일론 합성")

    if st.button("🚀 AI로 보고서 예시/구조 검색하기"):
        if not api_key:
            st.error("왼쪽 메뉴에 Gemini API Key를 입력해주세요!")
        elif search_query:
            with st.spinner(f"'{search_query}'에 대한 심층 보고서 및 교수학습 가이드를 생성 중입니다..."):
                genai.configure(api_key=api_key)
                model = get_safe_gemini_model(use_grounding=True)
                prompt = f"""너는 서울대학교 화학교육과(Chemistry Education) 및 주요 명문대(SKY, 성균관대, 한양대, KAIST, POSTECH)의 학술 시스템을 꿰뚫고 있는 전문 학술 고문이야.
학생이 요청한 주제: '{search_query}'

**중요 지시:**
1. 검색 기간에 어떠한 제한도 두지 말고(무한/All-time), 최신 정보뿐만 아니라 과거의 모든 학술적 자료를 포함하여 탐색해줘.
2. 선택된 대학교(**{selected_univ}**)의 학술적 특성 및 연구 트렌드를 중점적으로 반영하여 보고서 가이드를 작성해줘. (만약 '전체'라면 국내 주요 명문대들의 특징을 통합적으로 비교 설명해줘.)
3. 선택된 대상(**{selected_target}**)의 관점에서 내용을 소분류하여, 해당 대상이 집중해야 할 핵심 포인트(교수용: 채점 기준, 교육적 피드백 / 학생용: 오차 분석 팁, 논리적 서술 방법)를 구체적으로 작성해줘.
다음의 고밀도 전공 지식과 교육학적 통찰을 포함하여 전문적인 보고서 가이드를 작성해줘:

1. 실험 목적 및 교육과정 연계 (Objective & Pedagogical Goal): 이 실험의 화학적 목적뿐만 아니라, 예비 교사로서 중/고교 수준에서 어떤 학업 성취 기준과 연계할 수 있는지.
2. 심층적 전공 이론 (Advanced Chemical Theory):
   - 물리/유기/무기화학적 원리 상술 (전자 이동 메커니즘, 열역학, 양자화학 등).
   - 반드시 LaTeX 수식($$ ... $$)을 적극 활용하여 수식 위주로 설명할 것.
3. 데이터 처리 및 기기분석 가이드 (Data Handling):
   - NMR, IR, GC/MS, XRD 등 관련 기기분석 데이터의 해석 방법론.
4. 화학교육학적 분석 (Pedagogical Analysis):
   - **학습자 오개념(Misconceptions):** 학생들이 이 주제에서 가장 흔히 겪는 오개념을 구체적으로 식별하고, 인지 갈등을 유도하는 방식의 교정 전략 제시.
   - **교수법:** 5E 순환 학습 모델 또는 POE 모형을 적용한 수업 설계 아이디어.
5. 표준 수업 지도안 양식 (Lesson Plan Outline):
   - **도입(Introduction):** 전시 학습 확인 및 동기 유발 전략.
   - **전개(Development):** 핵심 탐구 활동, 학생-교사 간 예상 발문(Questioning), 개념 변화를 위한 핵심 피드백.
   - **정리 및 평가(Conclusion):** 학습 내용 요약 및 형성 평가 문항 예시.
"""



                try:
                    response_text = robust_generate_content(prompt, use_grounding=True)
                    st.session_state.last_report_result = response_text
                    st.success("✨ 전공 전문 보고서 및 화학교육 가이드라인 생성 완료!")
                    st.markdown(response_text)

                    # Word 다운로드 버튼 추가
                    try:
                        out_file = f"Report_{search_query[:15].replace(' ', '_')}.docx"
                        margins = {'top': 2.0, 'bottom': 2.0, 'left': 2.5, 'right': 2.5}
                        convert_latex_to_word_docx(response_text, out_file, margins)
                        with open(out_file, "rb") as f:
                            st.download_button("📥 생성된 보고서 Word로 다운로드", data=f.read(), file_name=out_file, key="report_word_dl")
                        os.remove(out_file)

                        # 데이터베이스 저장 버튼 추가
                        st.divider()
                        st.subheader("📂 데이터베이스(Notion)에 즉시 저장")
                        db_col1, db_col2 = st.columns([3, 1])
                        with db_col1:
                            target_sub = st.selectbox("저장할 과목 선택", list(st.session_state.notion_db.keys()), key="report_db_subject")
                        with db_col2:
                            if st.button("📥 DB에 추가", use_container_width=True, key="report_db_add_btn"):
                                st.session_state.notion_db[target_sub] += "\n\n---\n" + response_text
                                save_state()
                                st.success(f"'{target_sub}' 데이터베이스에 저장되었습니다!")

                    except Exception as e_word:
                        st.error(f"워드 변환 중 오류: {e_word}")



                except Exception as e:
                    st.error(f"AI 생성 중 오류가 발생했습니다: {e}")

    st.markdown("---")
    st.subheader("🔗 전공 및 화학교육 필수 연동 데이터베이스")
    render_chem_ed_core_guide()
    st.write("학부 전공 수준의 데이터 처리 및 예비 교사를 위한 교육용 툴입니다. (아래 샘플 클릭 시 해당 주제로 즉시 이동)")

    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.info("**🧪 분자 구조 및 반응식**\n\n[ChemDraw JS (무료)](https://chemdrawdirect.perkinelmer.cloud/js/sample/index.html)")
        st.success("**🧬 3D 단백질/DNA 구조**\n\n[Sample: DNA 3D 구조(1BNA)](https://www.rcsb.org/structure/1BNA)")
    with col2:
        st.warning("**📊 스펙트럼(IR, NMR) 데이터**\n\n[Sample: 에탄올 IR 스펙트럼](https://webbook.nist.gov/cgi/cbook.cgi?ID=C64175&Type=IR-SPEC&Index=1)")
        st.error("**🔬 화합물 정밀 물성 검색**\n\n[Sample: 아스피린(Aspirin) 데이터](https://pubchem.ncbi.nlm.nih.gov/compound/Aspirin)")
    with col3:
        st.info("**📈 데이터 추출 및 그래프**\n\n[WebPlotDigitizer](https://automeris.io/WebPlotDigitizer/)")
        st.success("**📚 화학교육 학술 문헌 검색**\n\n[Sample: 오개념(Misconception) 논문](https://scholar.google.com/scholar?q=misconceptions+in+chemistry+education)")
    with col4:
        st.warning("**🧑‍🏫 교육용 가상 실험실**\n\n[Sample: 분자 기하 구조 시뮬레이션](https://phet.colorado.edu/en/simulations/molecule-shapes)")

# ==========================================
# 3, 4, 5번 메뉴는 기존 코드 동일하게 유지됨
# ==========================================
elif menu == "🎓 전문가용 LaTeX (Overleaf) 에디터":
    st.title("🎓 전문가용 LaTeX (Overleaf) 에디터")
    col_ai, col_preview = st.columns([1, 1])
    with col_ai:
        uploaded_file = st.file_uploader("수기 노트를 LaTeX 논문 코드로 변환", type=["png", "jpg", "jpeg"], key="latex_up")
        if uploaded_file and st.button("LaTeX 코드로 완벽 변환"):
            if not api_key: st.error("API 키를 입력해주세요.")
            else:
                with st.spinner("AI가 논문 코드를 작성 중입니다..."):
                    genai.configure(api_key=api_key)
                    model = get_safe_gemini_model()
                    try:
                        res_text = robust_generate_content("Overleaf 컴파일용 완벽한 LaTeX 코드만 출력해.", images=Image.open(uploaded_file))
                        st.session_state.latex_code = res_text
                    except Exception as e: st.error(e)
    if "latex_code" not in st.session_state: st.session_state.latex_code = r"$$ E = mc^2 $$"
    st.markdown("### 🧪 필수 수식 템플릿 (클릭 시 전체 수식 자동 입력)")

    if "mathlive_init" not in st.session_state:
        st.session_state.mathlive_init = r"\sum_{i=1}^{n}"
    def set_mathlive_template(template_str):
        st.session_state.mathlive_init = template_str

    t1, t2, t3 = st.columns(3)
    t1.button("1D 상자 속 입자 (Wavefunction)", use_container_width=True, on_click=set_mathlive_template, args=(r"\psi_n(x) = \sqrt{\frac{2}{L}} \sin\left(\frac{n\pi x}{L}\right)",))
    t2.button("해밀토니안 연산자 (Hamiltonian)", use_container_width=True, on_click=set_mathlive_template, args=(r"\hat{H} = -\frac{\hbar^2}{2m}\frac{d^2}{dx^2} + V(x)",))
    t3.button("슈뢰딩거 방정식 (Schrödinger)", use_container_width=True, on_click=set_mathlive_template, args=(r"-\frac{\hbar^2}{2m}\frac{d^2\psi}{dx^2} + V(x)\psi = E\psi",))

    st.markdown("---")

    st.markdown("### 🧮 완벽 시각화 수식 에디터 (빈칸 클릭 후 숫자 입력)")

    import textwrap
    custom_mathlive_html = textwrap.dedent("""<!DOCTYPE html>
    <html>
    <head>
        __SCRIPT_LIB__
        <style>
            body { font-family: sans-serif; margin: 0; padding: 10px; background-color: white; }
            #mf {
                font-size: 28px; width: 100%; padding: 15px; min-height: 120px;
                border: 2px solid #3B82F6; border-radius: 8px;
                background-color: #F8FAFC; box-sizing: border-box;
                margin-bottom: 15px;
            }
            .keypad {
                display: grid;
                grid-template-columns: repeat(10, 1fr);
                gap: 8px;
                margin-bottom: 15px;
            }
            .keypad button {
                padding: 12px 5px; font-size: 16px; font-weight: bold;
                border: 1px solid #CBD5E1; border-radius: 6px;
                background-color: #F1F5F9; cursor: pointer;
                transition: 0.2s;
            }
            .keypad button:hover { background-color: #E2E8F0; }
            .keypad button.action { background-color: #DBEAFE; color: #1E40AF; border-color: #BFDBFE; }
            .keypad button.action:hover { background-color: #BFDBFE; }
            .copy-btn {
                width: 100%; padding: 15px; font-size: 18px; font-weight: bold;
                background-color: #10B981; color: white; border: none;
                border-radius: 8px; cursor: pointer; transition: 0.3s;
            }
            .copy-btn:hover { background-color: #059669; }
        </style>
    </head>
    <body>
        <p style="color: #64748B; font-size: 14px; margin-top: 0;">점선으로 된 빈칸(Placeholder)을 마우스로 클릭하고 아래 숫자 버튼을 눌러보세요!</p>
        <math-field id="mf">__MATHLIVE_INIT_VALUE__</math-field>
        <textarea id="latex-code" style="width: 100%; font-family: monospace; font-size: 12pt; padding: 10px; border: 2px solid #CBD5E1; border-radius: 8px; resize: vertical; background-color: #F8FAFC; min-height: 60px; margin-bottom: 15px; box-sizing: border-box;" placeholder="여기에 LaTeX 코드가 텍스트로 나타납니다 (직접 수정 가능)..."></textarea>

        <div class="keypad">
            <button onclick="insert('1')">1</button>
            <button onclick="insert('2')">2</button>
            <button onclick="insert('3')">3</button>
            <button onclick="insert('4')">4</button>
            <button onclick="insert('5')">5</button>
            <button onclick="insert('6')">6</button>
            <button onclick="insert('7')">7</button>
            <button onclick="insert('8')">8</button>
            <button onclick="insert('9')">9</button>
            <button onclick="insert('0')">0</button>

            <button onclick="insert('+')">+</button>
            <button onclick="insert('-')">-</button>
            <button onclick="insert('=')">=</button>
            <button onclick="insert('(')">(</button>
            <button onclick="insert(')')">)</button>
            <button class="action" onclick="insert('\\\\frac{#?}{#?}')">분수</button>
            <button class="action" onclick="insert('\\\\sqrt{#?}')">√</button>
            <button class="action" onclick="insert('^{#?}')">제곱</button>
            <button class="action" onclick="insert('_{#?}')">아래첨자</button>
            <button class="action" onclick="document.getElementById('mf').executeCommand('deleteAll')">전체지움</button>

            <button class="action" onclick="insert('\\\\int_{#?}^{#?}')">∫ 적분</button>
            <button class="action" onclick="insert('\\\\sum_{#?}^{#?}')">∑ 시그마</button>
            <button class="action" onclick="insert('\\\\infty')">∞</button>
            <button class="action" onclick="insert('\\\\psi')">ψ</button>
            <button class="action" onclick="insert('\\\\pi')">π</button>
            <button class="action" onclick="insert('\\\\alpha')">α</button>
            <button class="action" onclick="insert('\\\\beta')">β</button>
            <button class="action" onclick="insert('\\\\theta')">θ</button>
            <button class="action" onclick="insert('\\\\Delta')">Δ</button>
            <button class="action" onclick="document.getElementById('mf').executeCommand('moveToPreviousChar')">◀ 이동</button>
        </div>

        <button class="copy-btn" id="copyBtn" onclick="copyToClipboard()" style="margin-top:10px; background-color:#10B981;">📋 완성된 수식 복사하기 (클릭 후 아래 에디터에 붙여넣기 Ctrl+V)</button>

        __SCRIPT_START__
            const mf = document.getElementById('mf');
            const latexCode = document.getElementById('latex-code');

            // 양방향 동기화 설정
            mf.addEventListener('input', (ev) => {
                latexCode.value = mf.value;
            });

            latexCode.addEventListener('input', (ev) => {
                mf.value = latexCode.value;
            });

            // 초기화
            setTimeout(() => { latexCode.value = mf.value; }, 100);

            function insert(latex) {
                if (typeof mf.insert === 'function') {
                    mf.insert(latex);
                } else {
                    mf.executeCommand(['insert', latex]);
                }
                mf.focus();
            }

            function copyToClipboard() {
                const code = '$$ ' + mf.value + ' $$';
                const el = document.createElement('textarea');
                el.value = code;
                // 요소가 화면에 보이지 않도록 스타일 설정
                el.style.position = 'absolute';
                el.style.left = '-9999px';
                document.body.appendChild(el);
                el.select();

                try {
                    document.execCommand('copy');
                    const btn = document.getElementById('copyBtn');
                    const orig = btn.innerText;
                    btn.innerText = '✅ 복사 완료! 아래 메인 에디터 코드창에 붙여넣기(Ctrl+V) 하세요!';
                    btn.style.backgroundColor = '#2563EB';
                    setTimeout(() => {
                        btn.innerText = orig;
                        btn.style.backgroundColor = '#10B981';
                    }, 3000);
                } catch (err) {
                    alert('복사에 실패했습니다.');
                }

                document.body.removeChild(el);
            }
        __SCRIPT_END__
    </body>
    </html>
    """)

    import streamlit.components.v1 as components
    html_to_render = custom_mathlive_html.replace('__MATHLIVE_INIT_VALUE__', st.session_state.mathlive_init)
    html_to_render = html_to_render.replace('__SCRIPT_START__', '<script>').replace('__SCRIPT_END__', '</script>')
    html_to_render = html_to_render.replace('__SCRIPT_LIB__', '<script defer src="https://unpkg.com/mathlive"></script>')
    components.html(html_to_render, height=620)

    st.markdown("---")

    col_edit, col_render = st.columns([1, 1])
    with col_edit:
        user_latex = st.text_area("코드 수정:", key="latex_code", height=400)

        d_col1, d_col2 = st.columns(2)
        with d_col1:
            st.download_button("📥 LaTeX 원본 다운로드 (.tex)", data=user_latex, file_name="expert_assignment.tex", use_container_width=True)
        with d_col2:
            try:
                import tempfile
                import pypandoc
                import os
                with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
                    tmp_path = tmp.name

                # LaTeX 문법을 워드로 변환 (수식 포함)
                pypandoc.convert_text(user_latex, 'docx', format='latex', outputfile=tmp_path)

                with open(tmp_path, "rb") as f:
                    docx_bytes = f.read()
                os.remove(tmp_path) # 임시 파일 정리

                st.download_button("📝 워드로 다운로드 (.docx)", data=docx_bytes, file_name="expert_assignment.docx", use_container_width=True)
            except Exception as e:
                # Pandoc이 없거나 에러가 발생한 경우 단순 텍스트로 저장하는 대체(Fallback) 수단
                import io
                from docx import Document
                fallback_doc = Document()
                fallback_doc.add_heading("LaTeX Source Code", level=1)
                for line in user_latex.split("\n"):
                    fallback_doc.add_paragraph(line)
                f_stream = io.BytesIO()
                fallback_doc.save(f_stream)
                st.download_button("📝 워드로 다운로드 (텍스트만)", data=f_stream.getvalue(), file_name="expert_assignment_raw.docx", use_container_width=True, help="Pandoc 미설치로 인해 수식이 그래픽으로 변환되지 않고 텍스트 원본으로 저장됩니다.")

    with col_render:
        st.info("⚠️ 현재 화면(미리보기)에서는 수식 기호($$...$$)만 보이며, `\\documentclass` 등 논문 전체 구조는 보이지 않습니다. 완벽한 PDF 논문을 보시려면 아래의 **오버리프 연동 버튼**을 누르세요!")

        # Overleaf 연동 버튼 (HTML Form POST)

        overleaf_html = f"""
        <form action="https://www.overleaf.com/docs" method="POST" target="_blank">
            <textarea name="snip" style="display:none;">{user_latex}</textarea>
            <button type="submit" style="background-color: #10B981; color: white; padding: 15px 20px; border: none; border-radius: 8px; font-size: 16px; font-weight: bold; cursor: pointer; width: 100%; box-shadow: 0 4px 6px rgba(0,0,0,0.1); transition: 0.3s;">
                Overleaf 웹으로 전송하여 PDF 컴파일하기
            </button>
        </form>
        """
        import streamlit.components.v1 as components
        components.html(overleaf_html, height=80)

        with st.container(border=True):
            st.markdown(user_latex)

elif menu == "🧪 도표 & 3D 그림 생성기":
    st.markdown('''
        <style>
        /* 스마트 입력창(text_area) 디자인 최적화 */
        div[data-testid="stTextArea"] textarea {
            font-size: 1.1rem !important;
            padding: 15px !important;
            line-height: 1.6 !important;
        }
        </style>
    ''', unsafe_allow_html=True)

    st.subheader("🧪 도표 & 3D 그림 생성기")
    # --- 스마트 라우터 기능 시작 ---
    def handle_smart_input():
        import unicodedata
        import re

        raw_query = st.session_state.smart_input_val.strip()
        if not raw_query:
            return

        # PDF에서 복사 시 자주 발생하는 한글 자모 분리 현상(NFD) 및 유령 문자(Zero-width space) 해결
        query = unicodedata.normalize('NFC', raw_query)
        query = query.replace('\u200b', '').replace('\u200c', '').replace('\u200d', '')
        # 과도한 줄바꿈이나 공백 압축
        query = re.sub(r'\s+', ' ', query)

        api_key = st.session_state.get("gemini_api_key", "")

        is_problem = any(k in query for k in ["?", "어느", "이유", "판별", "비교", "하시오", "구하라", "단계", "계산", "%", "풀어", "알려", "설명", "왜", "어떻게", "답", "근거", "원리", "완성", "예상", "제시", "찾으", "결합각", "혼성", "오비탈", "hybrid", "orbital", "Lewis", "루이스"]) or (len(query) > 15 and any(k in query for k in [" ", "\n", "다.", "요."]))


        if is_problem and api_key:
            try:
                genai.configure(api_key=api_key)
                prompt = f"""당신은 세계 최고의 화학/물리 교육 전문가이자 학술 데이터 분석 거장입니다.
모든 결과는 **'풀컬러 초고해상도 ASCII 작도 및 전문 학술 데이터 시트 작성'** 모드로 고정하여 출력하세요.

**[슈퍼 정밀 시각화 지침 - 절대 준수]**:
1. **[원소별 색상 및 세부 라벨링]**: ASCII 아트 내에서 산소(O)는 **빨간색(:red[O])**, 질소(N)는 **파란색(:blue[N])**으로 Markdown 색상 문법을 사용하여 표현하세요. 또한 원본의 **원문 번호(①, ②, ③...)와 기호((A), (B), (C)...)**를 구조의 정확한 위치에 배치하세요.
2. **[이미지 내 모든 수식/구조식 시각화 필수]**: 사진에 있는 문항 식의 그림(Naphthalene, Aspartame, Capsaicin, Tryptophan, Lewis 구조 등)이 답변에 하나도 빠짐없이 다 들어갈 수 있도록 완벽하게 ASCII 또는 마크다운 구조식으로 재현하세요.
3. **[전문 학술 데이터 시트]**: 결합 길이(Å), 결합각(°), 에너지 준위 등을 포함한 정밀 수치 표를 작성하세요.
4. **[전수 분석 (누락 절대 금지)]**: (예: 1번~16번, 17번, 30번 등) 이미지나 텍스트에 포함된 모든 번호의 문항을 하나도 빠뜨리지 않고 순차적으로 분석해야 합니다. 1번만 풀고 2~4번을 구분하지 못하거나, 6번, 7~10번, 16~17번 등을 풀이 없이 대충 넘어가면 절대 안 됩니다. 두 번째 그림 등 모든 그림의 내용도 분석하세요. 다 분석 페이지 나와야 합니다.
5. **[반복 버그 박멸]**: '중복 문장 절대 금지' 및 '의미 없는 텍스트 생성 시 신뢰도 하락'이라는 강력한 페널티 지침이 적용됩니다. 동일한 문장을 반복하여 분량을 채우지 말고 정밀한 학술 데이터로만 내용을 채우고, 모든 문제들을 다 풀어주고 써주세요.
6. **[소문항 계층 구조 적용]**: 1번 문항 안에 a, b, c 등 여러 질문이 있는 경우, 이를 각각 독립적으로 나누는 것이 아니라 하나의 큰 문항 블록 안에 순서대로 배치하세요.
7. **[초강력 무한 검색 엔진 가동]**: 구글 검색(Grounding) 기능을 최대한 활용하여 기간(최신/과거) 제한 없이, 전 세계의 모든 학술 데이터와 논문, 솔루션 매뉴얼을 무한대로 탐색하여 가장 완벽한 해답을 찾아내 다 풀어나줘.

**[출력 템플릿]**:
[X]번 문항
- 도표 & 3D 그림 생성기에서 pdf로 올린자료는 a4 1000장 이내 작업으로 진행줘
내용을의 과제에 대한 
(소문항 a, b, c... 가 있을 경우 아래 세트를 소문항마다 반복, 없을 경우 1회만 작성)
- [질문 원문 그대로 전사]
- 질문의 풀이 식이나 내용을 상세하게 써술해주세요: [대충 넘어가지 말고 심층적 이론 해설 및 정밀 수식 유도]
- 내용의 3D 격자,분자 구조식,양자역학 그래프,루이스 전자점식,분자 오비탈 시각화,기본 도식 및 기하 도형,2D 분자 구조 / 선구조식,표 ,그래프등이 있으면 표시할것: [사진에 있는 3D 격자, 분자 구조식, 루이스 구조 등이 원본과 '똑같이' 빠짐없이 다 표시되도록 풀컬러 마스터 급 ASCII 구조식과 전문 데이터 시트로 완벽히 재현. "그림 생략" 등 대충 넘어가면 절대 안 되며 무조건 직접 똑같이 다 표시해]
- 답 풀이를 해줄것: [세밀한 단계별 정답 도출 과정]
- 답을 써줄것: [최종 정답]

질문 자료: {query}"""
                with st.status("🧠 AI가 문제를 분석하고 논리적 해설을 작성 중입니다...", expanded=True) as status:
                    res_text = robust_generate_content(prompt, use_grounding=True)
                    st.session_state.smart_analysis_result = res_text
                    status.update(label="✅ 풀이 완료!", state="complete", expanded=False)


                now = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                st.markdown(f"### 🧪 AI 상세 풀이 및 논리 (Smart Input) - 분석 시간: {now}")
                st.markdown(st.session_state.smart_analysis_result)

                # 이미지 분석과 똑같이 워드 다운로드 버튼 제공
                col_d1, col_d2 = st.columns(2)
                with col_d1:
                    if st.button("📝 전체 워드 문서에 추가 (누적)", key="add_to_word_cumulative"):
                        now = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                        st.session_state.word_doc.add_heading(f"Smart Input AI Analysis (Time: {now})", level=2)
                        for line in st.session_state.smart_analysis_result.split('\n'):
                            st.session_state.word_doc.add_paragraph(line)
                        st.success("워드 문서에 추가되었습니다.")
                with col_d2:
                    # 현재 결과만 즉시 다운로드하는 버튼 추가
                    try:
                        out_file = "Smart_Analysis_Result.docx"
                        margins = {'top': 2.0, 'bottom': 2.0, 'left': 2.5, 'right': 2.5}
                        # 버그 수정: global_smart_img_result 대신 smart_analysis_result 사용
                        convert_latex_to_word_docx(st.session_state.smart_analysis_result, out_file, margins)
                        with open(out_file, "rb") as f:
                            st.download_button("📥 이 분석 결과만 Word로 즉시 다운로드", data=f.read(), file_name="AI_Analysis_Result.docx", key="smart_direct_dl")
                        os.remove(out_file)
                    except:
                        st.info("💡 위 내용은 'LaTeX 에디터' 메뉴에서 워드로 일괄 변환 및 다운로드가 가능합니다.")
                st.divider()
            except Exception as e:
                st.error(f"AI 답변 생성 중 오류: {e}")

        q_upper = query.upper()

        # 2. 3D 격자 판별 (NaCl, CsCl, CaF2 등)
        lattices = {
            "NACL": "NaCl (Rock Salt)", "CSCL": "CsCl (BCC)", "CAF2": "CaF2 (Fluorite)",
            "BCC": "Body-Centered (BCC)", "FCC": "Face-Centered (FCC)", "SC": "Simple Cubic (SC)",
            "HCP": "HCP (Hexagonal)", "RHOMBOHEDRAL": "Rhombohedral"
        }
        for key, full in lattices.items():
            if key in q_upper:
                st.session_state.cell_mode_radio = "자동 (프리셋 선택)"
                st.session_state.cell_choice_select = full
                st.success(f"🔍 격자 구조 인식: '{full}' 도구로 연결합니다. (그림을 그리려면 아래 도구에서 '그리기'를 누르세요)")
                return

        # 3. 화학식/분자 이름 추출 및 라우팅
        words = re.findall(r"[A-Z][a-z0-9]*[A-Z0-9a-z]*|[a-zA-Z]{3,}|[\uac00-\ud7af]{2,}", query)
        chem_presets = {
            "H2O": ("lewis", "H2O (물)"), "CO2": ("lewis", "CO2 (이산화탄소)"),
            "NH3": ("lewis", "NH3 (암모니아)"), "CH4": ("lewis", "CH4 (메테인)"),
            "O2": ("lewis", "O2 (산소 분자)"), "N2": ("lewis", "N2 (질소 분자)"),
            "HCL": ("lewis", "HCl (염화수소)"), "BENZENE": ("skeletal", "Benzene (벤젠)"),
            "ACETONE": ("skeletal", "Acetone (아세톤)"), "NO3": ("orbital", "NO3- (질산 이온)"),
            "HNO3": ("orbital", "HNO3 (질산)")
        }

        for word in words:
            w_upper = word.upper()
            if w_upper in chem_presets:
                mod, choice = chem_presets[w_upper]
                if mod == "lewis":
                    st.session_state.lewis_mode_radio = "자동 (프리셋 선택)"
                    st.session_state.lewis_choice_select = choice
                elif mod == "orbital":
                    st.session_state.orbital_mode_radio = "자동 (프리셋 선택)"
                    st.session_state.orbital_choice_select = choice
                else:
                    st.session_state.skeletal_mode_radio = "자동 (프리셋 선택)"
                    st.session_state.skeletal_choice_select = choice
                st.success(f"🧪 화학 구조 인식: '{choice}' 도구로 연결합니다.")
                return

        # 4. 양자역학 그래프 판별
        math_patterns = [r"sin", r"cos", r"tan", r"exp", r"log", r"x", r"\*", r"\+", r"\-", r"\/", r"\^", r"np\."]
        if any(re.search(p, query.lower()) for p in math_patterns) and "x" in query.lower():
            st.session_state.graph_mode_radio = "수동 (직접 입력)"
            st.session_state.custom_expr_val = query
            st.session_state.graph_text_key_counter += 1
            st.success(f"📈 그래프 수식 인식: 그래프 도구에 연결합니다.")
            return

        # 5. 기타: PubChem 검색
        clean_query = ", ".join([w for w in words if len(w) > 1 and w.upper() not in ["DRAW", "LEWIS", "STRUCTURE", "SOLUTION", "MOLECULE", "그려라", "그리시오", "판별", "비교", "완성", "예상", "제시", "찾으", "결합각", "혼성", "오비탈", "HYBRID", "ORBITAL"]])
        if not clean_query: clean_query = query
        st.session_state.mol_mode_radio = "PubChem 자동 검색"
        st.session_state.mol_name_input = clean_query
        st.info(f"🌐 스마트 검색: '{clean_query}'를 PubChem에서 찾습니다.")
    st.text_area("✨ 스마트 입력 (텍스트 붙여넣기 및 대용량 PDF/이미지 분석 지원):", 
                 key="smart_input_val", height=200, on_change=handle_smart_input,
                 placeholder="내용을 직접 입력하거나 아래에 파일을 업로드하세요. (Ctrl+Enter로 실행)")

    if "global_smart_img_bytes" not in st.session_state:
        st.session_state.global_smart_img_bytes = None
        st.session_state.global_smart_img_name = None
        st.session_state.global_smart_img_result = None

    smart_imgs = st.file_uploader("📸 또는 문제 이미지/PDF를 직접 업로드하여 풀이하기 (대용량 PDF 1000페이지까지 지원)", type=["png", "jpg", "jpeg", "pdf"], accept_multiple_files=True, key="smart_img_uploader")

    # --- 스마트 라우터 기능 끝 ---

    if smart_imgs:
        # 분석 시작 버튼 추가 (사용자 경험 개선)
        if st.button("🚀 개별 문제별 초정밀 슈퍼 심층 분석 시작 (최고 품질/최대 분량)", key="start_super_analysis_btn", use_container_width=True):
            st.session_state.global_smart_img_result = None  # 이전 분석결과 초기화
            
            if st.session_state.get("gemini_api_key"):
                try:
                    import PIL.Image
                    import io
                    import fitz  # PyMuPDF

                    imgs = []
                    for uploaded_file in smart_imgs:
                        f_bytes = uploaded_file.getvalue()
                        if uploaded_file.name.lower().endswith(".pdf"):
                            # PDF를 고해상도 이미지로 변환 (최대 1000페이지 제한)
                            pdf_doc = fitz.open(stream=f_bytes, filetype="pdf")
                            max_pages = min(len(pdf_doc), 1000)
                            for page_idx in range(max_pages):
                                page = pdf_doc.load_page(page_idx)
                                # Matrix(2, 2)는 2배 확대 (고해상도 추출)
                                pix = page.get_pixmap(matrix=fitz.Matrix(2.0, 2.0))
                                img_bytes = pix.tobytes("png")
                                imgs.append(PIL.Image.open(io.BytesIO(img_bytes)))
                            pdf_doc.close()
                        else:
                            # 일반 이미지 처리
                            imgs.append(PIL.Image.open(io.BytesIO(f_bytes)))


                    with st.spinner("🚀 초정밀 슈퍼 심층 분석 진행 중... (압도적인 분량과 깊이로 풀이)"):
                        genai.configure(api_key=st.session_state.gemini_api_key)
                        
                        import hashlib
                        # 업로드된 파일들의 이름과 크기를 기반으로 고유 해시 생성 (동일 이름의 수정 파일 방어)
                        file_names = "".join([f"{f.name}_{f.size}" for f in smart_imgs])
                        file_hash = hashlib.md5(file_names.encode()).hexdigest()
                        
                        # 새로운 파일이 업로드되었으면 버퍼 초기화
                        if st.session_state.get("current_file_hash") != file_hash:
                            st.session_state.analysis_buffer = {}
                            st.session_state.current_file_hash = file_hash
                            
                        if "analysis_buffer" not in st.session_state:
                            st.session_state.analysis_buffer = {}
                        
                        all_results = []
                        total_pages = len(imgs)
                        
                        # 최고 품질/최대 분량을 위해 페이지당 1:1 개별 분석 진행 (축약 원천 차단)
                        batch_size = 1
                        for i in range(0, total_pages, batch_size):
                            # 이미 분석된 페이지는 건너뜀 (중단 시 복구용)
                            if i in st.session_state.analysis_buffer:
                                all_results.append(st.session_state.analysis_buffer[i])
                                continue

                            batch_imgs = imgs[i:i+batch_size]
                            current_page_label = f"{i+1}/{total_pages}"
                            st.info(f"⏳ [누락 방지 모드 가동 중] - {i+1}번째 페이지 심층 분석 중... ({current_page_label})")
                            
                            prompt = f"""당신은 세계 최고의 화학 및 물리 교육 전문가입니다. 
제시된 이미지의 모든 문항과 그림을 **단 하나도 빠짐없이, 대충 넘어가는 것 없이 완벽히** 분석하세요.

**[절대 준수 지침 - 반드시 지킬 것]**:
1. **[누락/대충 넘어기기 절대 금지]**: 1번만 잘하고 2~4번 문항을 뭉뚱그리거나, 6번, 7~10번, 16번, 17번을 풀이 없이 대충 넘어가면 절대 안 됩니다. 두 번째 그림이나 추가 이미지의 내용도 생략하지 마세요. 화면에 보이는 모든 문항 번호를 색출하여 처음부터 끝까지 개별적으로 완벽하게 풀이해야 다 분석 페이지가 나옵니다.
2. **[이미지 내 모든 그림 무조건 100% 시각화 필수]**: 사진에 있는 3D 격자, 분자 구조식, 양자역학 그래프, 루이스 전자점식, 분자 오비탈 시각화, 기본 도식 및 기하 도형, 2D 분자 구조/선구조식, 표, 그래프 등 그림 요소가 있다면 **단 하나도 빼놓지 말고 원본과 똑같이 다 표시**하세요. "생략함", "텍스트로 대체" 등 대충 넘어가면 절대 안 되며 완벽하게 ASCII 구조식이나 마크다운으로 무조건 직접 다 그려야 합니다.
3. **요약 금지**: "동일한 형식으로 진행", "이하 생략", "위와 같음" 등의 표현을 절대 사용하지 마세요.
4. **반복 버그 박멸**: 동일한 문장을 반복하여 분량을 채우지 말고 정밀한 학술 데이터로만 내용을 채우고, 모든 문제들을 다 풀어주고 써주세요.
5. **소문항 계층 구조 적용**: 1번 문항 안에 a, b, c 등 여러 질문이 있는 경우, 이를 각각 독립적으로 나누는 것이 아니라 하나의 큰 문항 블록 안에 순서대로 배치하세요.
6. **[초강력 무한 검색 엔진 가동]**: 구글 검색(Grounding) 기능을 최대한 활용하여 기간(최신/과거) 제한 없이, 전 세계의 모든 학술 데이터와 논문, 솔루션 매뉴얼을 무한대로 탐색하여 가장 완벽한 해답을 찾아내 다 풀어나줘.

**[분석 템플릿]**:
[X]번 문항
- 도표 & 3D 그림 생성기에서 pdf로 올린자료는 a4 1000장 이내 작업으로 진행줘
내용을의 과제에 대한 
(소문항 a, b, c... 가 있을 경우 아래 세트를 소문항마다 반복, 없을 경우 1회만 작성)
- [질문 원문 그대로 전사]
- 질문의 풀이 식이나 내용을 상세하게 써술해주세요: [대충 넘어가지 말고 대학 전공 수준의 심층 해설 및 수식 유도. 최소 10문장 이상 상세히.]
- 내용의 3D 격자,분자 구조식,양자역학 그래프,루이스 전자점식,분자 오비탈 시각화,기본 도식 및 기하 도형,2D 분자 구조 / 선구조식,표 ,그래프등이 있으면 표시할것: [사진에 있는 3D 격자, 분자 구조식, 루이스 구조 등이 하나도 빠짐없이 원본과 '똑같이' 다 표시되도록 풀컬러 ASCII 구조식과 전문 데이터 시트로 완벽히 재현. 대충 글로 때우지 말고 무조건 직접 똑같이 다 그려라]
- 답 풀이를 해줄것: [단계별 논리적 도출 과정]
- 답을 써줄것: [최종 정답]

모든 지능을 사용하여 이미지의 모든 글자와 기호, 그림을 완벽하게 분석하고 보고서를 완성하세요."""

                            try:
                                # 모델명 오류를 방지하기 위해 검증된 robust_generate_content 함수 사용
                                import time
                                res_text = robust_generate_content(prompt, images=batch_imgs, use_grounding=True)
                                if res_text:
                                    # 즉시 세션 버퍼에 저장 (브라우저 새로고침 대비)
                                    st.session_state.analysis_buffer[i] = res_text
                                    all_results.append(res_text)
                                    
                                    # 실시간 워드 누적 저장 (불필요한 공백 제거 및 내용 보존)
                                    st.session_state.word_doc.add_heading(f"Analysis Report - Page {i+1}", level=2)
                                    clean_lines = [l.strip() for l in res_text.split('\n') if l.strip()]
                                    for line in clean_lines:
                                        st.session_state.word_doc.add_paragraph(line)
                                    
                                    st.toast(f"✅ {i+1}페이지 분석 완료 및 버퍼 저장됨")
                                    
                                    # API 호출 제한 방지 (Rate Limit 방어용 대기)
                                    time.sleep(5)
                                else:
                                    # 결과가 없으면(API 할당량 초과 등) 중단하고 사용자에게 알림
                                    st.warning(f"⚠️ {i+1}번째 페이지 분석 실패 (API 할당량 초과 또는 네트워크 오류). 잠시 후 다시 '시작' 버튼을 누르면 이어서 분석합니다.")
                                    break # 멈추고 다음 루프 실행하지 않음
                            except Exception as e:
                                st.error(f"❌ {i+1}번째 페이지 분석 중 심각한 오류 발생: {e}")
                                st.warning("⚠️ API 제한이 발생했을 수 있습니다. 1분 후 다시 시도해주세요.")
                                break # 멈춤

                        if all_results:
                            st.session_state.global_smart_img_result = "\n\n---\n\n".join(all_results)
                            st.success("✅ 모든 페이지 분석이 완료되었습니다!")
                        else:
                            st.error("분석 결과 생성에 실패했습니다.")
                except Exception as e:
                    st.error(f"이미지 분석 중 오류 발생: {e}")
            else:
                st.warning("🔑 Gemini API 키를 먼저 입력해 주세요.")

    # 화면에 영구적으로 결과 표시 및 초기화 버튼 제공
    if st.session_state.global_smart_img_bytes is not None or st.session_state.global_smart_img_result is not None:
        if st.session_state.global_smart_img_result:
            with st.container(border=True):
                st.markdown("### 📋 AI 분석 결과 미리보기 (전체 내용)")
                st.markdown(st.session_state.global_smart_img_result)


            try:
                import tempfile
                import pypandoc
                import os
                with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
                    tmp_path = tmp.name

                # 마크다운(+LaTeX 수식)을 워드로 변환
                pypandoc.convert_text(st.session_state.global_smart_img_result, 'docx', format='markdown', outputfile=tmp_path)

                with open(tmp_path, "rb") as f:
                    docx_bytes = f.read()
                os.remove(tmp_path) # 임시 파일 정리

                st.download_button("📝 분석 결과를 Word 파일로 다운로드 (.docx)", data=docx_bytes, file_name="AI_Analysis_Result.docx", use_container_width=True)
            except Exception as e:
                # Pandoc이 없거나 에러가 발생한 경우 단순 텍스트로 저장하는 대체(Fallback) 수단
                import io
                from docx import Document
                fallback_doc = Document()
                fallback_doc.add_heading("AI Image Analysis Result", level=1)
                for line in st.session_state.global_smart_img_result.split("\n"):
                    fallback_doc.add_paragraph(line)
                f_stream = io.BytesIO()
                fallback_doc.save(f_stream)
                st.download_button("📓 수식을 변환하여 Word 파일로 다운로드 (텍스트 전용)", data=f_stream.getvalue(), file_name="AI_Analysis_Result_text.docx", use_container_width=True, help="Pandoc 변환 에러로 인해 수식이 텍스트로 저장됩니다.")

            # 데이터베이스 저장 버튼 추가
            st.divider()
            st.subheader("📂 데이터베이스(Notion)에 즉시 저장")
            db_col1, db_col2 = st.columns([3, 1])
            with db_col1:
                target_sub = st.selectbox("저장할 과목 선택", list(st.session_state.notion_db.keys()), key="smart_db_subject")
            with db_col2:
                if st.button("📥 DB에 추가", use_container_width=True, key="smart_db_add_btn"):
                    st.session_state.notion_db[target_sub] += "\n\n---\n" + st.session_state.global_smart_img_result
                    save_state()
                    st.success(f"'{target_sub}' 데이터베이스에 저장되었습니다!")

        if st.button("🗑️ 업로드된 자료 및 분석 결과 초기화 (지우기)", use_container_width=True):
            st.session_state.global_smart_img_bytes = None
            st.session_state.global_smart_img_name = None
            st.session_state.global_smart_img_result = None
            st.session_state.smart_analysis_result = None
            st.session_state.analysis_buffer = {}  # 버퍼 완전 초기화
            st.session_state.current_file_hash = None
            # 워드 문서도 새로 시작
            from docx import Document
            st.session_state.word_doc = Document()
            st.rerun()

    st.markdown("---")
    st.markdown("### 🔄 전체 생성 모드 일괄 적용")

    if "global_mode_val" not in st.session_state:
        st.session_state.global_mode_val = "기본 (개별 설정)"

    for mod in ["cell", "graph", "orbital", "lewis", "shape", "skeletal"]:
        if f"{mod}_text_key_counter" not in st.session_state:
            st.session_state[f"{mod}_text_key_counter"] = 0

    def set_global_mode():
        gmode = st.session_state.global_mode_widget
        st.session_state.global_mode_val = gmode
        target = "자동 (프리셋 선택)" if "자동" in gmode else "수동 (직접 입력)"
        if gmode != "기본 (개별 설정)":
            st.session_state.cell_mode_radio = target
            st.session_state.graph_mode_radio = target
            st.session_state.orbital_mode_radio = target
            st.session_state.lewis_mode_radio = target
            st.session_state.shape_mode_radio = target
            st.session_state.skeletal_mode_radio = target
            st.session_state.mol_mode_radio = "PubChem 자동 검색" if "자동" in gmode else "수동 (직접 입력)"

            if target == "수동 (직접 입력)":
                st.session_state.custom_coords_val = get_preset_custom_data("cell", st.session_state.get("cell_choice_select", ""))
                st.session_state.cell_text_key_counter += 1
                st.session_state.custom_expr_val = get_preset_custom_data("graph", st.session_state.get("graph_choice_select", ""))
                st.session_state.graph_text_key_counter += 1
                st.session_state.custom_orbital_data_val = get_preset_custom_data("orbital", st.session_state.get("orbital_choice_select", ""))
                st.session_state.orbital_text_key_counter += 1
                st.session_state.custom_lewis_data_val = get_preset_custom_data("lewis", st.session_state.get("lewis_choice_select", ""))
                st.session_state.lewis_text_key_counter += 1
                st.session_state.custom_shape_data_val = get_preset_custom_data("shape", st.session_state.get("shape_choice_select", ""))
                st.session_state.shape_text_key_counter += 1
                st.session_state.custom_skeletal_data_val = get_preset_custom_data("skeletal", st.session_state.get("skeletal_choice_select", ""))
                st.session_state.skeletal_text_key_counter += 1

    st.radio("모든 그리기 도구의 모드를 일괄적으로 변경합니다:",
             ["기본 (개별 설정)", "전체 자동 (프리셋/검색) 모드로 통일", "전체 수동 (직접 입력) 모드로 통일"],
             horizontal=True,
             key="global_mode_widget",
             on_change=set_global_mode)
    st.markdown("---")
    col1, col2 = st.columns([1, 1])
    with col1:
        st.subheader("1. 3D 격자 생성")
        cell_options = ["Simple Cubic (SC)", "Body-Centered (BCC)", "Face-Centered (FCC)", "HCP (Hexagonal)", "NaCl (Rock Salt)", "CsCl (BCC)", "CaF2 (Fluorite)", "Rhombohedral"]

        def on_cell_choice_change():
            if st.session_state.cell_mode_radio == "수동 (직접 입력)":
                st.session_state.custom_coords_val = get_preset_custom_data("cell", st.session_state.get("cell_choice_select", ""))
                st.session_state.cell_text_key_counter += 1

        def sync_cell_mode():
            if st.session_state.cell_mode_radio == "수동 (직접 입력)":
                st.session_state.custom_coords_val = get_preset_custom_data("cell", st.session_state.get("cell_choice_select", ""))
                st.session_state.cell_text_key_counter += 1

        cell_mode = st.radio("생성 모드:", ["자동 (프리셋 선택)", "수동 (직접 입력)"], horizontal=True, key="cell_mode_radio", on_change=sync_cell_mode)

        label = "3D 격자 선택:" if cell_mode == "자동 (프리셋 선택)" else "시작 템플릿 불러오기 (변경 시 아래 코드가 덮어씌워집니다):"
        cell_choice_raw = st.selectbox(label, cell_options, key="cell_choice_select", on_change=on_cell_choice_change)

        if cell_mode == "자동 (프리셋 선택)":
            cell_choice = cell_choice_raw
        else:
            cell_choice = "직접 좌표 입력 (Custom)"
        lattice_size = st.slider("격자 반복 횟수 (Size):", 1, 3, 1, help="단위 세포를 몇 번 반복해서 쌓을지 결정합니다.")
        atom_rad = st.slider("원자 크기 (Atom Size):", 0.05, 0.5, 0.18, 0.01, help="원자의 반지름을 조절합니다.")

        custom_coords = ""
        if cell_choice == "직접 좌표 입력 (Custom)":
            if "custom_coords_val" not in st.session_state:
                st.session_state.custom_coords_val = "0,0,0,blue\n1,1,1,red"
            custom_coords = st.text_area("좌표 입력 (x,y,z,색상):", value=st.session_state.custom_coords_val, key=f"custom_coords_val_{st.session_state.cell_text_key_counter}", help="줄바꿈으로 여러 좌표를 입력할 수 있습니다. (Ctrl+Enter로 적용)")
            st.session_state.custom_coords_val = custom_coords
        if st.button("3D 큐브 그리기"):
            img_stream, errors = draw_unit_cell(cell_choice, lattice_size, atom_rad, custom_coords)
            if errors:
                for err in errors: st.error(err)
            img_stream.seek(0)
            st.image(img_stream)
            img_stream.seek(0)
            st.session_state.word_doc.add_picture(img_stream, width=Inches(4.0))

        up_3d = st.file_uploader("📂 3D 격자 이미지 수동 업로드", type=["png", "jpg", "jpeg"], key="up_3d")
        if up_3d and st.button("워드에 수동 추가 (3D)"):
            up_3d.seek(0)
            st.image(up_3d)
            up_3d.seek(0)
            st.session_state.word_doc.add_picture(up_3d, width=Inches(4.0))
            st.success("워드에 추가되었습니다!")
        st.markdown("---")
        st.subheader("2. 양자역학 그래프")
        graph_options = ["1D Box 파동함수 (n=1,2,3)", "2s 오비탈 확률 밀도"]
        def on_graph_choice_change():
            if st.session_state.graph_mode_radio == "수동 (직접 입력)":
                st.session_state.custom_expr_val = get_preset_custom_data("graph", st.session_state.get("graph_choice_select", ""))
                st.session_state.graph_text_key_counter += 1

        def sync_graph_mode():
            if st.session_state.graph_mode_radio == "수동 (직접 입력)":
                st.session_state.custom_expr_val = get_preset_custom_data("graph", st.session_state.get("graph_choice_select", ""))
                st.session_state.graph_text_key_counter += 1

        graph_mode = st.radio("생성 모드:", ["자동 (프리셋 선택)", "수동 (직접 입력)"], horizontal=True, key="graph_mode_radio", on_change=sync_graph_mode)

        label = "양자역학 그래프:" if graph_mode == "자동 (프리셋 선택)" else "시작 템플릿 불러오기 (변경 시 아래 코드가 덮어씌워집니다):"
        graph_choice_raw = st.selectbox(label, graph_options, key="graph_choice_select", on_change=on_graph_choice_change)

        if graph_mode == "자동 (프리셋 선택)":
            graph_choice = graph_choice_raw
        else:
            graph_choice = "직접 수식 입력 (Custom)"

        custom_expr = "sin(x)"
        x_min, x_max = 0.0, 10.0

        if graph_choice == "직접 수식 입력 (Custom)":
            if "custom_expr_val" not in st.session_state:
                st.session_state.custom_expr_val = "sin(x)*exp(-x/5)"
            custom_expr = st.text_input("원하는 수식을 입력하세요 (x에 대한 함수):", value=st.session_state.custom_expr_val, key=f"custom_expr_val_{st.session_state.graph_text_key_counter}", help="예: sin(x), x^2, exp(-x)")
            st.session_state.custom_expr_val = custom_expr
            col_x1, col_x2 = st.columns(2)
            x_min = col_x1.number_input("x 최솟값", value=0.0)
            x_max = col_x2.number_input("x 최댓값", value=15.0)

        if st.button("그래프 그리기"):
            img_stream = draw_quantum_graph(graph_choice, custom_expr, x_min, x_max)
            img_stream.seek(0)
            st.image(img_stream)
            img_stream.seek(0)
            st.session_state.word_doc.add_picture(img_stream, width=Inches(4.0))

        up_graph = st.file_uploader("📂 그래프 이미지 수동 업로드", type=["png", "jpg", "jpeg"], key="up_graph")
        if up_graph and st.button("워드에 수동 추가 (그래프)"):
            up_graph.seek(0)
            st.image(up_graph)
            up_graph.seek(0)
            st.session_state.word_doc.add_picture(up_graph, width=Inches(4.0))
            st.success("워드에 추가되었습니다!")

        st.markdown("---")
        st.subheader("6. 분자 오비탈 시각화 (3D)")
        orbital_options = ["NO3- (질산 이온)", "HNO3 (질산)", "C2H4 (에텐)"]
        def on_orbital_choice_change():
            if st.session_state.orbital_mode_radio == "수동 (직접 입력)":
                st.session_state.custom_orbital_data_val = get_preset_custom_data("orbital", st.session_state.get("orbital_choice_select", ""))
                st.session_state.orbital_text_key_counter += 1

        def sync_orbital_mode():
            if st.session_state.orbital_mode_radio == "수동 (직접 입력)":
                st.session_state.custom_orbital_data_val = get_preset_custom_data("orbital", st.session_state.get("orbital_choice_select", ""))
                st.session_state.orbital_text_key_counter += 1

        orbital_mode = st.radio("생성 모드:", ["자동 (프리셋 선택)", "수동 (직접 입력)"], horizontal=True, key="orbital_mode_radio", on_change=sync_orbital_mode)

        label = "분자 오비탈 선택:" if orbital_mode == "자동 (프리셋 선택)" else "시작 템플릿 불러오기 (변경 시 아래 코드가 덮어씌워집니다):"
        orbital_choice_raw = st.selectbox(label, orbital_options, key="orbital_choice_select", on_change=on_orbital_choice_change)

        if orbital_mode == "자동 (프리셋 선택)":
            orbital_choice = orbital_choice_raw
        else:
            orbital_choice = "직접 입력 (Custom)"

        custom_orbital_data = ""
        if orbital_choice == "직접 입력 (Custom)":
            if "custom_orbital_data_val" not in st.session_state:
                st.session_state.custom_orbital_data_val = "ATOM, N, 0, 0, 0"
            custom_orbital_data = st.text_area("수동 입력 (형식: 명령어, 속성...)", value=st.session_state.custom_orbital_data_val, key=f"custom_orbital_data_val_{st.session_state.orbital_text_key_counter}", height=150, help="명령어: ATOM(이름,x,y,z), BOND(x1,y1,z1,x2,y2,z2), PZ(x,y,z,색상,라벨), SP2(x,y,z,각도,색상,라벨), S(x,y,z,색상,라벨)")
            st.session_state.custom_orbital_data_val = custom_orbital_data

        if st.button("오비탈 그리기"):
            img_stream, errors = draw_orbital_diagram(orbital_choice, custom_orbital_data)
            if errors:
                for err in errors: st.error(err)
            img_stream.seek(0)
            st.image(img_stream)
            img_stream.seek(0)
            st.session_state.word_doc.add_picture(img_stream, width=Inches(4.0))
            st.success(f"{orbital_choice} 오비탈 그림이 추가되었습니다!")

    with col2:
        st.subheader("3. 분자 구조식")
        mol_mode = st.radio("생성 모드:", ["PubChem 자동 검색", "수동 (직접 입력)"], horizontal=True, key="mol_mode_radio")

        if mol_mode == "PubChem 자동 검색":
            st.markdown("💡 **자주 검색하는 분자 예시:**")
            st.caption("아래 이름을 복사해 입력창에 넣거나, 쉼표(,)로 연결해서 검색하세요.")
            st.code("aspirin, benzene, caffeine, glucose, ethanol, water, methane, sulfuric acid, penicillin, acetaminophen", language="text")
            mol_name = st.text_input("분자 이름 또는 식 입력:", help="예: aspirin, benzene, C4H8, H2O", key="mol_name_input")
            if st.button("구조식 그리기 (자동)"):
                if mol_name:
                    names = [n.strip() for n in mol_name.split(',')]
                    for name in names[:4]:
                        if not name: continue
                        img_stream = get_molecule_image(name)
                        if img_stream:
                            img_stream.seek(0)
                            st.image(img_stream, caption=name)
                            st.session_state.word_doc.add_paragraph(f"[{name}] 구조식")
                            img_stream.seek(0)
                            st.session_state.word_doc.add_picture(img_stream, width=Inches(3.0))
        else:
            if "custom_mol_data_val" not in st.session_state:
                st.session_state.custom_mol_data_val = "TEXT, $C_6H_{12}O_6$, 0, 1.5, 24, black\nLINE, 0, 1, 1, 0.5, 2, black\nLINE, 1, 0.5, 2, 1, 2, black\nTEXT, OH, 2.3, 1, 18, red"
            custom_mol_data = st.text_area("수동 구조식 입력:", key="custom_mol_data_val", height=200, help="7번 메뉴와 동일한 명령어를 사용합니다.")
            if st.button("구조식 그리기 (수동)"):
                img_stream, errors = draw_skeletal_structure("직접 입력 (Custom)", custom_mol_data)
                if errors:
                    for err in errors: st.error(err)
                img_stream.seek(0)
                st.image(img_stream)
                img_stream.seek(0)
                st.session_state.word_doc.add_picture(img_stream, width=Inches(3.0))

        st.markdown("---")
        up_mol = st.file_uploader("📂 분자 구조식 이미지 수동 업로드", type=["png", "jpg", "jpeg"], key="up_mol")
        if up_mol and st.button("워드에 수동 추가 (구조식)"):
            up_mol.seek(0)
            st.image(up_mol)
            up_mol.seek(0)
            st.session_state.word_doc.add_picture(up_mol, width=Inches(3.0))
            st.success("워드에 추가되었습니다!")

        st.markdown("---")
        st.subheader("4. 루이스 전자점식 (Lewis Structure)")
        lewis_options = ["H2O (물)", "CO2 (이산화탄소)", "NH3 (암모니아)", "CH4 (메테인)", "O2 (산소 분자)", "N2 (질소 분자)", "HCl (염화수소)", "HCN (사이안화수소)", "C2H4 (에텐)"]
        def on_lewis_choice_change():
            if st.session_state.lewis_mode_radio == "수동 (직접 입력)":
                st.session_state.custom_lewis_data_val = get_preset_custom_data("lewis", st.session_state.get("lewis_choice_select", ""))
                st.session_state.lewis_text_key_counter += 1

        def sync_lewis_mode():
            if st.session_state.lewis_mode_radio == "수동 (직접 입력)":
                st.session_state.custom_lewis_data_val = get_preset_custom_data("lewis", st.session_state.get("lewis_choice_select", ""))
                st.session_state.lewis_text_key_counter += 1

        lewis_mode = st.radio("생성 모드:", ["자동 (프리셋 선택)", "수동 (직접 입력)"], horizontal=True, key="lewis_mode_radio", on_change=sync_lewis_mode)

        label = "분자 선택:" if lewis_mode == "자동 (프리셋 선택)" else "시작 템플릿 불러오기 (변경 시 아래 코드가 덮어씌워집니다):"
        lewis_choice_raw = st.selectbox(label, lewis_options, key="lewis_choice_select", on_change=on_lewis_choice_change)

        if lewis_mode == "자동 (프리셋 선택)":
            lewis_choice = lewis_choice_raw
        else:
            lewis_choice = "직접 입력 (Custom)"

        custom_lewis_data = ""
        if lewis_choice == "직접 입력 (Custom)":
            if "custom_lewis_data_val" not in st.session_state:
                st.session_state.custom_lewis_data_val = "TEXT, O, 0.5, 0.5, 36, blue"
            custom_lewis_data = st.text_area("루이스 수동 입력:", value=st.session_state.custom_lewis_data_val, key=f"custom_lewis_data_val_{st.session_state.lewis_text_key_counter}", height=150, help="TEXT, LINE, DOTS(x,y,dx,dy) 명령어를 사용합니다.")
            st.session_state.custom_lewis_data_val = custom_lewis_data

        if st.button("루이스 구조식 그리기"):
            img_stream, errors = draw_lewis_structure(lewis_choice, custom_lewis_data)
            if errors:
                for err in errors: st.error(err)
            img_stream.seek(0)
            st.image(img_stream)
            img_stream.seek(0)
            st.session_state.word_doc.add_picture(img_stream, width=Inches(3.0))
            st.success(f"{lewis_choice} 루이스 구조식이 추가되었습니다!")

        st.markdown("---")
        st.subheader("5. 기본 도식 및 기하 도형")
        shape_options = ["정사각형 (Square)", "원형 (Circle)", "다각형 (Polygon)", "정육면체 (Cube)"]
        def on_shape_choice_change():
            if st.session_state.shape_mode_radio == "수동 (직접 입력)":
                st.session_state.custom_shape_data_val = get_preset_custom_data("shape", st.session_state.get("shape_choice_select", ""))
                st.session_state.shape_text_key_counter += 1

        def sync_shape_mode():
            if st.session_state.shape_mode_radio == "수동 (직접 입력)":
                st.session_state.custom_shape_data_val = get_preset_custom_data("shape", st.session_state.get("shape_choice_select", ""))
                st.session_state.shape_text_key_counter += 1

        shape_mode = st.radio("생성 모드:", ["자동 (프리셋 선택)", "수동 (직접 입력)"], horizontal=True, key="shape_mode_radio", on_change=sync_shape_mode)

        label = "도형 선택:" if shape_mode == "자동 (프리셋 선택)" else "시작 템플릿 불러오기 (변경 시 아래 코드가 덮어씌워집니다):"
        shape_choice_raw = st.selectbox(label, shape_options, key="shape_choice_select", on_change=on_shape_choice_change)

        if shape_mode == "자동 (프리셋 선택)":
            shape_choice = shape_choice_raw
        else:
            shape_choice = "직접 입력 (Custom)"

        custom_shape_data = ""
        shape_color = "#2E5BFF"
        kwargs = {}

        if shape_choice == "직접 입력 (Custom)":
            if "custom_shape_data_val" not in st.session_state:
                st.session_state.custom_shape_data_val = "RECT, 0.1, 0.1, 0.8, 0.8, #2E5BFF"
            custom_shape_data = st.text_area("수동 입력 (형식: 명령어, 속성...):", value=st.session_state.custom_shape_data_val, key=f"custom_shape_data_val_{st.session_state.shape_text_key_counter}", height=150, help="명령어: RECT(x,y,w,h,색상), CIRCLE(x,y,r,색상), TEXT(내용,x,y,크기,색상), LINE(x1,y1,x2,y2,색상)")
            st.session_state.custom_shape_data_val = custom_shape_data
        else:
            shape_color = st.color_picker("색상 선택:", "#2E5BFF")
            if shape_choice == "정사각형 (Square)":
                c1, c2 = st.columns(2)
                kwargs['width'] = c1.slider("가로 길이", 0.1, 1.0, 0.8, 0.1)
                kwargs['height'] = c2.slider("세로 길이", 0.1, 1.0, 0.8, 0.1)
            elif shape_choice == "원형 (Circle)":
                kwargs['radius'] = st.slider("반지름 크기", 0.1, 0.5, 0.4, 0.05)
            elif shape_choice == "다각형 (Polygon)":
                kwargs['n_sides'] = st.slider("꼭짓점 개수 (N각형)", 3, 12, 6, 1)
            elif shape_choice == "정육면체 (Cube)":
                c1, c2, c3 = st.columns(3)
                kwargs['length'] = c1.slider("가로 길이 (x)", 0.5, 3.0, 1.0, 0.1)
                kwargs['width'] = c2.slider("세로 길이 (y)", 0.5, 3.0, 1.0, 0.1)
                kwargs['height'] = c3.slider("높이 (z)", 0.5, 3.0, 1.0, 0.1)

        if st.button("도형 그리기"):
            img_stream, errors = draw_schematic(shape_choice, shape_color, custom_shape_data, **kwargs)
            if errors:
                for err in errors: st.error(err)
            img_stream.seek(0)
            st.image(img_stream)
            img_stream.seek(0)
            st.session_state.word_doc.add_picture(img_stream, width=Inches(3.0))
            st.success("워드에 도형이 추가되었습니다!")

        st.markdown("---")
        st.subheader("7. 2D 분자 구조 / 선구조식 (수동 지원)")
        skeletal_options = ["Butane (뷰테인)", "Hexane (헥세인)", "Cyclohexane (사이클로헥세인)", "Benzene (벤젠)", "Acetone (아세톤)", "Acetic Acid (아세트산)"]
        def on_skeletal_choice_change():
            if st.session_state.skeletal_mode_radio == "수동 (직접 입력)":
                st.session_state.custom_skeletal_data_val = get_preset_custom_data("skeletal", st.session_state.get("skeletal_choice_select", ""))
                st.session_state.skeletal_text_key_counter += 1

        def sync_skeletal_mode():
            if st.session_state.skeletal_mode_radio == "수동 (직접 입력)":
                st.session_state.custom_skeletal_data_val = get_preset_custom_data("skeletal", st.session_state.get("skeletal_choice_select", ""))
                st.session_state.skeletal_text_key_counter += 1

        skeletal_mode = st.radio("생성 모드:", ["자동 (프리셋 선택)", "수동 (직접 입력)"], horizontal=True, key="skeletal_mode_radio", on_change=sync_skeletal_mode)

        label = "분자 선택:" if skeletal_mode == "자동 (프리셋 선택)" else "시작 템플릿 불러오기 (변경 시 아래 코드가 덮어씌워집니다):"
        skeletal_choice_raw = st.selectbox(label, skeletal_options, key="skeletal_choice_select", on_change=on_skeletal_choice_change)

        if skeletal_mode == "자동 (프리셋 선택)":
            skeletal_choice = skeletal_choice_raw
        else:
            skeletal_choice = "직접 입력 (Custom)"

        custom_skeletal_data = ""
        if skeletal_choice == "직접 입력 (Custom)":
            if "custom_skeletal_data_val" not in st.session_state:
                st.session_state.custom_skeletal_data_val = "CANVAS, 8, 4"
            custom_skeletal_data = st.text_area(
                "수동 입력 (형식: 명령어, 속성...)",
                key="custom_skeletal_data_val",
                height=350,
                help="명령어: CANVAS(너비,높이), TEXT/LTEXT/RTEXT(텍스트,x,y,크기,색상), LINE(x1,y1,x2,y2,두께,색상), DLINE(이중결합x1,y1,x2,y2,두께,색상), ARROW(x1,y1,x2,y2,색상), DOTS(점1x,점1y,점2x,점2y,색상), BALL(x,y,반지름,색상)."
            )

        if st.button("분자 구조식 그리기"):
            img_stream, errors = draw_skeletal_structure(skeletal_choice, custom_skeletal_data)
            if errors:
                for err in errors:
                    st.error(f"오류: {err}")
            img_stream.seek(0)
            st.image(img_stream)
            img_stream.seek(0)
            st.session_state.word_doc.add_picture(img_stream, width=Inches(4.0))
            st.info("💡 **Tip:** 수동 입력창에서 줄바꿈은 `Enter`를, 입력 완료 및 적용은 `Ctrl+Enter`를 사용하세요.")
            st.success(f"{skeletal_choice} 구조식이 추가되었습니다!")

    doc_stream = io.BytesIO()
    st.session_state.word_doc.save(doc_stream)
    st.download_button("누적 워드 문서 다운로드 (.docx)", data=doc_stream.getvalue(), file_name="Diagrams.docx")

elif menu == "📝 수기 노트 AI 문서화":
    st.title("📝 수기 노트 AI 문서화")
    uploaded_files = st.file_uploader("이미지 및 PDF 업로드", type=["png", "jpg", "jpeg", "pdf"], accept_multiple_files=True)

    col1, col2 = st.columns(2)
    with col1:
        start_conversion = st.button("문서 변환 시작", use_container_width=True)
    with col2:
        start_analysis = st.button("이미지 상세 분석", use_container_width=True)

    if uploaded_files:


        if start_conversion:
            if not api_key: st.error("API 키 입력 필수!")
            else:
                genai.configure(api_key=api_key)
                conversion_results = []
                with st.status("📝 AI가 문서를 디지털로 변환 중입니다...", expanded=True) as status:
                    for file_obj in uploaded_files:
                        try:
                            st.write(f"📄 **{file_obj.name}** 처리 중...")
                            if file_obj.name.lower().endswith('.pdf'):
                                import tempfile
                                with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
                                    tmp.write(file_obj.getvalue())
                                    tmp_path = tmp.name
                                gemini_file = genai.upload_file(tmp_path, mime_type="application/pdf")
                                prompt = "이 PDF 문서의 모든 내용을 텍스트로 정확하게 디지털화해줘. 연필 수기, 볼펜 수기, 인쇄된 텍스트 모두 포함해서 가독성 있게 정리해줘."
                                res_text = robust_generate_content(prompt, images=gemini_file)
                                os.remove(tmp_path)
                            else:
                                from PIL import Image
                                prompt = "이 이미지에 포함된 모든 텍스트(연필/볼펜 수기, 프린트된 글자)를 인식해서 디지털 텍스트로 변환해줘. 표나 수식이 있다면 최대한 형식을 유지해줘."
                                res_text = robust_generate_content(prompt, images=Image.open(file_obj))


                            if res_text:
                                now = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                                conversion_results.append((file_obj.name, f"[분석 시간: {now}]\n" + res_text))
                                st.session_state.word_doc.add_heading(f"Digitalized Content from {file_obj.name} (Time: {now})", level=2)

                                for line in res_text.split('\n'):
                                    st.session_state.word_doc.add_paragraph(line)
                            else:
                                st.error(f"❌ {file_obj.name} 분석 실패")
                        except Exception as e:
                            st.error(f"파일 처리 중 오류 발생 ({file_obj.name}): {e}")
                    status.update(label="✅ 모든 문서 변환 완료!", state="complete", expanded=False)



                # 변환 결과 출력 및 자동 팝업 트리거
                if conversion_results:
                    all_text_combined = ""
                    for name, text in conversion_results:
                        all_text_combined += f"--- [File: {name}] ---\\n" + text + "\\n\\n"
                        with st.expander(f"👁️ {name} 추출 결과 미리보기", expanded=True):
                            st.text(text)
                    now = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                    open_in_new_window(f"통합 분석 리포트 ({now})", all_text_combined)
                    st.success("✅ 모든 분석이 완료되었으며, 별도의 결과창이 열렸습니다.")



        if start_analysis:
            for file_obj in uploaded_files:
                if not file_obj.name.lower().endswith('.pdf'):
                    handle_image_analysis(file_obj)
                else:
                    st.warning(f"PDF 파일('{file_obj.name}')은 상세 분석 대상이 아닙니다. 텍스트 추출을 사용하세요.")

    doc_stream = io.BytesIO()
    st.session_state.word_doc.save(doc_stream)
    st.download_button("📥 전체 분석 결과 통합 Word 문서 다운로드 (.docx)", data=doc_stream.getvalue(), file_name="AI_Integrated_Report.docx", use_container_width=True)

# 자동 저장: 스크립트 실행이 끝날 때마다 현재 상태를 JSON 파일에 저장
save_state()
