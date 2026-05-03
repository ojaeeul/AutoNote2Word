import re
with open("app.py", "r", encoding="utf-8") as f:
    text = f.read()
    
# Extract the broken load_local_academic_db function and replace it
def_start = text.find("def load_local_academic_db():")
def_end = text.find("def open_in_new_window(")

new_func = """def load_local_academic_db():
    \"\"\"로컬 폴더의 과제, 채점기준표, 교과서 데이터를 로드하여 컨텍스트로 제공\"\"\"
    db_text = ""
    try:
        import glob
        import os
        import fitz  # PyMuPDF
        from docx import Document
        import zipfile
        import tempfile
        import shutil

        # [보안 로직] 암호화된 ZIP 파일이 존재하면 임시 폴더에 해제
        temp_extract_dir = None
        if os.path.exists("oje_secure.zip"):
            try:
                temp_extract_dir = tempfile.mkdtemp()
                with zipfile.ZipFile("oje_secure.zip", 'r') as zf:
                    zf.extractall(path=temp_extract_dir, pwd=b"AIs3cr3t!")
            except Exception as e:
                print(f"Zip extraction error: {e}")

        # 검색 대상 패턴 확장
        search_dirs = ["oje", os.path.join(temp_extract_dir, "oje")] if temp_extract_dir else ["oje"]
        potential_files = []
        for d in search_dirs:
            if not os.path.exists(d): continue
            potential_files.extend(glob.glob(f"{d}/*.pdf"))
            potential_files.extend(glob.glob(f"{d}/*.docx"))
            potential_files.extend(glob.glob(f"{d}/*.md"))
        
        potential_files.extend(glob.glob("*HW2-1*.pdf") + glob.glob("*과제*.pdf") + glob.glob("*채점기준표*.docx") + glob.glob("*결과보고서*.pdf"))
        
        unique_files = list(set([os.path.abspath(f) for f in potential_files]))
        
        for fpath in unique_files:
            if not os.path.exists(fpath): continue
            fname = os.path.basename(fpath)
            db_text += "\\n\\n=========================================\\n"
            if "채점기준표" in fname:
                db_text += f"🎯 [절대 준수: 채점기준표 (Rubric)]: {fname}\\n"
            else:
                db_text += f"📄 [우수 결과보고서 및 과제 해설 데이터]: {fname}\\n"
            db_text += "\\n=========================================\\n"
            
            if fpath.lower().endswith(".pdf"):
                doc = fitz.open(fpath)
                file_text = ""
                for page in doc:
                    file_text += page.get_text() + "\\n"
                    if len(file_text) > 150000: break # 파일당 최대 15만자 제한
                db_text += file_text
                doc.close()
            elif fpath.lower().endswith(".docx"):
                doc = Document(fpath)
                file_text = ""
                for para in doc.paragraphs:
                    file_text += para.text + "\\n"
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            file_text += cell.text + " | "
                        file_text += "\\n"
                db_text += file_text[:150000] # 파일당 최대 15만자 제한
            elif fpath.lower().endswith(".md"):
                with open(fpath, "r", encoding="utf-8") as f:
                    db_text += f.read() + "\\n"
                    
        # [보안 로직] 임시 폴더 즉시 파기 (보안 유지)
        if temp_extract_dir and os.path.exists(temp_extract_dir):
            shutil.rmtree(temp_extract_dir)

    except Exception as e:
        print(f"Error loading local academic db: {e}")
    return db_text

"""

with open("app.py", "w", encoding="utf-8") as f:
    f.write(text[:def_start] + new_func + text[def_end:])
