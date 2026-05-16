import datetime
# Last Updated: 2026-05-04 07:59 (LaTeX Button Labels Refined)
import re
import streamlit as st
import platform
import subprocess

def open_file_in_os(filepath):
    """지정된 파일을 해당 OS의 기본 프로그램(MS Word 등)으로 엽니다."""
    try:
        import subprocess
        import os
        if platform.system() == "Darwin":       # macOS
            subprocess.Popen(["open", filepath])
        elif platform.system() == "Windows":    # Windows
            os.startfile(filepath)
        else:                                   # linux variants
            # Streamlit Cloud 등 헤드리스 환경에서는 xdg-open이 없을 수 있음
            try:
                subprocess.Popen(["xdg-open", filepath])
            except FileNotFoundError:
                # 에러를 표시하지 않고 무시 (웹 환경이므로 자동 열기가 불가능한 것이 정상)
                return False
        return True
    except Exception as e:
        # 웹 환경에서의 일반적인 상황이므로 심각한 에러로 취급하지 않음
        return False

def render_download_and_open_buttons(title, content_bytes, filename, key_suffix="", dl_label=None, open_label=None):
    """다운로드 버튼과 (데스크탑 전용) 즉시 열기 버튼을 나란히 배치합니다."""
    col1, col2 = st.columns(2)
    with col1:
        st.download_button(
            label=dl_label if dl_label else f"💾 {title} 저장 (Save to PC)",
            data=content_bytes,
            file_name=filename,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
            type="primary",
            key=f"dl_btn_{key_suffix}"
        )
    with col2:
        if st.button(open_label if open_label else f"🚀 {title} 즉시 열기 (Desktop 전용)", use_container_width=True, key=f"open_btn_{key_suffix}"):
            import os
            out_file = os.path.join(os.getcwd(), f"Direct_{filename}")
            try:
                with open(out_file, "wb") as f:
                    f.write(content_bytes)
                if not open_file_in_os(out_file):
                    st.info("ℹ️ 웹 브라우저 환경(Streamlit Cloud)에서는 자동 열기가 지원되지 않습니다. 왼쪽 '저장' 버튼을 이용해 주세요!")
            except Exception as e:
                st.error(f"파일 준비 중 오류가 발생했습니다: {e}")

def open_any_word_direct(title, content, filename_prefix="Sample", session_key="any_word_bytes"):
    """마크다운 콘텐츠를 워드로 변환하고 세션에 저장합니다."""
    with st.spinner(f"{title} 문서를 워드로 변환 중..."):
        import time, os
        out_file = os.path.join(os.getcwd(), f"{filename_prefix}_{int(time.time())}.docx")
        margins = {"top": 2.0, "bottom": 2.0, "left": 2.5, "right": 2.5}
        if convert_latex_to_word_docx(f"# {title}\n\n" + content, out_file, margins):
            with open(out_file, "rb") as f:
                st.session_state[session_key] = f.read()
                st.session_state[f"{session_key}_filename"] = f"{filename_prefix}.docx"
            st.success(f"🎉 {title} 워드 변환이 완료되었습니다! 아래에서 확인하세요.")





# MUST BE THE FIRST STREAMLIT COMMAND
st.set_page_config(page_title="SNU Chem-Ed Studio Pro", page_icon="🧪", layout="wide")

# --- Authentication System ---
def check_password():
    # 1. URL 파라미터에서 영구 로그인 상태 확인 (새로고침 방어)
    if st.query_params.get("auth") == "ojaeeul_verified":
        st.session_state["password_correct"] = True
        return True
        
    def password_entered():
        # 기본값은 사용자가 요청한 ID/PW. 추후 Streamlit Secrets에서 오버라이드 가능.
        expected_id = st.secrets.get("admin_id", "ojaeeul")
        expected_pw = st.secrets.get("admin_pw", "calzone2@")
        
        if st.session_state["username"] == expected_id and st.session_state["password"] == expected_pw:
            st.session_state["password_correct"] = True
            if st.session_state.get("auto_login"):
                st.query_params["auth"] = "ojaeeul_verified" # 새로고침해도 풀리지 않도록 URL에 토큰 삽입
            elif "auth" in st.query_params:
                del st.query_params["auth"]
            del st.session_state["password"]  # 보안을 위해 세션에서 비밀번호 삭제
            del st.session_state["username"]
        else:
            st.session_state["password_correct"] = False

    if st.session_state.get("password_correct", False):
        return True

    # Show login UI
    st.markdown("<h1 style='text-align: center; color: #4F46E5; margin-top: 10vh;'>🔒 보안 접속</h1>", unsafe_allow_html=True)
    st.markdown("<p style='text-align: center; margin-bottom: 2rem;'>접근 권한이 필요합니다. 아이디와 비밀번호를 입력해주세요.</p>", unsafe_allow_html=True)
    
    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:
        with st.form(key='login_form'):
            st.text_input("아이디 (ID)", key="username")
            st.text_input("비밀번호 (Password)", type="password", key="password")
            st.checkbox("☑️ 자동 로그인 유지 (즐겨찾기용 주소 생성)", value=True, key="auto_login")
            st.form_submit_button("로그인", on_click=password_entered, use_container_width=True)
            
        if "password_correct" in st.session_state and not st.session_state["password_correct"]:
            st.error("아이디 또는 비밀번호가 일치하지 않습니다.")
    return False

if not check_password():
    st.stop()

import time
import openai
import streamlit.components.v1 as components

# --- Premium UI/UX Design System Injection ---
def inject_premium_design():
    st.markdown("""
        <style>
        @import url('https://fonts.googleapis.com/css2?family=Outfit:wght@300;400;600;700&display=swap');

        /* Global Styles */
        * { font-family: 'Outfit', sans-serif; }
        
        /* Sidebar Glassmorphism */
        section[data-testid="stSidebar"] {
            background: rgba(255, 255, 255, 0.05);
            backdrop-filter: blur(10px);
            border-right: 1px solid rgba(0, 0, 0, 0.1);
        }

        /* Gradient Headers */
        .main-header {
            background: linear-gradient(135deg, #6366F1 0%, #3B82F6 100%);
            -webkit-background-clip: text;
            -webkit-text-fill-color: transparent;
            font-weight: 700;
            font-size: 3rem;
            margin-bottom: 1rem;
        }

        /* Card Design */
        div.stButton > button {
            border-radius: 12px;
            background: linear-gradient(135deg, #4F46E5 0%, #3B82F6 100%);
            color: white;
            border: none;
            padding: 0.6rem 1.2rem;
            font-weight: 600;
            transition: all 0.3s cubic-bezier(0.4, 0, 0.2, 1);
            box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.1), 0 2px 4px -1px rgba(0, 0, 0, 0.06);
        }
        div.stButton > button:hover {
            transform: translateY(-2px);
            box-shadow: 0 10px 15px -3px rgba(79, 70, 229, 0.4);
            color: white;
        }

        /* Status Indicator Pulse */
        div[data-testid="stStatus"] {
            background-color: #F0F9FF;
            border: 1px solid #BAE6FD;
            border-radius: 12px;
            animation: pulse 2s infinite;
        }
        @keyframes pulse {
            0% { transform: scale(1); }
            50% { transform: scale(1.02); }
            100% { transform: scale(1); }
        }

        /* Text Area Focus Effect */
        textarea {
            border-radius: 12px !important;
            border: 1px solid #E2E8F0 !important;
            transition: border-color 0.3s ease;
        }
        textarea:focus {
            border-color: #3B82F6 !important;
            box-shadow: 0 0 0 3px rgba(59, 130, 246, 0.1) !important;
        }
        </style>
    """, unsafe_allow_html=True)

inject_premium_design()

if "gemini_api_key" not in st.session_state:
    # 보안 금고(secrets.toml)에서 키를 안전하게 가져옵니다.
    st.session_state.gemini_api_key = st.secrets.get("gemini_api_key", "")
import google.generativeai as genai
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import io
import os
import urllib.request
from PIL import Image
import matplotlib.pyplot as plt
import numpy as np
from mpl_toolkits.mplot3d import Axes3D
import ssl

# Fix for macOS SSL certificate verify failed error
try:
    _create_unverified_https_context = ssl._create_unverified_context
except AttributeError:
    pass
else:
    ssl._create_default_https_context = _create_unverified_https_context

def get_safe_gemini_model(use_grounding=False):
    """
    미래에 모델명이 변경되더라도 시스템이 중단되지 않도록 가용한 최신 모델을 자동 탐색하는 로직
    """
    tools = []
    if use_grounding:
        try:
            tools = [genai.Tool(google_search_retrieval=genai.GoogleSearchRetrieval())]
        except:
            pass

    try:
        # 현재 API 키로 사용 가능한 모델 목록 확보
        available_models = [m.name for m in genai.list_models() if 'generateContent' in m.supported_generation_methods]
        
        # 선호도 순위 (최신 및 안정성 기준)
        preferences = [
            'models/gemini-1.5-pro-002',
            'models/gemini-1.5-pro', 
            'models/gemini-2.0-flash', 
            'models/gemini-1.5-flash-002',
            'models/gemini-1.5-flash',
            'models/gemini-pro'
        ]
        
        for pref in preferences:
            if pref in available_models:
                return genai.GenerativeModel(pref, tools=tools)
        
        # 만약 선호 모델이 하나도 없다면, 목록 중 첫 번째 사용 가능 모델 선택 (중단 방지)
        if available_models:
            return genai.GenerativeModel(available_models[0], tools=tools)
            
        return genai.GenerativeModel('gemini-pro-latest', tools=tools)
    except Exception as e:
        # 최후의 보루: 기본 모델 반환
        return genai.GenerativeModel('gemini-pro-latest', tools=tools)

@st.cache_data(show_spinner=False)
def load_local_academic_db():
    """로컬 폴더의 과제, 채점기준표, 교과서 데이터를 로드하여 컨텍스트로 제공"""
    db_text = ""
    try:
        import glob
        import os
        import fitz  # PyMuPDF
        from docx import Document
        import zipfile
        import tempfile
        import shutil

        # [보안 로직] 암호화된 ZIP 파일이 존재하면 임시 폴더에 해제
        temp_extract_dir = None
        if os.path.exists("oje_secure.zip"):
            try:
                temp_extract_dir = tempfile.mkdtemp()
                with zipfile.ZipFile("oje_secure.zip", 'r') as zf:
                    zf.extractall(path=temp_extract_dir, pwd=b"AIs3cr3t!")
            except Exception as e:
                print(f"Zip extraction error: {e}")

        # 검색 대상 패턴 확장
        search_dirs = ["oje", os.path.join(temp_extract_dir, "oje")] if temp_extract_dir else ["oje"]
        potential_files = []
        for d in search_dirs:
            if not os.path.exists(d): continue
            potential_files.extend(glob.glob(f"{d}/*.pdf"))
            potential_files.extend(glob.glob(f"{d}/*.docx"))
            potential_files.extend(glob.glob(f"{d}/*.md"))
        
        potential_files.extend(glob.glob("*HW2-1*.pdf") + glob.glob("*과제*.pdf") + glob.glob("*채점기준표*.docx") + glob.glob("*결과보고서*.pdf"))
        
        unique_files = list(set([os.path.abspath(f) for f in potential_files]))
        
        for fpath in unique_files:
            if not os.path.exists(fpath): continue
            fname = os.path.basename(fpath)
            db_text += "\n\n=========================================\n"
            if "채점기준표" in fname:
                db_text += f"🎯 [절대 준수: 채점기준표 (Rubric)]: {fname}\n"
            else:
                db_text += f"📄 [우수 결과보고서 및 과제 해설 데이터]: {fname}\n"
            db_text += "\n=========================================\n"
            
            if fpath.lower().endswith(".pdf"):
                doc = fitz.open(fpath)
                file_text = ""
                for page in doc:
                    file_text += page.get_text() + "\n"
                    if len(file_text) > 150000: break # 파일당 최대 15만자 제한
                db_text += file_text
                doc.close()
            elif fpath.lower().endswith(".docx"):
                doc = Document(fpath)
                file_text = ""
                for para in doc.paragraphs:
                    file_text += para.text + "\n"
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            file_text += cell.text + " | "
                        file_text += "\n"
                db_text += file_text[:150000] # 파일당 최대 15만자 제한
            elif fpath.lower().endswith(".md"):
                with open(fpath, "r", encoding="utf-8") as f:
                    db_text += f.read() + "\n"
                    
        # [보안 로직] 임시 폴더 즉시 파기 (보안 유지)
        if temp_extract_dir and os.path.exists(temp_extract_dir):
            shutil.rmtree(temp_extract_dir)

    except Exception as e:
        print(f"Error loading local academic db: {e}")
    return db_text

def open_in_new_window(title, content):
    import streamlit.components.v1 as components
    import json
    # 줄바꿈 및 특수문자 처리를 위한 JSON 인코딩
    safe_content = json.dumps(content)
    html_code = f"""
    <script>
    function openWindow() {{
        const win = window.open("", "_blank");
        win.document.write(`
            <html>
            <head>
                <title>{title}</title>
                <script src="https://cdn.jsdelivr.net/npm/marked/marked.min.js"></script>
                <script src="https://polyfill.io/v3/polyfill.min.js?features=es6"></script>
                <script id="MathJax-script" async src="https://cdn.jsdelivr.net/npm/mathjax@3/es5/tex-mml-chtml.js"></script>
                <style>
                    body {{ 
                        font-family: 'Segoe UI', Roboto, Helvetica, Arial, sans-serif; 
                        padding: 40px; 
                        line-height: 1.8; 
                        color: #1e293b; 
                        background-color: #f8fafc;
                        max-width: 1000px;
                        margin: 0 auto;
                        word-wrap: break-word;
                        overflow-wrap: break-word;
                    }}
                    .container {{
                        background: white;
                        padding: 40px;
                        border-radius: 15px;
                        box-shadow: 0 10px 15px -3px rgba(0, 0, 0, 0.1);
                        border: 1px solid #e2e8f0;
                    }}
                    h1 {{ 
                        color: #2563eb; 
                        border-bottom: 3px solid #3b82f6; 
                        padding-bottom: 15px;
                        margin-top: 0;
                        font-size: 1.8rem;
                    }}
                    #rendered-content {{
                        font-size: 1.1rem;
                    }}
                    pre {{ 
                        white-space: pre-wrap; 
                        word-wrap: break-word; 
                        background: #f1f5f9; 
                        padding: 20px; 
                        border-radius: 8px; 
                        font-family: 'Fira Code', 'Courier New', monospace;
                        font-size: 0.95em;
                        border: 1px solid #e2e8f0;
                        overflow-x: auto;
                    }}
                    code {{
                        background: #f1f5f9;
                        padding: 2px 5px;
                        border-radius: 4px;
                        font-family: monospace;
                    }}
                    blockquote {{
                        border-left: 5px solid #3b82f6;
                        padding: 10px 20px;
                        margin: 20px 0;
                        background: #eff6ff;
                        color: #1e40af;
                        border-radius: 0 8px 8px 0;
                    }}
                    table {{
                        border-collapse: collapse;
                        width: 100%;
                        margin: 20px 0;
                        border-radius: 8px;
                        overflow: hidden;
                        box-shadow: 0 1px 3px rgba(0,0,0,0.1);
                    }}
                    th, td {{
                        border: 1px solid #e2e8f0;
                        padding: 15px;
                        text-align: left;
                    }}
                    th {{
                        background-color: #f1f5f9;
                        font-weight: 600;
                        color: #475569;
                    }}
                    .no-print {{
                        margin-top: 40px;
                        text-align: center;
                    }}
                    .btn {{
                        padding: 15px 30px; 
                        background: #2563eb; 
                        color: white; 
                        border: none; 
                        border-radius: 8px; 
                        cursor: pointer;
                        font-weight: bold;
                        font-size: 1rem;
                        transition: all 0.2s;
                        box-shadow: 0 4px 6px rgba(37, 99, 235, 0.2);
                    }}
                    .btn:hover {{ 
                        background: #1d4ed8; 
                        transform: translateY(-2px);
                        box-shadow: 0 6px 8px rgba(37, 99, 235, 0.3);
                    }}
                    @media print {{
                        .no-print {{ display: none; }}
                        body {{ padding: 0; background: white; }}
                        .container {{ box-shadow: none; border: none; padding: 0; }}
                    }}
                </style>
            </head>
            <body>
                <div class="container">
                    <h1>{title}</h1>
                    <div id="rendered-content">분석 결과를 렌더링 중입니다...</div>
                    <div class="no-print">
                        <button class="btn" onclick="window.print()">📥 결과 보고서 PDF / 인쇄로 저장</button>
                    </div>
                </div>
                <script>
                    setTimeout(() => {{
                        const rawContent = ${{safe_content}};
                        document.getElementById('rendered-content').innerHTML = marked.parse(rawContent);
                        
                        // MathJax 렌더링 재실행
                        if (window.MathJax) {{
                            window.MathJax.typesetPromise();
                        }}
                    }}, 100);
                </script>
            </body>
            </html>
        `);
        win.document.close();
    }}
    openWindow();
    </script>
    """
    components.html(html_code, height=0)



def handle_image_analysis(file_obj):
    """
    이미지 내의 복잡한 문제, 그래프, 구조식을 전문가 수준으로 정밀 분석하는 기능
    """
    from PIL import Image
    import streamlit as st



    with st.status(f"🔍 '{file_obj.name}' 상세 분석 중...", expanded=True) as status:
        prompt = "당신은 화학/물리 전문가입니다. 이 이미지에 있는 과학 문제, 그래프, 혹은 구조식을 분석하여 '상세한 풀이 과정'과 '논리적 근거', 그리고 '최종 답'을 설명해 주세요. 수식은 LaTeX 형식($$ ... $$)을 사용하세요."
        res_text = robust_generate_content(prompt, images=Image.open(file_obj))

        if res_text:
            now = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            st.markdown(f"### 📑 {file_obj.name} 상세 분석 결과 (분석 시간: {now})")
            st.markdown(res_text)
            open_in_new_window(f"상세 분석 결과: {file_obj.name}", res_text)

            # 워드에 추가
            now = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            st.session_state.word_doc.add_heading(f"Detailed Analysis of {file_obj.name} (Time: {now})", level=2)
            for line in res_text.split('\n'):
                st.session_state.word_doc.add_paragraph(line)

            st.success("분석 결과가 워드 문서에 추가되었습니다.")
            status.update(label="✅ 상세 분석 완료", state="complete", expanded=False)
        else:
            st.error("분석에 실패했습니다.")


def deduplicate_text(text):
    if not text: return text
    import re
    orig_len = len(text)
    # 1. 대규모 중복 패턴 강제 제거 (10자 이상의 문장/문단 반복)
    for _ in range(5): 
        text = re.sub(r'(.{15,})\1{2,}', r'\1', text, flags=re.DOTALL)
    
    # 2. 중간/짧은 단어 폭주(Model Collapse) 방어
    for _ in range(3):
        text = re.sub(r'(.{1,15})(\s?\1){4,}', r'\1', text)
        
    text = text.strip()
    # 3. 심각한 반복 버그(Model Collapse) 치명적 감지 로직
    # 만약 중복 제거 후 텍스트 길이가 원본 대비 비정상적으로 줄어들었다면(반복이 절반 이상을 차지함)
    if orig_len > 1000 and len(text) < orig_len * 0.4:
        raise Exception("Model Collapse Detected (심각한 문장 반복 버그 감지됨). 모델 붕괴로 인해 현재 결과를 폐기합니다.")
        
    return text

def robust_generate_content(prompt, images=None, use_grounding=False):
    """
    재시도 간격(Sleep)을 추가하고 모든 변종 모델명을 시도하는 최후의 철벽 로직
    """
    import streamlit as st


    import google.generativeai as genai
    import time

    # Gemini API 키 설정 확인 및 적용
    g_api_key = st.session_state.get("gemini_api_key")
    if g_api_key:
        genai.configure(api_key=g_api_key)

    tools = None
    if use_grounding:
        import google.ai.generativelanguage as gl
        tools = [gl.Tool({'google_search': {}})]

    # 1. 동적 목록 확보
    dynamic_models = []
    try:
        for m in genai.list_models():
            if 'generateContent' in m.supported_generation_methods:
                dynamic_models.append(m.name)
    except: pass

    # 2. 최신 고성능 모델 전수 동원 (초고지능 모드: Gemini + GPT-4o 하이브리드)
    # Grounding을 사용할 때는 최신 모델로 제한하여 에러 방지
    if use_grounding:
        manual_candidates = [
            'gemini-pro-latest', 'models/gemini-pro-latest',
            'gemini-flash-latest', 'models/gemini-flash-latest',
            'gemini-2.5-pro', 'models/gemini-2.5-pro',
            'gemini-2.5-flash', 'models/gemini-2.5-flash'
        ]
        dynamic_models = [] # grounding은 동적 모델에서 에러 잦음
    else:
        manual_candidates = [
            'gpt-4o', 'gpt-4o-2024-05-13', 
            'gemini-pro-latest', 'models/gemini-pro-latest',
            'gemini-flash-latest', 'models/gemini-flash-latest',
            'gemini-2.5-pro', 'models/gemini-2.5-pro',
            'gemini-2.5-flash', 'models/gemini-2.5-flash',
            'gemini-3.1-pro-preview', 'models/gemini-3.1-pro-preview'
        ]

    model_candidates = []
    for m in manual_candidates + dynamic_models:
        if m not in model_candidates: model_candidates.append(m)
    
    # 최대 시도 모델 수를 제한하여 무한 루프 멈춤 현상 방지
    model_candidates = model_candidates[:8]

    last_er = None
    inputs = [prompt]
    if images:
        if isinstance(images, list): inputs.extend(images)
        else: inputs.append(images)

    # 3. 안전 필터 완전 해제 및 전수 동원 루프
    safety_settings = [
        {"category": "HARM_CATEGORY_HARASSMENT", "threshold": "BLOCK_NONE"},
        {"category": "HARM_CATEGORY_HATE_SPEECH", "threshold": "BLOCK_NONE"},
        {"category": "HARM_CATEGORY_SEXUALLY_EXPLICIT", "threshold": "BLOCK_NONE"},
        {"category": "HARM_CATEGORY_DANGEROUS_CONTENT", "threshold": "BLOCK_NONE"},
    ]

    retry_delay = 3
    
    # [무한 우회: 다중 API Key 로테이션 시스템]
    g_api_keys = st.session_state.get("gemini_api_key", "").split(",")
    g_api_keys = [k.strip() for k in g_api_keys if k.strip()]
    key_idx = 0
    if g_api_keys: genai.configure(api_key=g_api_keys[0])

    # [고급 고도화] 전체 AI 군단(Arsenal)을 3번의 큰 파도(Wave)로 순환시키며 파상공세
    for wave in range(3):
        for idx, model_name in enumerate(model_candidates):
            try:
                    # [ChatGPT 모델 처리]
                if model_name.startswith('gpt'):
                    o_api_key = st.session_state.get("openai_api_key")
                    if not o_api_key: continue # 키가 없으면 건너뜀
                    
                    client = openai.OpenAI(api_key=o_api_key)
                    messages = [{"role": "user", "content": [{"type": "text", "text": prompt}]}]
                    
                    res_gpt = client.chat.completions.create(
                        model=model_name,
                        messages=messages,
                        max_tokens=4096,
                        temperature=0.1, # 0.0보다 0.1이 모델 루프(반복) 방지에 더 유리함
                        presence_penalty=1.0, # 강한 억제
                        frequency_penalty=1.0 # 강한 억제
                    )
                    if res_gpt.choices[0].message.content:
                        return deduplicate_text(res_gpt.choices[0].message.content)

                # [Gemini 모델 처리]
                else:
                    model = genai.GenerativeModel(model_name, tools=tools if use_grounding else None)
                    res = model.generate_content(
                        inputs, 
                        generation_config=genai.GenerationConfig(
                            max_output_tokens=8192, 
                            temperature=0.3, # 0.1보다 0.3이 지적 다양성과 반복 방지 사이의 최적 균형점
                            top_p=0.8, # 상위 80% 확률 내에서 선택하여 창의성과 안정성 확보
                            top_k=40   # 후보군을 40개로 제한하여 엉뚱한 토큰 루핑 방지
                        ),
                        safety_settings=safety_settings
                    )
                    if res and res.text: return deduplicate_text(res.text)
                    
            except Exception as e:
                last_er = e
                err_str = str(e).lower()
                # 429 과부하, 초과 횟수, 타임아웃, 403 권한 거부(Suspended) 등에 대한 자동 수위 조절 로직 (Arsenal Fallback)
                if "429" in err_str or "quota" in err_str or "rate limit" in err_str or "overloaded" in err_str or "timeout" in err_str or "403" in err_str or "permission denied" in err_str or "suspended" in err_str:
                    # [API Key Rotation] 보조 키가 존재할 경우 구글의 물리적 차단벽을 즉시 회피
                    if len(g_api_keys) > 1 and key_idx < len(g_api_keys) - 1:
                        key_idx += 1
                        genai.configure(api_key=g_api_keys[key_idx])
                        st.toast(f"🗝️ 구글 한도/권한 차단벽 도달! 즉각 {key_idx+1}번째 보조 API 키로 교체하여 물리적 제한을 우회합니다!")
                        continue # 쉬지 않고 곧바로 다음 모델로 우회 타격
                    
                    # 보조 키가 없거나 모두 소진된 상태에서 429/403을 맞았다면, 내부 루프를 끊고 정확한 대기 시간을 바깥으로 던집니다.
                    import re
                    match = re.search(r'retry in (\d+\.?\d*)s', err_str)
                    wait_sec = float(match.group(1)) + 3.0 if match else 65.0
                    raise Exception(f"QUOTA_EXHAUSTED:{wait_sec}")
                else:
                    # 400 등 구문 오류는 재시도해도 의미가 없으므로 즉시 다음 모델로 패스
                    continue

    if last_er and "429" in str(last_er):
        pass # 429 에러는 바깥쪽에서 쿨타임(휴식) 타이머로 예쁘게 처리하므로 흉측한 로우 에러는 숨깁니다.
    else:
        with st.expander("❌ 전 세계 AI 군단 응답 실패 (자세히 보기)"):
            st.error(f"마지막 에러: {last_er}")
    return None








# st.set_page_config removed (called at the beginning)

import streamlit.components.v1 as components
components.html("""
<script>
    const parent = window.parent;
    const mainContainer = parent.document.querySelector('.main') || parent.document.querySelector('section.main');

    if (mainContainer) {
        mainContainer.addEventListener('scroll', () => {
            parent.localStorage.setItem('st_scroll_pos', mainContainer.scrollTop);
        });

        setTimeout(() => {
            const scrollPos = parent.localStorage.getItem('st_scroll_pos');
            if (scrollPos) {
                mainContainer.scrollTop = parseInt(scrollPos);
            }
        }, 300);
    }
</script>
""", height=0, width=0)

import json
STATE_FILE = "workspace_state.json"

def load_state():
    if "state_loaded" not in st.session_state:
        if os.path.exists(STATE_FILE):
            try:
                import base64
                with open(STATE_FILE, "r", encoding="utf-8") as f:
                    saved = json.load(f)
                for k, v in saved.items():
                    if isinstance(v, dict) and v.get("__type__") == "bytes":
                        st.session_state[k] = base64.b64decode(v["data"])
                    else:
                        st.session_state[k] = v
            except:
                pass
        st.session_state.state_loaded = True

def save_state():
    state_to_save = {}
    for k in ["smart_target_pages", "smart_target_questions", "notion_db", "latex_code", "mathlive_init", "menu_selection", "global_smart_img_result", "global_smart_img_name", "analysis_buffer", "curent_file_hash", "smart_analysis_result", "hw_analysis_buffer", "hw_curent_file_hash", "smart_analysis_active", "smart_page_map", "smart_file_paths", "smart_q_structures", "last_report_result", "last_report_query", "last_report_safe_word", "last_report_safe_name", "last_report_hq_word", "last_report_hq_name", "ex_label_1", "ex_label_2", "ex_label_3", "ex_label_4", "report_search_query", "global_trash_bin", "report_gen_active", "report_gen_buffer", "report_gen_phases"]:
        if k in st.session_state:
            state_to_save[k] = st.session_state[k]
    try:
        import base64
        encoded_state = {}
        for k, v in state_to_save.items():
            if isinstance(v, bytes):
                encoded_state[k] = {"__type__": "bytes", "data": base64.b64encode(v).decode('utf-8')}
            else:
                encoded_state[k] = v
                
        with open(STATE_FILE, "w", encoding="utf-8") as f:
            json.dump(encoded_state, f, ensure_ascii=False, indent=2)
    except:
        pass

# --- 백그라운드 분석 관리 시스템 ---
if "pending_tasks" not in st.session_state:
    st.session_state.pending_tasks = []

if "global_trash_bin" not in st.session_state:
    st.session_state.global_trash_bin = st.session_state.get("report_trash_bin", [])

if "cell_text_key_counter" not in st.session_state:
    st.session_state.cell_text_key_counter = 0
    st.session_state.graph_text_key_counter = 0
    st.session_state.orbital_text_key_counter = 0
    st.session_state.lewis_text_key_counter = 0
    st.session_state.shape_text_key_counter = 0
    st.session_state.skeletal_text_key_counter = 0

def add_analysis_task(prompt, images, target_subject, file_name):
    st.session_state.pending_tasks.append({
        "prompt": prompt,
        "images": images,
        "subject": target_subject,
        "file_name": file_name,
        "status": "pending"
    })

def process_background_tasks():
    if st.session_state.pending_tasks:
        for task in st.session_state.pending_tasks:
            if task["status"] == "pending":
                with st.sidebar:
                    with st.status(f"⏳ '{task['file_name']}' 분석 중...", expanded=False):
                        task["status"] = "processing"
                        try:
                            res = robust_generate_content(task["prompt"], images=task["images"])
                            if res:
                                now = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                                header = f"\n\n--- [AI 자동추출: {task['file_name']} ({now})] ---\n"
                                if task["subject"] in st.session_state.notion_db:
                                    st.session_state.notion_db[task["subject"]] += header + res
                                    st.toast(f"✅ {task['file_name']} 분석 완료 및 '{task['subject']}'에 저장됨!")
                                task["status"] = "completed"
                                st.rerun()
                        except Exception as e:
                            st.error(f"오류: {e}")
                            task["status"] = "failed"
        st.session_state.pending_tasks = [t for t in st.session_state.pending_tasks if t["status"] not in ["completed", "failed"]]

# 앱 초기화 및 상태 로드
load_state()
process_background_tasks()


if "word_doc" not in st.session_state:
    st.session_state.word_doc = Document()
    st.session_state.word_doc.add_heading("SNU Chem-Ed Analysis Report", level=0)

if "analysis_buffer" not in st.session_state:
    st.session_state.analysis_buffer = {}


# 과제 샘플 데이터 (TOP 5) - 최적화 (캐싱 적용)
@st.cache_data(show_spinner=False)
def load_all_samples():
    # Cache busted by adding this comment
    samples_dict = {}
    subjects_list = ["물리화학", "유기화학", "분석화학", "무기화학", "화학교육"]
    for subj in subjects_list:
        try:
            import os
            sample_path = os.path.join(os.path.dirname(__file__), "samples", f"{subj}.md")
            if os.path.exists(sample_path):
                with open(sample_path, "r", encoding="utf-8") as f:
                    content = f.read()
                    # LaTeX 수식 제어 문자 복구 (백슬래시 이스케이프 오류 방지)
                    content = content.replace('\x0c', '\\f').replace('\x0d', '\\r').replace('\x09', '\\t').replace('\x08', '\\b').replace('\x07', '\\a')
                    # 깨진 수식(줄바꿈으로 분리된 경우) 추가 보정
                    content = content.replace('\n' + 'u ', '\\nu ').replace('\n' + 'abla', '\\nabla').replace('\n' + 'ho', '\\rho')
                    samples_dict[subj] = content
            else:
                samples_dict[subj] = f"샘플 파일({subj}.md)이 존재하지 않습니다."
        except Exception as e:
            samples_dict[subj] = f"샘플 데이터를 불러올 수 없습니다: {e}"
    return samples_dict

SAMPLES = load_all_samples()

if "notion_db" not in st.session_state:
    st.session_state.notion_db = {
        "물리화학": SAMPLES.get("물리화학", "데이터 없음"),
        "유기화학": SAMPLES.get("유기화학", "데이터 없음"),
        "무기화학": SAMPLES.get("무기화학", "데이터 없음"),
        "분석화학": SAMPLES.get("분석화학", "데이터 없음"),
        "화학교육": SAMPLES.get("화학교육", "데이터 없음")
    }

    st.markdown(f"""
<style>
    :root {{
        --bg-color: #F8FAFC;
        --sidebar-bg: #FFFFFF;
    }}
    .stApp {{ background-color: var(--bg-color); font-family: 'Inter', sans-serif; }}
    h1 {{ font-size: 2.2rem !important; font-weight: 700 !important; color: #0F172A !important; margin-bottom: 1rem !important; }}
    .main {{ overflow: auto !important; height: 100vh !important; }} /* 추가된 스크롤 고정 */
    .block-container {{ padding: 2rem 4rem !important; max-width: 100% !important; }}

    div.stButton > button {{
        border-radius: 10px;
        transition: all 0.3s ease;
        font-weight: 600;

        background: linear-gradient(135deg, #2E5BFF 0%, #6366F1 100%);
        color: white;
        border: none;
    }}

    div.stButton > button:hover {{ transform: translateY(-2px); box-shadow: 0 10px 15px -3px rgba(46, 91, 255, 0.4); }}

    .block-container {{ padding: 2rem 4rem !important; }}
</style>

""", unsafe_allow_html=True)

import platform
if platform.system() == 'Windows':
    plt.rcParams['font.family'] = 'Malgun Gothic'
elif platform.system() == 'Darwin':
    plt.rcParams['font.family'] = 'AppleGothic'
else:
    import os
    import matplotlib.font_manager as fm
    font_url = "https://github.com/google/fonts/raw/main/ofl/nanumgothic/NanumGothic-Regular.ttf"
    font_path = "NanumGothic.ttf"
    if not os.path.exists(font_path):
        import urllib.request
        urllib.request.urlretrieve(font_url, font_path)
    fm.fontManager.addfont(font_path)
    font_prop = fm.FontProperties(fname=font_path)
    plt.rcParams['font.family'] = font_prop.get_name()
plt.rcParams['axes.unicode_minus'] = False

# Floating Scroll Buttons
st.markdown("""
<style>
    #scroll-btns {
        position: fixed;
        bottom: 30px;
        right: 30px;
        z-index: 9999999;
        display: flex;
        flex-direction: column;
        gap: 15px;
    }
    .scroll-button {
        width: 60px;
        height: 60px;
        border-radius: 50%;
        cursor: pointer;
        display: flex;
        flex-direction: column;
        align-items: center;
        justify-content: center;
        font-weight: bold;
        font-size: 12px;
        transition: all 0.3s cubic-bezier(0.4, 0, 0.2, 1);
        user-select: none;
    }
    .scroll-button:hover {
        transform: scale(1.15) translateY(-5px);
        box-shadow: 0 8px 20px rgba(0,0,0,0.3) !important;
    }
    .scroll-up {
        background-color: white !important;
        color: #2E5BFF !important;
        border: 3px solid #2E5BFF !important;
        box-shadow: 0 4px 12px rgba(0,0,0,0.15);
    }
    .scroll-down {
        background-color: #F59E0B !important;
        color: white !important;
        border: 3px solid white !important;
        box-shadow: 0 4px 12px rgba(0,0,0,0.25);
    }
</style>

<div id="scroll-btns">
    <div class="scroll-button scroll-up" id="btn-scroll-up" title="맨 위로 이동">
        <span style="font-size: 20px; line-height: 1;">↑</span>위로
    </div>
    <div class="scroll-button scroll-down" id="btn-scroll-down" title="맨 아래로 이동">
        아래로<span style="font-size: 20px; line-height: 1;">↓</span>
    </div>
</div>
""", unsafe_allow_html=True)

import streamlit.components.v1 as components
components.html("""
<script>
    const parentDoc = window.parent.document;
    const btnUp = parentDoc.getElementById('btn-scroll-up');
    const btnDown = parentDoc.getElementById('btn-scroll-down');
    
    if (btnUp) {
        btnUp.onclick = function() {
            const scroller = parentDoc.querySelector('.main') || parentDoc.querySelector('[data-testid="stMain"]') || parentDoc.querySelector('.stAppViewContainer');
            if (scroller) {
                scroller.scrollTo({ top: 0, behavior: 'smooth' });
            } else {
                window.parent.scrollTo({ top: 0, behavior: 'smooth' });
            }
        };
    }
    
    if (btnDown) {
        btnDown.onclick = function() {
            const scroller = parentDoc.querySelector('.main') || parentDoc.querySelector('[data-testid="stMain"]') || parentDoc.querySelector('.stAppViewContainer');
            if (scroller) {
                scroller.scrollTo({ top: scroller.scrollHeight, behavior: 'smooth' });
            } else {
                window.parent.scrollTo({ top: 999999, behavior: 'smooth' });
            }
        };
    }
</script>
""", height=0, width=0)

# ==========================================
# 도식/그래프 함수들
# ==========================================

def get_preset_custom_data(module, choice):
    if module == "cell":
        if choice == "Simple Cubic (SC)":
            return "0,0,0,blue\n1,0,0,blue\n0,1,0,blue\n1,1,0,blue\n0,0,1,blue\n1,0,1,blue\n0,1,1,blue\n1,1,1,blue"
        elif choice == "Body-Centered (BCC)":
            return "0,0,0,blue\n1,0,0,blue\n0,1,0,blue\n1,1,0,blue\n0,0,1,blue\n1,0,1,blue\n0,1,1,blue\n1,1,1,blue\n0.5,0.5,0.5,red"
        elif choice == "Face-Centered (FCC)":
            return "0,0,0,blue\n1,0,0,blue\n0,1,0,blue\n1,1,0,blue\n0,0,1,blue\n1,0,1,blue\n0,1,1,blue\n1,1,1,blue\n0.5,0.5,0,green\n0.5,0.5,1,green\n0.5,0,0.5,green\n0.5,1,0.5,green\n0,0.5,0.5,green\n1,0.5,0.5,green"
        else:
            return "0,0,0,blue\n1,1,1,red"

    elif module == "graph":
        if choice == "1D Box 파동함수 (n=1,2,3)":
            return "np.sin(np.pi*x)*np.sqrt(2)"
        elif choice == "2s 오비탈 확률 밀도":
            return "(x*(2-x)*np.exp(-x/2))**2"
        return "sin(x)*exp(-x/5)"
        if choice == "Atkins: 2s 오비탈 노드 분석":
            return "(x*(2-x)*np.exp(-x/2))**2"
        if choice == "Atkins: 아레니우스 에너지 장벽":
            return "np.exp(-10/x)"


    elif module == "orbital":
        if choice == "NO3- (질산 이온)":
            return "ATOM, N, 0, 0, 0\nATOM, O, 0, 1.2, 0\nATOM, O, 1.04, -0.6, 0\nATOM, O, -1.04, -0.6, 0\nBOND, 0, 0, 0, 0, 1.2, 0\nBOND, 0, 0, 0, 1.04, -0.6, 0\nBOND, 0, 0, 0, -1.04, -0.6, 0\nPZ, 0, 0, 0, #60A5FA, N_pz\nPZ, 0, 1.2, 0, #9CA3AF, O_pz\nPZ, 1.04, -0.6, 0, #9CA3AF, O_pz\nPZ, -1.04, -0.6, 0, #9CA3AF, O_pz"
        elif choice == "HNO3 (질산)":
            return "ATOM, N, 0, 0, 0\nATOM, O, 0, 1.2, 0\nATOM, O, 1.04, -0.6, 0\nATOM, O, -1.04, -0.6, 0\nATOM, H, -1.8, -1.0, 0\nBOND, 0, 0, 0, 0, 1.2, 0\nBOND, 0, 0, 0, 1.04, -0.6, 0\nBOND, 0, 0, 0, -1.04, -0.6, 0\nBOND, -1.04, -0.6, 0, -1.8, -1.0, 0\nPZ, 0, 0, 0, #60A5FA, N_pz\nPZ, 0, 1.2, 0, #9CA3AF, O_pz\nPZ, 1.04, -0.6, 0, #9CA3AF, O_pz\nSP2, -1.04, -0.6, 0, 210, #4ADE80, O_sp2\nS, -1.8, -1.0, 0, #F87171, H_s"
        elif choice == "C2H4 (에텐)":
            return "ATOM, C1, -0.6, 0, 0\nATOM, C2, 0.6, 0, 0\nBOND, -0.6, 0, 0, 0.6, 0, 0\nPZ, -0.6, 0, 0, #60A5FA, C1_pz\nPZ, 0.6, 0, 0, #60A5FA, C2_pz\nATOM, H, -1.2, 0.8, 0\nATOM, H, -1.2, -0.8, 0\nATOM, H, 1.2, 0.8, 0\nATOM, H, 1.2, -0.8, 0\nBOND, -0.6, 0, 0, -1.2, 0.8, 0\nBOND, -0.6, 0, 0, -1.2, -0.8, 0\nBOND, 0.6, 0, 0, 1.2, 0.8, 0\nBOND, 0.6, 0, 0, 1.2, -0.8, 0\nS, -1.2, 0.8, 0, #9CA3AF, H_s\nS, -1.2, -0.8, 0, #9CA3AF, H_s\nS, 1.2, 0.8, 0, #9CA3AF, H_s\nS, 1.2, -0.8, 0, #9CA3AF, H_s"
        return "ATOM, X, 0, 0, 0"

    elif module == "lewis":
        if choice == "H2O (물)":
            return "TEXT, O, 0.5, 0.5, 36, #1E3A8A\nTEXT, H, 0.2, 0.5, 36, #1E3A8A\nTEXT, H, 0.8, 0.5, 36, #1E3A8A\nLINE, 0.3, 0.5, 0.4, 0.5, 2, k\nLINE, 0.6, 0.5, 0.7, 0.5, 2, k"



def draw_unit_cell(cell_type, lattice_size=1, atom_rad=0.15, custom_coords="", elev=20, azim=45):
    from matplotlib.figure import Figure
    errors = []
    fig = Figure(figsize=(10, 10))
    ax = fig.add_subplot(111, projection='3d')

    # 사실적인 구(Sphere)를 그리는 함수
    def draw_sphere(ax, center, radius, color):
        u, v = np.mgrid[0:2*np.pi:15j, 0:np.pi:10j]
        x = center[0] + radius * np.cos(u) * np.sin(v)
        y = center[1] + radius * np.sin(u) * np.sin(v)
        z = center[2] + radius * np.cos(v)
        # shade=True와 rstride/cstride 조절로 입체감 극대화
        ax.plot_surface(x, y, z, color=color, alpha=1.0, linewidth=0, shade=True, rstride=1, cstride=1)

    # 1. 격자 선 생성 (수동 모드 포함 모든 모드에서 상 표시)
    atoms = [] # (pos, color, radius)
    for i in range(lattice_size + 1):
        for j in range(lattice_size + 1):
            # X축 방향 선
            ax.plot([0, lattice_size], [i, i], [j, j], color='gray', lw=1, alpha=0.4)
            # Y축 방향 선
            ax.plot([i, i], [0, lattice_size], [j, j], color='gray', lw=1, alpha=0.4)
            # Z축 방향 선
            ax.plot([i, i], [j, j], [0, lattice_size], color='gray', lw=1, alpha=0.4)

    # 2. 원자 위치 데이터 생성
    if cell_type != "직접 좌표 입력 (Custom)":
        rad = atom_rad # 사용자 지정 반지름

        if cell_type == "Simple Cubic (SC)":
            for i in range(lattice_size + 1):
                for j in range(lattice_size + 1):
                    for k in range(lattice_size + 1):
                        draw_sphere(ax, (i, j, k), rad, '#D32F2F') # Corner: Red

        elif cell_type == "Body-Centered (BCC)":
            # Corners
            for i in range(lattice_size + 1):
                for j in range(lattice_size + 1):
                    for k in range(lattice_size + 1):
                        draw_sphere(ax, (i, j, k), rad, '#D32F2F') # Corner: Red
            # Centers
            for i in range(lattice_size):
                for j in range(lattice_size):
                    for k in range(lattice_size):
                        draw_sphere(ax, (i+0.5, j+0.5, k+0.5), rad, '#1976D2') # Body: Blue

        elif cell_type == "Face-Centered (FCC)":
            # Corners
            for i in range(lattice_size + 1):
                for j in range(lattice_size + 1):
                    for k in range(lattice_size + 1):
                        draw_sphere(ax, (i, j, k), rad, '#D32F2F') # Corner: Red
            # Faces
            for i in range(lattice_size):
                for j in range(lattice_size):
                    for k in range(lattice_size):
                        off = np.array([i, j, k])
                        # XY, XZ, YZ faces
                        for p, c in [(np.array([0.5,0.5,0]), '#388E3C'), (np.array([0.5,0,0.5]), '#1976D2'), (np.array([0,0.5,0.5]), '#FBC02D')]:
                            draw_sphere(ax, p + off, rad, c)

        elif cell_type == "HCP (Hexagonal)":
            c_a = 1.633 # c/a ratio
            for i in range(lattice_size + 1):
                for j in range(lattice_size + 1):
                    for layer in range(lattice_size * 2 + 1):
                        z = layer * (c_a / 2)
                        # Hexagonal points
                        angles = np.linspace(0, 2*np.pi, 7)
                        for ang in angles:
                            x, y = np.cos(ang) + i*1.5, np.sin(ang) + j*np.sqrt(3)
                            if layer % 2 == 0:
                                draw_sphere(ax, (x, y, z), rad, '#D32F2F')
                            else:
                                # Middle layer offset
                                draw_sphere(ax, (x+0.5, y+0.3, z), rad, '#1976D2')
                        # Draw vertical lines
                        if layer < lattice_size * 2:
                            for ang in angles:
                                x, y = np.cos(ang) + i*1.5, np.sin(ang) + j*np.sqrt(3)
                                ax.plot([x, x], [y, y], [z, z + c_a/2], color='gray', lw=1, alpha=0.3)

        elif cell_type == "CsCl (BCC)":
            # Corners: Cl-
            for i in range(lattice_size + 1):
                for j in range(lattice_size + 1):
                    for k in range(lattice_size + 1):
                        draw_sphere(ax, (i, j, k), rad, '#D32F2F') # 빨강 (Cl-)
            # Centers: Cs+
            for i in range(lattice_size):
                for j in range(lattice_size):
                    for k in range(lattice_size):
                        draw_sphere(ax, (i+0.5, j+0.5, k+0.5), rad*1.2, '#1976D2') # 파랑 (Cs+)

        elif cell_type == "CaF2 (Fluorite)":
            # Ca2+: FCC
            for i in range(lattice_size + 1):
                for j in range(lattice_size + 1):
                    for k in range(lattice_size + 1):
                        draw_sphere(ax, (i, j, k), rad, '#1976D2') # 파랑 (Ca2+)
            # F-: 8 sites in unit cell
            for i in range(lattice_size):
                for j in range(lattice_size):
                    for k in range(lattice_size):
                        off = np.array([i, j, k])
                        f_pts = [np.array([0.25,0.25,0.25]), np.array([0.75,0.25,0.25]), np.array([0.25,0.75,0.25]), np.array([0.75,0.75,0.25]),
                                 np.array([0.25,0.25,0.75]), np.array([0.75,0.25,0.75]), np.array([0.25,0.75,0.75]), np.array([0.75,0.75,0.75])]
                        for p in f_pts: draw_sphere(ax, p + off, rad*0.6, '#FBC02D') # 노랑 (F-)

        elif cell_type == "NaCl (Rock Salt)":
            for i in range(lattice_size + 1):
                for j in range(lattice_size + 1):
                    for k in range(lattice_size + 1):
                        # Cl-
                        if (i+j+k) % 2 == 0: draw_sphere(ax, (i, j, k), rad*1.1, '#D32F2F')
                        # Na+
                        else: draw_sphere(ax, (i, j, k), rad*0.7, '#E0E0E0')

        elif cell_type == "Rhombohedral":
            alpha = np.radians(60)
            a1, a2, a3 = np.array([1,0,0]), np.array([np.cos(alpha), np.sin(alpha), 0]), np.array([np.cos(alpha), 0.5, 0.8])
            for i in range(lattice_size + 1):
                for j in range(lattice_size + 1):
                    for k in range(lattice_size + 1):
                        pos = i*a1 + j*a2 + k*a3
                        draw_sphere(ax, pos, rad, '#C2185B')
                        if i < lattice_size: ax.plot([pos[0], (pos+a1)[0]], [pos[1], (pos+a1)[1]], [pos[2], (pos+a1)[2]], color='gray', lw=1, alpha=0.3)
                        if j < lattice_size: ax.plot([pos[0], (pos+a2)[0]], [pos[1], (pos+a2)[1]], [pos[2], (pos+a2)[2]], color='gray', lw=1, alpha=0.3)
                        if k < lattice_size: ax.plot([pos[0], (pos+a3)[0]], [pos[1], (pos+a3)[1]], [pos[2], (pos+a3)[2]], color='gray', lw=1, alpha=0.3)

        # 3. 단위 세포 강조박스 (Dashed Box) - 이미지 스타일 재현
        if lattice_size > 1:
            for start, end in [([0,0,0],[1,0,0]), ([1,0,0],[1,1,0]), ([1,1,0],[0,1,0]), ([0,1,0],[0,0,0]),
                               ([0,0,1],[1,0,1]), ([1,0,1],[1,1,1]), ([1,1,1],[0,1,1]), ([0,1,1],[0,0,1]),
                               ([0,0,0],[0,0,1]), ([1,0,0],[1,0,1]), ([1,1,0],[1,1,1]), ([0,1,0],[0,1,1])]:
                ax.plot([start[0], end[0]], [start[1], end[1]], [start[2], end[2]], color='#FFEB3B', lw=3, ls='--', alpha=0.9, zorder=10)

    else: # Custom Coords
        if custom_coords:
            # Support both newline and semicolon as delimiters
            data = custom_coords.replace(';', '\n')
            for line_idx, p in enumerate(data.split('\n')):
                if p.strip():
                    try:
                        vals = [v.strip() for v in p.split(',')]
                        if len(vals) >= 3:
                            x, y, z = float(vals[0]), float(vals[1]), float(vals[2])
                            c = vals[3] if len(vals) > 3 else 'red'
                            draw_sphere(ax, (x, y, z), 0.15, c)
                    except Exception as e:
                        errors.append(f"{line_idx+1}번째 줄 오류: {e}")

    ax.set_title(f"Detailed {cell_type} Lattice", fontsize=16, fontweight='bold')
    ax.axis('off')

    # 뷰 각도 조절 (더 입체적으로 보이게)
    ax.view_init(elev=elev, azim=azim)

    # 배경을 흰색으로 강제 (워드 다운로드 시 검은 배경 방지)
    fig.patch.set_facecolor('white')
    ax.set_facecolor('white')

    img_stream = io.BytesIO()
    fig.savefig(img_stream, format='png', bbox_inches='tight', dpi=200, facecolor='white')
    img_stream.seek(0)
    return img_stream, errors

def draw_quantum_graph(graph_type, custom_expr="", x_min=0.0, x_max=10.0):
    from matplotlib.figure import Figure
    fig = Figure(figsize=(6, 4))
    ax = fig.subplots()
    if graph_type == "1D Box 파동함수 (n=1,2,3)":
        x = np.linspace(0, 3, 100)
        ax.plot(x, np.sin(1 * np.pi * x / 3) + 4, label='n=1')
        ax.plot(x, np.sin(2 * np.pi * x / 3) + 2, label='n=2')
        ax.plot(x, np.sin(3 * np.pi * x / 3), label='n=3')
        ax.axis('off')
        ax.legend()
    elif graph_type == "2s 오비탈 확률 밀도":
        r = np.linspace(0, 15, 200)
        R2s = (1/(2*np.sqrt(2))) * (1)**(1.5) * (2 - r) * np.exp(-r/2)
        ax.plot(r, r**2 * R2s**2, color='purple', lw=2)
    elif graph_type == "직접 수식 입력 (Custom)":
        x = np.linspace(x_min, x_max, 300)
        safe_dict = {
            'x': x, 'np': np, 'sin': np.sin, 'cos': np.cos, 'tan': np.tan,
            'exp': np.exp, 'log': np.log, 'sqrt': np.sqrt, 'pi': np.pi, 'e': np.e,
            'abs': np.abs
        }
        try:
            # Replace common math terms to numpy terms for convenience
            expr = custom_expr.replace('^', '**').replace('sin', 'np.sin').replace('cos', 'np.cos').replace('exp', 'np.exp').replace('np.np.', 'np.')
            y = eval(expr, {"__builtins__": None}, safe_dict)
            ax.plot(x, y, color='blue', lw=2, label=custom_expr)
            ax.legend()
            ax.grid(True, linestyle='--', alpha=0.6)
        except Exception as e:
            ax.text(0.5, 0.5, f"수식 오류!\n{e}", ha='center', va='center', color='red', transform=ax.transAxes)

    img_stream = io.BytesIO()
    fig.savefig(img_stream, format='png', bbox_inches='tight', facecolor='white')
    img_stream.seek(0)
    return img_stream

def get_molecule_image(name):
    name = name.strip()
    urls = [
        f"https://pubchem.ncbi.nlm.nih.gov/rest/pug/compound/name/{name}/PNG?record_type=2d&image_size=large",
        f"https://pubchem.ncbi.nlm.nih.gov/rest/pug/compound/fastformula/{name}/PNG?record_type=2d&image_size=large"
    ]
    for url in urls:
        try:
            req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
            with urllib.request.urlopen(req) as response:
                return io.BytesIO(response.read())
        except:
            continue
    st.error(f"'{name}'에 대한 이미지를 PubChem에서 찾을 수 없습니다. (이름이나 분자식을 확인해주세요)")
    return None

def draw_schematic(shape_type, color='#2E5BFF', custom_data="", **kwargs):
    from matplotlib.figure import Figure
    errors = []
    fig = Figure(figsize=(4, 4))

    if shape_type == "직접 입력 (Custom)":
        ax = fig.add_subplot(111)
        ax.set_aspect('equal')
        ax.axis('off')
        ax.set_xlim(0, 10)
        ax.set_ylim(0, 10)
        for line_idx, line in enumerate(custom_data.split('\n')):
            parts = [p.strip() for p in line.split(',')]
            if not parts or not parts[0]: continue
            cmd = parts[0].upper()
            try:
                if cmd == 'RECT' and len(parts) >= 5:
                    x, y, w, h = map(float, parts[1:5])
                    c = parts[5] if len(parts) > 5 else color
                    rect = plt.Rectangle((x, y), w, h, fill=True, color=c, alpha=0.7, ec='k', lw=2)
                    ax.add_patch(rect)
                elif cmd == 'CIRCLE' and len(parts) >= 4:
                    x, y, r = map(float, parts[1:4])
                    c = parts[4] if len(parts) > 4 else color
                    circle = plt.Circle((x, y), r, fill=True, color=c, alpha=0.7, ec='k', lw=2)
                    ax.add_patch(circle)
                elif cmd == 'TEXT' and len(parts) >= 4:
                    t, x, y = parts[1], float(parts[2]), float(parts[3])
                    s = int(parts[4]) if len(parts) > 4 else 12
                    c = parts[5] if len(parts) > 5 else 'black'
                    ax.text(x, y, t, fontsize=s, color=c, ha='center', va='center')
                elif cmd == 'LINE' and len(parts) >= 5:
                    x1, y1, x2, y2 = map(float, parts[1:5])
                    c = parts[5] if len(parts) > 5 else 'black'
                    ax.plot([x1, x2], [y1, y2], color=c, lw=2)
            except Exception as e:
                errors.append(f"{line_idx+1}번째 줄 오류: {e}")

    if shape_type == "정육면체 (Cube)":
        ax = fig.add_subplot(111, projection='3d')
        length = kwargs.get('length', 1.0)
        width = kwargs.get('width', 1.0)
        height = kwargs.get('height', 1.0)

        r_x = [0, length]
        r_y = [0, width]
        X, Y = np.meshgrid(r_x, r_y)

        z_top = np.full((2,2), height)
        z_bot = np.zeros((2,2))

        from matplotlib.colors import LightSource
        ls = LightSource(azdeg=315, altdeg=45)

        ax.plot_surface(X, Y, z_top, alpha=0.9, color=color, edgecolor='none', shade=True, lightsource=ls, antialiased=True)
        ax.plot_surface(X, Y, z_bot, alpha=0.9, color=color, edgecolor='none', shade=True, lightsource=ls, antialiased=True)

        X_xz, Z_xz = np.meshgrid(r_x, [0, height])
        ax.plot_surface(X_xz, np.zeros((2,2)), Z_xz, alpha=0.9, color=color, edgecolor='none', shade=True, lightsource=ls, antialiased=True)
        ax.plot_surface(X_xz, np.full((2,2), width), Z_xz, alpha=0.9, color=color, edgecolor='none', shade=True, lightsource=ls, antialiased=True)

        Y_yz, Z_yz = np.meshgrid(r_y, [0, height])
        ax.plot_surface(np.zeros((2,2)), Y_yz, Z_yz, alpha=0.9, color=color, edgecolor='none', shade=True, lightsource=ls, antialiased=True)
        ax.plot_surface(np.full((2,2), length), Y_yz, Z_yz, alpha=0.9, color=color, edgecolor='none', shade=True, lightsource=ls, antialiased=True)

        ax.set_axis_off()
        ax.view_init(elev=20, azim=30)
    else:
        ax = fig.add_subplot(111)
        ax.set_aspect('equal')
        ax.axis('off')

        import matplotlib.patheffects as pe
        shadow = [pe.SimplePatchShadow(shadow_color='#94a3b8', offset=(3,-3), alpha=0.5), pe.Normal()]
        
        if shape_type == "정사각형 (Square)":
            w = kwargs.get('width', 0.8)
            h = kwargs.get('height', 0.8)
            rect = plt.Rectangle(((1-w)/2, (1-h)/2), w, h, fill=True, color=color, alpha=0.85, ec='none', path_effects=shadow)
            ax.add_patch(rect)
            ax.set_xlim(0, 1)
            ax.set_ylim(0, 1)
        elif shape_type == "원형 (Circle)":
            r = kwargs.get('radius', 0.4)
            circle = plt.Circle((0.5, 0.5), r, fill=True, color=color, alpha=0.85, ec='none', path_effects=shadow)
            ax.add_patch(circle)
            ax.set_xlim(0, 1)
            ax.set_ylim(0, 1)
        elif shape_type == "다각형 (Polygon)":
            n = kwargs.get('n_sides', 6)
            from matplotlib.patches import Polygon
            angles = np.linspace(0, 2 * np.pi, n, endpoint=False)
            points = np.column_stack((0.5 + 0.4 * np.cos(angles), 0.5 + 0.4 * np.sin(angles)))
            poly = Polygon(points, fill=True, color=color, alpha=0.85, ec='none', path_effects=shadow)
            ax.add_patch(poly)
            ax.set_xlim(0, 1)
            ax.set_ylim(0, 1)

    img_stream = io.BytesIO()
    fig.savefig(img_stream, format='png', bbox_inches='tight', facecolor='white')
    img_stream.seek(0)
    return img_stream, errors

def draw_lewis_structure(molecule, custom_data=""):
    from matplotlib.figure import Figure
    errors = []
    fig = Figure(figsize=(4, 4))
    ax = fig.subplots()
    ax.set_aspect('equal')
    ax.axis('off')
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)

    def draw_pair(x, y, dx, dy):
        # 전자쌍 글로우(Glow) 효과 추가
        ax.plot([x - dx, x + dx], [y - dy, y + dy], 'o', color='#3B82F6', markersize=10, alpha=0.3)
        ax.plot([x - dx, x + dx], [y - dy, y + dy], 'o', color='#1E3A8A', markersize=5)

    def text_sym(x, y, sym):
        import matplotlib.patheffects as pe
        ax.text(x, y, sym, fontsize=38, ha='center', va='center', fontweight='900', color='#0F172A',
                path_effects=[pe.withStroke(linewidth=4, foreground='white'), pe.SimpleLineShadow(shadow_color='#cbd5e1', offset=(2,-2)), pe.Normal()])

    if molecule == "H2O (물)":
        text_sym(0.5, 0.5, "O")
        text_sym(0.2, 0.5, "H")
        text_sym(0.8, 0.5, "H")
        ax.plot([0.3, 0.4], [0.5, 0.5], 'k-', lw=2)
        ax.plot([0.6, 0.7], [0.5, 0.5], 'k-', lw=2)
        draw_pair(0.5, 0.65, 0.04, 0)
        draw_pair(0.5, 0.35, 0.04, 0)

    elif molecule == "CO2 (이산화탄소)":
        text_sym(0.5, 0.5, "C")
        text_sym(0.2, 0.5, "O")
        text_sym(0.8, 0.5, "O")
        ax.plot([0.3, 0.4], [0.53, 0.53], 'k-', lw=2)
        ax.plot([0.3, 0.4], [0.47, 0.47], 'k-', lw=2)
        ax.plot([0.6, 0.7], [0.53, 0.53], 'k-', lw=2)
        ax.plot([0.6, 0.7], [0.47, 0.47], 'k-', lw=2)
        draw_pair(0.2, 0.65, 0.04, 0)
        draw_pair(0.2, 0.35, 0.04, 0)
        draw_pair(0.8, 0.65, 0.04, 0)
        draw_pair(0.8, 0.35, 0.04, 0)

    elif molecule == "NH3 (암모니아)":
        text_sym(0.5, 0.5, "N")
        text_sym(0.2, 0.5, "H")
        text_sym(0.8, 0.5, "H")
        text_sym(0.5, 0.2, "H")
        ax.plot([0.3, 0.4], [0.5, 0.5], 'k-', lw=2)
        ax.plot([0.6, 0.7], [0.5, 0.5], 'k-', lw=2)
        ax.plot([0.5, 0.5], [0.4, 0.3], 'k-', lw=2)
        draw_pair(0.5, 0.65, 0.04, 0)

    elif molecule == "CH4 (메테인)":
        text_sym(0.5, 0.5, "C")
        text_sym(0.2, 0.5, "H")
        text_sym(0.8, 0.5, "H")
        text_sym(0.5, 0.8, "H")
        text_sym(0.5, 0.2, "H")
        ax.plot([0.3, 0.4], [0.5, 0.5], 'k-', lw=2)
        ax.plot([0.6, 0.7], [0.5, 0.5], 'k-', lw=2)
        ax.plot([0.5, 0.5], [0.6, 0.7], 'k-', lw=2)
        ax.plot([0.5, 0.5], [0.4, 0.3], 'k-', lw=2)

    elif molecule == "O2 (산소 분자)":
        text_sym(0.35, 0.5, "O")
        text_sym(0.65, 0.5, "O")
        ax.plot([0.45, 0.55], [0.53, 0.53], 'k-', lw=2)
        ax.plot([0.45, 0.55], [0.47, 0.47], 'k-', lw=2)
        draw_pair(0.35, 0.65, 0.04, 0)
        draw_pair(0.35, 0.35, 0.04, 0)
        draw_pair(0.65, 0.65, 0.04, 0)
        draw_pair(0.65, 0.35, 0.04, 0)

    elif molecule == "N2 (질소 분자)":
        text_sym(0.35, 0.5, "N")
        text_sym(0.65, 0.5, "N")
        ax.plot([0.45, 0.55], [0.55, 0.55], 'k-', lw=2)
        ax.plot([0.45, 0.55], [0.50, 0.50], 'k-', lw=2)
        ax.plot([0.45, 0.55], [0.45, 0.45], 'k-', lw=2)
        draw_pair(0.2, 0.5, 0, 0.04)
        draw_pair(0.8, 0.5, 0, 0.04)

    elif molecule == "HCl (염화수소)":
        text_sym(0.3, 0.5, "H")
        text_sym(0.7, 0.5, "Cl")
        ax.plot([0.4, 0.55], [0.5, 0.5], 'k-', lw=2)
        draw_pair(0.7, 0.65, 0.04, 0)
        draw_pair(0.7, 0.35, 0.04, 0)
        draw_pair(0.85, 0.5, 0, 0.04)

    elif molecule == "HCN (사이안화수소)":
        text_sym(0.2, 0.5, "H")
        text_sym(0.5, 0.5, "C")
        text_sym(0.8, 0.5, "N")
        ax.plot([0.3, 0.4], [0.5, 0.5], 'k-', lw=2)
        ax.plot([0.6, 0.7], [0.55, 0.55], 'k-', lw=2)
        ax.plot([0.6, 0.7], [0.50, 0.50], 'k-', lw=2)
        ax.plot([0.6, 0.7], [0.45, 0.45], 'k-', lw=2)
        draw_pair(0.95, 0.5, 0, 0.04)

    elif molecule == "C2H4 (에텐)":
        text_sym(0.35, 0.5, "C")
        text_sym(0.65, 0.5, "C")
        text_sym(0.15, 0.7, "H")
        text_sym(0.15, 0.3, "H")
        text_sym(0.85, 0.7, "H")
        text_sym(0.85, 0.3, "H")
        ax.plot([0.45, 0.55], [0.53, 0.53], 'k-', lw=2)
        ax.plot([0.45, 0.55], [0.47, 0.47], 'k-', lw=2)
        ax.plot([0.25, 0.3], [0.6, 0.55], 'k-', lw=2)
        ax.plot([0.25, 0.3], [0.4, 0.45], 'k-', lw=2)
        ax.plot([0.75, 0.7], [0.6, 0.55], 'k-', lw=2)
        ax.plot([0.75, 0.7], [0.4, 0.45], 'k-', lw=2)

    elif molecule == "직접 입력 (Custom)":
        for line_idx, line in enumerate(custom_data.split('\n')):
            if not line.strip(): continue
            parts = [p.strip() for p in line.split(',')]
            if not parts or not parts[0]: continue
            cmd = parts[0].upper()
            try:
                if cmd == 'TEXT' and len(parts) >= 4:
                    t, x, y = parts[1], float(parts[2]), float(parts[3])
                    fs = int(parts[4]) if len(parts) > 4 else 36
                    color = parts[5] if len(parts) > 5 else '#1E3A8A'
                    ax.text(x, y, t, fontsize=fs, ha='center', va='center', fontweight='bold', color=color)
                elif cmd == 'LINE' and len(parts) >= 5:
                    x1, y1, x2, y2 = map(float, parts[1:5])
                    lw = float(parts[5]) if len(parts) > 5 else 2
                    color = parts[6] if len(parts) > 6 else 'k'
                    ax.plot([x1, x2], [y1, y2], '-', color=color, lw=lw)
                elif cmd == 'DOTS' and len(parts) >= 5:
                    x, y, dx, dy = map(float, parts[1:5])
                    draw_pair(x, y, dx, dy)
            except Exception as e:
                errors.append(f"{line_idx+1}번째 줄 오류: {e}")

    img_stream = io.BytesIO()
    fig.savefig(img_stream, format='png', bbox_inches='tight', facecolor='white', dpi=150)
    img_stream.seek(0)
    return img_stream, errors

def draw_orbital_diagram(molecule_type, custom_data="", elev=25, azim=-55):
    from matplotlib.figure import Figure
    errors = []
    fig = Figure(figsize=(8, 6))
    ax = fig.add_subplot(111, projection='3d')
    ax.set_axis_off()

    # 고도화: 배경 제거 및 프리미엄 컬러 세팅
    fig.patch.set_facecolor('#ffffff')
    ax.set_facecolor('#ffffff')

    def draw_pz(center, color='#3B82F6', label=''):
        size = 0.3
        u = np.linspace(0, 2 * np.pi, 30)
        v = np.linspace(0, np.pi, 20)
        x = size * 0.8 * np.outer(np.cos(u), np.sin(v)) + center[0]
        y = size * 0.8 * np.outer(np.sin(u), np.sin(v)) + center[1]

        from matplotlib.colors import LightSource
        ls = LightSource(azdeg=315, altdeg=45)

        z_top = size * 1.5 * np.outer(np.ones(np.size(u)), np.cos(v)) + center[2] + size*1.2
        ax.plot_surface(x, y, z_top, color=color, alpha=0.85, edgecolor='none', antialiased=True, shade=True, lightsource=ls)

        z_bot = size * 1.5 * np.outer(np.ones(np.size(u)), np.cos(v)) + center[2] - size*1.2
        ax.plot_surface(x, y, z_bot, color=color, alpha=0.85, edgecolor='none', antialiased=True, shade=True, lightsource=ls)

        if label:
            ax.text(center[0], center[1]+0.2, center[2]+size*2.5, label, color=color, fontsize=12, fontweight='bold')

    def draw_py(center, color='#F59E0B', label=''):
        size = 0.3
        u = np.linspace(0, 2 * np.pi, 30)
        v = np.linspace(0, np.pi, 20)
        x = size * 0.8 * np.outer(np.sin(u), np.sin(v)) + center[0]
        z = size * 0.8 * np.outer(np.cos(u), np.sin(v)) + center[2]
        
        from matplotlib.colors import LightSource
        ls = LightSource(azdeg=315, altdeg=45)

        y_top = size * 1.5 * np.outer(np.ones(np.size(u)), np.cos(v)) + center[1] + size*1.2
        ax.plot_surface(x, y_top, z, color=color, alpha=0.85, edgecolor='none', antialiased=True, shade=True, lightsource=ls)

        y_bot = size * 1.5 * np.outer(np.ones(np.size(u)), np.cos(v)) + center[1] - size*1.2
        ax.plot_surface(x, y_bot, z, color=color, alpha=0.85, edgecolor='none', antialiased=True, shade=True, lightsource=ls)

        if label:
            ax.text(center[0], center[1]+size*2.5, center[2]+0.2, label, color=color, fontsize=12, fontweight='bold')

    def draw_sp2(center, angle_deg, color, label=''):
        size = 0.25
        u = np.linspace(0, 2 * np.pi, 30)
        v = np.linspace(0, np.pi, 20)
        rad = np.radians(angle_deg)

        X = size * 1.5 * np.outer(np.cos(u), np.sin(v))
        Y = size * 0.6 * np.outer(np.sin(u), np.sin(v))
        Z = size * 0.6 * np.outer(np.ones(np.size(u)), np.cos(v))

        X_rot = X * np.cos(rad) - Y * np.sin(rad)
        Y_rot = X * np.sin(rad) + Y * np.cos(rad)

        offset = size * 1.0
        cx = center[0] + offset * np.cos(rad)
        cy = center[1] + offset * np.sin(rad)
        cz = center[2]

        from matplotlib.colors import LightSource
        ls = LightSource(azdeg=315, altdeg=45)
        ax.plot_surface(X_rot + cx, Y_rot + cy, Z + cz, color=color, alpha=0.85, edgecolor='none', antialiased=True, shade=True, lightsource=ls)
        if label:
            ax.text(cx + 0.3*np.cos(rad), cy + 0.3*np.sin(rad), cz, label, color=color, fontsize=10)

    def draw_s(center, color='#9CA3AF', label=''):
        size = 0.3
        u = np.linspace(0, 2 * np.pi, 30)
        v = np.linspace(0, np.pi, 20)
        X = size * np.outer(np.cos(u), np.sin(v)) + center[0]
        Y = size * np.outer(np.sin(u), np.sin(v)) + center[1]
        Z = size * np.outer(np.ones(np.size(u)), np.cos(v)) + center[2]
        
        from matplotlib.colors import LightSource
        ls = LightSource(azdeg=315, altdeg=45)
        ax.plot_surface(X, Y, Z, color=color, alpha=0.85, edgecolor='none', antialiased=True, shade=True, lightsource=ls)
        if label:
            ax.text(center[0], center[1], center[2]+0.4, label, color='gray', fontsize=12)

    # Base coords
    N_pos = (0, 0, 0)
    Oa_pos = (0, 1.2, 0)
    Ob_pos = (-1.04, -0.6, 0) # 210 deg
    Oc_pos = (1.04, -0.6, 0)  # 330 deg

    if molecule_type == "NO3- (질산 이온)":
        ax.text(N_pos[0]+0.1, N_pos[1], 0, 'N', fontsize=16, fontweight='bold')
        ax.text(Oa_pos[0]+0.1, Oa_pos[1], 0, 'O_a', fontsize=14)
        ax.text(Ob_pos[0]-0.3, Ob_pos[1]-0.2, 0, 'O_b', fontsize=14)
        ax.text(Oc_pos[0]+0.1, Oc_pos[1]-0.2, 0, 'O_c', fontsize=14)

        ax.plot([0, Oa_pos[0]], [0, Oa_pos[1]], [0, 0], 'k-', lw=2)
        ax.plot([0, Ob_pos[0]], [0, Ob_pos[1]], [0, 0], 'k-', lw=2)
        ax.plot([0, Oc_pos[0]], [0, Oc_pos[1]], [0, 0], 'k-', lw=2)

        draw_pz(N_pos, color='#60A5FA', label='N 2p_z')
        draw_pz(Oa_pos, color='#93C5FD', label='O 2p_z')
        draw_pz(Ob_pos, color='#93C5FD')
        draw_pz(Oc_pos, color='#93C5FD')

        draw_sp2(N_pos, 90, '#4ADE80', 'N 2sp²')
        draw_sp2(N_pos, 210, '#4ADE80')
        draw_sp2(N_pos, 330, '#4ADE80')

        draw_sp2(Oa_pos, 270, '#F87171', 'O 2sp²')
        draw_sp2(Ob_pos, 30, '#F87171')
        draw_sp2(Oc_pos, 150, '#F87171')
        ax.set_title("Molecular Orbitals of NO3-", fontsize=16, y=0.95)

    elif molecule_type == "HNO3 (질산)":
        ax.text(N_pos[0]+0.1, N_pos[1], 0, 'N', fontsize=16, fontweight='bold')
        ax.text(Oa_pos[0]+0.1, Oa_pos[1], 0, 'O_a', fontsize=14)
        ax.text(Ob_pos[0]-0.3, Ob_pos[1]-0.2, 0, 'O_b', fontsize=14)
        ax.text(Oc_pos[0]+0.1, Oc_pos[1]-0.2, 0, 'O_c', fontsize=14)

        ax.plot([0, Oa_pos[0]], [0, Oa_pos[1]], [0, 0], 'k-', lw=2)
        ax.plot([0, Ob_pos[0]], [0, Ob_pos[1]], [0, 0], 'k-', lw=2)
        ax.plot([0, Oc_pos[0]], [0, Oc_pos[1]], [0, 0], 'k-', lw=2)

        draw_pz(N_pos, color='#60A5FA', label='N 2p_z')
        draw_pz(Oa_pos, color='#93C5FD', label='O 2p_z')
        draw_pz(Oc_pos, color='#93C5FD')

        H_pos = (-1.8, -0.6, 0)
        ax.text(H_pos[0]-0.2, H_pos[1], 0, 'H', fontsize=14)
        ax.plot([Ob_pos[0], H_pos[0]], [Ob_pos[1], H_pos[1]], [0, 0], 'k-', lw=2)

        draw_s(H_pos, label='H 1s')

        draw_sp2(N_pos, 90, '#4ADE80', 'N 2sp²')
        draw_sp2(N_pos, 210, '#4ADE80')
        draw_sp2(N_pos, 330, '#4ADE80')

        draw_sp2(Oa_pos, 270, '#F87171', 'O 2sp²')
        draw_sp2(Ob_pos, 30, '#F87171')
        draw_sp2(Ob_pos, 180, '#F87171')
        draw_sp2(Oc_pos, 150, '#F87171')
        ax.set_title("Molecular Orbitals of HNO3", fontsize=16, y=0.95)

    elif molecule_type in ["C2H4 (에텐)", "에텐 (C2H4)"]:
        C1_pos = (-0.7, 0, 0)
        C2_pos = (0.7, 0, 0)
        H1_pos = (-1.3, 0.8, 0)
        H2_pos = (-1.3, -0.8, 0)
        H3_pos = (1.3, 0.8, 0)
        H4_pos = (1.3, -0.8, 0)

        ax.text(C1_pos[0]+0.1, C1_pos[1], 0, 'C', fontsize=16, fontweight='bold')
        ax.text(C2_pos[0]+0.1, C2_pos[1], 0, 'C', fontsize=16, fontweight='bold')

        ax.plot([C1_pos[0], C2_pos[0]], [C1_pos[1], C2_pos[1]], [0, 0], 'k-', lw=2)
        ax.plot([C1_pos[0], H1_pos[0]], [C1_pos[1], H1_pos[1]], [0, 0], 'k-', lw=2)
        ax.plot([C1_pos[0], H2_pos[0]], [C1_pos[1], H2_pos[1]], [0, 0], 'k-', lw=2)
        ax.plot([C2_pos[0], H3_pos[0]], [C2_pos[1], H3_pos[1]], [0, 0], 'k-', lw=2)
        ax.plot([C2_pos[0], H4_pos[0]], [C2_pos[1], H4_pos[1]], [0, 0], 'k-', lw=2)

        draw_pz(C1_pos, color='#60A5FA', label='C 2p_z')
        draw_pz(C2_pos, color='#60A5FA', label='C 2p_z')

        draw_sp2(C1_pos, 0, '#4ADE80')
        draw_sp2(C2_pos, 180, '#4ADE80')
        draw_sp2(C1_pos, 125, '#4ADE80')
        draw_sp2(C1_pos, 235, '#4ADE80')
        draw_sp2(C2_pos, 55, '#4ADE80')
        draw_sp2(C2_pos, 305, '#4ADE80')

        draw_s(H1_pos, label='H 1s')
        draw_s(H2_pos)
        draw_s(H3_pos)
        draw_s(H4_pos)
        ax.set_title("Molecular Orbitals of Ethene (C2H4)", fontsize=16, y=0.95)

    elif molecule_type == "에타인 (C2H2)":
        C1_pos = (-0.7, 0, 0)
        C2_pos = (0.7, 0, 0)
        H1_pos = (-1.8, 0, 0)
        H2_pos = (1.8, 0, 0)

        ax.text(C1_pos[0]+0.1, C1_pos[1], 0, 'C', fontsize=16, fontweight='bold')
        ax.text(C2_pos[0]-0.2, C2_pos[1], 0, 'C', fontsize=16, fontweight='bold')

        ax.plot([C1_pos[0], C2_pos[0]], [C1_pos[1], C2_pos[1]], [0, 0], 'k-', lw=2)
        ax.plot([C1_pos[0], H1_pos[0]], [C1_pos[1], H1_pos[1]], [0, 0], 'k-', lw=2)
        ax.plot([C2_pos[0], H2_pos[0]], [C2_pos[1], H2_pos[1]], [0, 0], 'k-', lw=2)

        # pi bonds (pz and py)
        draw_pz(C1_pos, color='#60A5FA', label='C 2p_z')
        draw_pz(C2_pos, color='#60A5FA')
        draw_py(C1_pos, color='#F59E0B', label='C 2p_y')
        draw_py(C2_pos, color='#F59E0B')

        # sp orbitals (along x-axis)
        draw_sp2(C1_pos, 0, '#4ADE80', 'C sp')
        draw_sp2(C1_pos, 180, '#4ADE80')
        draw_sp2(C2_pos, 0, '#4ADE80')
        draw_sp2(C2_pos, 180, '#4ADE80')

        draw_s(H1_pos, label='H 1s')
        draw_s(H2_pos)
        ax.set_title("Molecular Orbitals of Ethyne (C2H2)", fontsize=16, y=0.95)

    elif molecule_type in ["물 (H2O)", "H2O (물)"]:
        O_pos = (0, 0, 0)
        H1_pos = (-1.2, -1.0, 0)
        H2_pos = (1.2, -1.0, 0)

        ax.text(O_pos[0]+0.1, O_pos[1], 0, 'O', fontsize=16, fontweight='bold')
        ax.text(H1_pos[0]-0.2, H1_pos[1]-0.2, 0, 'H', fontsize=14)
        ax.text(H2_pos[0]+0.1, H2_pos[1]-0.2, 0, 'H', fontsize=14)

        ax.plot([O_pos[0], H1_pos[0]], [O_pos[1], H1_pos[1]], [0, 0], 'k-', lw=2)
        ax.plot([O_pos[0], H2_pos[0]], [O_pos[1], H2_pos[1]], [0, 0], 'k-', lw=2)

        # O sp3 hybrid orbitals (approximated for visual)
        draw_sp2(O_pos, 220, '#F87171', 'O sp³') 
        draw_sp2(O_pos, 320, '#F87171')
        
        # Lone pairs
        draw_sp2(O_pos, 90, '#FCD34D', 'Lone Pair') 
        draw_sp2(O_pos, 50, '#FCD34D')

        draw_s(H1_pos, label='H 1s')
        draw_s(H2_pos, label='H 1s')
        ax.set_title("Molecular Orbitals of Water (H2O)", fontsize=16, y=0.95)

    elif molecule_type == "직접 입력 (Custom)":
        for line_idx, line in enumerate(custom_data.split('\n')):
            parts = [p.strip() for p in line.split(',')]
            if not parts or not parts[0]: continue
            cmd = parts[0].upper()
            try:
                if cmd == 'ATOM' and len(parts) >= 5:
                    label, x, y, z = parts[1], float(parts[2]), float(parts[3]), float(parts[4])
                    ax.text(x+0.1, y, z, label, fontsize=14, fontweight='bold')
                elif cmd == 'BOND' and len(parts) >= 7:
                    x1, y1, z1, x2, y2, z2 = map(float, parts[1:7])
                    ax.plot([x1, x2], [y1, y2], [z1, z2], 'k-', lw=2)
                elif cmd == 'PZ' and len(parts) >= 4:
                    x, y, z = map(float, parts[1:4])
                    color = parts[4] if len(parts) > 4 else '#3B82F6'
                    label = parts[5] if len(parts) > 5 else ''
                    draw_pz((x, y, z), color, label)
                elif cmd == 'SP2' and len(parts) >= 5:
                    x, y, z, angle = map(float, parts[1:5])
                    color = parts[5] if len(parts) > 5 else '#4ADE80'
                    label = parts[6] if len(parts) > 6 else ''
                    draw_sp2((x, y, z), angle, color, label)
                elif cmd == 'S' and len(parts) >= 4:
                    x, y, z = map(float, parts[1:4])
                    color = parts[4] if len(parts) > 4 else '#9CA3AF'
                    label = parts[5] if len(parts) > 5 else ''
                    draw_s((x, y, z), color, label)
            except Exception as e:
                errors.append(f"{line_idx+1}번째 줄 오류: {e}")
        ax.set_title("Custom Molecular Orbitals", fontsize=16, y=0.95)

    ax.view_init(elev=elev, azim=azim)

    img_stream = io.BytesIO()
    fig.savefig(img_stream, format='png', bbox_inches='tight', facecolor='white', dpi=200)
    img_stream.seek(0)
    return img_stream, errors

def draw_skeletal_structure(molecule, custom_data=""):
    from matplotlib.figure import Figure
    errors = []
    fig = Figure(figsize=(6, 4))
    ax = fig.subplots()
    ax.set_aspect('equal')
    ax.axis('off')

    import matplotlib.patheffects as pe
    text_pe = [pe.withStroke(linewidth=3, foreground='white')]
    
    def plot_bond(x, y, is_double=False):
        ax.plot(x, y, '-', color='#1E293B', lw=4.5, solid_capstyle='round')

    if molecule == "Butane (뷰테인)":
        plot_bond([0, 1, 2, 3], [0, 0.866, 0, 0.866])
    elif molecule == "Hexane (헥세인)":
        plot_bond([0, 1, 2, 3, 4, 5], [0, 0.866, 0, 0.866, 0, 0.866])
    elif molecule == "Cyclohexane (사이클로헥세인)":
        angles = np.linspace(0, 2 * np.pi, 7)
        plot_bond(np.cos(angles), np.sin(angles))
    elif molecule == "Benzene (벤젠)":
        angles = np.linspace(0, 2 * np.pi, 7)
        plot_bond(np.cos(angles), np.sin(angles))
        circle = plt.Circle((0, 0), 0.65, fill=False, color='#1E293B', lw=4.5)
        ax.add_patch(circle)
    elif molecule == "Acetone (아세톤)":
        plot_bond([0, 1, 2], [0, 0.866, 0])
        # Double bond to O
        plot_bond([0.9, 0.9], [0.866, 1.866])
        plot_bond([1.1, 1.1], [0.866, 1.866])
        ax.text(1, 2.15, 'O', fontsize=26, ha='center', va='center', fontweight='bold', color='#B91C1C', path_effects=text_pe)
    elif molecule == "Acetic Acid (아세트산)":
        plot_bond([0, 1], [0, 0.866]) # C-C
        plot_bond([1, 1.866], [0.866, 0.366]) # C-OH
        ax.text(2.35, 0.1, 'OH', fontsize=26, ha='center', va='center', fontweight='bold', color='#B91C1C', path_effects=text_pe)
        ax.plot([0.9, 0.9], [0.866, 1.866]) # C=O
        plot_bond([1.1, 1.1], [0.866, 1.866])
        ax.text(1, 2.15, 'O', fontsize=26, ha='center', va='center', fontweight='bold', color='#B91C1C', path_effects=text_pe)
    elif molecule == "Aspirin (아스피린)":
        # Benzene ring center at (0,0)
        angles = np.linspace(0, 2 * np.pi, 7)
        plot_bond(np.cos(angles), np.sin(angles))
        circle = plt.Circle((0, 0), 0.65, fill=False, color='#1E293B', lw=4.5)
        ax.add_patch(circle)
        # Acetyl group at top
        plot_bond([0, 0], [1, 2])
        ax.text(0, 2.2, 'O', fontsize=26, ha='center', va='center', fontweight='bold', color='#B91C1C', path_effects=text_pe)
        plot_bond([0.2, 1], [2.2, 2.6])
        plot_bond([1, 1.8], [2.6, 2.2])
        # C=O
        plot_bond([0.9, 0.9], [2.6, 3.4])
        plot_bond([1.1, 1.1], [2.6, 3.4])
        ax.text(1, 3.6, 'O', fontsize=26, ha='center', va='center', fontweight='bold', color='#B91C1C', path_effects=text_pe)
        # Carboxylic acid group at bottom right
        ax.text(0.866, -0.6, 'COOH', fontsize=22, ha='left', va='top', fontweight='bold', color='#B91C1C', path_effects=text_pe)
    elif molecule == "Caffeine (카페인)":
        # Simplified purine ring core
        # Pyrimidine part (left)
        plot_bond([-1, -1, 0, 1, 1, 0, -1], [1, -1, -2, -1, 1, 2, 1])
        ax.text(-1, 1, 'N', fontsize=26, ha='center', va='center', fontweight='bold', color='#1E3A8A', path_effects=text_pe)
        ax.text(-1, -1, 'N', fontsize=26, ha='center', va='center', fontweight='bold', color='#1E3A8A', path_effects=text_pe)
        plot_bond([-1.2, -2], [1.2, 1.8]) # Methyl
        plot_bond([-1.2, -2], [-1.2, -1.8]) # Methyl
        # C=O groups
        plot_bond([0, -0.2], [-2, -2.8])
        plot_bond([0.2, 0], [-2, -2.8])
        ax.text(0, -3.1, 'O', fontsize=26, ha='center', va='center', fontweight='bold', color='#B91C1C', path_effects=text_pe)
        plot_bond([0, -0.2], [2, 2.8])
        plot_bond([0.2, 0], [2, 2.8])
        ax.text(0, 3.1, 'O', fontsize=26, ha='center', va='center', fontweight='bold', color='#B91C1C', path_effects=text_pe)
        # Imidazole part (right)
        plot_bond([1, 2.5, 2.5, 1], [1, 1.5, -1.5, -1])
        ax.text(2.5, 1.5, 'N', fontsize=26, ha='center', va='center', fontweight='bold', color='#1E3A8A', path_effects=text_pe)
        ax.text(2.5, -1.5, 'N', fontsize=26, ha='center', va='center', fontweight='bold', color='#1E3A8A', path_effects=text_pe)
        plot_bond([2.7, 3.5], [1.7, 2]) # Methyl


    elif molecule == "직접 입력 (Custom)":
        for line in custom_data.split('\n'):
            parts = [p.strip() for p in line.split(',')]
            if not parts or not parts[0]: continue
            cmd = parts[0].upper()
            try:
                if cmd in ['TEXT', 'LTEXT', 'RTEXT'] and len(parts) >= 4:
                    t, x, y = parts[1], float(parts[2]), float(parts[3])
                    fs = int(parts[4]) if len(parts) > 4 else 18
                    color = parts[5] if len(parts) > 5 else 'k'
                    ha = 'center'
                    if cmd == 'LTEXT': ha = 'left'
                    elif cmd == 'RTEXT': ha = 'right'
                    ax.text(x, y, t, fontsize=fs, color=color, ha=ha, va='center', fontfamily='serif')
                elif cmd == 'LINE' and len(parts) >= 5:
                    x1, y1, x2, y2 = map(float, parts[1:5])
                    lw = float(parts[5]) if len(parts) > 5 else 2
                    color = parts[6] if len(parts) > 6 else 'k'
                    ax.plot([x1, x2], [y1, y2], '-', color=color, lw=lw, zorder=2)
                elif cmd == 'DLINE' and len(parts) >= 5:
                    x1, y1, x2, y2 = map(float, parts[1:5])
                    lw = float(parts[5]) if len(parts) > 5 else 2
                    color = parts[6] if len(parts) > 6 else 'k'
                    dx, dy = x2 - x1, y2 - y1
                    dist = np.hypot(dx, dy)
                    nx, ny = -dy/dist * 0.08, dx/dist * 0.08
                    ax.plot([x1+nx, x2+nx], [y1+ny, y2+ny], '-', color=color, lw=lw, zorder=2)
                    ax.plot([x1-nx, x2-nx], [y1-ny, y2-ny], '-', color=color, lw=lw, zorder=2)
                elif cmd == 'ARROW' and len(parts) >= 5:
                    x1, y1, x2, y2 = map(float, parts[1:5])
                    color = parts[5] if len(parts) > 5 else 'k'
                    ax.annotate('', xy=(x2, y2), xytext=(x1, y1), arowprops=dict(arowstyle="->", lw=1.5, color=color))
                elif cmd == 'DOTS' and len(parts) >= 5:
                    x1, y1, x2, y2 = map(float, parts[1:5])
                    color = parts[5] if len(parts) > 5 else 'k'
                    ax.plot([x1, x2], [y1, y2], 'o', color=color, markersize=3)
                elif cmd == 'BALL' and len(parts) >= 5:
                    x, y, r = map(float, parts[1:4])
                    color = parts[4]
                    base_circle = plt.Circle((x, y), r, color=color, ec='black', lw=1.5, zorder=3)
                    ax.add_patch(base_circle)
                    hl_circle = plt.Circle((x - r*0.3, y + r*0.3), r*0.3, color='white', alpha=0.4, zorder=4, edgecolor='none')
                    ax.add_patch(hl_circle)
            except Exception as e:
                errors.append(f"{line} 처리 중 오류: {e}")

    img_stream = io.BytesIO()
    fig.savefig(img_stream, format='png', bbox_inches='tight', facecolor='white', dpi=150)
    img_stream.seek(0)
    return img_stream, errors

# ==========================================
# 워드 수식(OMML) 자동 변환 처리용 Pandoc 엔진
# ==========================================
def convert_latex_to_word_docx(markdown_text, output_filename, margins):
    import pypandoc
    import os
    import uuid
    import re

    # --- 1. 철저한 텍스트 정화 (Word 충돌 유발 문자 제거) ---
    # 제어 문자 및 비정상적인 유니코드 문자 제거 (한글/공백/일반기호 유지)
    text_clean = "".join(c for c in markdown_text if c.isprintable() or c in "\n\r\t")
    
    # --- 1-1. 마크다운 기호(**, *) 제거 요청 반영 ---
    # 수식 내의 별표(*)는 유지하기 위해, 텍스트 강조용으로 쓰인 패턴만 제거합니다.
    # 1) 볼드 기호 (**) 제거
    text_clean = re.sub(r'\*\*(.*?)\*\*', r'\1', text_clean)
    # 2) 이탤릭 기호 (*) 제거 (글자 양옆에 붙어있는 경우만)
    text_clean = re.sub(r'(^|[^a-zA-Z0-9])\*(.*?)\*([^a-zA-Z0-9]|$)', r'\1\2\3', text_clean)
    
    sanitized_text = text_clean
    
    # --- 2. 고유 임시 파일 이름 생성 (충돌 방지) ---
    temp_id = str(uuid.uuid4())[:8]
    temp_md = f"temp_{temp_id}.md"
    
    try:
        # Pandoc 설치 확인
        try:
            pypandoc.get_pandoc_version()
        except:
            pypandoc.ensure_pandoc_installed()

        with open(temp_md, "w", encoding="utf-8") as f:
            f.write(sanitized_text)

        # --- 3. Pandoc 자체 여백 설정 사용 (파일 손상 방지 핵심) ---
        # python-docx로 다시 열어 저장하는 과정을 생략하고 Pandoc 단계에서 모든 레이아웃을 결정합니다.
        # 이 방식이 맥/윈도우 공용으로 가장 안정적입니다.
        extra_args = [
            '--standalone',
            f'--variable=margin-top:{margins.get("top", 2.0)}cm',
            f'--variable=margin-bottom:{margins.get("bottom", 2.0)}cm',
            f'--variable=margin-left:{margins.get("left", 2.5)}cm',
            f'--variable=margin-right:{margins.get("right", 2.5)}cm'
        ]

        pypandoc.convert_file(temp_md, 'docx', outputfile=output_filename, extra_args=extra_args)
        
        if os.path.exists(output_filename) and os.path.getsize(output_filename) > 0:
            return True
        else:
            raise Exception("File creation failed or empty.")

    except Exception as e:
        # --- 4. 최후의 수단: 텍스트 보존 모드 (python-docx) ---
        try:
            from docx import Document
            doc = Document()
            doc.add_heading("[안전 복구 모드] 문서 내용", level=1)
            for line in sanitized_text.split('\n'):
                doc.add_paragraph(line)
            doc.save(output_filename)
            return True
        except:
            return False
    finally:
        if os.path.exists(temp_md):
            try: os.remove(temp_md)
            except: pass

# ==========================================
# 사이드바 & 워크스페이스 설정
# ==========================================
with st.sidebar:
    st.markdown("<h2 style='text-align: center; color: #1E3A8A;'>🧪 Chem-Ed Studio</h2>", unsafe_allow_html=True)
    st.caption("Professional AI for Chemistry Education")
    st.markdown("---")

    menu_options = [
        "📓 Notion / MS Word 스타일 매니저 (추천)",
        "🔬 실험 보고서 AI 도우미",
        "🎓 전문가용 LaTeX (Overleaf) 에디터",
        "🧪 도표 & 3D 그림 생성기 / 80페이지+ 초정밀 분석",
        "📝 수기 노트 AI 문서화",
        "💬 실시간 AI 학술 상담 (ChatGPT 스타일)"
    ]
    
    # 세션 상태에서 이전 선택 메뉴를 찾기 위한 유연한 로직
    saved_menu = st.session_state.get("menu_selection", menu_options[0])
    default_idx = 0
    for i, opt in enumerate(menu_options):
        if saved_menu == opt:
            default_idx = i
            break
        # 이름이 변경되었더라도 키워드가 맞으면 해당 메뉴로 복구
        elif ("도표" in saved_menu and "3D" in saved_menu) and ("도표" in opt):
            default_idx = i
            break
        elif ("수기" in saved_menu) and ("수기" in opt):
            default_idx = i
            break
        elif ("실시간" in saved_menu) and ("실시간" in opt):
            default_idx = i
            break

    # StreamlitAPIException 방지: index 파라미터 대신 세션 상태를 직접 동기화
    st.session_state["menu_selection"] = menu_options[default_idx]

    menu = st.radio(
        "🏠 Workspace 메뉴", 
        menu_options, 
        key="menu_selection",
        on_change=save_state # 메뉴를 클릭하는 즉시 파일에 영구 저장
    )

    st.markdown("---")
    # Gemini API 철벽 보안 시스템 (Secrets 가동 중)
    secret_key = st.secrets.get("gemini_api_key", "")
    if secret_key:
        st.session_state.gemini_api_key = secret_key
        # [철벽 방어] 화면이나 로그에 키가 절대 노출되지 않도록 가짜 문자로 마스킹 처리 (UI 알림은 숨김)
        pass
    else:
        st.session_state.gemini_api_key = st.text_input("🔑 Gemini API Key 입력 (쉼표(,)로 여러 개 입력 시 한도 우회!)", type="password", value=st.session_state.get("gemini_api_key", ""))
        if not st.session_state.gemini_api_key:
            st.warning("분석을 위해 Gemini API Key가 필요합니다.")
            
    api_key = st.session_state.gemini_api_key

    # OpenAI (ChatGPT) 설정 추가
    with st.expander("🤖 ChatGPT (OpenAI) 연동 설정"):
        openai_key = st.text_input("OpenAI API Key 입력", type="password", key="openai_api_key_input")
        if openai_key:
            st.session_state.openai_api_key = openai_key
            st.success("✅ ChatGPT 군단 합류 완료!")
        else:
            st.caption("GPT-4o를 분석에 동원하려면 키를 입력하세요.")

    st.markdown("---")
    st.subheader("📥 통합 보고서 다운로드")

    # 임시 파일 생성을 위한 메모리 버퍼
    doc_stream = io.BytesIO()
    st.session_state.word_doc.save(doc_stream)

    st.download_button(
        label="📝 전체 작업 내용 Word로 받기",
        data=doc_stream.getvalue(),
        file_name="SNU_Chem_Report_Total.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True,
        type="primary",
        help="지금까지 '워드에 추가' 버튼을 눌러 쌓인 모든 분석 결과를 하나의 문서로 저장합니다."
    )

    st.markdown("---")
    st.subheader("🗑️ 통합 휴지통")
    trash_count = len(st.session_state.get("global_trash_bin", []))
    with st.expander(f"최근 삭제 항목 ({trash_count}/5)", expanded=False):
        if trash_count == 0:
            st.caption("휴지통이 비어 있습니다.")
        else:
            for i, item in enumerate(st.session_state.global_trash_bin):
                st.markdown(f"**[{i+1}] {item.get('type', '항목')}**")
                st.caption(f"{item['query']} ({item['time']})")
                
                # 복구 버튼
                if st.button("♻️ 이 항목 복구하기", key=f"global_restore_{i}", use_container_width=True):
                    if item.get("type") == "보고서":
                        st.session_state.last_report_result = item['content']
                        st.session_state.last_report_query = item['query']
                        st.session_state["menu_selection"] = "💬 실시간 AI 학술 상담 (ChatGPT 스타일)"
                    st.rerun()
                st.divider()

    st.caption("v2.5.0 Professional Edition")


# 화학교육(Chem-Ed) 교수학습 상세 가이드 데이터베이스
# 화학교육 가이드 데이터 로드 함수
def load_chem_ed_guides():
    base_dir = os.path.join(os.path.dirname(__file__), "chem_ed_guides")
    mapping = {
        "주요 오개념": "misconception.md",
        "화학 결합": "bonding.md",
        "동적 평형": "equilibrium.md",
        "권장 교수학습 모델": "models.md",
        "5E 순환 학습": "5e_model.md",
        "POE 모형": "poe_model.md",
        "평가": "assessment.md"
    }
    data = {}
    for key, filename in mapping.items():
        path = os.path.join(base_dir, filename)
        if os.path.exists(path):
            with open(path, "r", encoding="utf-8") as f:
                data[key] = f.read()
        else:
            # 기본 데이터 유지 (파일이 없을 경우 대비)
            data[key] = f"# {key}\n상세 자료를 불러올 수 없습니다."
    return data

CHEM_ED_GUIDE_DATA = load_chem_ed_guides()
def show_sample_dialog(title, content, target_subject):
    st.markdown("---")
    st.subheader(f"🎓 {title} 미리보기")


    sample_key = f"sample_bytes_{title}"
    if st.button("⚡ 이 샘플을 MS Word 문서로 변환하기 (수식 완벽 지원)", type="primary", use_container_width=True):
        with st.spinner("샘플 수식을 완벽히 변환 중입니다..."):
            import time
            out_file = os.path.join(os.getcwd(), f"Sample_{int(time.time())}.docx")
            full_markdown = f"# {title}\n\n" + content
            margins = {"top": 2.0, "bottom": 2.0, "left": 2.5, "right": 2.5}
            if convert_latex_to_word_docx(full_markdown, out_file, margins):
                with open(out_file, "rb") as f:
                    st.session_state[sample_key] = f.read()
                os.remove(out_file)
                st.success("변환 완료! 아래 버튼을 클릭하여 여세요.")
            else:
                st.error("샘플 변환 중 오류가 발생했습니다.")

    if st.session_state.get(sample_key):
        render_download_and_open_buttons(title, st.session_state[sample_key], f"{title}_Sample.docx", f"sc_{title}")

    import json
    content_escaped = json.dumps(content)
    copy_code = f"""
    <button onclick="copySample(event)" style="background:#0F172A; color:white; border:none; padding:12px 15px; border-radius:8px; cursor:pointer; font-weight:bold; width:100%; font-size:15px; transition:0.2s; box-shadow: 0 2px 4px rgba(0,0,0,0.1);">
        📓 Notion / 옵시디언 등 마크다운 노트에 샘플 복사
    </button>
    <script>
    function copySample(event) {{
        const text = {content_escaped};
        const el = document.createElement('textarea');
        el.value = text;
        document.body.appendChild(el);
        el.select();
        document.execCommand('copy');
        document.body.removeChild(el);
        const btn = event.target;
        const orig = btn.innerText;
        btn.innerText = '✅ 복사 완료! (Notion에서 Ctrl+V 하세요)';
        btn.style.background = '#047857';
        setTimeout(() => {{ btn.innerText = orig; btn.style.background = '#0F172A'; }}, 2500);
    }}
    </script>
    """
    import streamlit.components.v1 as components
    components.html(copy_code, height=60)

    st.markdown(content)

    # 덮어쓰기 방지를 위해 이어붙이기로 변경
    if st.button(f"📥 이 내용을 '{target_subject}' 에디터 아래에 추가하기 (이어붙이기)", use_container_width=True):
        st.session_state.notion_db[target_subject] += "" + content
        if f"editor_{target_subject}" in st.session_state: del st.session_state[f"editor_{target_subject}"]
        st.success("에디터에 성공적으로 추가되었습니다!")
        st.rerun()

def render_chem_ed_core_guide():
    with st.expander("🎓 화학교육(Chem-Ed) 교수학습 핵심 가이드 (예비교사 필독)"):
        st.write("과제나 지도안 작성 시 다음의 교육학적 요소를 점검해 보세요. 각 버튼을 클릭하면 **즉시 내 컴퓨터의 MS Word**로 열립니다.")
        
        ce_cols = st.columns(7)
        items = [
            ("🧠 오개념", "주요 오개념"),
            ("💎 화학결합", "화학 결합"),
            ("⚖️ 동적평형", "동적 평형"),
            ("🧑‍🏫 교수모델", "권장 교수학습 모델"),
            ("🔄 5E순환", "5E 순환 학습"),
            ("🧪 POE모형", "POE 모형"),
            ("📊 평가설계", "평가")
        ]
        
        for i, (label, key) in enumerate(items):
            with ce_cols[i]:
                if st.button(label, key=f"ce_main_{i}", use_container_width=True, type="secondary"):
                    content = CHEM_ED_GUIDE_DATA.get(key, "내용 없음")
                    open_any_word_direct(key, content, "ChemEd_Guide", f"ce_main_dl_{i}")
                
                if st.session_state.get(f"ce_main_dl_{i}"):
                    render_download_and_open_buttons(key, st.session_state[f"ce_main_dl_{i}"], st.session_state[f"ce_main_dl_{i}_filename"], f"ce_main_{i}_persistent")

        st.markdown("---")
        st.markdown("### 📥 화학교육 가이드(.docx) 직접 다운로드 센터")
        st.caption("가이드 내용을 오프라인에서 확인하거나 과제에 활용할 수 있도록 워드 문서로 제공합니다.")
        
        d_col1, d_col2 = st.columns(2)
        for i, (label, key) in enumerate(items):
            target_col = d_col1 if i % 2 == 0 else d_col2
            with target_col:
                if st.button(f"📝 {key} (Word 변환 시작)", key=f"ce_dl_{i}", use_container_width=True):
                    with st.spinner(f"{key} 가이드 변환 중..."):
                        import time
                        out_file = os.path.join(os.getcwd(), f"ChemEd_{i}_{int(time.time())}.docx")
                        content = CHEM_ED_GUIDE_DATA.get(key, "내용 없음")
                        margins = {"top": 2.0, "bottom": 2.0, "left": 2.5, "right": 2.5}
                        if convert_latex_to_word_docx(f"# {key} 상세 가이드\n\n" + content, out_file, margins):
                            with open(out_file, "rb") as f:
                                st.session_state[f"ce_guide_bytes_{i}"] = f.read()
                            st.success(f"✅ {key} 가이드가 준비되었습니다!")
                
                if st.session_state.get(f"ce_guide_bytes_{i}"):
                    render_download_and_open_buttons(key, st.session_state[f"ce_guide_bytes_{i}"], f"ChemEd_{key}.docx", f"ce_guide_{i}")
        
        items = list(CHEM_ED_GUIDE_DATA.items())
        for i in range(len(items)):
            key, content = items[i]
            col = d_col1 if i % 2 == 0 else d_col2
            with col:
                if st.button(f"📝 {key} (Word 변환 시작)", key=f"dl_trigger_{i}", use_container_width=True):
                    out_file = f"ChemEd_Guide_{key.replace(' ', '_')}.docx"
                    margins = (2.54, 2.54, 2.54, 2.54)
                    try:
                        convert_latex_to_word_docx(content, out_file, margins)
                        with open(out_file, "rb") as f:
                            st.session_state[f"chemed_guide_{i}_bytes"] = f.read()
                            st.session_state[f"chemed_guide_{i}_name"] = out_file
                    except Exception as e:
                        st.error(f"변환 중 오류: {e}")
                
                if st.session_state.get(f"chemed_guide_{i}_bytes"):
                    st.download_button(
                        label=f"💾 {key}.docx 다운로드",
                        data=st.session_state[f"chemed_guide_{i}_bytes"],
                        file_name=st.session_state[f"chemed_guide_{i}_name"],
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key=f"dl_btn_{i}_persistent",
                        use_container_width=True
                    )

        st.write("")
        if st.button("📚 화학교육 가이드 전체 통합본 변환 시작", use_container_width=True, type="primary"):
            full_content = "# 화학교육(Chem-Ed) 교수학습 핵심 가이드 통합본\n\n"
            for key, content in CHEM_ED_GUIDE_DATA.items():
                full_content += f"\n\n---\n\n{content}\n"
            
            out_file = "ChemEd_Full_Guide.docx"
            margins = (2.54, 2.54, 2.54, 2.54)
            try:
                convert_latex_to_word_docx(full_content, out_file, margins)
                with open(out_file, "rb") as f:
                    st.session_state["chemed_full_guide_bytes"] = f.read()
            except Exception as e:
                st.error(f"통합본 생성 중 오류: {e}")

        if st.session_state.get("chemed_full_guide_bytes"):
            st.download_button(
                label="💾 통합 가이드.docx 다운로드",
                data=st.session_state["chemed_full_guide_bytes"],
                file_name="ChemEd_Full_Guide.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="dl_full_persistent",
                use_container_width=True,
                type="primary"
            )

# ==========================================
# 1. Notion 스타일 관리 및 여백 조절 + 네이티브 워드 수식
# ==========================================
if menu == "📓 Notion / MS Word 스타일 매니저 (추천)":

    # 상단 컨트롤 패널 (모든 요소를 가로 1줄로 쫙 펼치기)
    c1, c2, c3, c4, c5 = st.columns([1.5, 1.5, 1.0, 3.5, 1.5])

    with c1:
        if "current_subject_tab" not in st.session_state:
            st.session_state.current_subject_tab = list(st.session_state.notion_db.keys())[0]
        
        subjects_list = list(st.session_state.notion_db.keys())
        idx = subjects_list.index(st.session_state.current_subject_tab) if st.session_state.current_subject_tab in subjects_list else 0
        
        subject = st.selectbox("과목", subjects_list, index=idx, label_visibility="collapsed")
        st.session_state.current_subject_tab = subject

        if st.button("📝 샘플 불러오기", help="이 과목의 표준 과제 샘플을 에디터로 불러옵니다.", use_container_width=True):
            if subject in SAMPLES:
                st.session_state.notion_db[subject] = SAMPLES[subject]
                if f"editor_{subject}" in st.session_state: del st.session_state[f"editor_{subject}"]
                st.success(f"✅ {subject} 샘플 로드 완료!")
                st.rerun()
    with c2:
        new_subject = st.text_input("추가", placeholder="새 과목 이름", label_visibility="collapsed")
    with c3:
        if st.button("추가", use_container_width=True):
            if new_subject:
                st.session_state.notion_db[new_subject] = ""
                st.rerun()

    with c4:
        upload_img = st.file_uploader("사진", type=["png", "jpg", "jpeg"], label_visibility="collapsed")
        # 강력한 JavaScript 주입을 통해 업로드 창의 텍스트를 강제로 숨기고 높이를 압축합니다.
        import streamlit.components.v1 as components
        components.html("""
        <script>
        // 부모 창(Streamlit 메인 창)의 DOM에 접근
        const doc = window.parent.document;

        function flattenUploader() {
            const dropzones = doc.querySelectorAll('[data-testid="stFileUploadDropzone"]');
            dropzones.forEach(dz => {
                // 박스 자체의 여백과 높이 최소화
                dz.style.padding = '2px 5px';
                dz.style.minHeight = '35px';
                dz.style.height = '35px';
                dz.style.display = 'flex';
                dz.style.alignItems = 'center';
                dz.style.justifyContent = 'center';

                // 내부에 있는 'Drag and drop file here' 등 모든 글자와 아이콘 찾기
                const innerElements = dz.querySelectorAll('div, span, small, p, svg');
                innerElements.forEach(el => {
                    // 버튼(Browse files)이 아닌 텍스트 요소들은 모조리 숨김 처리
                    if(el.tagName !== 'BUTTON' && !el.closest('button')) {
                        el.style.display = 'none';
                    }
                });
            });
        }

        // 요소가 렌더링될 때까지 반복해서 시도
        flattenUploader();
        setTimeout(flattenUploader, 500);
        setTimeout(flattenUploader, 1000);
        </script>
        """, height=0)

    with c5:
        if upload_img and st.button("🚀 AI 자동추출"):
            if not api_key: st.error("키를 입력하세요!")
            else:
                with st.spinner("분석 중..."):
                    genai.configure(api_key=api_key)
                    model = get_safe_gemini_model()
                    try:
                        from PIL import Image
                        add_analysis_task("이 이미지 안의 텍스트와 수식을 LaTeX 형식($$ ... $$)을 사용해서 정확히 추출해줘.", Image.open(upload_img), subject, upload_img.name)
                        st.success(f"🚀 {upload_img.name} 분석 작업이 백그라운드에서 시작되었습니다. 다른 메뉴를 보셔도 분석은 계속됩니다!")
                    except Exception as e:
                        st.error(f"오류: {e}")

    st.markdown("---")

    # ==============================
    # 🚀 MS Word / Notion 바로 복사 기능
    # ==============================
    st.subheader("🚀 원클릭 바로 쓰기 (다운로드 없이 복사)")
    st.write("번거로운 파일 변환/다운로드 과정 없이, 작성하신 노트를 클릭 한 번으로 복사하여 원하는 곳에 바로 붙여넣으세요!")
    st.caption(r"팁: 수식을 작성할 때는 기호 앞뒤를 $$ 로 감싸주세요! (예: $$ \Delta G $$)")

    # 최신 텍스트 상태 가져오기
    curent_text = st.session_state.get(f"editor_{subject}", st.session_state.notion_db[subject])

    import json
    md_escaped = json.dumps(curent_text)

    # 최신 텍스트 상태 가져오기
    curent_text = st.session_state.get(f"editor_{subject}", st.session_state.notion_db[subject])

    # --- Word 다운로드 및 복사 섹션 ---
    st.subheader("📥 결과물 내보내기 (Download & Copy)")
    
    col_dl1, col_dl2 = st.columns(2)
    with col_dl1:
        if st.button("📄 MS Word 파일로 저장하기 (추천)", use_container_width=True, type="primary"):
            with st.spinner("고해상도 수식을 포함하여 Word 문서를 생성 중입니다..."):
                import time
                out_file = os.path.join(os.getcwd(), f"Report_{subject}_{int(time.time())}.docx")
                full_markdown = f"# {subject}\n\n" + curent_text
                margins = {'top': 2.0, 'bottom': 2.0, 'left': 2.5, 'right': 2.5}
                if convert_latex_to_word_docx(full_markdown, out_file, margins):
                    with open(out_file, "rb") as f:
                        st.session_state[f"last_word_{subject}"] = f.read()
                    st.success("🎉 Word 생성이 완료되었습니다! 아래 버튼을 눌러 저장하세요.")
                else:
                    st.error("워드 파일 생성 중 오류가 발생했습니다.")
    
    with col_dl2:
        if st.button("🚀 워드로 즉시 열기 (Desktop 앱 전용)", use_container_width=True):
            import time
            out_file = os.path.join(os.getcwd(), f"Direct_{subject}_{int(time.time())}.docx")
            full_markdown = f"# {subject}\n\n" + curent_text
            margins = {'top': 2.0, 'bottom': 2.0, 'left': 2.5, 'right': 2.5}
            if convert_latex_to_word_docx(full_markdown, out_file, margins):
                open_file_in_os(out_file)
                st.info("ℹ️ 로컬 환경에서는 Word가 자동 실행되지만, 웹(Streamlit Cloud) 버전에서는 작동하지 않을 수 있습니다. 왼쪽 '파일로 저장' 버튼을 이용해 주세요!")

    if st.session_state.get(f"last_word_{subject}"):
        st.download_button(
            label=f"💾 생성된 {subject} 워드 파일 컴퓨터에 저장하기",
            data=st.session_state[f"last_word_{subject}"],
            file_name=f"{subject}_Report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
            key=f"dl_btn_{subject}"
        )

    copy_html = f"""
    <style>
        .btn {{
            border: none;
            border-radius: 8px;
            padding: 12px 20px;
            font-size: 15px;
            font-weight: 700;
            font-family: 'Inter', sans-serif;
            cursor: pointer;
            width: 100%;
            transition: all 0.3s;
            color: white;
            box-shadow: 0 4px 6px rgba(0,0,0,0.1);
        }}
        .btn-notion {{ background-color: #0F172A; }}
        .btn-notion:hover {{ background-color: #334155; transform: translateY(-2px); }}
        .copy-container {{ display: flex; flex-direction: row; align-items: center; justify-content: flex-start; padding: 5px 0px; }}
    </style>

    <div class="copy-container">
        <button class="btn btn-notion" id="btn-notion" onclick='copyNotion()'>
            📓 Notion / 옵시디언 등 마크다운 노트에 바로 복사
        </button>
        <button class="btn" id="openKbBtn" onclick="window.mathVirtualKeyboard ? window.mathVirtualKeyboard.show() : alert('키보드를 로드 중입니다.')"
            style="margin-left: 10px; background-color:#3B82F6;"
            >🔢 숫자 입력 키보드 열기</button>
    </div>

    <script>
    function copyNotion() {{
        const text = {md_escaped};
        const el = document.createElement('textarea');
        el.value = text;
        document.body.appendChild(el);
        el.select();
        document.execCommand('copy');
        document.body.removeChild(el);

        const btn = document.getElementById('btn-notion');
        const originalText = btn.innerText;
        btn.innerText = '✅ 복사 완료! (Notion에서 Ctrl+V 하세요)';
        setTimeout(() => btn.innerText = originalText, 2500);
    }}
    </script>
    """
    import streamlit.components.v1 as components
    components.html(copy_html, height=60)

    def render_copyable_math(latex_display, latex_code):
        import json
        import streamlit.components.v1 as components
        safe_display = latex_display.replace("\\", "\\\\").replace("\"", "\\\"")
        safe_code = latex_code.replace("\\", "\\\\").replace("\"", "\\\"")
        
        components.html(f"""
            <link rel="stylesheet" href="https://cdn.jsdelivr.net/npm/katex@0.16.8/dist/katex.min.css">
            <script src="https://cdn.jsdelivr.net/npm/katex@0.16.8/dist/katex.min.js"></script>
            <style>
                .math-box {{ background: #ffffff; border: 1px solid #e2e8f0; border-radius: 8px; padding: 20px; margin: 10px 0; cursor: pointer; transition: all 0.2s; text-align: center; position: relative; display: flex; flex-direction: column; align-items: center; justify-content: center; }}
                .math-box:hover {{ border-color: #3b82f6; background: #f8fafc; box-shadow: 0 4px 12px rgba(59, 130, 246, 0.1); }}
                .copy-icon {{ position: absolute; top: 10px; right: 10px; color: #94a3b8; opacity: 0.6; }}
                .copy-hint {{ font-size: 11px; color: #64748b; margin-top: 10px; font-weight: 500; }}
                #toast {{ visibility: hidden; min-width: 120px; background-color: #1e293b; color: #fff; text-align: center; border-radius: 6px; padding: 10px; position: fixed; z-index: 100; bottom: 20px; left: 50%; transform: translateX(-50%); font-size: 13px; }}
                #toast.show {{ visibility: visible; animation: fadein 0.5s, fadeout 0.5s 1.5s; }}
                @keyframes fadein {{ from {{bottom: 0; opacity: 0;}} to {{bottom: 20px; opacity: 1;}} }}
                @keyframes fadeout {{ from {{bottom: 20px; opacity: 1;}} to {{bottom: 0; opacity: 0;}} }}
            </style>
            <div class="math-box" onclick="copyCode()">
                <div class="copy-icon"><svg width="20" height="20" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><rect x="9" y="9" width="13" height="13" rx="2" ry="2"></rect><path d="M5 15H4a2 2 0 0 1-2-2V4a2 2 0 0 1 2-2h9a2 2 0 0 1 2 2v1"></path></svg></div>
                <div id="math-content"></div>
                <div class="copy-hint">📋 클릭하여 LaTeX 코드 복사</div>
            </div>
            <div id="toast">✅ 복사 완료!</div>
            <script>
                const display = "{safe_display}";
                const code = "{safe_code}";
                const el = document.getElementById('math-content');
                katex.render(display, el, {{ throwOnError: false, displayMode: true }});
                function copyCode() {{
                    const temp = document.createElement('textarea');
                    temp.value = code;
                    document.body.appendChild(temp);
                    temp.select();
                    document.execCommand('copy');
                    document.body.removeChild(temp);
                    const toast = document.getElementById('toast');
                    toast.className = "show";
                    setTimeout(() => {{ toast.className = ""; }}, 2000);
                }}
            </script>
        """, height=130)

    # --- 추가된 전공 샘플 워드 다운로드 센터 ---
    with st.expander("📥 전공 과제 샘플(.docx) 직접 다운로드 센터", expanded=False):
        st.caption("팝업 확인 없이 바로 워드 파일로 소장하고 싶으실 때 사용하세요.")
        sd_col1, sd_col2 = st.columns(2)
        sample_list = [("물리화학", "⚛️"), ("유기화학", "🧪"), ("분석화학", "📊"), ("무기화학", "💎"), ("화학교육", "🧑‍🏫")]
        for i, (name, icon) in enumerate(sample_list):
            target_col = sd_col1 if i % 2 == 0 else sd_col2
            with target_col:
                if st.button(f"{icon} {name} 샘플 생성 (Save to PC)", key=f"sd_btn_{name}", use_container_width=True):
                    with st.spinner(f"{name} 샘플 변환 중..."):
                        import time
                        out_file = os.path.join(os.getcwd(), f"{name}_Sample_{int(time.time())}.docx")
                        content = SAMPLES.get(name, "데이터 없음")
                        margins = {"top": 2.0, "bottom": 2.0, "left": 2.5, "right": 2.5}
                        if convert_latex_to_word_docx(f"# {name} 샘플\n\n" + content, out_file, margins):
                            # 세션에 개별 샘플 데이터 저장
                            with open(out_file, "rb") as f:
                                st.session_state[f"sample_file_{name}"] = f.read()
                            st.success(f"✅ {name} 샘플이 생성되었습니다!")

                if st.button(f"📝 {name} 샘플 에디터로 불러오기 (미리보기)", key=f"preview_btn_{name}", use_container_width=True):
                    if name in SAMPLES:
                        st.session_state.notion_db[name] = SAMPLES[name]
                        st.session_state.current_subject_tab = name
                        if f"editor_{name}" in st.session_state: del st.session_state[f"editor_{name}"]
                        st.success(f"✅ {name} 샘플이 에디터에 로드되었습니다! 위쪽 에디터와 미리보기에서 확인하세요.")
                        st.rerun()
                
                # 생성된 샘플이 있으면 다운로드 버튼 표시
                if st.session_state.get(f"sample_file_{name}"):
                    st.download_button(
                        f"💾 {name} 파일 직접 다운로드", 
                        st.session_state[f"sample_file_{name}"], 
                        f"{name}_Sample.docx", 
                        use_container_width=True,
                        key=f"dl_sample_{name}"
                    )

    # 텍스트 에디터 및 미리보기 (전체 너비 사용)
    st.subheader(f"📝 {subject} 노트 에디터")
    st.info("이곳의 에디터는 화면 전체를 넓게 사용합니다. 내용이 길어도 편하게 작성하세요!")

    with st.expander("📌 자주 쓰는 수학/화학 기호 사전 (클릭해서 열기)"):
        st.info("각 코드 박스 우측의 아이콘을 클릭하여 복사(Copy)한 뒤 에디터에 붙여넣으세요! (LaTeX 형식)")

        # 가로 중첩 방지를 위해 세로로 순차 배치
        st.markdown("**1. 그리스 문자**")
        render_copyable_math(r"\alpha, \beta, \gamma, \delta, \Delta, \pi, \sigma, \theta, \psi, \Psi, \phi, \omega", r"\alpha, \beta, \gamma, \delta, \Delta, \pi, \sigma, \theta, \psi, \Psi, \phi, \omega")
        st.code(r"\alpha, \beta, \gamma, \delta, \Delta, \pi, \sigma, \theta, \psi, \Psi, \phi, \omega", language="latex")
        
        st.markdown("---")
        st.markdown("**2. 기본 수식 기호**")
        render_copyable_math(r"x^2, y_{i}, \sqrt{x}, \sqrt[n]{x}, \pm, \mp, \times, \div, \approx, \neq, \propto, \infty", r"x^2, y_{i}, \sqrt{x}, \sqrt[n]{x}, \pm, \mp, \times, \div, \approx, \neq, \propto, \infty")
        st.code(r"x^2, y_{i}, \sqrt{x}, \sqrt[n]{x}, \pm, \mp, \times, \div, \approx, \neq, \propto, \infty", language="latex")
        
        st.markdown("---")
        st.markdown("**3. 화학 반응/화살표**")
        render_copyable_math(r"A \rightarrow B, C \rightleftharpoons D, E \xrightarrow{H^+} F, G \uparrow, H \downarrow", r"\rightarrow, \rightleftharpoons, \xrightarrow{H^+}, \xleftarrow[temp]{cat}, \uparrow, \downarrow, \leftrightarrow")
        st.code(r"\rightarrow, \rightleftharpoons, \xrightarrow{H^+}, \xleftarrow[temp]{cat}, \uparrow, \downarrow, \leftrightarrow", language="latex")

        st.markdown("---")
        st.markdown("**4. 열역학/속도론**")
        render_copyable_math(r"\Delta G = \Delta H - T\Delta S, \quad k = A e^{-\frac{E_a}{RT}}, \quad PV = nRT, \quad \ln K = -\frac{\Delta G^\circ}{RT}", r"\Delta G = \Delta H - T\Delta S, k = A e^{-\frac{E_a}{RT}}, PV = nRT, \ln K = -\frac{\Delta G^\circ}{RT}")
        st.code(r"\Delta G = \Delta H - T\Delta S, k = A e^{-\frac{E_a}{RT}}, PV = nRT, \ln K = -\frac{\Delta G^\circ}{RT}", language="latex")
        
        st.markdown("---")
        st.markdown("**5. 양자역학/오비탈**")
        render_copyable_math(r"\hat{H}\psi = E\psi, \quad \lambda = \frac{h}{p}, \quad \psi_{n,l,m}, \quad \nabla^2, \quad \hbar, \quad \int |\psi|^2 d\tau = 1", r"\hat{H}\psi = E\psi, \lambda = \frac{h}{p}, \psi_{n,l,m}, \nabla^2, \hbar, \int |\psi|^2 d\tau = 1")
        st.code(r"\hat{H}\psi = E\psi, \lambda = \frac{h}{p}, \psi_{n,l,m}, \nabla^2, \hbar, \int |\psi|^2 d\tau = 1", language="latex")
        
        st.markdown("---")
        st.markdown("**6. 고급 수학 (미적분/행렬)**")
        render_copyable_math(r"\frac{dy}{dx}, \quad \frac{\partial f}{\partial x}, \quad \int_a^b f(x) dx, \quad \sum_{i=1}^n x_i, \quad \begin{pmatrix} a & b \\ c & d \end{pmatrix}", r"\frac{dy}{dx}, \frac{\partial f}{\partial x}, \int_a^b, \sum_{i=1}^n, \prod, \lim_{x \to 0}, \begin{pmatrix} a & b \\ c & d \end{pmatrix}")
        st.code(r"\frac{dy}{dx}, \frac{\partial f}{\partial x}, \int_a^b, \sum_{i=1}^n, \prod, \lim_{x \to 0}, \begin{pmatrix} a & b \\ c & d \end{pmatrix}", language="latex")

        st.caption(r"팁: 수식을 작성할 때는 기호 앞뒤를 $$ 로 감싸주세요! (예: $$ \Delta G $$)")


    render_chem_ed_core_guide()

    with st.expander("📊 스마트 데이터 표(Table) & 화학 그래프 빌더"):
        st.write("데이터를 입력하면 에디터에 맞는 Markdown 표나 그래프를 생성해 줍니다.")
        t1, t2 = st.tabs(["📋 표 빌더", "📈 그래프 생성"])

        with t1:
            st.markdown("#### 엑셀형 데이터 표 에디터")
            st.caption("마우스로 셀을 클릭하여 직접 수정하거나, 표 아래를 클릭하여 행을 추가하세요. **(행 삭제: 왼쪽 끝 체크박스 선택 후 우측 상단의 휴지통 🗑️ 아이콘 클릭)**")

            # 프리미엄 논문용 템플릿 선택
            t_col1, t_col2 = st.columns([2, 1])
            with t_col1:
                table_template = st.selectbox("📝 논문용 프리미엄 표 양식 선택:", [
                    "빈 표 (기본)",
                    "기술 통계 요약표 (Descriptive)",
                    "상관관계 분석표 (Corelation)",
                    "분산분석표 (ANOVA)",
                    "이화학 실험 결과 요약표",
                    "반응 속도 데이터",
                    "산-염기 적정 데이터"
                ])
            with t_col2:
                t_color = st.color_picker("🎨 표 헤더 색상", "#334155")

            import pandas as pd
            # Ensure the table builder dataframe exists and matches the current template
            if "table_builder_df" not in st.session_state or st.session_state.get("table_template_prev") != table_template:
                if table_template == "기술 통계 요약표 (Descriptive)":
                    st.session_state.table_builder_df = pd.DataFrame([{"변수 (Variables)": "연령 (Age)", "표본 수 (N)": 120, "평균 (Mean)": 25.4, "표준편차 (SD)": 3.2}, {"변수 (Variables)": "시험 점수 (Score)", "표본 수 (N)": 120, "평균 (Mean)": 82.1, "표준편차 (SD)": 7.5}, {"변수 (Variables)": "학습 시간 (Hours)", "표본 수 (N)": 120, "평균 (Mean)": 4.5, "표준편차 (SD)": 1.1}])
                elif table_template == "상관관계 분석표 (Corelation)":
                    st.session_state.table_builder_df = pd.DataFrame([{"변수명": "1. 학습 동기", "1": "1.00", "2": "0.45**", "3": "0.38*"}, {"변수명": "2. 자기효능감", "1": "-", "2": "1.00", "3": "0.62**"}, {"변수명": "3. 학업 성취도", "1": "-", "2": "-", "3": "1.00"}])
                elif table_template == "분산분석표 (ANOVA)":
                    st.session_state.table_builder_df = pd.DataFrame([{"요인 (Source)": "집단 간 (Between)", "제곱합 (SS)": 450.2, "자유도 (df)": 2, "평균제곱 (MS)": 225.1, "F 값": 5.42, "p-value": "0.008**"}, {"요인 (Source)": "집단 내 (Within)", "제곱합 (SS)": 1205.5, "자유도 (df)": 29, "평균제곱 (MS)": 41.5, "F 값": "-", "p-value": "-"}, {"요인 (Source)": "전체 (Total)", "제곱합 (SS)": 1655.7, "자유도 (df)": 31, "평균제곱 (MS)": "-", "F 값": "-", "p-value": "-"}])
                elif table_template == "이화학 실험 결과 요약표":
                    st.session_state.table_builder_df = pd.DataFrame([{"실험군 (Group)": "대조군 (Control)", "이론 수득량 (g)": 5.0, "실제 수득량 (g)": 4.2, "수득률 (%)": 84.0, "오차율 (%)": 16.0}, {"실험군 (Group)": "촉매 A 사용", "이론 수득량 (g)": 5.0, "실제 수득량 (g)": 4.8, "수득률 (%)": 96.0, "오차율 (%)": 4.0}])
                elif table_template == "반응 속도 데이터":
                    st.session_state.table_builder_df = pd.DataFrame([{"시간 (s)": 0, "농도 (M)": 1.00}, {"시간 (s)": 10, "농도 (M)": 0.85}, {"시간 (s)": 20, "농도 (M)": 0.72}])
                elif table_template == "산-염기 적정 데이터":
                    st.session_state.table_builder_df = pd.DataFrame([{"적정액 부피 (mL)": 0.0, "pH": 2.5}, {"적정액 부피 (mL)": 5.0, "pH": 3.8}, {"적정액 부피 (mL)": 10.0, "pH": 7.0}])
                else:
                    st.session_state.table_builder_df = pd.DataFrame([{"컬럼 1": "", "컬럼 2": "", "컬럼 3": ""} for _ in range(3)])
                st.session_state.table_template_prev = table_template

            try:
                # Use data_editor directly. If it fails, fallback to something else, though data_editor is standard in newer Streamlit versions.
                edited_df = st.data_editor(
                    st.session_state.table_builder_df,
                    num_rows="dynamic",
                    use_container_width=True,
                    key="table_editor_ui"
                )
            except Exception as e:
                st.error(f"데이터 에디터 로딩 중 오류가 발생했습니다: {e}")
                edited_df = st.session_state.table_builder_df

            t_btn1, t_btn2, t_btn3 = st.columns([1.5, 1.5, 1.5])
            with t_btn1:
                insert_table = st.button("✨ Markdown 표로 에디터에 삽입", use_container_width=True)
            with t_btn2:
                generate_table_word = st.button("💾 단독 워드 파일 생성", use_container_width=True)
            
            t_dl_placeholder = t_btn3.empty()

            if insert_table:
                if not edited_df.empty:
                    md_table = f"\n<table style='width:100%; border-collapse: collapse; text-align: center;' border='1'>\n"
                    md_table += f"  <thead style='background-color: {t_color}; color: white;'>\n    <tr>\n"
                    for c in edited_df.columns:
                        md_table += f"      <th style='padding: 8px;'>{c}</th>\n"
                    md_table += "    </tr>\n  </thead>\n  <tbody>\n"
                    for _, row in edited_df.iterrows():
                        md_table += "    <tr>\n"
                        for val in row.values:
                            md_table += f"      <td style='padding: 8px;'>{val}</td>\n"
                        md_table += "    </tr>\n"
                    md_table += "  </tbody>\n</table>\n<br>\n"

                    st.session_state.notion_db[subject] += "\n" + md_table
                    if f"editor_{subject}" in st.session_state: del st.session_state[f"editor_{subject}"]
                    st.success("고급스러운 표가 성공적으로 에디터에 추가되었습니다!")
                    st.rerun()

            if generate_table_word:
                if not edited_df.empty:
                    with st.spinner("워드 문서 생성 중..."):
                        md_table = f"\n<table style='width:100%; border-collapse: collapse; text-align: center;' border='1'>\n"
                        md_table += f"  <thead style='background-color: {t_color}; color: white;'>\n    <tr>\n"
                        for c in edited_df.columns:
                            md_table += f"      <th style='padding: 8px;'>{c}</th>\n"
                        md_table += "    </tr>\n  </thead>\n  <tbody>\n"
                        for _, row in edited_df.iterrows():
                            md_table += "    <tr>\n"
                            for val in row.values:
                                md_table += f"      <td style='padding: 8px;'>{val}</td>\n"
                            md_table += "    </tr>\n"
                        md_table += "  </tbody>\n</table>\n<br>\n"
                        import time
                        out_file = f"Table_{int(time.time())}.docx"
                        if convert_latex_to_word_docx(md_table, out_file, {'top': 2.0, 'bottom': 2.0, 'left': 2.5, 'right': 2.5}):
                            with open(out_file, "rb") as f:
                                st.session_state["table_dl_bytes"] = f.read()
                            st.success("준비 완료! 우측 다운로드 버튼을 눌러주세요.")

            if st.session_state.get("table_dl_bytes"):
                t_dl_placeholder.download_button("📥 생성된 표 다운로드 (.docx)", st.session_state["table_dl_bytes"], file_name="Data_Table.docx", use_container_width=True, type="primary", key="table_dl_btn")

        with t2:
            st.markdown("#### 프리미엄 화학/물리 데이터 시각화")
            g_col1, g_col2 = st.columns([1, 2])

            with g_col1:
                g_type = st.radio("📈 시각화 유형", [
                    "선 그래프 (Line)",
                    "이중선 그래프 (Dual Line)",
                    "영역 그래프 (Area)",
                    "막대 그래프 (Bar)",
                    "산점도 (Scatter)",
                    "박스 플롯 (Box Plot)",
                    "방사형 차트 (Radar)",
                    "단일 적정 곡선 (Single Spline)",
                    "이중 적정 곡선 (Dual Spline)",
                    "파이 차트 (Pie)",
                    "3D 산점도 (3D Scatter)",
                    "3D 표면 (3D Surface)"
                ])
                g_color = st.color_picker("🎨 그래프 주요 색상", "#3B82F6")
                g_title = st.text_input("📝 그래프 제목", f"{subject} 데이터 분석")
                g_xlabel = st.text_input("📝 표의 X열 이름 / X축 라벨 변경", "")
                g_ylabel = st.text_input("📝 표의 Y열 이름 / Y축 라벨 변경", "")
                if "3D" in g_type:
                    g_zlabel = st.text_input("📝 표의 Z열 이름 / Z축 라벨 변경", "")
                elif "이중선" in g_type or "이중 적정" in g_type:
                    g_zlabel = st.text_input("📝 표의 Y2열 이름 / 대조군 라벨 변경", "")

            with g_col2:
                st.caption("그래프에 표시할 데이터를 입력하세요. 3D 차트는 X, Y, Z 세 열이 필요합니다. **(행 삭제: 우측 상단 🗑️ 클릭)**")
                import pandas as pd
                if "graph_builder_df" not in st.session_state or st.session_state.get("g_type_prev") != g_type:
                    if "3D" in g_type:
                        st.session_state.graph_builder_df = pd.DataFrame([{"X": 1, "Y": 2, "Z": 3}, {"X": 2, "Y": 5, "Z": 5}, {"X": 3, "Y": 1, "Z": 2}, {"X": 4, "Y": 6, "Z": 8}, {"X": 5, "Y": 3, "Z": 4}])
                    elif "이중선" in g_type:
                        st.session_state.graph_builder_df = pd.DataFrame([{"X (시간/조건)": 0, "Y1 (실험군)": 2.5, "Y2 (대조군)": 2.0}, {"X (시간/조건)": 10, "Y1 (실험군)": 5.8, "Y2 (대조군)": 3.5}, {"X (시간/조건)": 20, "Y1 (실험군)": 12.0, "Y2 (대조군)": 4.1}, {"X (시간/조건)": 30, "Y1 (실험군)": 15.5, "Y2 (대조군)": 5.0}])
                    elif "이중 적정 곡선" in g_type:
                        st.session_state.graph_builder_df = pd.DataFrame([{"X (부피)": 0, "Y1 (pH-A)": 2.5, "Y2 (pH-B)": 3.0}, {"X (부피)": 10, "Y1 (pH-A)": 3.8, "Y2 (pH-B)": 4.5}, {"X (부피)": 20, "Y1 (pH-A)": 7.0, "Y2 (pH-B)": 7.2}, {"X (부피)": 25, "Y1 (pH-A)": 11.2, "Y2 (pH-B)": 10.5}])
                    elif "단일 적정 곡선" in g_type:
                        st.session_state.graph_builder_df = pd.DataFrame([{"X (적정 부피)": 0, "Y (pH)": 2.5}, {"X (적정 부피)": 10, "Y (pH)": 3.8}, {"X (적정 부피)": 20, "Y (pH)": 7.0}, {"X (적정 부피)": 25, "Y (pH)": 11.2}])
                    elif "선 그래프" in g_type or "영역 그래프" in g_type or "산점도" in g_type:
                        st.session_state.graph_builder_df = pd.DataFrame([{"X (시간/농도 등)": 10, "Y 값": 2.5}, {"X (시간/농도 등)": 20, "Y 값": 5.8}, {"X (시간/농도 등)": 30, "Y 값": 12.0}, {"X (시간/농도 등)": 40, "Y 값": 15.5}])
                    elif "박스 플롯" in g_type:
                        st.session_state.graph_builder_df = pd.DataFrame([{"집단 (X)": "대조군", "측정값 (Y)": 45.2}, {"집단 (X)": "대조군", "측정값 (Y)": 48.1}, {"집단 (X)": "대조군", "측정값 (Y)": 44.5}, {"집단 (X)": "실험군", "측정값 (Y)": 85.3}, {"집단 (X)": "실험군", "측정값 (Y)": 92.1}, {"집단 (X)": "실험군", "측정값 (Y)": 88.0}])
                    else: # 막대, 방사형, 파이
                        st.session_state.graph_builder_df = pd.DataFrame([{"항목 (X 범주)": "A", "수치 (Y)": 10.0}, {"항목 (X 범주)": "B", "수치 (Y)": 25.5}, {"항목 (X 범주)": "C", "수치 (Y)": 15.2}, {"항목 (X 범주)": "D", "수치 (Y)": 30.1}])
                    st.session_state.g_type_prev = g_type

                try:
                    display_df = st.session_state.graph_builder_df.copy()
                    if g_xlabel:
                        display_df.rename(columns={display_df.columns[0]: g_xlabel}, inplace=True)
                    if g_ylabel and len(display_df.columns) > 1:
                        display_df.rename(columns={display_df.columns[1]: g_ylabel}, inplace=True)
                    if 'g_zlabel' in locals() and g_zlabel and len(display_df.columns) > 2:
                        display_df.rename(columns={display_df.columns[2]: g_zlabel}, inplace=True)

                    graph_df = st.data_editor(
                        display_df,
                        num_rows="dynamic",
                        use_container_width=True,
                        key="graph_editor_ui"
                    )
                    
                    # Update session state with edited values but keep original internal column names
                    updated_df = graph_df.copy()
                    updated_df.columns = st.session_state.graph_builder_df.columns
                    st.session_state.graph_builder_df = updated_df

                except Exception as e:
                    st.error(f"데이터 에디터 로딩 중 오류가 발생했습니다: {e}")
                    graph_df = st.session_state.graph_builder_df

            g_btn1, g_btn2 = st.columns(2)
            with g_btn1:
                generate_graph = st.button("🖼️ 고해상도 그래프 생성 (전체 워드 추가)", use_container_width=True)
            
            g_dl_placeholder = g_btn2.empty()

            if True: # 항상 실시간 미리보기를 위해 렌더링
                try:
                    import pandas as pd
                    import matplotlib.pyplot as plt
                    import numpy as np

                    clean_df = graph_df.dropna()
                    
                    is_3d = "3D" in g_type
                    is_dual = "이중선" in g_type or "이중 적정" in g_type
                    if is_3d or is_dual:
                        if clean_df.shape[1] < 3:
                            st.error(f"{g_type}를 그리려면 최소 3개의 열이 필요합니다.")
                            raise ValueError("Insufficient columns")
                        
                        try:
                            x_vals = pd.to_numeric(clean_df.iloc[:, 0]).tolist()
                        except Exception:
                            x_vals = clean_df.iloc[:, 0].astype(str).tolist()

                        y_vals = pd.to_numeric(clean_df.iloc[:, 1], errors='coerce').tolist()
                        z_vals = pd.to_numeric(clean_df.iloc[:, 2], errors='coerce').tolist()
                        
                        valid_triplets = [(x, y, z) for x, y, z in zip(x_vals, y_vals, z_vals) if not (pd.isna(x) or pd.isna(y) or pd.isna(z))]
                        x = [p[0] for p in valid_triplets]
                        y = [p[1] for p in valid_triplets]
                        z = [p[2] for p in valid_triplets]
                    else:
                        # X값 자동 형변환
                        try:
                            x_vals = pd.to_numeric(clean_df.iloc[:, 0]).tolist()
                        except Exception:
                            x_vals = clean_df.iloc[:, 0].astype(str).tolist()
                        y_vals = pd.to_numeric(clean_df.iloc[:, 1], errors='coerce').tolist()
                        
                        valid_pairs = [(x, y) for x, y in zip(x_vals, y_vals) if not (pd.isna(x) or pd.isna(y))]
                        x = [p[0] for p in valid_pairs]
                        y = [p[1] for p in valid_pairs]

                    if not is_3d and ("적정 곡선" in g_type and len(x) > 0 and not all(isinstance(v, (int, float, np.integer, np.floating)) for v in x)):
                        st.error("적정 곡선(Spline)은 X축 값이 반드시 '숫자'여야 합니다.")
                    elif not is_3d and (not x or not y or len(x) != len(y)):
                        st.error("데이터 형식이 올바르지 않거나 Y값에 문자가 섞여있습니다.")
                    elif is_3d and (not x or not y or not z):
                        st.error("3D 데이터 형식이 올바르지 않거나 숫자가 아닌 값이 섞여있습니다.")
                    else:
                        plt.style.use('seaborn-v0_8-whitegrid')
                        
                        import platform
                        if platform.system() == 'Darwin':
                            plt.rc('font', family='AppleGothic')
                        elif platform.system() == 'Windows':
                            plt.rc('font', family='Malgun Gothic')
                        else:
                            plt.rc('font', family='NanumGothic')
                        plt.rcParams['axes.unicode_minus'] = False

                        if "방사형 차트" in g_type:
                            fig, ax = plt.subplots(figsize=(5, 5), dpi=300, subplot_kw=dict(polar=True))
                        elif is_3d:
                            fig = plt.figure(figsize=(7, 5), dpi=300)
                            ax = fig.add_subplot(111, projection='3d')
                        else:
                            fig, ax = plt.subplots(figsize=(6, 4), dpi=300)

                        fig.patch.set_facecolor('#FFFFFF')
                        ax.set_facecolor('#FFFFFF')

                        if is_3d:
                            if "3D 산점도" in g_type:
                                ax.scatter(x, y, z, color=g_color, s=80, alpha=0.8, edgecolor='white', linewidth=1)
                            elif "3D 표면" in g_type:
                                try:
                                    ax.plot_trisurf(x, y, z, color=g_color, alpha=0.8, linewidth=0.2, edgecolor='grey')
                                except RuntimeError as e:
                                    if "qhull" in str(e).lower() or "singular" in str(e).lower():
                                        st.error("⚠️ 3D 표면 그래프 오류: X, Y 데이터가 모두 일직선상에 있어 면(표면)을 만들 수 없습니다. X와 Y값을 서로 다르게 흩어지도록 수정해주세요.")
                                        st.stop()
                                    else:
                                        raise e
                            ax.set_xlabel(g_xlabel if g_xlabel else (clean_df.columns[0] if len(clean_df.columns) > 0 else 'X'), fontweight='bold', labelpad=10)
                            ax.set_ylabel(g_ylabel if g_ylabel else (clean_df.columns[1] if len(clean_df.columns) > 1 else 'Y'), fontweight='bold', labelpad=10)
                            ax.set_zlabel(g_zlabel if 'g_zlabel' in locals() and g_zlabel else (clean_df.columns[2] if len(clean_df.columns) > 2 else 'Z'), fontweight='bold', labelpad=10)
                            ax.view_init(elev=20, azim=30)
                        elif "이중선 그래프" in g_type:
                            ax.plot(x, y, marker='o', linestyle='-', color=g_color, linewidth=2.5, markersize=8, label=clean_df.columns[1] if len(clean_df.columns) > 1 else 'Y1')
                            ax.plot(x, z, marker='s', linestyle='--', color='#EF4444', linewidth=2.5, markersize=8, label=clean_df.columns[2] if len(clean_df.columns) > 2 else 'Y2')
                            ax.legend(frameon=True, shadow=True)
                        elif "선 그래프" in g_type:
                            ax.plot(x, y, marker='o', linestyle='-', color=g_color, linewidth=2.5, markersize=8, markerfacecolor='white', markeredgecolor=g_color, markeredgewidth=2)
                        elif "영역 그래프" in g_type:
                            ax.fill_between(x, y, color=g_color, alpha=0.3)
                            ax.plot(x, y, marker='o', color=g_color, linewidth=2)
                        elif "산점도" in g_type:
                            ax.scatter(x, y, color=g_color, s=80, alpha=0.8, edgecolor='white', linewidth=1)
                        elif "막대 그래프" in g_type:
                            ax.bar(x, y, color=g_color, alpha=0.85, edgecolor='black', linewidth=0.5, width=0.5)
                        elif "박스 플롯" in g_type:
                            unique_x = list(dict.fromkeys(x))
                            data_groups = [[y[i] for i in range(len(x)) if x[i] == ux] for ux in unique_x]
                            ax.boxplot(data_groups, labels=unique_x, patch_artist=True, boxprops=dict(facecolor=g_color, color='black', alpha=0.7), medianprops=dict(color='#EF4444', linewidth=2))
                        elif "적정 곡선" in g_type:
                            def plot_spline(x_arr, y_arr, color, label):
                                unique_xy = {}
                                for px, py in sorted(zip(x_arr, y_arr), key=lambda p: p[0]):
                                    if px in unique_xy:
                                        unique_xy[px].append(py)
                                    else:
                                        unique_xy[px] = [py]

                                x_sorted = list(unique_xy.keys())
                                y_sorted = [np.mean(unique_xy[px]) for px in x_sorted]

                                if len(x_sorted) < 2:
                                    st.error(f"{label} 곡선을 그리려면 고유한 X값이 최소 2개 이상 필요합니다.")
                                    return
                                
                                p = np.poly1d(np.polyfit(x_sorted, y_sorted, min(3, len(x_sorted)-1)))
                                x_new = np.linspace(min(x_sorted), max(x_sorted), 300)
                                y_new = p(x_new)
                                ax.plot(x_new, y_new, color=color, lw=3, label=label + ' (곡선)')
                                ax.scatter(x_sorted, y_sorted, color=color, s=40, zorder=5, label=label + ' (데이터)')

                            if is_dual:
                                plot_spline(x, y, g_color, clean_df.columns[1] if len(clean_df.columns) > 1 else 'Y1')
                                plot_spline(x, z, '#EF4444', clean_df.columns[2] if len(clean_df.columns) > 2 else 'Y2')
                                ax.legend(frameon=True, shadow=True)
                            else:
                                plot_spline(x, y, g_color, clean_df.columns[1] if len(clean_df.columns) > 1 else 'Data')
                                ax.legend(frameon=True, shadow=True)
                        elif "파이 차트" in g_type:
                            colors = [g_color, '#FCA5A5', '#FCD34D', '#6EE7B7', '#93C5FD', '#C4B5FD', '#F9A8D4', '#D1D5DB']
                            ax.pie(y, labels=x, autopct='%1.1f%%', startangle=90, colors=colors, wedgeprops={'edgecolor': 'white', 'linewidth': 1.5})
                            ax.axis('equal')
                        elif "방사형 차트" in g_type:
                            angles = np.linspace(0, 2 * np.pi, len(x), endpoint=False).tolist()
                            y_rad = y + [y[0]]
                            angles += [angles[0]]

                            ax.plot(angles, y_rad, color=g_color, linewidth=2)
                            ax.fill(angles, y_rad, color=g_color, alpha=0.25)
                            ax.set_xticks(angles[:-1])
                            ax.set_xticklabels(x, fontsize=11, fontweight='bold')
                            ax.tick_params(axis='x', pad=10)

                        if not is_3d and g_type not in ["파이 차트 (Pie)", "방사형 차트 (Radar)"]:
                            ax.set_xlabel(g_xlabel if g_xlabel else clean_df.columns[0], fontweight='bold', color='#334155', fontsize=11)
                            ax.set_ylabel(g_ylabel if g_ylabel else clean_df.columns[1], fontweight='bold', color='#334155', fontsize=11)
                            ax.spines['top'].set_visible(False)
                            ax.spines['right'].set_visible(False)
                            ax.spines['left'].set_linewidth(1.5)
                            ax.spines['bottom'].set_linewidth(1.5)
                            ax.tick_params(axis='both', which='major', labelsize=10)

                        ax.set_title(g_title, fontsize=15, fontweight='bold', pad=20, color='#1E293B')

                        st.pyplot(fig) # 실시간 라이브 미리보기

                        if generate_graph:
                            buf = io.BytesIO()
                            fig.savefig(buf, format='png', bbox_inches='tight', dpi=300, facecolor='white')
                            buf.seek(0)
                            st.session_state.word_doc.add_picture(buf, width=Inches(5.0))
                            st.success("🎉 세련된 그래프가 생성되어 전체 워드 문서에 성공적으로 추가되었습니다!")

                            # 단독 다운로드를 위한 독립 워드 문서 생성
                            from docx import Document
                            from docx.shared import Inches
                            temp_doc = Document()
                            temp_doc.add_heading(g_title if g_title else "그래프", level=1)
                            buf.seek(0)
                            temp_doc.add_picture(buf, width=Inches(6.0))
                            doc_io = io.BytesIO()
                            temp_doc.save(doc_io)
                            doc_io.seek(0)
                            st.session_state.graph_dl_bytes = doc_io.read()
                            st.info("💡 우측 상단의 '단독 워드 다운로드' 버튼이 활성화되었습니다!")
                except Exception as e:
                    st.error(f"데이터 형식을 확인하거나 다른 템플릿을 선택하세요: {e}")

            if st.session_state.get("graph_dl_bytes"):
                g_dl_placeholder.download_button("📥 이 그래프 단독 워드(.docx) 다운로드", st.session_state.graph_dl_bytes, file_name="Data_Graph.docx", use_container_width=True, type="primary", key="graph_dl_btn")

    # ==============================================================
    # 💡 양방향 자동화: 수동/자동 모드 스위치(Toggle)
    # ==============================================================
    col_t1, col_t2 = st.columns([1, 4])
    with col_t1:
        auto_mode = st.toggle("🤖 AI 자동 완성 모드 켜기")
    with col_t2:
        if auto_mode:
            st.info("현재 **자동(AI) 모드**입니다. 아래에 지시어를 입력하면 에디터 맨 밑에 내용이 자동으로 작성됩니다.")
        else:
            st.write("현재 **수동 모드**입니다. 자유롭게 키보드로 과제를 타이핑하세요.")

    if auto_mode:
        with st.container(border=True):
            ai_c1, ai_c2 = st.columns([4, 1.2])
            with ai_c1:
                ai_prompt = st.text_input("AI에게 작성시킬 내용을 입력하세요", placeholder="예: 살리실산과 아세트산 무수물의 에스터화 반응 메커니즘을 상세히 설명해줘", label_visibility="collapsed")
            with ai_c2:
                if st.button("✨ 에디터에 이어붙이기", use_container_width=True):
                    if not api_key:
                        st.error("API 키를 입력하세요!")
                    elif ai_prompt:
                        with st.spinner("AI가 전공 수준의 지식으로 내용을 작성하여 에디터에 붙여넣고 있습니다..."):
                            genai.configure(api_key=api_key)
                            model = get_safe_gemini_model()
                            try:
                                sys_prompt = "너는 서울대학교 화학교육과 조교야. 학생이 과제 레포트 작성 중 너에게 다음 부분의 대필을 요청했어. 전공 수준의 정확한 화학/물리 지식(LaTeX 수식 $$ $$ 적극 활용)과 논리적인 문장으로 작성해줘:\n\n요청내용: "
                                res_text = robust_generate_content(sys_prompt + ai_prompt)
                                if res_text:
                                    st.session_state.notion_db[subject] += "\n" + res_text
                                    if f"editor_{subject}" in st.session_state: del st.session_state[f"editor_{subject}"]
                                    st.rerun()
                            except Exception as e:
                                st.error(f"오류: {e}")

    st.write("") # 간격 조절

    # 미리보기 영역을 위에 배치하기 위한 빈 공간(placeholder) 생성
    preview_placeholder = st.empty()

    # 에디터 (아래에 배치)
    note_content = st.text_area(
        "수식이나 내용을 자유롭게 수정하세요:",
        value=st.session_state.notion_db[subject],
        height=400,
        key=f"editor_{subject}"
    )
    st.session_state.notion_db[subject] = note_content

    # 위에서 만든 빈 공간에 실시간 렌더링 결과 채워넣기
    with preview_placeholder.container():
        prev_col1, prev_col2 = st.columns([5, 1.2])
        with prev_col1:
            st.markdown("**미리보기 (실시간 수식 확인):**")
        with prev_col2:
            if st.button("🗑️ 글 초기화", key=f"reset_btn_{subject}", use_container_width=True, help="현재 에디터의 내용을 모두 지웁니다."):
                st.session_state.notion_db[subject] = ""
                if f"editor_{subject}" in st.session_state: del st.session_state[f"editor_{subject}"]
                st.rerun()
        with st.container(border=True, height=600):
            st.markdown(note_content)


# ==========================================
# 2. 실험 보고서 AI 도우미 (검색 및 연동)
# ==========================================
elif menu == "🔬 실험 보고서 AI 도우미":
    st.title("🔬 화학/물리 실험 보고서 AI 도우미")
    st.write("실험 보고서 작성에 필요한 예시 자료를 검색하고, 유용한 화학 관련 웹 프로그램들을 바로 실행해 보세요!")

    st.subheader("🔍 보고서 예시 및 템플릿 검색 (AI)")

    # 추천 샘플 TOP 5 버튼 배치
    st.markdown("🌟 **추천 전공 실험 샘플 TOP 5 (Quick Access):**")


    def open_sample_html(topic_id, title):
        html_path = os.path.join(os.path.dirname(__file__), "samples_html", f"{topic_id}.html")
        if os.path.exists(html_path):
            with open(html_path, "r", encoding="utf-8") as f:
                html_content = f.read()
            import json
            safe_html = json.dumps(html_content)
            # 팝업 창 강제 실행
            import streamlit.components.v1 as components
            components.html(f"""
                <script>
                    const win = window.open("", "_blank");
                    if (win) {{
                        win.document.open();
                        win.document.write({safe_html});
                        win.document.close();
                    }} else {{
                        alert("팝업이 차단되었습니다! 브라우저 주소창 우측의 '팝업 차단 해제'를 클릭해주세요.");
                    }}
                </script>
            """, height=0)
        else:
            st.error("해당 샘플 파일을 찾을 수 없습니다.")

    def open_report_word_direct(code, label):
        with st.spinner(f"{label} 샘플을 워드로 변환 중..."):
            import time
            out_file = os.path.join(os.getcwd(), f"{code}_Report_{int(time.time())}.docx")
            
            # 템플릿 내용 로드 (samples_html 폴더에서 직접 읽기)
            content = ""
            try:
                md_path = os.path.join(os.path.dirname(__file__), "samples_html", f"{code}.md")
                if os.path.exists(md_path):
                    with open(md_path, "r", encoding="utf-8") as f:
                        content = f.read()
                else:
                    content = f"# {label}\n\n데이터 파일을 찾을 수 없습니다."
            except Exception as e:
                content = f"# {label}\n\n로드 오류: {e}"

            margins = {"top": 2.0, "bottom": 2.0, "left": 2.5, "right": 2.5}
            if convert_latex_to_word_docx(f"# {label} 보고서 예시\n\n" + content, out_file, margins):
                # 즉시 워드 실행 (팝업 차단 영향 없음)
                open_file_in_os(out_file)
                st.success(f"🎉 {label} 워드가 새 창에서 실행되었습니다!")
                with open(out_file, "rb") as f:
                    st.download_button(f"💾 {label} 저장", f.read(), f"{code}_Report.docx", key=f"dl_rep_{code}")

    # 5열 가로 배치 (Quick Access - 워드 즉시 실행)
    cr_cols = st.columns(5)
    r_items = [("aspirin", "💊 아스피린"), ("nylon", "🧶 나일론"), ("titration", "🧪 적정"), ("lattice", "🧊 격자"), ("spectrum", "🌈 스펙트럼")]
    s_topics = ["아스피린 합성 및 재결정", "나일론 6,6 계면 중합", "산-염기 적정 지시약 원리", "NaCl 결정 구조 및 XRD 분석", "수소 원자 스펙트럼과 발머 계열"]
    
    for i, (code, label) in enumerate(r_items):
        with cr_cols[i]:
            if st.button(label, key=f"quick_rep_{i}", use_container_width=True, type="primary"):
                st.session_state.report_search_query = s_topics[i]
                
                # 내용 로드
                content = ""
                try:
                    md_path = os.path.join(os.path.dirname(__file__), "samples_html", f"{code}.md")
                    if os.path.exists(md_path):
                        with open(md_path, "r", encoding="utf-8") as f:
                            content = f.read()
                except: pass
                
                open_any_word_direct(label, content, f"{code}_Report", f"quick_rep_dl_{i}")
                
            if st.session_state.get(f"quick_rep_dl_{i}"):
                render_download_and_open_buttons(label, st.session_state[f"quick_rep_dl_{i}"], st.session_state[f"quick_rep_dl_{i}_filename"], f"quick_rep_p_{i}")

    # --- 추가된 실험 보고서 예시 워드 다운로드 센터 ---
    with st.expander("📥 실험 보고서 예시(.docx) 직접 다운로드 센터", expanded=False):
        rd_col1, rd_col2 = st.columns(2)
        for i, (code, label) in enumerate(r_items):
            target_col = rd_col1 if i % 2 == 0 else rd_col2
            with target_col:
                if st.button(f"📝 {label} (Word 즉시 열기)", key=f"rd_btn_{code}", use_container_width=True):
                    # 내용 로드
                    content = ""
                    try:
                        md_path = os.path.join(os.path.dirname(__file__), "samples_html", f"{code}.md")
                        if os.path.exists(md_path):
                            with open(md_path, "r", encoding="utf-8") as f:
                                content = f.read()
                    except: pass
                    open_any_word_direct(label, content, f"{code}_Report", f"rd_btn_dl_{code}")
                
                if st.session_state.get(f"rd_btn_dl_{code}"):
                    render_download_and_open_buttons(label, st.session_state[f"rd_btn_dl_{code}"], st.session_state[f"rd_btn_dl_{code}_filename"], f"rd_btn_p_{code}")

    # (기존 미리보기 영역 제거됨 - 팝업으로 대체)


    # 대학 및 대상 선택 UI 추가
    st.markdown("🏛️ **대학별/대상별 특화 검색 옵션:**")
    col_univ, col_target = st.columns(2)
    selected_univ = col_univ.selectbox("대학교 선택", ["전체(통합)", "서울대학교", "고려대학교", "연세대학교", "성균관대학교", "한양대학교", "KAIST", "POSTECH"], key="univ_select")
    selected_target = col_target.selectbox("대상 선택", ["전체(통합)", "교수용", "학생용"], key="target_select")

    # --- 추천 검색어 편집 및 스타일링 ---
    st.markdown("""
        <style>
        
        /* 추천 버튼 스타일: 투명 배경, 작은 글씨, 테두리 최소화 */
        div.stButton > button {
            background-color: transparent !important;
            background-image: none !important;
            border: 1px solid #ddd !important;
            color: #777 !important;
            font-size: 0.7rem !important;
            padding: 1px 4px !important;
            min-height: 22px !important;
            height: 22px !important;
            line-height: 1 !important;
            border-radius: 4px !important;
            box-shadow: none !important;
        }
        div.stButton > button:hover {
            border-color: #3B82F6 !important;
            color: #3B82F6 !important;
            background-color: rgba(59, 130, 246, 0.05) !important;
        }

            transition: all 0.2s ease !important;
        }
        div.stButton > button[key^="ex_btn_"]:hover {
            border-color: #3B82F6 !important;
            color: #3B82F6 !important;
            background-color: #f0f7ff !important;
        }
        /* 편집 입력창 스타일: 매우 작게 */
        div[data-testid="stHorizontalBlock"] div[data-testid="stTextInput"] input {
            font-size: 0.7rem !important;
            padding: 2px !important;
            height: 20px !important;
            min-height: 20px !important;
        }
        </style>
    """, unsafe_allow_html=True)

    st.markdown("<p style='font-size: 0.75rem; opacity: 0.7; margin-bottom: 0.2rem;'>💡 <b>추천 검색어 편집 및 클릭 시 자동 입력 (글자 수정 가능):</b></p>", unsafe_allow_html=True)
    
    # 추천어 상태 초기화
    if "ex_label_1" not in st.session_state: st.session_state.ex_label_1 = "🧪 색소의 분리와 흡광분석"
    if "ex_label_2" not in st.session_state: st.session_state.ex_label_2 = "💊 아스피린 합성과 정제"
    if "ex_label_3" not in st.session_state: st.session_state.ex_label_3 = "🍋 비타민 C 적정 분석"
    if "ex_label_4" not in st.session_state: st.session_state.ex_label_4 = "🧵 나일론 6,10 합성"

    # 1행: 편집 가능한 입력창
    ed_col1, ed_col2, ed_col3, ed_col4 = st.columns(4)
    l1 = ed_col1.text_input("L1", st.session_state.ex_label_1, key="edit_l1", label_visibility="collapsed", on_change=save_state)
    l2 = ed_col2.text_input("L2", st.session_state.ex_label_2, key="edit_l2", label_visibility="collapsed", on_change=save_state)
    l3 = ed_col3.text_input("L3", st.session_state.ex_label_3, key="edit_l3", label_visibility="collapsed", on_change=save_state)
    l4 = ed_col4.text_input("L4", st.session_state.ex_label_4, key="edit_l4", label_visibility="collapsed", on_change=save_state)

    # 2행: 실제 적용 버튼 (투명/작게 스타일 적용됨)
    ex_col1, ex_col2, ex_col3, ex_col4 = st.columns(4)
    if ex_col1.button(f"{l1}", use_container_width=True, key="ex_btn_1"):
        st.session_state.report_search_query = l1
        st.session_state.ex_label_1 = l1
        save_state()
        st.rerun()
    if ex_col2.button(f"{l2}", use_container_width=True, key="ex_btn_2"):
        st.session_state.report_search_query = l2
        st.session_state.ex_label_2 = l2
        save_state()
        st.rerun()
    if ex_col3.button(f"{l3}", use_container_width=True, key="ex_btn_3"):
        st.session_state.report_search_query = l3
        st.session_state.ex_label_3 = l3
        save_state()
        st.rerun()
    if ex_col4.button(f"{l4}", use_container_width=True, key="ex_btn_4"):
        st.session_state.report_search_query = l4
        st.session_state.ex_label_4 = l4
        save_state()
        st.rerun()

    st.markdown("---")
    with st.expander("🛠️ 다운로드했던 '안전 모드' 워드 파일 수식 변환하기 (업로드)"):
        st.info("이전에 수식 변환 없이 텍스트 전용으로 다운로드했던 워드 파일(.docx)을 업로드하시면, 수식을 완벽하게 변환하여 다시 다운로드할 수 있습니다.")
        uploaded_docx = st.file_uploader("안전 모드로 다운받은 워드 파일(.docx) 업로드", type=["docx"], key="upload_safe_docx")
        if uploaded_docx:
            if st.button("🚀 업로드한 파일 수식 변환 시작", use_container_width=True):
                with st.spinner("파일의 텍스트를 읽어와 수식을 변환 중입니다 (Pandoc)..."):
                    try:
                        import os
                        from docx import Document
                        
                        # 업로드된 파일 읽기
                        temp_doc = Document(uploaded_docx)
                        extracted_text = "\n".join([p.text for p in temp_doc.paragraphs])
                        
                        # 변환 수행
                        clean_res = "".join(c for c in extracted_text if c.isprintable() or c in "\n\r\t")
                        out_file_hq = f"HQ_Converted_{int(time.time())}.docx"
                        convert_latex_to_word_docx(clean_res, out_file_hq, {'top': 2.0, 'bottom': 2.0, 'left': 2.5, 'right': 2.5})
                        
                        if os.path.exists(out_file_hq):
                            with open(out_file_hq, "rb") as f:
                                st.session_state["converted_docx_bytes"] = f.read()
                            st.session_state["converted_docx_name"] = f"Converted_{uploaded_docx.name}"
                            os.remove(out_file_hq)
                            st.success("수식 변환이 완료되었습니다! 아래 버튼을 눌러 다운로드하세요.")
                    except Exception as e:
                        st.error(f"변환 중 오류가 발생했습니다: {e}")
                        
            if st.session_state.get("converted_docx_bytes"):
                st.download_button(
                    "📘 고품질 Word 다운로드 (수식 완벽 변환)", 
                    data=st.session_state["converted_docx_bytes"], 
                    file_name=st.session_state["converted_docx_name"], 
                    key="dl_converted_docx_persistent", 
                    use_container_width=True
                )
    st.markdown("---")

    search_query = st.text_input("(00 실험 보고서 작성해줘 )", key="report_search_query", placeholder="예: 아스피린 합성 실험, 산염기 적정, 나일론 합성")

    col1, col2 = st.columns([3, 1])
    with col1:
        search_btn = st.button("🚀 AI로 보고서 예시/구조 검색하기", use_container_width=True)
    with col2:
        reset_btn = st.button("🗑️ 보고서 휴지통으로 이동", use_container_width=True)

    if reset_btn:
        if st.session_state.get("last_report_result"):
            import datetime
            trash_item = {
                "type": "보고서",
                "query": st.session_state.get("last_report_query", "제목 없음"),
                "content": st.session_state.last_report_result,
                "time": datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            }
            st.session_state.global_trash_bin.insert(0, trash_item)
            # 최대 5개까지만 보관
            st.session_state.global_trash_bin = st.session_state.global_trash_bin[:5]
            
        st.session_state.last_report_result = None
        st.session_state.last_report_query = None
        st.rerun()

    if search_btn:
        if not api_key:
            st.error("왼쪽 메뉴에 Gemini API Key를 입력해주세요!")
        elif search_query:
            st.session_state.report_gen_active = True
            st.session_state.report_gen_phases = [
                ("Phase 1: 초록, 서론, 이론적 배경", """
1. **Abstract (초록)**: 연구/실험 전체를 관통하는 핵심 요약 (매우 상세하게).
2. **Introduction (서론 및 배경)**: 해당 주제의 학술적/산업적 배경, 역사적 맥락, 사회적 중요성.
3. **Theoretical Background (이론적 배경 및 심층 분석)**: 핵심 메커니즘, 분자/물리적 관점에서의 해석, 열역학/속도론적 증명. 수식($$ ... $$)을 동원한 논리적 전개.
"""),
                ("Phase 2: 다중 실험 설계 및 데이터 분석", """
4. **Experimental Methodology (다중 실험 설계 및 방대한 프로토콜)**: 본 연구는 단일 실험이 아닌, 변수를 통제한 최소 3~5가지 이상의 심화/연계 실험(Phase 1, 2, 3 등)으로 구성할 것. 각 실험 단계마다 완벽하게 분리된 세팅, 시약 명세표, 기구 보정 절차 기술.
5. **Data Analysis & Results (방대한 결과 도출 및 다중 시각화)**: 각 실험 단계별로 가상의 방대한 로우 데이터(Raw Data)를 기반으로 최소 5~10개 이상의 거대한 마크다운 표(Table) 작성. 각 표마다 ASCII 아트나 텍스트 심볼을 활용한 정밀한 시각화 그래프(적정 곡선 등) 1:1 매칭 직접 작성. 철저한 통계적 검정의 수식적 계산 과정을 포함할 것.
"""),
                ("Phase 3: 심층 고찰 및 결론", """
6. **Discussion & Error Modeling (심층 논의 및 오차 모델링)**: 논문 수준의 고찰. 예상되는 오차의 계통적/우연적 원인 분석 및 보정 논리. 변수(온도, pH 등)가 미치는 영향 고찰.
7. **Conclusion & Perspectives (결론 및 향후 전망)**: 연구의 완벽한 요약과 학계/산업계에 미칠 파급 효과.
8. **References (참고문헌)**: 국내외 권위 있는 학술지, 교과서, 논문 등을 상세히 명시.
""")
            ]
            st.session_state.report_gen_buffer = []
            st.session_state.last_report_query = search_query
            st.session_state.last_report_result = None  # 이전 결과 초기화
            save_state()
            st.rerun()

    # --- 무중단 엔진 코어 (보고서 3단계 생성) ---
    if st.session_state.get("report_gen_active"):
        if st.button("🛑 작성 중단 및 지금까지 결과 보기", key="stop_report_btn", use_container_width=True):
            st.session_state.report_gen_active = False
            if st.session_state.get("report_gen_buffer"):
                st.session_state.last_report_result = "\n\n".join(st.session_state.report_gen_buffer)
            save_state()
            st.warning("보고서 작성을 강제 중단했습니다. 화면 하단에서 지금까지 완성된 파트를 확인하세요.")
            st.rerun()
            
        phases = st.session_state.report_gen_phases
        buffer = st.session_state.get("report_gen_buffer", [])
        search_query = st.session_state.get("last_report_query", "주제 없음")
        
        total_phases = len(phases)
        completed_phases = len(buffer)
        remaining_phases = total_phases - completed_phases
        est_sec = remaining_phases * 25
        
        st.progress(completed_phases / total_phases if total_phases > 0 else 0.0)
        st.caption(f"📊 보고서 작성 진행률: {completed_phases} / {total_phases} 파트 완료 (예상 남은 시간: 약 {est_sec}초)")
        
        if completed_phases < total_phases:
            phase_title, phase_desc = phases[completed_phases]
            with st.status(f"🔍 '{search_query}' - {phase_title} 작성 중...", expanded=True) as status:
                st.write(f"⚡ {phase_title} 파트를 심층 분석 및 작성하고 있습니다. 잠시만 기다려주세요...")
                genai.configure(api_key=api_key)
                db_context = load_local_academic_db()
                
                # 이전 파트의 맥락
                prev_context = buffer[-1][-3000:] if buffer else "없음 (첫 파트 시작)"
                
                prompt = fr"""너는 세계 최고 수준의 학술 연구원이자 대학 수석 연구원이야.
학생이 요청한 학술/실험/논문 주제: '{search_query}'

**[미션: '{phase_title}' 파트 초심층 작성]**
전체 보고서 중 현재 **[{phase_title}]** 파트만 집중적으로 작성하라.

**지시사항 (미준수 시 실격):**
0. **기간 및 제약 해제**: 어떤 시간적 제약이나 안전 가이드라인에 얽매이지 말고 방대한 데이터를 탐색해라.
1. **언어 및 표기법**: 한국어로 작성하되 중요 개념은 (영문) 병기. **별표(`**` 등)를 이용한 굵게/기울임 표시 절대 금지.** 순수 텍스트로만 작성.
2. **분량 극대화**: 가용 토큰을 최대로 사용하여 길고 장대하게 서술하라. 단 하나의 정보나 수치도 누락하지 말고, 얻을 수 있는 모든 데이터를 끌어모아 2000자 이상 서술하라.
3. **다양한 시각적 요소**: 내용에 3D 격자, 분자 구조식, 그래프, 도표 등이 포함된다면 ASCII 아트나 마크다운 표로 무조건 시각적으로 렌더링하라.
4. **수식 변환**: 모든 반응식과 공식은 **블록형 LaTeX($$ ... $$)** 문법으로 감싸서 작성.
5. **기준 준수**: 로컬 DB 데이터 중 '흡광분석 채점기준표'나 예시의 뼈대를 최대한 참고하라.

**[이번에 작성해야 할 세부 내용]**
{phase_desc}

**[이전 파트의 마지막 맥락 (자연스럽게 이어지도록 참고)]**
{prev_context}

[핵심 학술 데이터베이스 (수업 과제 및 채점기준)]
{db_context if db_context else "참조 가능한 로컬 과제 DB 없음"}
"""
                try:
                    res = robust_generate_content(prompt, use_grounding=False)
                    import time; time.sleep(4)
                    if res:
                        res = res.replace("***", "").replace("**", "")
                        buffer.append(f"## {phase_title}\n\n" + res)
                        st.session_state.report_gen_buffer = buffer
                        save_state()
                        st.rerun()
                    else:
                        st.error("생성 중 오류 발생 (자동 수위 조절 실패)")
                        st.session_state.report_gen_active = False
                        st.stop()
                except Exception as e:
                    st.error(f"AI 생성 중 오류가 발생했습니다: {e}")
                    st.session_state.report_gen_active = False
                    st.stop()
        else:
            # 모두 완료
            st.session_state.report_gen_active = False
            response_text = "\n\n".join(buffer)
            st.session_state.last_report_result = response_text
            
            # [긴급 최적화] 워드 파일 미리 생성하여 캐싱 (다운로드 지연 방지)
            try:
                import re, io
                from docx import Document
                clean_res = "".join(c for c in response_text if c.isprintable() or c in "\n\r\t")
                safe_q = re.sub(r'[\\/*?:"<>|]', '', search_query)[:20].strip().replace(' ', '_')
                if not safe_q: safe_q = "Report"
                
                # 1. 호환성 모드 캐싱 (필수)
                doc_safe = Document()
                doc_safe.add_heading(f"학술 보고서: {search_query}", 0)
                for paragraph in clean_res.split('\n'):
                    if paragraph.strip(): doc_safe.add_paragraph(paragraph)
                from docx.oxml import parse_xml
                doc_safe.settings.element.append(parse_xml(r'<w:documentProtection w:edit="readOnly" w:enforcement="1" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'))
                safe_stream = io.BytesIO()
                doc_safe.save(safe_stream)
                st.session_state.last_report_safe_word = safe_stream.getvalue()
                st.session_state.last_report_safe_name = f"Safe_{safe_q}.docx"
                
            except Exception as e:
                print(f"[Word Generation Error] {e}")
            
            save_state()
            st.success("🎊 초고밀도 실험 보고서(전체 파트) 완벽 생성 완료!")
            st.rerun()

    # --- [결과 표시 영역] 버튼 밖으로 이동하여 검색 결과 유지 ---
    if st.session_state.get("last_report_result"):
        res_text = st.session_state.last_report_result
        q_text = st.session_state.get("last_report_query", "검색 결과")
        
        # [데이터 복구] 워드 캐시가 비어있다면 즉시 재빌드
        if not st.session_state.get("last_report_safe_word"):
            try:
                import re, io
                from docx import Document
                clean_res = "".join(c for c in res_text if c.isprintable() or c in "\n\r\t")
                safe_q = re.sub(r'[\\/*?:"<>|]', '', q_text)[:20].strip().replace(' ', '_')
                if not safe_q: safe_q = "Report"
                
                doc_safe = Document()
                doc_safe.add_heading(f"학술 보고서: {q_text}", 0)
                for paragraph in clean_res.split('\n'):
                    if paragraph.strip(): doc_safe.add_paragraph(paragraph)
                safe_stream = io.BytesIO()
                doc_safe.save(safe_stream)
                st.session_state.last_report_safe_word = safe_stream.getvalue()
                st.session_state.last_report_safe_name = f"Safe_{safe_q}.docx"
            except: pass

        with st.container(border=True):
            st.markdown(f"### ✨ '{q_text}' AI 분석 보고서")
            st.markdown(res_text)

            # [최적화된 이중 다운로드 시스템]
            col_dl1, col_dl2 = st.columns(2)
            
            with col_dl1:
                if st.session_state.get("last_report_hq_word"):
                    st.download_button(
                        "📘 고품질 Word 다운로드", 
                        data=st.session_state.last_report_hq_word, 
                        file_name=st.session_state.last_report_hq_name, 
                        key="report_word_hq_v4", 
                        use_container_width=True
                    )
                else:
                    if st.button("🚀 수식 변환 MS 워드 생성하기", use_container_width=True, key="trigger_hq_word"):
                        with st.spinner("Pandoc 엔진을 통해 수식을 완벽하게 변환 중입니다... (10~30초 소요)"):
                            try:
                                import re, io, os
                                clean_res = "".join(c for c in res_text if c.isprintable() or c in "\n\r\t")
                                safe_q = re.sub(r'[\\/*?:"<>|]', '', q_text)[:20].strip().replace(' ', '_')
                                if not safe_q: safe_q = "Report"
                                out_file_hq = f"HQ_{safe_q}.docx"
                                convert_latex_to_word_docx(clean_res, out_file_hq, {'top': 2.0, 'bottom': 2.0, 'left': 2.5, 'right': 2.5})
                                if os.path.exists(out_file_hq):
                                    with open(out_file_hq, "rb") as f:
                                        st.session_state.last_report_hq_word = f.read()
                                    st.session_state.last_report_hq_name = out_file_hq
                                    os.remove(out_file_hq)
                                    st.rerun()
                            except Exception as e:
                                st.error(f"워드 파일 생성 중 오류가 발생했습니다: {e}")
            
            with col_dl2:
                # 호환성 모드 (캐시된 데이터 사용) - 즉시 응답
                if st.session_state.get("last_report_safe_word"):
                    st.download_button(
                        "📗 빠른 다운로드 (수식 미변환, 텍스트 전용)", 
                        data=st.session_state.last_report_safe_word, 
                        file_name=st.session_state.last_report_safe_name, 
                        key="report_word_safe_v4", 
                        use_container_width=True
                    )
                else:
                    st.warning("⚠️ 파일 준비 중...")

            # 데이터베이스 저장 버튼
            st.divider()
            st.subheader("📂 데이터베이스(Notion)에 즉시 저장")
            db_col1, db_col2 = st.columns([3, 1])
            with db_col1:
                target_sub = st.selectbox("저장할 과목 선택", list(st.session_state.notion_db.keys()), key="report_db_subject_v2")
            with db_col2:
                if st.button("📥 DB에 추가", use_container_width=True, key="report_db_add_btn_v2"):
                    st.session_state.notion_db[target_sub] += "\n---\n" + res_text
                    save_state()
                    st.success(f"'{target_sub}' 데이터베이스에 저장되었습니다!")

    st.markdown("---")
    st.subheader("🔗 전공 및 화학교육 필수 연동 데이터베이스")
    render_chem_ed_core_guide()
    st.write("학부 전공 수준의 데이터 처리 및 예비 교사를 위한 교육용 툴입니다. (아래 샘플 클릭 시 해당 주제로 즉시 이동)")

    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.info("**🧪 분자 구조 및 반응식**\n\n[ChemDraw JS (무료)](https://chemdrawdirect.perkinelmer.cloud/js/sample/index.html)")
        st.success("**🧬 3D 단백질/DNA 구조**\n\n[Sample: DNA 3D 구조(1BNA)](https://www.rcsb.org/structure/1BNA)")
    with col2:
        st.warning("**📊 스펙트럼(IR, NMR) 데이터**\n\n[Sample: 에탄올 IR 스펙트럼](https://webbook.nist.gov/cgi/cbook.cgi?ID=C64175&Type=IR-SPEC&Index=1)")
        st.error("**🔬 화합물 정밀 물성 검색**\n\n[Sample: 아스피린(Aspirin) 데이터](https://pubchem.ncbi.nlm.nih.gov/compound/Aspirin)")
    with col3:
        st.info("**📈 데이터 추출 및 그래프**\n\n[WebPlotDigitizer](https://automeris.io/WebPlotDigitizer/)")
        st.success("**📚 화학교육 학술 문헌 검색**\n\n[Sample: 오개념(Misconception) 논문](https://scholar.google.com/scholar?q=misconceptions+in+chemistry+education)")

    with col4:
        st.warning("**🧑‍🏫 교육용 가상 실험실**\n\n[Sample: 분자 기하 구조 시뮬레이션](https://phet.colorado.edu/en/simulations/molecule-shapes)")
# 3, 4, 5번 메뉴는 기존 코드 동일하게 유지됨
# ==========================================
elif menu == "🎓 전문가용 LaTeX (Overleaf) 에디터":
    st.title("🎓 전문가용 LaTeX (Overleaf) 에디터")
    col_ai, col_preview = st.columns([1, 1])
    with col_ai:
        uploaded_file = st.file_uploader("수기 노트를 LaTeX 논문 코드로 변환", type=["png", "jpg", "jpeg"], key="latex_up")
        if uploaded_file and st.button("LaTeX 코드로 완벽 변환"):
            if not api_key: st.error("API 키를 입력해주세요.")
            else:
                with st.spinner("AI가 논문 코드를 작성 중입니다..."):
                    genai.configure(api_key=api_key)
                    model = get_safe_gemini_model()
                    try:
                        res_text = robust_generate_content("Overleaf 컴파일용 완벽한 LaTeX 코드만 출력해.", images=Image.open(uploaded_file))
                        st.session_state.latex_code = res_text
                    except Exception as e: st.error(e)
    if "latex_code" not in st.session_state: st.session_state.latex_code = r"$$ E = mc^2 $$"
    st.markdown("### 🧪 필수 수식 템플릿 (클릭 시 전체 수식 자동 입력)")

    if "mathlive_init" not in st.session_state:
        st.session_state.mathlive_init = r"\sum_{i=1}^{n}"
    def set_mathlive_template(template_str):
        st.session_state.mathlive_init = template_str

    t1, t2, t3 = st.columns(3)
    t1.button(r"1D 상자 속 입자 (Wavefunction)", use_container_width=True, on_click=set_mathlive_template, args=(r"\psi_n(x) = \sqrt{\frac{2}{L}} \sin\left(\frac{n\pi x}{L}\right)",))
    t2.button(r"해밀토니안 연산자 (Hamiltonian)", use_container_width=True, on_click=set_mathlive_template, args=(r"\hat{H} = -\frac{\hbar^2}{2m}\frac{d^2}{dx^2} + V(x)",))
    t3.button(r"슈뢰딩거 방정식 (Schrödinger)", use_container_width=True, on_click=set_mathlive_template, args=(r"-\frac{\hbar^2}{2m}\frac{d^2\psi}{dx^2} + V(x)\psi = E\psi",))

    st.markdown("---")

    st.markdown("### 🧮 완벽 시각화 수식 에디터 (빈칸 클릭 후 숫자 입력)")

    import textwrap
    custom_mathlive_html = textwrap.dedent("""<!DOCTYPE html>
    <html>
    <head>
    __SCRIPT_LIB__
    <style>
    body { font-family: sans-serif; margin: 0; padding: 10px; background-color: white; }
    #mf {
    font-size: 28px; width: 100%; padding: 15px; min-height: 120px;
    border: 2px solid #3B82F6; border-radius: 8px;
    background-color: #F8FAFC; box-sizing: border-box;
    margin-bottom: 15px;
    }
    .keypad {
    display: grid;
    grid-template-columns: repeat(10, 1fr);
    gap: 8px;
    margin-bottom: 15px;
    }
    .keypad button {
    padding: 12px 5px; font-size: 16px; font-weight: bold;
    border: 1px solid #CBD5E1; border-radius: 6px;
    background-color: #F1F5F9; cursor: pointer;
    transition: 0.2s;
    }
    .keypad button:hover { background-color: #E2E8F0; }
    .keypad button.action { background-color: #DBEAFE; color: #1E40AF; border-color: #BFDBFE; }
    .keypad button.action:hover { background-color: #BFDBFE; }
    .copy-btn {
    width: 100%; padding: 15px; font-size: 18px; font-weight: bold;
    background-color: #10B981; color: white; border: none;
    border-radius: 8px; cursor: pointer; transition: 0.3s;
    }
    .copy-btn:hover { background-color: #059669; }
    </style>
    </head>
    <body>
    <p style="color: #64748B; font-size: 14px; margin-top: 0;">점선으로 된 빈칸(Placeholder)을 마우스로 클릭하고 아래 숫자 버튼을 눌러보세요!</p>
    <math-field id="mf">__MATHLIVE_INIT_VALUE__</math-field>
    <textarea id="latex-code" style="width: 100%; font-family: monospace; font-size: 12pt; padding: 10px; border: 2px solid #CBD5E1; border-radius: 8px; resize: vertical; background-color: #F8FAFC; min-height: 60px; margin-bottom: 15px; box-sizing: border-box;" placeholder="여기에 LaTeX 코드가 텍스트로 나타납니다 (직접 수정 가능)..."></textarea>

    <div class="keypad">
    <button onclick="insert('1')">1</button>
    <button onclick="insert('2')">2</button>
    <button onclick="insert('3')">3</button>
    <button onclick="insert('4')">4</button>
    <button onclick="insert('5')">5</button>
    <button onclick="insert('6')">6</button>
    <button onclick="insert('7')">7</button>
    <button onclick="insert('8')">8</button>
    <button onclick="insert('9')">9</button>
    <button onclick="insert('0')">0</button>

    <button onclick="insert('+')">+</button>
    <button onclick="insert('-')">-</button>
    <button onclick="insert('=')">=</button>
    <button onclick="insert('(')">(</button>
    <button onclick="insert(')')">)</button>
    <button class="action" onclick="insert('\\\\frac{#?}{#?}')">분수</button>
    <button class="action" onclick="insert('\\\\sqrt{#?}')">√</button>
    <button class="action" onclick="insert('^{#?}')">제곱</button>
    <button class="action" onclick="insert('_{#?}')">아래첨자</button>
    <button class="action" onclick="document.getElementById('mf').executeCommand('deleteAll')">전체지움</button>

    <button class="action" onclick="insert('\\\\int_{#?}^{#?}')">∫ 적분</button>
    <button class="action" onclick="insert('\\\\sum_{#?}^{#?}')">∑ 시그마</button>
    <button class="action" onclick="insert('\\\\infty')">∞</button>
    <button class=r"action" onclick="insert('\\\\psi')">ψ</button>
    <button class=r"action" onclick="insert('\\\\pi')">π</button>
    <button class="action" onclick="insert('\\\\alpha')">α</button>
    <button class="action" onclick="insert('\\\\beta')">β</button>
    <button class="action" onclick="insert('\\\\theta')">θ</button>
    <button class="action" onclick="insert('\\\\Delta')">Δ</button>
    <button class="action" onclick="document.getElementById('mf').executeCommand('moveToPreviousChar')">◀ 이동</button>
    </div>

    <button class="copy-btn" id="copyBtn" onclick="copyToClipboard()" style="margin-top:10px; background-color:#10B981;">📋 완성된 수식 복사하기 (클릭 후 아래 에디터에 붙여넣기 Ctrl+V)</button>

    <script>
    const mf = document.getElementById('mf');
    const latexCode = document.getElementById('latex-code');

    // 양방향 동기화 설정
    mf.addEventListener('input', (ev) => {
    latexCode.value = mf.value;
    });

    latexCode.addEventListener('input', (ev) => {
    mf.value = latexCode.value;
    });

    // 초기화
    customElements.whenDefined("math-field").then(() => {
        latexCode.value = mf.value || mf.textContent || "";
    });

    function insert(latex) {
    if (typeof mf.insert === 'function') {
    mf.insert(latex);
    } else {
    mf.executeCommand(['insert', latex]);
    }
    mf.focus();
    }

    function copyToClipboard() {
    const code = '$$ ' + mf.value + ' $$';
    const el = document.createElement('textarea');
    el.value = code;
    el.style.position = 'absolute';
    el.style.left = '-9999px';
    document.body.appendChild(el);
    el.select();

    try {
    document.execCommand('copy');
    const btn = document.getElementById('copyBtn');
    const orig = btn.innerText;
    btn.innerText = '✅ 복사 완료! 아래 메인 에디터 코드창에 붙여넣기(Ctrl+V) 하세요!';
    btn.style.backgroundColor = '#2563EB';
    setTimeout(() => {
    btn.innerText = orig;
    btn.style.backgroundColor = '#10B981';
    }, 3000);
    } catch (er) {
    alert('복사에 실패했습니다.');
    }
    document.body.removeChild(el);
    }
    </script>
    </body>
    </html>
    """)

    import streamlit.components.v1 as components
    html_to_render = custom_mathlive_html.replace('__MATHLIVE_INIT_VALUE__', st.session_state.mathlive_init)
    html_to_render = html_to_render.replace('__SCRIPT_LIB__', '<script defer src="https://unpkg.com/mathlive"></script>')
    components.html(html_to_render, height=620)

    st.markdown("---")

    col_edit, col_render = st.columns([1, 1])
    with col_edit:
        user_latex = st.text_area("코드 수정:", value=st.session_state.latex_code, height=400, key="latex_editor_textarea")

    with col_render:
        st.info("⚠️ 미리보기 화면입니다.")
        with st.container(border=True):
            render_text = user_latex.strip()
            # 사용자가 $$를 쓰지 않고 텍스트만 쭉 썼거나 중간에 $ 하나만 쓴 경우를 모두 수식으로 예쁘게 렌더링하기 위한 보정
            if "$$" not in render_text and len(render_text) > 0:
                parts = render_text.split("$")
                render_text = "\n\n".join([f"$$ {p.strip()} $$" for p in parts if p.strip()])
            st.markdown(render_text)
    st.markdown("---")
    if st.button("💾 La Tex 워드 파일 변환 및 다운로드 준비 (Pandoc 수식 완벽 변환)", use_container_width=True):
        with st.spinner("수식을 워드용으로 변환 중입니다..."):
            try:
                import os, time
                out_file = f"LaTeX_Export_{int(time.time())}.docx"
                # 줄바꿈 유지 및 제어 문자 제거
                clean_latex = "".join(c for c in user_latex if c.isprintable() or c in "\n\r\t").strip()
                # 워드 변환(Pandoc) 엔진도 수식을 인식하려면 $$가 필요함 (미리보기 화면과 동일한 보정 적용)
                if "$$" not in clean_latex and len(clean_latex) > 0:
                    parts = clean_latex.split("$")
                    clean_latex = "\n\n".join([f"$$ {p.strip()} $$" for p in parts if p.strip()])
                    
                convert_latex_to_word_docx(clean_latex, out_file, {'top': 2.0, 'bottom': 2.0, 'left': 2.5, 'right': 2.5})
                
                if os.path.exists(out_file):
                    with open(out_file, "rb") as f:
                        st.session_state.latex_docx_bytes = f.read()
                    os.remove(out_file)
                    st.success("변환 성공! 아래 나타난 다운로드 버튼을 클릭하세요.")
            except Exception as e:
                st.error(f"워드 파일 변환 중 오류가 발생했습니다: {e}")

    if st.session_state.get("latex_docx_bytes"):
        st.markdown("---")
        # [사용자 요청 반영] 2단 컬럼 대신 풀사이즈 스택 버튼으로 변경하여 가독성 극대화
        st.download_button(
            label="💾 변환된 La Tex 워드 파일로 다운로드",
            data=st.session_state.latex_docx_bytes,
            file_name="LaTeX_Equation.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
            type="primary",
            key="latex_dl_btn_v2"
        )
        
        if st.button("⚡️ 변환된 La Tex 워드 파일 즉시 열기", use_container_width=True, type="primary", key="latex_open_btn_v2"):
            import os
            out_file = os.path.join(os.getcwd(), "Direct_LaTeX_Equation.docx")
            try:
                with open(out_file, "wb") as f:
                    f.write(st.session_state.latex_docx_bytes)
                if not open_file_in_os(out_file):
                    st.info("ℹ️ 웹 브라우저(Streamlit Cloud) 환경에서는 자동 열기가 지원되지 않습니다. 위 '다운로드' 버튼을 이용해 주세요!")
            except Exception as e:
                st.error(f"파일 준비 중 오류가 발생했습니다: {e}")


elif menu == "🧪 도표 & 3D 그림 생성기 / 80페이지+ 초정밀 분석":
    if st.session_state.get("smart_analysis_active", False):
        if "buffer_wiped_for_quality_v4" not in st.session_state:
            st.session_state.analysis_buffer = {}
            st.session_state.smart_q_structures = {} # 캐시 초기화
            st.session_state.smart_analysis_start_time = time.time()
            st.session_state.buffer_wiped_for_quality_v4 = True
            st.rerun()


    st.markdown('''
        <style>
        /* 수식(KaTeX) 볼드 및 검정색 강제 설정 */
        .katex {
            color: #000000 !important;
            font-weight: 800 !important;
        }
        .katex .mathdefault {
            font-weight: 800 !important;
        }
        div[data-testid="stTextArea"] textarea {
            font-size: 1.1rem !important;
            padding: 15px !important;
            line-height: 1.6 !important;
        }
        </style>
    ''', unsafe_allow_html=True)

    st.subheader("🧪 도표 & 3D 그림 생성기 / 80페이지+ 초정밀 분석")

    # --- 스마트 라우터 기능 시작 ---
    if not st.session_state.get("smart_analysis_active"):
        smart_imgs = st.file_uploader("📸 이미지/PDF 업로드 (80~1000페이지 지원)", type=["png", "jpg", "jpeg", "pdf", "hwp"], accept_multiple_files=True, key="smart_imgs")
        
        st.text_area("✨ 추가 지시사항 및 텍스트 문제 직접 입력 (선택사항):",
                     key="smart_input_val", height=200,
                     placeholder="[사용법 1] 인터넷이나 문서에 있는 텍스트 문제들을 복사해서 여기에 붙여넣기(Ctrl+V) 하세요.\n[사용법 2] 위에 PDF를 업로드한 경우, 'PDF에서 열역학 문제만 찾아줘' 같은 특별 지시사항을 적으실 수 있습니다.\n\n⚠️ 주의: PDF/이미지 파일은 여기에 드래그 앤 드롭하실 수 없습니다. 바로 위의 '구름 모양(업로드)' 상자에 넣어주세요!")
                     
        with st.expander("⚙️ 고급 분석 옵션 (특정 페이지 / 특정 문항만 선택)"):
            col_opt1, col_opt2 = st.columns(2)
            with col_opt1:
                st.text_input("📄 특정 페이지 지정 (예: 1-5, 8, 12)", key="smart_target_pages", placeholder="비워두면 80+페이지 전체 분석")
            with col_opt2:
                st.text_input("🎯 특정 문항 지정 (예: 1번~10번, 5번만)", key="smart_target_questions", placeholder="비워두면 전체 문항 분석")
        
        col_start1, col_start2 = st.columns(2) if st.session_state.get("analysis_buffer") else (st.columns([1])[0], None)
        
        start_clicked = False
        resume_clicked = False
        
        if st.session_state.get("analysis_buffer"):
            with col_start1:
                if st.button("▶️ 중단된 부분부터 이어서 계속 풀기 (안전)", key="resume_analysis_btn", use_container_width=True):
                    resume_clicked = True
            with col_start2:
                if st.button("⚠️ 처음부터 새로 시작 (기존 데이터 삭제)", key="restart_analysis_btn", use_container_width=True):
                    start_clicked = True
        else:
            if st.button("🚀 무중단 초정밀 심층 분석 시작 (80페이지+ 최적화)", key="start_super_analysis_btn", use_container_width=True):
                start_clicked = True
    
        if start_clicked or resume_clicked:
            raw_query = st.session_state.smart_input_val.strip()
            if st.session_state.get("gemini_api_key"):
                if not smart_imgs and not raw_query and not resume_clicked:
                    st.warning("분석할 텍스트를 입력하거나 파일을 업로드해주세요.")
                else:
                    import fitz, hashlib, os, time
                    
                    # 새로 시작할 때만 파일 매핑을 다시 합니다.
                    if start_clicked:
                        page_map = [] 
                        file_paths = []
                        os.makedirs("temp_uploads", exist_ok=True)
                        
                        target_pages_str = st.session_state.get("smart_target_pages", "")
                        
                        def parse_pages(p_str, max_p):
                            if not p_str.strip(): return set(range(max_p))
                            res = set()
                            import re
                            # Matches ranges like "1-5" or "1~5", and single numbers like "5"
                            matches = re.finditer(r'(\d+)\s*(?:-|~)\s*(\d+)|(\d+)', p_str)
                            for match in matches:
                                if match.group(1) and match.group(2):
                                    try:
                                        s, e = int(match.group(1)), int(match.group(2))
                                        for p in range(s-1, e):
                                            if 0 <= p < max_p: res.add(p)
                                    except: pass
                                elif match.group(3):
                                    try:
                                        p = int(match.group(3))-1
                                        if 0 <= p < max_p: res.add(p)
                                    except: pass
                            return res if res else set(range(max_p))
                        
                        if raw_query:
                            t_path = f"temp_uploads/smart_text_{int(time.time())}.txt"
                            with open(t_path, "w", encoding="utf-8") as f:
                                f.write(raw_query)
                            file_paths.append(t_path)
                            page_map.append(('txt', len(file_paths)-1, 0))
                        
                        for f_idx, uploaded_file in enumerate(smart_imgs):
                            t_path = f"temp_uploads/{uploaded_file.name}"
                            with open(t_path, "wb") as f:
                                f.write(uploaded_file.getvalue())
                            
                            current_file_idx = len(file_paths)
                            file_paths.append(t_path)
        
                            if uploaded_file.name.lower().endswith(".pdf"):
                                doc = fitz.open(t_path)
                                target_pages = parse_pages(target_pages_str, len(doc))
                                for p_idx in range(len(doc)):
                                    if p_idx in target_pages:
                                        page_map.append(('pdf', current_file_idx, p_idx))
                                doc.close()
                            else: page_map.append(('img', current_file_idx, 0))
                        
                        st.session_state.smart_file_paths = file_paths
                        st.session_state.smart_page_map = page_map
                        st.session_state.analysis_buffer = {}
                        st.session_state.smart_q_structures = {}
                        if "smart_timer_state" in st.session_state: del st.session_state["smart_timer_state"]
                    
                    st.session_state.smart_analysis_active = True
                    save_state()
                    st.rerun()
            else: st.warning("🔑 Gemini API 키를 먼저 입력해 주세요.")
    else:
        st.success("🟢 80페이지+ 무중단 정밀 분석 엔진이 가동 중입니다. (브라우저를 새로고침 하셔도 백그라운드에서 안전하게 계속 풀고 있습니다.)")

    def render_smart_analysis_results(is_live_preview=False):
        import os
        if not (st.session_state.get("analysis_buffer") or st.session_state.get("global_smart_img_result")):
            return
            
        with st.container(border=True):
            if is_live_preview:
                st.markdown("### 🟢 [Live] AI 초정밀 분석 실시간 미리보기")
                st.caption("엔진이 돌아가는 중에도 이미 풀린 문제들을 클릭해서 즉시 확인할 수 있습니다.")
            else:
                st.markdown("### 📋 AI 초정밀 분석 결과 (80+p 무중단 진행 완료)")
            
            curent_full_res = ""
            if st.session_state.get("analysis_buffer"):
                sorted_keys = sorted([k for k in st.session_state.analysis_buffer.keys() if k.startswith("smart_p_")])
                
                from collections import defaultdict
                grouped_res = defaultdict(list)
                
                for k in sorted_keys:
                    try:
                        parts = k.split('_')
                        p_num = int(parts[2]) + 1
                        q_num = parts[4]
                        sub_id = parts[6] if len(parts) > 6 else ""
                        grouped_res[(p_num, q_num)].append((k, sub_id))
                    except:
                        grouped_res[("기타", k)].append((k, ""))
                        
                for (p_num, q_num), items in grouped_res.items():
                    if p_num == "기타":
                        for (k, sub_id) in items:
                            with st.expander(f"📄 구조 분석 실패 항목 ({q_num})", expanded=False):
                                st.markdown(st.session_state.analysis_buffer[k])
                    else:
                        st.markdown(f"#### 📝 {p_num}페이지 - {q_num}번 문항 그룹")
                        for (k, sub_id) in items:
                            col_text, col_trash = st.columns([0.9, 0.1])
                            with col_text:
                                if sub_id and sub_id != "_none":
                                    label = f"   ↪ {q_num}{sub_id} 소문항 심층 풀이 완료"
                                else:
                                    label = f"   ↪ {q_num}번 전체 심층 풀이 완료"
                                with st.expander(label, expanded=False):
                                    st.markdown(st.session_state.analysis_buffer[k])
                            with col_trash:
                                if st.button("🗑️", key=f"del_{k}"):
                                    del st.session_state.analysis_buffer[k]
                                    save_state(); st.rerun()
                curent_full_res = "\n---\n".join([st.session_state.analysis_buffer[k] for k in sorted_keys])

            if curent_full_res:
                col_w1, col_w2 = st.columns(2)
                with col_w1:
                    if not st.session_state.get(f"word_ready_pandoc_{is_live_preview}"):
                        if st.button("⚡ 완벽 변환 Word 준비 (클릭)" + (" (현재까지)" if is_live_preview else ""), use_container_width=True, key=f"word_export_{is_live_preview}"):
                            try:
                                with st.spinner("Word 파일 생성 중... 잠시만 기다려주세요."):
                                    import pypandoc, tempfile, os, shutil
                                    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
                                        tmp_path = tmp.name
                                    pypandoc.convert_text(curent_full_res, 'docx', format='markdown', outputfile=tmp_path)
                                    out_file = os.path.join(os.getcwd(), "AI_Analysis_Result.docx")
                                    shutil.copy(tmp_path, out_file)
                                    os.remove(tmp_path)
                                    st.session_state[f"word_ready_pandoc_{is_live_preview}"] = True
                                    st.rerun()
                            except Exception as e:
                                st.error(f"Pandoc 변환 실패: {e}")
                    else:
                        out_file = os.path.join(os.getcwd(), "AI_Analysis_Result.docx")
                        try:
                            with open(out_file, "rb") as f:
                                st.download_button("📥 완벽 변환 Word 다운로드", f, file_name="AI_Analysis_Result.docx", use_container_width=True, key=f"dl_pandoc_{is_live_preview}")
                        except: pass
                        if st.button("🔄 파일 다시 생성하기", use_container_width=True, key=f"regen_{is_live_preview}"):
                            st.session_state[f"word_ready_pandoc_{is_live_preview}"] = False
                            st.rerun()
                            
                with col_w2:
                    if not st.session_state.get(f"word_ready_txt_{is_live_preview}"):
                        if st.button("⚡ 텍스트 전용 Word 준비 (클릭)" + (" (현재까지)" if is_live_preview else ""), use_container_width=True, key=f"word_export_txt_{is_live_preview}"):
                            try:
                                with st.spinner("Word 파일 생성 중... 잠시만 기다려주세요."):
                                    import os
                                    from docx import Document
                                    out_file = os.path.join(os.getcwd(), "AI_Result_Text.docx")
                                    doc = Document()
                                    for line in curent_full_res.split('\n'): doc.add_paragraph(line)
                                    doc.save(out_file)
                                    st.session_state[f"word_ready_txt_{is_live_preview}"] = True
                                    st.rerun()
                            except Exception as e:
                                st.error(f"Word 변환 실패: {e}")
                    else:
                        out_file = os.path.join(os.getcwd(), "AI_Result_Text.docx")
                        try:
                            with open(out_file, "rb") as f:
                                st.download_button("📥 텍스트 전용 Word 다운로드", f, file_name="AI_Result_Text.docx", use_container_width=True, key=f"dl_txt_{is_live_preview}")
                        except: pass
                        if st.button("🔄 텍스트 파일 다시 생성하기", use_container_width=True, key=f"regen_txt_{is_live_preview}"):
                            st.session_state[f"word_ready_txt_{is_live_preview}"] = False
                            st.rerun()

            st.divider()
            st.subheader("📂 데이터베이스(Notion)에 즉시 저장")
            db_col1, db_col2 = st.columns([3, 1])
            with db_col1: target_sub = st.selectbox("저장할 과목 선택", list(st.session_state.notion_db.keys()), key=f"smart_db_subject_{is_live_preview}")
            with db_col2:
                if st.button("📥 DB에 추가", use_container_width=True, key=f"db_add_{is_live_preview}"):
                    st.session_state.notion_db[target_sub] += "\n---\n" + curent_full_res
                    save_state(); st.success("DB 저장 완료!")

        if not is_live_preview and st.button("🗑️ 모든 분석 결과 초기화 (새로 시작)", use_container_width=True):
            st.session_state.analysis_buffer = {}
            st.session_state.smart_analysis_result = None
            save_state(); st.rerun()

    # --- 무중단 엔진 코어 ---
    if st.session_state.get("smart_analysis_active") and st.session_state.get("smart_page_map"):
        if st.button("🛑 분석 중단 및 지금까지 결과 보기", key="stop_analysis_engine", use_container_width=True):
            st.session_state.smart_analysis_active = False
            if not st.session_state.analysis_buffer:
                st.warning("분석을 강제 중단했습니다. (아직 첫 번째 문항이 완료되지 않아 저장된 결과가 없습니다. 결과를 보려면 최소 1개 항목이 완료될 때까지 기다려 주세요.)")
            else:
                st.warning("분석을 강제 중단했습니다. 화면 아래에서 지금까지의 결과를 확인하세요.")
            st.rerun()
            
        page_map = st.session_state.smart_page_map
        buffer = st.session_state.analysis_buffer
        
        # [남은 시간 예측 로직: 대문항/소문항 갯수 기반 초정밀 동적 타이머]
        total_pages = len(page_map)
        
        # 1. 스캔 상태 파악
        scanned_structs = st.session_state.get("smart_q_structures", {})
        scanned_pages_count = len(scanned_structs)
        unscanned_pages = max(0, total_pages - scanned_pages_count)
        
        # 2. 스캔된 페이지의 정확한 총 소문항 갯수 계산
        total_scanned_subs = 0
        for p_key, struct in scanned_structs.items():
            for q in struct.get("questions", []):
                subs = q.get("subs", [])
                total_scanned_subs += len(subs) if subs else 1
                
        # 3. 완료된 소문항 갯수
        completed_subs = len([k for k in buffer.keys() if "_q_" in k])
        remaining_scanned_subs = max(0, total_scanned_subs - completed_subs)
        
        # 4. 소문항 기반 정밀 잔여 시간 계산 (문항당 45초 + 미스캔 페이지당 180초)
        target_remaining_sec = (remaining_scanned_subs * 45) + (unscanned_pages * 180)
        
        import time
        if "smart_timer_state" not in st.session_state:
            st.session_state.smart_timer_state = {"target_end_time": time.time() + target_remaining_sec}
            
        # 스캔 직후 문항이 폭증하거나, 예측 시간 차이가 40초 이상 나면 동적 타이머 보정
        current_target = st.session_state.smart_timer_state["target_end_time"]
        if abs(target_remaining_sec - (current_target - time.time())) > 40:
            st.session_state.smart_timer_state["target_end_time"] = time.time() + target_remaining_sec
            
        remaining_sec = max(5, int(st.session_state.smart_timer_state["target_end_time"] - time.time()))
        
        # 5. 진행률 프로그레스 바 계산 (발견된 문항 기준)
        progress_ratio = min(1.0, completed_subs / total_scanned_subs if total_scanned_subs > 0 else 0.0)
        st.progress(progress_ratio)
        
        status_text = f"현재 발견된 문항: 총 {total_scanned_subs}개 중 {completed_subs}개 완료"
        if unscanned_pages > 0:
            status_text += f" (아직 읽지 않은 미탐색 페이지 {unscanned_pages}장 대기 중)"
        else:
            status_text += " (모든 페이지 탐색 완료)"

        html_ticker = f"""
        <style>body {{ background: transparent !important; margin: 0; padding: 0; }}</style>
        <div style="font-family: 'Inter', sans-serif; font-size: 0.85rem; color: #555; display: flex; align-items: center;">
            📊 진행 상태: {status_text} | 예상 남은 시간: 약 <span id="time-ticker" style="font-weight: bold; margin-left: 4px; color: #1d4ed8;">계산 중...</span>
        </div>
        <script>
            let timeLeft = {remaining_sec};
            const display = document.getElementById('time-ticker');
            function render() {{
                if (timeLeft <= 0) {{ display.innerText = "답변 생성 마무리 중 (AI 초정밀 렌더링 대기)..."; return; }}
                const m = Math.floor(timeLeft / 60);
                const s = timeLeft % 60;
                display.innerText = m > 0 ? m + "분 " + s + "초" : s + "초";
            }}
            render();
            setInterval(() => {{
                if (timeLeft > 0) {{ timeLeft--; render(); }}
            }}, 1000);
        </script>
        """
        components.html(html_ticker, height=30)

        # [고도화] 엔진이 도는 와중에도 이미 풀린 문제들을 바로 열어볼 수 있도록 실시간 렌더링 호출
        render_smart_analysis_results(is_live_preview=True)

        target_found = False
        for p_idx, (f_type, file_idx, p_in_file) in enumerate(page_map):
            page_base_key = f"smart_p_{p_idx}"
            
            try:
                import PIL.Image, io, fitz
                # 지속성 파일 경로에서 이미지 로드
                f_path = st.session_state.smart_file_paths[file_idx]
                cur_img = None
                cur_text = ""
                if f_type == 'txt':
                    with open(f_path, "r", encoding="utf-8") as tf:
                        cur_text = tf.read()
                elif f_type == 'pdf':
                    doc = fitz.open(f_path)
                    page = doc.load_page(p_in_file)
                    pix = page.get_pixmap(matrix=fitz.Matrix(3.0, 3.0))
                    cur_img = PIL.Image.open(io.BytesIO(pix.tobytes("png")))
                    doc.close()
                else: cur_img = PIL.Image.open(f_path)
                
                # [0단계: 페이지 완료 여부 사전 체크 (고속 패스)]
                if "smart_q_structures" not in st.session_state:
                    st.session_state.smart_q_structures = {}
                
                cached_structure = st.session_state.smart_q_structures.get(page_base_key)
                if cached_structure:
                    # 이미 스캔된 구조가 있다면, 모든 문항이 완료되었는지 즉시 확인
                    all_done = True
                    for q_item in cached_structure.get("questions", []):
                        q_id = q_item.get("id", "전체")
                        subs = q_item.get("subs", [])
                        target_subs = subs if subs else ["_none"]
                        for sub_id in target_subs:
                            if f"{page_base_key}_q_{q_id}_sub_{sub_id}" not in buffer:
                                all_done = False; break
                        if not all_done: break
                    if all_done: continue # 모든 문항 완료 시 즉시 다음 페이지로 점프 (1초 컷)

                # [1단계: 계층적 구조 스캔] 구조가 없거나 미완료된 경우에만 실행
                if not cached_structure:
                    target_q_str = st.session_state.get("smart_target_questions", "").strip()
                    if target_q_str:
                        target_instruction = f"사용자가 특정 문항({target_q_str})만 분석하길 원합니다. 문서에 해당 문항 번호가 존재한다면 오직 그 문항들만 추출하여 JSON 형식으로 답변하고, 해당하지 않는 문항들은 완벽하게 무시(제외)하십시오."
                    else:
                        target_instruction = "문서에 존재하는 모든 문항 번호를 무조건 전부 배열에 개별적으로 분리해서 넣어야 합니다."

                    scan_prompt = f"""제공된 자료(이미지 또는 텍스트)에서 독립적인 대문항 번호(예: 1, 2, 3...)와 그에 속한 소문항(예: a, b, c, i, ii...)을 스캔하여 반드시 다음 JSON 형식으로만 답변하세요. 다른 설명은 일절 추가하지 마십시오.
                    [JSON 예시]
                    {{"questions": [{{"id": "1", "subs": ["a", "b", "c"]}}, {{"id": "2", "subs": []}}]}}
                    **[초비상 절대 엄수 사항]** 소문항이 5a부터 5p까지 16개가 있든, a부터 z까지 무수히 많든, **단 하나의 문항이나 소문항도 절대로 생략하거나 축약(건너뛰기)하지 마십시오!** 대문항과 하위 소문항을 완벽히 그룹화하세요. 번호가 도저히 없으면 {{"questions": [{{"id": "전체", "subs": []}}]}} 로 답하세요.
                    **[타겟 필터링 지시]** {target_instruction}
                    {"[분석할 텍스트]: " + cur_text if cur_text else ""}"""

                    scan_res_raw = robust_generate_content(scan_prompt, images=[cur_img] if cur_img else None)
                    if not scan_res_raw:
                        raise Exception("QUOTA_EXHAUSTED")
                    try:
                        import json
                        json_match = re.search(r'\{.*\}', scan_res_raw, re.DOTALL)
                        q_structure = json.loads(json_match.group()) if json_match else {"questions": [{"id": "전체", "subs": []}]}
                    except:
                        q_structure = {"questions": [{"id": "전체", "subs": []}]}
                    
                    # 구조 캐싱
                    st.session_state.smart_q_structures[page_base_key] = q_structure
                else:
                    q_structure = cached_structure

                # [2단계: 계층적 분석 루프 실행]
                page_results_found = False
                actual_page_num = p_in_file + 1 if f_type == 'pdf' else p_idx + 1
                with st.status(f"🚀 [원본 {actual_page_num}페이지] 계층적 정밀 분석 중...", expanded=True) as status:
                    for q_item in q_structure.get("questions", []):
                        q_id = q_item.get("id", "전체")
                        subs = q_item.get("subs", [])

                        target_subs = subs if subs else ["_none"]
                        for sub_id in target_subs:
                            sub_suffix = f"({sub_id})" if sub_id != "_none" else ""
                            q_key = f"{page_base_key}_q_{q_id}_sub_{sub_id}"
                            if q_key in buffer: continue

                            st.write(f"⚡ {q_id}번{sub_suffix} 문항 정밀 분석 중...")
                            if f_type == 'txt' and cur_text:
                                preview = cur_text[:100].replace('\n', ' ')
                                st.caption(f"📝 **현재 분석 대상 내용:** '{preview}...'")
                            elif f_path:
                                import os
                                st.caption(f"📄 **현재 분석 대상 파일:** {os.path.basename(f_path)} (실제 위치: {actual_page_num}페이지)")


                            # [지능형 자동 답안지/참조 탐색 시스템]
                            import os, glob
                            auto_ref_data = ""
                            try:
                                potential_refs = glob.glob("*.pdf") + glob.glob("temp_uploads/*.pdf")
                                curent_fname = st.session_state.smart_file_paths[file_idx]
                                for ref_path in potential_refs:
                                    if os.path.abspath(ref_path) == os.path.abspath(curent_fname): continue
                                    if any(kw in ref_path for kw in ["답", "Answer", "Solution", "정답", "해설"]):
                                        import fitz
                                        doc_ref = fitz.open(ref_path)
                                        auto_ref_data += f"\n--- [{os.path.basename(ref_path)} 참조 정답지] ---\n"
                                        auto_ref_data += "".join([p.get_text() for p in doc_ref])[:5000]
                                        doc_ref.close()
                            except: pass
                            
                            global_db = load_local_academic_db()
                            if global_db:
                                auto_ref_data += f"\n[글로벌 학술 데이터베이스 (수업 전체 참고)]\n{global_db}\n"
                            # [최종 고품질 프롬프트 생성]
                            target_label = f"{q_id}번 대문항" if sub_id == "_none" else f"{q_id}번 대문항의 소문항 ({sub_id})"
                            solve_prompt = fr"""당신은 세계 최고의 물리/화학 석박사급 학술 분석 전문가입니다. (Task ID: {time.time()})
                **[미션: 제공된 참조 데이터 및 당신의 지능을 결합한 무결점 {target_label} 풀이]**
                제공된 자료에서 **[{target_label}]**을 완벽하게 분석하십시오.
                {"[분석할 텍스트 대상]: " + cur_text if cur_text else ""}

                [참조 데이터]
                {auto_ref_data if auto_ref_data else "직접 지적 분석 수행"}

                [절대 엄수 가이드라인]
                1. **순차적 논리**: {target_label}에만 집중하여 가장 정밀한 정답과 해설을 도출하십시오.
                2. **참조 데이터 기준 압도적 확장**: 제공된 참조 정답지의 정답(수치, 결론)은 100% 일치시키되, 해설의 깊이와 분량, 배경지식, 시각적 시뮬레이션 데이터는 참조 정답지를 아득히 초월할 정도로 훨씬 더 상세하고 방대하게 추가 작성하십시오.
                3. **외부 링크 절대 금지**: 모든 외부 URL을 배제하십시오.
                4. **수식 렌더링 에러(빨간 글씨) 원천 차단 규칙**: 
                   - 수학/화학 수식은 `$$ ... $$` 기호로 감싸되, **수식 내부에 한글을 절대 포함하지 마십시오.** (`\text{{한글}}`, `\fbox{{한글}}` 사용 금지. 한글 설명은 수식 밖 일반 텍스트로 빼세요.)
                   - **절대 금지**: `\begin{{tikzpicture}}`, `\begin{{gathered}}`, 복잡한 `\begin{{array}}` 등 웹에서 호환되지 않는 고급 LaTeX 패키지는 절대 사용하지 마십시오.
                5. **시각 자료 완벽 복원 및 강제 렌더링 (Visual Rendering)**: 풀이 내용이나 대상에 **3D 격자, 분자 구조식, 양자역학 그래프, 루이스 전자점식, 분자 오비탈 시각화, 기본 도식 및 기하 도형, 2D 분자 구조 / 선구조식, 표, 그래프** 중 하나라도 포함된다면 절대 생략하지 마십시오!
                   **[초비상 절대 엄수: 양자역학 파동함수 및 각종 그래프 크기 대폭 확대]** 양자역학의 파동함수(Wavefunction), 1차원 상자 속 입자(Particle in a box) 모형, 파동/함수 그래프 등을 ASCII 아트로 그릴 때는 **상하좌우 크기를 기존보다 최소 2~3배 이상 대폭 확장하여 매우 큼직하고 시원하게(Large-scale) 그리십시오.** 사용자가 진폭과 노드의 위치를 직관적으로 바로 파악할 수 있도록 x축, y축의 간격을 넓게 잡아야 하며, 너무 작고 옹졸하게 그리는 것은 절대 금지됩니다.
                   **[초비상 절대 엄수: 분자 구조 및 결합 표시 방식]**
                   - 2D 분자 구조, 선구조식, 루이스 전자점식, 화학 결합을 그릴 때 **절대로 Mermaid graph(노드와 엣지를 연결하는 방식, 예: A---B, subgraph 등)를 사용하지 마십시오.** 화학 구조를 Mermaid 네트워크 그래프로 그리는 것은 절대 금지됩니다.
                   - 대신 반드시 **정교한 텍스트 기반 ASCII 아트**를 사용하여 원자들의 실제 2D 배치와 결합(단일, 이중 결합 등), 그리고 비공유 전자쌍을 시각적으로 정확하게 직관적으로 그려내야 합니다.
                   - **[ASCII 선구조식 예시]**
                        Cl  H
                         \ /
                          C ㅡ C ㅡ H
                         / \
                        F   H
                   - **[ASCII 루이스 전자점식 예시 (비공유 전자쌍 표시 필수)]**
                         ..
                       : O :
                         ||
                       : O :
                         ..
                   - 표, 데이터 플롯 등 화학 구조가 아닌 일반 도표의 경우에만 마크다운 표나 필요한 경우 Mermaid 등을 적절히 사용하십시오.
                6. **출력 구조 및 문제 원본 그림 복원 필수**: 
                   - **[질문 전사 및 원본 그림 복원]**: 대문항이나 소문항의 원본 문제 자체에 **3D 격자, 분자 구조식, 양자역학 그래프, 루이스 전자점식, 분자 오비탈 시각화, 기본 도식 및 기하 도형, 2D 분자 구조 / 선구조식, 표, 그래프 등 그림이 포함되어 있다면, 풀이를 시작하기 전에 반드시 가장 먼저 100% 똑같이 그려서 복원해야 합니다.** (위에서 제시한 ASCII 아트 또는 마크다운 표 등을 적극 활용하여 질문 내용과 그림을 함께 표시하세요.)
                   - **[심층 학술 해설(LaTeX)]**: 수식을 포함한 매우 상세한 풀이 과정.
                   - **[시각화 가이드 및 풀이용 그림]**: 풀이 과정에서 추가로 필요한 도식, 반응 메커니즘, 그래프, 분자 구조 변화 등이 있다면 여기서 **꼭 추가로 그려주세요.**
                   - **[최종 정답]**: 명확한 결론 도출.
                7. **초고도화 분량 극대화 및 반복 버그 절대 금지**: 모든 분석은 석/박사 학위 논문 수준으로 얻을 수 있는 모든 물리/화학적 데이터, 개념, 배경 지식을 끌어모아 **매우 길고 깊이 있게(최소 3000자 이상) 해설**하십시오. 단, **동일한 문장이나 단어가 비정상적으로 반복되는 기계적 붕괴(Model Collapse) 현상이 발생하면 즉시 실패 처리되므로**, 반드시 창의적이고 다양한 고급 학술 어휘를 사용하여 풍부하고 생동감 있게 작성하십시오.
                8. **시각적/구조적 고도화**: 단순 텍스트 나열을 피하고, 다층적 구조(서론-본론-결론, 단계별 유도 과정) 및 고급 시각적 포맷(마크다운 표, 수식 배열, 다이어그램)을 복합적으로 사용하여 **가장 압도적이고 정밀한 퀄리티의 풀이 리포트**를 완성하십시오.
                """
                            res = robust_generate_content(solve_prompt, images=[cur_img] if cur_img else None, use_grounding=False)
                            if not res:
                                raise Exception("QUOTA_EXHAUSTED")
                            
                            if res:
                                buffer[q_key] = res
                                page_results_found = True
                                st.toast(f"✅ {q_id}{sub_suffix} 완료")
                                save_state() # 소문항 단위 실시간 저장
                                st.rerun() # 소문항 풀이 즉시 화면에 반영 (Live 업데이트)

                    status.update(label=f"✅ {p_idx+1}페이지 모든 문항 완료", state="complete", expanded=False)

                target_found = True
                break
            except Exception as e:
                err_msg = str(e)
                if "QUOTA_EXHAUSTED" in err_msg:
                    import time
                    import re
                    match = re.search(r'QUOTA_EXHAUSTED:(\d+\.?\d*)', err_msg)
                    wait_time = int(float(match.group(1))) if match else 65
                    
                    msg = st.empty()
                    for i in range(wait_time, 0, -1):
                        msg.warning(f"🚫 구글 API 한도 도달. 누락 방지를 위해 {i}초간 강제 휴식 후, 스스로 앞선 문항들을 재점검하고 자동 진격합니다...")
                        time.sleep(1)
                    msg.empty()
                    st.toast("🔄 휴식 완료! 누락된 문항 점검 및 다음 문제 풀이를 자동으로 재개합니다 🚀")
                    st.rerun()
                
                # [생명력 유지 로직] 오류 발생 시 전체 시스템을 죽이지 않고 해당 페이지만 영구 스킵 처리하여 무한 루프 탈출
                st.error(f"⚠️ {p_idx+1}페이지 처리 중 치명적 오류 발생 (해당 페이지 스킵 후 다음으로 강제 진행): {e}")
                # 다음 새로고침 시 이 페이지를 완료된 것으로 강제 인식시켜 무한 루프 및 멈춤 방지
                buffer[f"{page_base_key}_q_Error_sub__none"] = f"⚠️ {p_idx+1}페이지는 심각한 구조적 오류 또는 반복 버그로 인해 스킵되었습니다.\n\n에러 내용: {e}"
                st.session_state.smart_q_structures[page_base_key] = {"questions": [{"id": "Error", "subs": []}]}
                save_state()
                continue # 전체 프로세스를 정지(st.stop)하지 않고 즉시 다음 페이지로 넘어가서 생명력 유지
        if not target_found:
            st.session_state.smart_analysis_active = False
            st.success("🎊 모든 페이지(80+p) 분석 완료!")
            st.rerun()

    # --- [상시 결과 표시 및 관리 섹션 (엔진 종료 후)] ---
    if not st.session_state.get("smart_analysis_active"):
        render_smart_analysis_results(is_live_preview=False)

    st.markdown("---")
    st.markdown("### 🔄 전체 생성 모드 일괄 적용")
    def set_global_mode():
        gmode = st.session_state.global_mode_widget
        target = "자동 (프리셋 선택)" if "자동" in gmode else "수동 (직접 입력)"
        if gmode != "기본 (개별 설정)":
            st.session_state.cell_mode_radio = target
            st.session_state.graph_mode_radio = target
            st.session_state.orbital_mode_radio = target
            st.session_state.lewis_mode_radio = target
            st.session_state.shape_mode_radio = target
            st.session_state.skeletal_mode_radio = target
            if target == "수동 (직접 입력)":
                st.session_state.cell_text_key_counter += 1
                st.session_state.graph_text_key_counter += 1
                st.session_state.orbital_text_key_counter += 1
                st.session_state.lewis_text_key_counter += 1
                st.session_state.shape_text_key_counter += 1
                st.session_state.skeletal_text_key_counter += 1

    st.radio("모든 그리기 도구의 모드를 일괄 변경:", ["기본 (개별 설정)", "전체 자동 모드로 통일", "전체 수동 모드로 통일"], horizontal=True, key="global_mode_widget", on_change=set_global_mode)
    st.markdown("---")
    
    st.markdown("### 📸 AI 비전 자동 역설계 (Vision-to-Plot)")
    st.markdown("이미지나 PDF를 업로드하면, AI가 형태를 인식하여 가장 적절한 화학/물리 그래픽으로 자동 복원하고 워드 파일에 추가합니다.")
    
    if "vision_history" not in st.session_state:
        st.session_state.vision_history = []
        
    vision_upload = st.file_uploader("이미지 또는 PDF 파일 업로드", type=["png", "jpg", "jpeg", "pdf"], key="vision_uploader_unified")
    
    col1, col2 = st.columns([1, 1])
    with col1:
        split_mode = st.radio("✂️ 원본 수동 분할 옵션", ["AI 자동 판단", "가로 균등 분할", "가로+세로 격자 분할", "원본 그대로 보존"], index=0, horizontal=False)
    with col2:
        if split_mode == "가로 균등 분할":
            split_cols = st.number_input("몇 등분 할까요?", min_value=2, max_value=10, value=3)
            split_rows = 1
        elif split_mode == "가로+세로 격자 분할":
            subcol1, subcol2 = st.columns(2)
            with subcol1:
                split_cols = st.number_input("가로 등분", min_value=1, max_value=10, value=4)
            with subcol2:
                split_rows = st.number_input("세로 등분", min_value=1, max_value=10, value=4)
        else:
            split_cols = 1
            split_rows = 1
            
    # 기존 작업 내역 렌더링 (새로고침/페이지 이동 시 유지용)
    if st.session_state.vision_history:
        with st.expander(f"🗂️ 기존 변환 내역 보기 ({len(st.session_state.vision_history)}개)", expanded=False):
            for idx, item in enumerate(reversed(st.session_state.vision_history)):
                with st.container(border=True):
                    st.success(f"✅ 기존 변환 내역: {item['category']} ({item['reasoning']})")
                    st.image(item['image'])
                    
    if vision_upload and st.button("🚀 AI 분석 및 워드에 자동 복원", use_container_width=True):
        if not st.session_state.get("gemini_api_key"):
            st.error("Gemini API 키가 필요합니다. 왼쪽 사이드바 하단에 키를 입력해주세요.")
        else:
            with st.spinner("AI가 이미지를 분석하여 코드로 역설계 중입니다..."):
                try:
                    import google.generativeai as genai
                    import json
                    from PIL import Image
                    import io
                    import fitz
                    
                    genai.configure(api_key=st.session_state.gemini_api_key)
                    
                    if vision_upload.name.lower().endswith(".pdf"):
                        doc_pdf = fitz.open(stream=vision_upload.read(), filetype="pdf")
                        page = doc_pdf.load_page(0)
                        pix = page.get_pixmap(dpi=150)
                        img_bytes = pix.tobytes("png")
                        img = Image.open(io.BytesIO(img_bytes))
                    else:
                        img = Image.open(vision_upload)
                        
                    prompt = """
                    You are an expert scientific diagram analyzer. 
                    Look at this image and determine which of the following 6 categories it belongs to, and extract the necessary parameters to recreate it.
                    
                    Categories & Required Parameters:
                    1. "3D 격자 생성": parameters: {"cell_type": "Face-Centered (FCC)" | "Body-Centered (BCC)" | "Simple Cubic (SC)" | "HCP (Hexagonal)" | "NaCl (Rock Salt)", "lattice_size": 1, "atom_rad": 0.18}
                    2. "양자역학 그래프": parameters: {"graph_type": "1D Box 파동함수 (n=1,2,3)" | "2s 오비탈 확률 밀도"}
                    3. "루이스 전자점식": parameters: {"molecule": e.g., "O2", "N2", "H2O"}
                    4. "분자 오비탈 시각화": parameters: {"molecule_type": "에텐 (C2H4)" | "에타인 (C2H2)" | "물 (H2O)"}
                    5. "기본 도식 및 기하 도형": parameters: {"shape_type": "정사면체 (Tetrahedral)" | "옥타헤드론 (Octahedral)" | "삼각쌍뿔 (Trigonal Bipyramidal)"}
                    6. "2D 분자 구조 / 선구조식": parameters: {"molecule": e.g., "벤젠", "아스피린"}
                    7. "표 / 기타 그래프": parameters: {"type": "Table", "description": "설명"}
                    
                    Return ONLY a valid JSON string (no markdown ticks) with keys:
                    {
                        "category": "<one of the 6 categories above>",
                        "reasoning": "<short explanation>",
                        "parameters": { ... },
                        "crop_boxes": [ [ymin, xmin, ymax, xmax], ... ], // MUST use this if the image contains MULTIPLE independent molecules or diagrams side-by-side (e.g., 3 different lattice drawings). Provide a list of bounding boxes for EACH object. Values MUST be normalized between 0 and 1000. If there is only one single diagram, leave it empty [].
                        "has_complex_annotations": false // MUST BE TRUE if the image contains any custom arrows, distance labels (e.g., r_0), descriptive text (e.g., 'Nearest'), or custom structures. If TRUE, the system will preserve the original image.
                    }
                    """
                    
                    if split_mode != "AI 자동 판단":
                        # 수동 분할 모드일 경우 AI 호출을 아예 생략 (시간 단축 및 오류 방지)
                        data = {
                            "category": "수동 분할 이미지",
                            "reasoning": f"사용자가 '{split_mode}' 옵션을 선택하여 원본 이미지를 수동 처리합니다.",
                            "parameters": {}
                        }
                    else:
                        res_text = robust_generate_content(prompt, images=[img], use_grounding=False)
                        if not res_text:
                            raise Exception("무료 할당량(Quota)이 완전히 소진되었거나 접근 가능한 모델이 없습니다.")
                        
                        import re
                        match = re.search(r'\{.*\}', res_text, re.DOTALL)
                        if match:
                            res_text = match.group(0)
                            
                        try:
                            data = json.loads(res_text)
                        except Exception:
                            # AI가 JSON 포맷을 어기거나 파싱 실패 시 원본 강제 보존
                            data = {
                                "category": "알 수 없음 (분석 실패)",
                                "reasoning": "AI가 이미지 형식을 분석하지 못해 원본을 보존합니다.",
                                "parameters": {},
                                "has_complex_annotations": True,
                                "crop_boxes": []
                            }
                        
                    st.success(f"✅ AI 판독 완료: {data['category']} ({data['reasoning']})")
                    
                    cat = data["category"]
                    params = data["parameters"]
                    
                    img_streams = []
                    errors = []
                    if split_mode == "가로 균등 분할":
                        data["has_complex_annotations"] = True
                        data["crop_boxes"] = []
                        step = 1000.0 / split_cols
                        for i in range(split_cols):
                            data["crop_boxes"].append([0, i*step, 1000.0, (i+1)*step])
                    elif split_mode == "가로+세로 격자 분할":
                        data["has_complex_annotations"] = True
                        data["crop_boxes"] = []
                        step_x = 1000.0 / split_cols
                        step_y = 1000.0 / split_rows
                        for j in range(split_rows):
                            for i in range(split_cols):
                                data["crop_boxes"].append([j*step_y, i*step_x, (j+1)*step_y, (i+1)*step_x])
                    elif split_mode == "원본 그대로 보존":
                        data["has_complex_annotations"] = True
                        data["crop_boxes"] = []
                        
                    crop_boxes = data.get("crop_boxes", [])
                    
                    is_complex = data.get("has_complex_annotations", False)
                    # "can_recreate_perfectly"가 남아있을 수 있으니 하위 호환성 체크
                    if not data.get("can_recreate_perfectly", True):
                        is_complex = True
                        
                    if is_complex:
                        # AI가 화살표나 복잡한 텍스트 기호가 있어서 파이썬 코드로 완벽 복원이 불가능하다고 판단한 경우 -> 원본 복사
                        img_stream = io.BytesIO()
                        img.save(img_stream, format='PNG')
                        img_streams.append(img_stream)
                        # 만약 복원이 불가능한데 여러 개체라면, 원본 분할 복사를 위해 아래 조건문을 타도록 수정
                        if len(crop_boxes) > 1:
                            img_streams = [] # 초기화하고 아래 분할 로직으로 넘어감
                            pass 
                        
                    if len(crop_boxes) > 1 and not img_streams:
                        # 여러 독립된 개체가 발견된 경우, 원본 이미지를 여러 개로 분할 복사
                        w, h = img.size
                        for box in crop_boxes:
                            if len(box) == 4:
                                ymin, xmin, ymax, xmax = box
                                left = max(0, (xmin / 1000.0) * w)
                                right = min(w, (xmax / 1000.0) * w)
                                top = max(0, (ymin / 1000.0) * h)
                                bottom = min(h, (ymax / 1000.0) * h)
                                
                                # 수동 분할은 정확히 자르고, AI 분할은 여백 5% 추가
                                if split_mode == "AI 자동 판단":
                                    pw, ph = right - left, bottom - top
                                    left = int(max(0, left - pw * 0.05))
                                    right = int(min(w, right + pw * 0.05))
                                    top = int(max(0, top - ph * 0.05))
                                    bottom = int(min(h, bottom + ph * 0.05))
                                else:
                                    left = int(left)
                                    right = int(right)
                                    top = int(top)
                                    bottom = int(bottom)
                                
                                # 좌우/상하 역전 방지
                                if right <= left: right = left + 1
                                if bottom <= top: bottom = top + 1
                                
                                cropped = img.crop((left, top, right, bottom))
                                stream = io.BytesIO()
                                cropped.save(stream, format='PNG')
                                img_streams.append(stream)
                    if not img_streams:
                        img_stream = None
                        if "3D 격자" in cat:
                            img_stream, errors = draw_unit_cell(params.get("cell_type", "Face-Centered (FCC)"), params.get("lattice_size", 1), params.get("atom_rad", 0.18))
                        elif "양자역학" in cat:
                            img_stream, errors = draw_quantum_graph(params.get("graph_type", "1D Box 파동함수 (n=1,2,3)"))
                        elif "루이스" in cat:
                            img_stream, errors = draw_lewis_structure(params.get("molecule", "O2"))
                        elif "오비탈" in cat:
                            img_stream, errors = draw_orbital_diagram(params.get("molecule_type", "에텐 (C2H4)"))
                        elif "기하 도형" in cat or "기본 도식" in cat:
                            img_stream, errors = draw_schematic(params.get("shape_type", "정사면체 (Tetrahedral)"))
                        elif "2D 분자" in cat or "선구조식" in cat:
                            molecule_name = str(params.get("molecule", ""))
                            supported = ["Butane (뷰테인)", "Hexane (헥세인)", "Cyclohexane (사이클로헥세인)", "Benzene (벤젠)", "Acetone (아세톤)", "Acetic Acid (아세트산)"]
                            matched = next((s for s in supported if molecule_name.lower() in s.lower() or s.lower() in molecule_name.lower()), None)
                            
                            if matched:
                                img_stream, errors = draw_skeletal_structure(matched)
                            else:
                                img_stream = io.BytesIO()
                                img.save(img_stream, format='PNG')
                        elif "표" in cat or "그래프" in cat or "기타" in cat:
                            img_stream = io.BytesIO()
                            img.save(img_stream, format='PNG')
                        else:
                            img_stream = io.BytesIO()
                            img.save(img_stream, format='PNG')
                            
                        if img_stream:
                            img_streams.append(img_stream)
                            
                    if errors:
                        for e in errors: st.error(e)
                    
                    if img_streams:
                        from docx.shared import Inches
                        for stream in img_streams:
                            stream.seek(0)
                            st.image(stream)
                            
                            stream.seek(0)
                            st.session_state.vision_history.append({
                                "category": data['category'] + (" (분할 크롭)" if len(img_streams)>1 else ""),
                                "reasoning": data['reasoning'],
                                "image": stream.getvalue()
                            })
                            
                            stream.seek(0)
                            st.session_state.word_doc.add_picture(stream, width=Inches(4.0))
                            
                        st.success(f"📝 MS Word 파일에 {len(img_streams)}개의 이미지가 성공적으로 추가되었습니다! 아래 버튼을 눌러 바로 다운로드하세요.")
                        
                        import time
                        doc_stream = io.BytesIO()
                        st.session_state.word_doc.save(doc_stream)
                        st.session_state.vision_docx_bytes = doc_stream.getvalue()
                        

                except Exception as e:
                    st.error(f"처리 중 오류 발생: {str(e)}")
                    
    if st.session_state.get("vision_docx_bytes"):
        render_download_and_open_buttons("AI 비전 전체 작업 내용", st.session_state.vision_docx_bytes, "SNU_Chem_Report_AI_Vision.docx", "vision_dl_persistent")

    st.markdown("---")

    col1, col2 = st.columns([1, 1])
    with col1:
        st.subheader("1. 3D 격자 생성")
        cell_options = ["Simple Cubic (SC)", "Body-Centered (BCC)", "Face-Centered (FCC)", "HCP (Hexagonal)", "NaCl (Rock Salt)", "CsCl (BCC)", "CaF2 (Fluorite)", "Rhombohedral"]

        def on_cell_choice_change():
            if st.session_state.cell_mode_radio == "수동 (직접 입력)":
                st.session_state.custom_coords_val = get_preset_custom_data("cell", st.session_state.get("cell_choice_select", ""))
                st.session_state.cell_text_key_counter += 1

        def sync_cell_mode():
            if st.session_state.cell_mode_radio == "수동 (직접 입력)":
                st.session_state.custom_coords_val = get_preset_custom_data("cell", st.session_state.get("cell_choice_select", ""))
                st.session_state.cell_text_key_counter += 1

        cell_mode = st.radio("생성 모드:", ["자동 (프리셋 선택)", "수동 (직접 입력)"], horizontal=True, key="cell_mode_radio", on_change=sync_cell_mode)

        label = "3D 격자 선택:" if cell_mode == "자동 (프리셋 선택)" else "시작 템플릿 불러오기 (변경 시 아래 코드가 덮어씌워집니다):"
        cell_choice_raw = st.selectbox(label, cell_options, key="cell_choice_select", on_change=on_cell_choice_change)

        if cell_mode == "자동 (프리셋 선택)":
            cell_choice = cell_choice_raw
        else:
            cell_choice = "직접 좌표 입력 (Custom)"
        lattice_size = st.slider("격자 반복 횟수 (Size):", 1, 3, 1, help="단위 세포를 몇 번 반복해서 쌓을지 결정합니다.")
        atom_rad = st.slider("원자 크기 (Atom Size):", 0.05, 0.5, 0.18, 0.01, help="원자의 반지름을 조절합니다.")

        custom_coords = ""
        if cell_choice == "직접 좌표 입력 (Custom)":
            if "custom_coords_val" not in st.session_state:
                st.session_state.custom_coords_val = "0,0,0,blue\n1,1,1,red"
            custom_coords = st.text_area("좌표 입력 (x,y,z,색상):", value=st.session_state.custom_coords_val, key=f"custom_coords_val_{st.session_state.cell_text_key_counter}", help="줄바꿈으로 여러 좌표를 입력할 수 있습니다. (Ctrl+Enter로 적용)")
            st.session_state.custom_coords_val = custom_coords
        if st.button("3D 큐브 그리기"):
            img_stream, errors = draw_unit_cell(cell_choice, lattice_size, atom_rad, custom_coords)
            if errors:
                for er in errors: st.error(er)
            img_stream.seek(0)
            st.image(img_stream)
            img_stream.seek(0)
            st.session_state.word_doc.add_picture(img_stream, width=Inches(4.0))

        up_3d = st.file_uploader("📂 3D 격자 이미지 수동 업로드", type=["png", "jpg", "jpeg"], key="up_3d")
        if up_3d and st.button("워드에 수동 추가 (3D)"):
            up_3d.seek(0)
            st.image(up_3d)
            up_3d.seek(0)
            st.session_state.word_doc.add_picture(up_3d, width=Inches(4.0))
            st.success("워드에 추가되었습니다!")
        st.markdown("---")
        st.subheader("2. 양자역학 그래프")
        graph_options = ["1D Box 파동함수 (n=1,2,3)", "2s 오비탈 확률 밀도"]
        def on_graph_choice_change():
            if st.session_state.graph_mode_radio == "수동 (직접 입력)":
                st.session_state.custom_expr_val = get_preset_custom_data("graph", st.session_state.get("graph_choice_select", ""))
                st.session_state.graph_text_key_counter += 1

        def sync_graph_mode():
            if st.session_state.graph_mode_radio == "수동 (직접 입력)":
                st.session_state.custom_expr_val = get_preset_custom_data("graph", st.session_state.get("graph_choice_select", ""))
                st.session_state.graph_text_key_counter += 1

        graph_mode = st.radio("생성 모드:", ["자동 (프리셋 선택)", "수동 (직접 입력)"], horizontal=True, key="graph_mode_radio", on_change=sync_graph_mode)

        label = "양자역학 그래프:" if graph_mode == "자동 (프리셋 선택)" else "시작 템플릿 불러오기 (변경 시 아래 코드가 덮어씌워집니다):"
        graph_choice_raw = st.selectbox(label, graph_options, key="graph_choice_select", on_change=on_graph_choice_change)

        if graph_mode == "자동 (프리셋 선택)":
            graph_choice = graph_choice_raw
        else:
            graph_choice = "직접 수식 입력 (Custom)"

        custom_expr = "sin(x)"
        x_min, x_max = 0.0, 10.0

        if graph_choice == "직접 수식 입력 (Custom)":
            if "custom_expr_val" not in st.session_state:
                st.session_state.custom_expr_val = "sin(x)*exp(-x/5)"
            custom_expr = st.text_input("원하는 수식을 입력하세요 (x에 대한 함수):", value=st.session_state.custom_expr_val, key=f"custom_expr_val_{st.session_state.graph_text_key_counter}", help="예: sin(x), x^2, exp(-x)")
            st.session_state.custom_expr_val = custom_expr
            col_x1, col_x2 = st.columns(2)
            x_min = col_x1.number_input("x 최솟값", value=0.0)
            x_max = col_x2.number_input("x 최댓값", value=15.0)

        if st.button("그래프 그리기"):
            img_stream = draw_quantum_graph(graph_choice, custom_expr, x_min, x_max)
            img_stream.seek(0)
            st.image(img_stream)
            img_stream.seek(0)
            st.session_state.word_doc.add_picture(img_stream, width=Inches(4.0))

        up_graph = st.file_uploader("📂 그래프 이미지 수동 업로드", type=["png", "jpg", "jpeg"], key="up_graph")
        if up_graph and st.button("워드에 수동 추가 (그래프)"):
            up_graph.seek(0)
            st.image(up_graph)
            up_graph.seek(0)
            st.session_state.word_doc.add_picture(up_graph, width=Inches(4.0))
            st.success("워드에 추가되었습니다!")

        st.markdown("---")
        st.subheader("6. 분자 오비탈 시각화 (3D)")
        orbital_options = ["NO3- (질산 이온)", "HNO3 (질산)", "C2H4 (에텐)"]
        def on_orbital_choice_change():
            if st.session_state.orbital_mode_radio == "수동 (직접 입력)":
                st.session_state.custom_orbital_data_val = get_preset_custom_data("orbital", st.session_state.get("orbital_choice_select", ""))
                st.session_state.orbital_text_key_counter += 1

        def sync_orbital_mode():
            if st.session_state.orbital_mode_radio == "수동 (직접 입력)":
                st.session_state.custom_orbital_data_val = get_preset_custom_data("orbital", st.session_state.get("orbital_choice_select", ""))
                st.session_state.orbital_text_key_counter += 1

        orbital_mode = st.radio("생성 모드:", ["자동 (프리셋 선택)", "수동 (직접 입력)", "이미지 직접 업로드"], horizontal=True, key="orbital_mode_radio", on_change=sync_orbital_mode)

        if orbital_mode == "이미지 직접 업로드":
            uploaded_orbital = st.file_uploader("분자 오비탈 이미지 파일 업로드", type=["png", "jpg", "jpeg"], key="orbital_upload")
            if st.button("업로드한 이미지 워드에 추가", key="orbital_upload_btn"):
                if uploaded_orbital is not None:
                    import io
                    img_stream = io.BytesIO(uploaded_orbital.read())
                    st.image(img_stream)
                    img_stream.seek(0)
                    from docx.shared import Inches
                    st.session_state.word_doc.add_picture(img_stream, width=Inches(4.0))
                    st.success("업로드한 이미지가 워드에 추가되었습니다!")
                else:
                    st.error("이미지를 업로드해주세요.")
        else:
            label = "분자 오비탈 선택:" if orbital_mode == "자동 (프리셋 선택)" else "시작 템플릿 불러오기 (변경 시 아래 코드가 덮어씌워집니다):"
            orbital_choice_raw = st.selectbox(label, orbital_options, key="orbital_choice_select", on_change=on_orbital_choice_change)
    
            if orbital_mode == "자동 (프리셋 선택)":
                orbital_choice = orbital_choice_raw
            else:
                orbital_choice = "직접 입력 (Custom)"
    
            custom_orbital_data = ""
            if orbital_choice == "직접 입력 (Custom)":
                if "custom_orbital_data_val" not in st.session_state:
                    st.session_state.custom_orbital_data_val = "ATOM, N, 0, 0, 0"
                custom_orbital_data = st.text_area("수동 입력 (형식: 명령어, 속성...)", value=st.session_state.custom_orbital_data_val, key=f"custom_orbital_data_val_{st.session_state.orbital_text_key_counter}", height=150, help="명령어: ATOM(이름,x,y,z), BOND(x1,y1,z1,x2,y2,z2), PZ(x,y,z,색상,라벨), SP2(x,y,z,각도,색상,라벨), S(x,y,z,색상,라벨)")
                st.session_state.custom_orbital_data_val = custom_orbital_data
    
            if st.button("오비탈 그리기"):
                img_stream, errors = draw_orbital_diagram(orbital_choice, custom_orbital_data)
                if errors:
                    for er in errors: st.error(er)
                img_stream.seek(0)
                st.image(img_stream)
                img_stream.seek(0)
                from docx.shared import Inches
                st.session_state.word_doc.add_picture(img_stream, width=Inches(4.0))
                st.success(f"{orbital_choice} 오비탈 그림이 추가되었습니다!")

    with col2:
        st.subheader("3. 분자 구조식")
        mol_mode = st.radio("생성 모드:", ["PubChem 자동 검색", "수동 (직접 입력)"], horizontal=True, key="mol_mode_radio")

        if mol_mode == "PubChem 자동 검색":
            st.markdown("💡 **자주 검색하는 분자 예시:**")
            st.caption("아래 이름을 복사해 입력창에 넣거나, 쉼표(,)로 연결해서 검색하세요.")
            st.code("aspirin, benzene, caffeine, glucose, ethanol, water, methane, sulfuric acid, penicillin, acetaminophen", language="text")
            mol_name = st.text_input("분자 이름 또는 식 입력:", help="예: aspirin, benzene, C4H8, H2O", key="mol_name_input")
            if st.button("구조식 그리기 (자동)"):
                if mol_name:
                    names = [n.strip() for n in mol_name.split(',')]
                    for name in names[:4]:
                        if not name: continue
                        img_stream = get_molecule_image(name)
                        if img_stream:
                            img_stream.seek(0)
                            st.image(img_stream, caption=name)
                            st.session_state.word_doc.add_paragraph(f"[{name}] 구조식")
                            img_stream.seek(0)
                            st.session_state.word_doc.add_picture(img_stream, width=Inches(3.0))
        else:
            if "custom_mol_data_val" not in st.session_state:
                st.session_state.custom_mol_data_val = "TEXT, $C_6H_{12}O_6$, 0, 1.5, 24, black\nLINE, 0, 1, 1, 0.5, 2, black"
            if st.button("구조식 그리기 (수동)"):
                img_stream, errors = draw_skeletal_structure("직접 입력 (Custom)", custom_mol_data)
                if errors:
                    for er in errors: st.error(er)
                img_stream.seek(0)
                st.image(img_stream)
                img_stream.seek(0)
                st.session_state.word_doc.add_picture(img_stream, width=Inches(3.0))

        st.markdown("---")
        up_mol = st.file_uploader("📂 분자 구조식 이미지 수동 업로드", type=["png", "jpg", "jpeg"], key="up_mol")
        if up_mol and st.button("워드에 수동 추가 (구조식)"):
            up_mol.seek(0)
            st.image(up_mol)
            up_mol.seek(0)
            st.session_state.word_doc.add_picture(up_mol, width=Inches(3.0))
            st.success("워드에 추가되었습니다!")

        st.markdown("---")
        st.subheader("4. 루이스 전자점식 (Lewis Structure)")
        lewis_options = ["H2O (물)", "CO2 (이산화탄소)", "NH3 (암모니아)", "CH4 (메테인)", "O2 (산소 분자)", "N2 (질소 분자)", "HCl (염화수소)", "HCN (사이안화수소)", "C2H4 (에텐)"]
        def on_lewis_choice_change():
            if st.session_state.lewis_mode_radio == "수동 (직접 입력)":
                st.session_state.custom_lewis_data_val = get_preset_custom_data("lewis", st.session_state.get("lewis_choice_select", ""))
                st.session_state.lewis_text_key_counter += 1

        def sync_lewis_mode():
            if st.session_state.lewis_mode_radio == "수동 (직접 입력)":
                st.session_state.custom_lewis_data_val = get_preset_custom_data("lewis", st.session_state.get("lewis_choice_select", ""))
                st.session_state.lewis_text_key_counter += 1

        lewis_mode = st.radio("생성 모드:", ["자동 (프리셋 선택)", "수동 (직접 입력)", "이미지 직접 업로드"], horizontal=True, key="lewis_mode_radio", on_change=sync_lewis_mode)

        if lewis_mode == "이미지 직접 업로드":
            uploaded_lewis = st.file_uploader("루이스 전자점식 이미지 파일 업로드", type=["png", "jpg", "jpeg"], key="lewis_upload")
            if st.button("업로드한 이미지 워드에 추가", key="lewis_upload_btn"):
                if uploaded_lewis is not None:
                    import io
                    img_stream = io.BytesIO(uploaded_lewis.read())
                    st.image(img_stream)
                    img_stream.seek(0)
                    from docx.shared import Inches
                    st.session_state.word_doc.add_picture(img_stream, width=Inches(3.0))
                    st.success("업로드한 이미지가 워드에 추가되었습니다!")
                else:
                    st.error("이미지를 업로드해주세요.")
        else:
            label = "분자 선택:" if lewis_mode == "자동 (프리셋 선택)" else "시작 템플릿 불러오기 (변경 시 아래 코드가 덮어씌워집니다):"
            lewis_choice_raw = st.selectbox(label, lewis_options, key="lewis_choice_select", on_change=on_lewis_choice_change)
    
            if lewis_mode == "자동 (프리셋 선택)":
                lewis_choice = lewis_choice_raw
            else:
                lewis_choice = "직접 입력 (Custom)"
    
            custom_lewis_data = ""
            if lewis_choice == "직접 입력 (Custom)":
                if "custom_lewis_data_val" not in st.session_state:
                    st.session_state.custom_lewis_data_val = "TEXT, O, 0.5, 0.5, 36, blue"
                custom_lewis_data = st.text_area("루이스 수동 입력:", value=st.session_state.custom_lewis_data_val, key=f"custom_lewis_data_val_{st.session_state.lewis_text_key_counter}", height=150, help="TEXT, LINE, DOTS(x,y,dx,dy) 명령어를 사용합니다.")
                st.session_state.custom_lewis_data_val = custom_lewis_data
    
            if st.button("루이스 구조식 그리기"):
                img_stream, errors = draw_lewis_structure(lewis_choice, custom_lewis_data)
                if errors:
                    for er in errors: st.error(er)
                img_stream.seek(0)
                st.image(img_stream)
                img_stream.seek(0)
                from docx.shared import Inches
                st.session_state.word_doc.add_picture(img_stream, width=Inches(3.0))
                st.success(f"{lewis_choice} 루이스 구조식이 추가되었습니다!")

        st.markdown("---")
        st.subheader("5. 기본 도식 및 기하 도형")
        shape_options = ["정사각형 (Square)", "원형 (Circle)", "다각형 (Polygon)", "정육면체 (Cube)"]
        def on_shape_choice_change():
            if st.session_state.shape_mode_radio == "수동 (직접 입력)":
                st.session_state.custom_shape_data_val = get_preset_custom_data("shape", st.session_state.get("shape_choice_select", ""))
                st.session_state.shape_text_key_counter += 1

        def sync_shape_mode():
            if st.session_state.shape_mode_radio == "수동 (직접 입력)":
                st.session_state.custom_shape_data_val = get_preset_custom_data("shape", st.session_state.get("shape_choice_select", ""))
                st.session_state.shape_text_key_counter += 1

        shape_mode = st.radio("생성 모드:", ["자동 (프리셋 선택)", "수동 (직접 입력)", "이미지 직접 업로드"], horizontal=True, key="shape_mode_radio", on_change=sync_shape_mode)

        if shape_mode == "이미지 직접 업로드":
            uploaded_shape = st.file_uploader("도식 및 기하 도형 이미지 파일 업로드", type=["png", "jpg", "jpeg"], key="shape_upload")
            if st.button("업로드한 이미지 워드에 추가", key="shape_upload_btn"):
                if uploaded_shape is not None:
                    import io
                    img_stream = io.BytesIO(uploaded_shape.read())
                    st.image(img_stream)
                    img_stream.seek(0)
                    from docx.shared import Inches
                    st.session_state.word_doc.add_picture(img_stream, width=Inches(3.0))
                    st.success("업로드한 이미지가 워드에 추가되었습니다!")
                else:
                    st.error("이미지를 업로드해주세요.")
        else:
            label = "도형 선택:" if shape_mode == "자동 (프리셋 선택)" else "시작 템플릿 불러오기 (변경 시 아래 코드가 덮어씌워집니다):"
            shape_choice_raw = st.selectbox(label, shape_options, key="shape_choice_select", on_change=on_shape_choice_change)
    
            if shape_mode == "자동 (프리셋 선택)":
                shape_choice = shape_choice_raw
            else:
                shape_choice = "직접 입력 (Custom)"
    
            custom_shape_data = ""
            shape_color = "#2E5BFF"
            kwargs = {}
    
            if shape_choice == "직접 입력 (Custom)":
                if "custom_shape_data_val" not in st.session_state:
                    st.session_state.custom_shape_data_val = "RECT, 0.1, 0.1, 0.8, 0.8, #2E5BFF"
                custom_shape_data = st.text_area("수동 입력 (형식: 명령어, 속성...):", value=st.session_state.custom_shape_data_val, key=f"custom_shape_data_val_{st.session_state.shape_text_key_counter}", height=150, help="명령어: RECT(x,y,w,h,색상), CIRCLE(x,y,r,색상), TEXT(내용,x,y,크기,색상), LINE(x1,y1,x2,y2,색상)")
                st.session_state.custom_shape_data_val = custom_shape_data
            else:
                shape_color = st.color_picker("색상 선택:", "#2E5BFF")
                if shape_choice == "정사각형 (Square)":
                    c1, c2 = st.columns(2)
                    kwargs['width'] = c1.slider("가로 길이", 0.1, 1.0, 0.8, 0.1)
                    kwargs['height'] = c2.slider("세로 길이", 0.1, 1.0, 0.8, 0.1)
                elif shape_choice == "원형 (Circle)":
                    kwargs['radius'] = st.slider("반지름 크기", 0.1, 0.5, 0.4, 0.05)
                elif shape_choice == "다각형 (Polygon)":
                    kwargs['n_sides'] = st.slider("꼭짓점 개수 (N각형)", 3, 12, 6, 1)
                elif shape_choice == "정육면체 (Cube)":
                    c1, c2, c3 = st.columns(3)
                    kwargs['length'] = c1.slider("가로 길이 (x)", 0.5, 3.0, 1.0, 0.1)
                    kwargs['width'] = c2.slider("세로 길이 (y)", 0.5, 3.0, 1.0, 0.1)
                    kwargs['height'] = c3.slider("높이 (z)", 0.5, 3.0, 1.0, 0.1)
    
            if st.button("도형 그리기"):
                img_stream, errors = draw_schematic(shape_choice, shape_color, custom_shape_data, **kwargs)
                if errors:
                    for er in errors: st.error(er)
                img_stream.seek(0)
                st.image(img_stream)
                img_stream.seek(0)
                from docx.shared import Inches
                st.session_state.word_doc.add_picture(img_stream, width=Inches(3.0))
                st.success("워드에 도형이 추가되었습니다!")

        st.markdown("---")
        st.subheader("7. 2D 분자 구조 / 선구조식 (수동 지원)")
        skeletal_options = ["Butane (뷰테인)", "Hexane (헥세인)", "Cyclohexane (사이클로헥세인)", "Benzene (벤젠)", "Acetone (아세톤)", "Acetic Acid (아세트산)"]
        def on_skeletal_choice_change():
            if st.session_state.skeletal_mode_radio == "수동 (직접 입력)":
                st.session_state.custom_skeletal_data_val = get_preset_custom_data("skeletal", st.session_state.get("skeletal_choice_select", ""))
                st.session_state.skeletal_text_key_counter += 1

        def sync_skeletal_mode():
            if st.session_state.skeletal_mode_radio == "수동 (직접 입력)":
                st.session_state.custom_skeletal_data_val = get_preset_custom_data("skeletal", st.session_state.get("skeletal_choice_select", ""))
                st.session_state.skeletal_text_key_counter += 1

        skeletal_mode = st.radio("생성 모드:", ["자동 (프리셋 선택)", "수동 (직접 입력)", "이미지 직접 업로드"], horizontal=True, key="skeletal_mode_radio", on_change=sync_skeletal_mode)

        if skeletal_mode == "이미지 직접 업로드":
            uploaded_skeletal = st.file_uploader("2D 분자 구조 / 선구조식 이미지 파일 업로드", type=["png", "jpg", "jpeg"], key="skeletal_upload")
            if st.button("업로드한 이미지 워드에 추가", key="skeletal_upload_btn"):
                if uploaded_skeletal is not None:
                    import io
                    img_stream = io.BytesIO(uploaded_skeletal.read())
                    st.image(img_stream)
                    img_stream.seek(0)
                    from docx.shared import Inches
                    st.session_state.word_doc.add_picture(img_stream, width=Inches(4.0))
                    st.success("업로드한 이미지가 워드에 추가되었습니다!")
                else:
                    st.error("이미지를 업로드해주세요.")
        else:
            label = "분자 선택:" if skeletal_mode == "자동 (프리셋 선택)" else "시작 템플릿 불러오기 (변경 시 아래 코드가 덮어씌워집니다):"
            skeletal_choice_raw = st.selectbox(label, skeletal_options, key="skeletal_choice_select", on_change=on_skeletal_choice_change)
    
            if skeletal_mode == "자동 (프리셋 선택)":
                skeletal_choice = skeletal_choice_raw
            else:
                skeletal_choice = "직접 입력 (Custom)"
    
            custom_skeletal_data = ""
            if skeletal_choice == "직접 입력 (Custom)":
                if "custom_skeletal_data_val" not in st.session_state:
                    st.session_state.custom_skeletal_data_val = "CANVAS, 8, 4"
                custom_skeletal_data = st.text_area(
                    "수동 입력 (형식: 명령어, 속성...)",
                    key="custom_skeletal_data_val",
                    height=350,
                    help="명령어: CANVAS(너비,높이), TEXT/LTEXT/RTEXT(텍스트,x,y,크기,색상), LINE(x1,y1,x2,y2,두께,색상), DLINE(이중결합x1,y1,x2,y2,두께,색상), ARROW(x1,y1,x2,y2,색상), DOTS(점1x,점1y,점2x,점2y,색상), BALL(x,y,반지름,색상)."
                )
    
            if st.button("분자 구조식 그리기"):
                img_stream, errors = draw_skeletal_structure(skeletal_choice, custom_skeletal_data)
                if errors:
                    for er in errors:
                        st.error(f"오류: {er}")
                img_stream.seek(0)
                st.image(img_stream)
                img_stream.seek(0)
                from docx.shared import Inches
                st.session_state.word_doc.add_picture(img_stream, width=Inches(4.0))
                st.info("💡 **Tip:** 수동 입력창에서 줄바꿈은 `Enter`를, 입력 완료 및 적용은 `Ctrl+Enter`를 사용하세요.")
                st.success(f"{skeletal_choice} 구조식이 추가되었습니다!")

    doc_stream = io.BytesIO()
    st.session_state.word_doc.save(doc_stream)
    
    if st.button("🚀 누적된 모든 도표/그림 워드로 변환하기", type="primary", use_container_width=True):
        doc_stream = io.BytesIO()
        st.session_state.word_doc.save(doc_stream)
        st.session_state.diagram_docx_bytes = doc_stream.getvalue()
        st.success("✅ 모든 도표가 워드 파일로 변환되었습니다! 아래에서 다운로드하세요.")

    if st.session_state.get("diagram_docx_bytes"):
        render_download_and_open_buttons("도표/그림 모음", st.session_state.diagram_docx_bytes, "Diagrams.docx", "diagrams")

    st.markdown("---")
    st.subheader("🎨 통합 수기 도표 보드 (그림판)")
    st.markdown("아이콘을 선택하고 캔버스에 직접 도표나 그림을 수기로 그려보세요. AI가 형태를 인식해 깔끔한 프로그래밍 이미지로 변환 후 워드에 추가합니다.")

    try:
        # streamlit-drawable-canvas + 최신 Streamlit(1.30+) 호환성 패치
        import streamlit.elements.image as st_image
        if not hasattr(st_image, 'image_to_url'):
            def patched_image_to_url(image, *args, **kwargs):
                import io
                import base64
                buf = io.BytesIO()
                image.save(buf, format="PNG")
                b64 = base64.b64encode(buf.getvalue()).decode("utf-8")
                return f"data:image/png;base64,{b64}"
            st_image.image_to_url = patched_image_to_url

        from streamlit_drawable_canvas import st_canvas
        canvas_supported = True
    except ImportError:
        canvas_supported = False
        st.warning("⚠️ `streamlit-drawable-canvas` 패키지가 설치되지 않았습니다. 터미널에서 `pip install streamlit-drawable-canvas`를 실행해주세요.")

    if canvas_supported:
        draw_cat = st.radio("그릴 도표 유형 선택:", ["🧊 3D 격자", "🔮 분자 오비탈", "⚛️ 루이스 전자점식", "🔺 기하 도형", "⌬ 2D 선구조식", "📊 표/그래프"], horizontal=True, key="canvas_category_select")
        
        with st.expander(f"💡 [{draw_cat}] 빠른 자동 생성 (프리셋)", expanded=True):
            st.markdown("### ✨ 빠른 자동 생성 (프리셋)")
            st.caption("수기로 그리기 어렵다면 아래 버튼을 눌러 캔버스로 불러온 뒤 수정하세요.")
            
            # --- 2D Vectorization Engine ---
            def generate_fabric_preset(cat, val):
                objects = []
                def add_text(text, left, top, color="#000000", size=24):
                    objects.append({
                        "type": "text", "version": "4.4.0", "originX": "center", "originY": "center",
                        "left": left, "top": top, "fill": color, "text": text,
                        "fontSize": size, "fontFamily": "sans-serif", "fontWeight": "bold",
                        "angle": 0, "scaleX": 1, "scaleY": 1,
                        "selectable": True, "evented": True, "hasControls": True, "hasBorders": True
                    })
                def add_line(x1, y1, x2, y2, color="#000000", width=4):
                    objects.append({
                        "type": "line", "version": "4.4.0", "originX": "center", "originY": "center",
                        "left": (x1 + x2) / 2, "top": (y1 + y2) / 2,
                        "x1": x1 - (x1 + x2)/2, "y1": y1 - (y1 + y2)/2,
                        "x2": x2 - (x1 + x2)/2, "y2": y2 - (y1 + y2)/2,
                        "stroke": color, "strokeWidth": width, "fill": "",
                        "angle": 0, "scaleX": 1, "scaleY": 1,
                        "selectable": True, "evented": True, "hasControls": True, "hasBorders": True
                    })
                def add_circle(left, top, radius=3, color="#000000"):
                    objects.append({
                        "type": "circle", "version": "4.4.0", "originX": "center", "originY": "center",
                        "left": left, "top": top, "radius": radius, "fill": color,
                        "angle": 0, "scaleX": 1, "scaleY": 1,
                        "selectable": True, "evented": True, "hasControls": True, "hasBorders": True
                    })
                
                if val == "H2O (물)":
                    add_text("O", 300, 180, "#B91C1C", 32)
                    add_text("H", 240, 240, "#1E293B", 24)
                    add_text("H", 360, 240, "#1E293B", 24)
                    add_line(290, 190, 250, 230)
                    add_line(310, 190, 350, 230)
                    # Lone pairs
                    add_circle(290, 150)
                    add_circle(310, 150)
                    add_circle(280, 160)
                    add_circle(320, 160)
                elif val == "CO2 (이산화탄소)":
                    add_text("O", 180, 200, "#B91C1C", 32)
                    add_text("C", 300, 200, "#1E293B", 32)
                    add_text("O", 420, 200, "#B91C1C", 32)
                    add_line(210, 192, 275, 192)
                    add_line(210, 208, 275, 208)
                    add_line(325, 192, 390, 192)
                    add_line(325, 208, 390, 208)
                    add_circle(170, 175)
                    add_circle(190, 175)
                    add_circle(170, 225)
                    add_circle(190, 225)
                    add_circle(410, 175)
                    add_circle(430, 175)
                    add_circle(410, 225)
                    add_circle(430, 225)
                elif val == "NH3 (암모니아)":
                    add_text("N", 300, 180, "#2563EB", 32)
                    add_text("H", 240, 250, "#1E293B", 24)
                    add_text("H", 300, 270, "#1E293B", 24)
                    add_text("H", 360, 250, "#1E293B", 24)
                    add_line(290, 195, 250, 240)
                    add_line(300, 195, 300, 255)
                    add_line(310, 195, 350, 240)
                    add_circle(290, 150)
                    add_circle(310, 150)
                elif val == "O2 (산소 분자)":
                    add_text("O", 250, 200, "#B91C1C", 32)
                    add_text("O", 350, 200, "#B91C1C", 32)
                    add_line(280, 192, 320, 192)
                    add_line(280, 208, 320, 208)
                    add_circle(230, 175); add_circle(250, 170)
                    add_circle(230, 225); add_circle(250, 230)
                    add_circle(350, 170); add_circle(370, 175)
                    add_circle(350, 230); add_circle(370, 225)
                elif val == "Benzene (벤젠)":
                    # Hexagon
                    import math
                    cx, cy, r = 300, 200, 80
                    pts = [(cx + r*math.cos(a), cy + r*math.sin(a)) for a in [math.pi/6, math.pi/2, 5*math.pi/6, 7*math.pi/6, 3*math.pi/2, 11*math.pi/6]]
                    for i in range(6):
                        add_line(pts[i][0], pts[i][1], pts[(i+1)%6][0], pts[(i+1)%6][1], width=5)
                    # Double bonds
                    add_line(cx + r*0.7*math.cos(math.pi/6), cy + r*0.7*math.sin(math.pi/6), cx + r*0.7*math.cos(math.pi/2), cy + r*0.7*math.sin(math.pi/2), width=3)
                    add_line(cx + r*0.7*math.cos(5*math.pi/6), cy + r*0.7*math.sin(5*math.pi/6), cx + r*0.7*math.cos(7*math.pi/6), cy + r*0.7*math.sin(7*math.pi/6), width=3)
                    add_line(cx + r*0.7*math.cos(3*math.pi/2), cy + r*0.7*math.sin(3*math.pi/2), cx + r*0.7*math.cos(11*math.pi/6), cy + r*0.7*math.sin(11*math.pi/6), width=3)
                elif val == "Acetic Acid (아세트산)":
                    add_line(250, 200, 310, 170) # C-C
                    add_line(310, 170, 370, 200) # C-O
                    add_line(305, 170, 305, 110) # C=O double
                    add_line(315, 170, 315, 110)
                    add_text("O", 310, 90, "#B91C1C", 28)
                    add_text("OH", 390, 210, "#B91C1C", 28)
                else:
                    return None
                return {"version": "4.4.0", "objects": objects}

            preset_clicked = None
            
            if "3D 격자" in draw_cat:
                col1, col2, col3 = st.columns(3)
                with col1:
                    if st.button("🧊 FCC 격자"): preset_clicked = ("3D 격자", "Face-Centered (FCC)")
                    if st.button("🧊 HCP 격자"): preset_clicked = ("3D 격자", "HCP (Hexagonal)")
                with col2:
                    if st.button("🧊 BCC 격자"): preset_clicked = ("3D 격자", "Body-Centered (BCC)")
                    if st.button("🧊 NaCl 구조"): preset_clicked = ("3D 격자", "NaCl (Rock Salt)")
                with col3:
                    if st.button("🧊 SC 격자"): preset_clicked = ("3D 격자", "Simple Cubic (SC)")
            elif "분자 오비탈" in draw_cat:
                col1, col2, col3 = st.columns(3)
                with col1:
                    if st.button("🔮 물(H2O) 오비탈"): preset_clicked = ("분자 오비탈", "물 (H2O)")
                with col2:
                    if st.button("🔮 에텐(C2H4) sp2"): preset_clicked = ("분자 오비탈", "에텐 (C2H4)")
                with col3:
                    if st.button("🔮 에타인(C2H2) sp"): preset_clicked = ("분자 오비탈", "에타인 (C2H2)")
            elif "루이스" in draw_cat:
                col1, col2, col3, col4 = st.columns(4)
                with col1:
                    if st.button("⚛️ H2O 구조"): preset_clicked = ("루이스", "H2O (물)")
                with col2:
                    if st.button("⚛️ CO2 구조"): preset_clicked = ("루이스", "CO2 (이산화탄소)")
                with col3:
                    if st.button("⚛️ NH3 구조"): preset_clicked = ("루이스", "NH3 (암모니아)")
                with col4:
                    if st.button("⚛️ O2 구조"): preset_clicked = ("루이스", "O2 (산소 분자)")
            elif "선구조식" in draw_cat:
                col1, col2, col3, col4 = st.columns(4)
                with col1:
                    if st.button("⌬ 벤젠 고리"): preset_clicked = ("선구조식", "Benzene (벤젠)")
                with col2:
                    if st.button("⌬ 아세트산"): preset_clicked = ("선구조식", "Acetic Acid (아세트산)")
                with col3:
                    if st.button("⌬ 아스피린"): preset_clicked = ("선구조식", "Aspirin (아스피린)")
                with col4:
                    if st.button("⌬ 카페인"): preset_clicked = ("선구조식", "Caffeine (카페인)")
            elif "기하 도형" in draw_cat:
                col1, col2 = st.columns(2)
                with col1:
                    if st.button("🔺 3D 정육면체 자동 완성"): preset_clicked = ("기하 도형", "정육면체 (Cube)")
                with col2:
                    if st.button("🌐 대화형 3D 분자 조립기 (PhET)"): preset_clicked = ("기하 도형", "PhET")
            else:
                st.info("이 카테고리는 프리셋을 제공하지 않습니다.")
                
            if preset_clicked:
                cat, val = preset_clicked
                
                if cat in ["3D 격자", "분자 오비탈"]:
                    st.session_state["active_3d_preset_cat"] = cat
                    st.session_state["active_3d_preset_val"] = val
                    st.session_state["preset_elev"] = 25 if cat == "분자 오비탈" else 20
                    st.session_state["preset_azim"] = -55 if cat == "분자 오비탈" else 45
                else:
                    st.session_state.pop("active_3d_preset_cat", None)
                    st.session_state.pop("active_3d_preset_val", None)

                if cat in ["루이스", "선구조식"] and val in ["H2O (물)", "CO2 (이산화탄소)", "NH3 (암모니아)", "O2 (산소 분자)", "Benzene (벤젠)", "Acetic Acid (아세트산)"]:
                    # 벡터화 가능한 2D 프리셋인 경우
                    fabric_json = generate_fabric_preset(cat, val)
                    if fabric_json:
                        if "global_canvas_bg" in st.session_state:
                            del st.session_state["global_canvas_bg"] # 배경 제거
                        st.session_state["canvas_initial_drawing"] = fabric_json
                        st.session_state["canvas_tool_index"] = 4 # 자동으로 '이동/크기조절' 도구 선택
                        import uuid
                        st.session_state["canvas_key_id"] = "canvas_" + str(uuid.uuid4())
                        st.success(f"✅ '{val}' 도표가 [편집 가능한 조각(벡터)] 형태로 캔버스로 불러와졌습니다! 이제 원자나 결합 선을 마우스로 잡고 이동하거나 형태를 변형해 보세요.")
                elif val == "PhET":
                    st.session_state["show_phet_iframe"] = True
                    st.success("✅ 대화형 3D 분자 조립기(PhET)가 열렸습니다. 마우스로 분자를 직접 조립하고 입체적으로 돌려보세요!")
                else:
                    # 기존 이미지 방식 (3D 격자, 분자 오비탈 등)
                    st.session_state.pop("show_phet_iframe", None)
                    p_stream = None
                    p_errors = []
                    if cat == "3D 격자":
                        p_stream, p_errors = draw_unit_cell(val, elev=st.session_state["preset_elev"], azim=st.session_state["preset_azim"])
                    elif cat == "분자 오비탈":
                        p_stream, p_errors = draw_orbital_diagram(val, elev=st.session_state["preset_elev"], azim=st.session_state["preset_azim"])
                    elif cat == "루이스":
                        p_stream, p_errors = draw_lewis_structure(val)
                    elif cat == "선구조식":
                        p_stream, p_errors = draw_skeletal_structure(val)
                    elif cat == "기하 도형":
                        p_stream, p_errors = draw_schematic(val)
                        
                    if p_stream:
                        from PIL import Image
                        st.session_state["global_canvas_bg"] = Image.open(p_stream).copy()
                        import uuid
                        st.session_state["canvas_key_id"] = "canvas_" + str(uuid.uuid4())
                        st.success(f"✅ '{val}' 도표가 캔버스로 불러와졌습니다! 이제 아래 캔버스에서 수기로 수정해 보세요.")
        
        save_canvas_btn = False
        ai_transform_btn = False
        if st.session_state.get("show_phet_iframe", False):
            st.markdown("### 🌐 대화형 3D 분자 조립기 (PhET)")
            st.caption("참고: 이 모드에서는 3D 상호작용이 완벽히 지원되지만, 펜으로 덧그리는 '수기 필기' 기능은 사용할 수 없습니다. 필기가 필요하시다면 다른 프리셋을 사용해 주세요.")
            import streamlit.components.v1 as components
            components.iframe("https://phet.colorado.edu/sims/html/molecule-shapes/latest/molecule-shapes_all.html", height=600, width=800)
            if st.button("❌ PhET 닫기 및 그림판으로 돌아가기"):
                st.session_state.pop("show_phet_iframe", None)
                st.rerun()
        else:
            col_c1, col_c2 = st.columns([1, 4])
            with col_c1:
                stroke_width = st.slider("펜 굵기", 1, 15, 3)
            stroke_color = st.color_picker("펜 색상", "#000000")
            bg_color = st.color_picker("배경 색상", "#ffffff")
            
            tool_mapping = {
                "✍️ 펜 (태블릿 글쓰기/자유그리기)": "freedraw",
                "📏 직선": "line",
                "⬜ 사각형": "rect",
                "⭕ 원": "circle",
                "🖐️ 캔버스 이동/크기조절": "transform"
            }
            
            if "canvas_tool_index" not in st.session_state:
                st.session_state["canvas_tool_index"] = 0
                
            tool_names = list(tool_mapping.keys())
            current_index = st.session_state.get("canvas_tool_index", 0)
            if current_index >= len(tool_names):
                current_index = 0
                
            selected_tool = st.selectbox("도구 선택 (글쓰기 등)", tool_names, index=current_index)
            st.session_state["canvas_tool_index"] = tool_names.index(selected_tool)
            drawing_mode = tool_mapping[selected_tool]
            if st.button("🗑️ 캔버스 템플릿(배경) 완전히 지우기", use_container_width=True, help="캔버스 아래의 작은 휴지통 아이콘은 '직접 그린 펜 자국'만 지웁니다. 템플릿 전체를 지우려면 이 버튼을 누르세요."):
                if "global_canvas_bg" in st.session_state:
                    del st.session_state["global_canvas_bg"]
                if "canvas_initial_drawing" in st.session_state:
                    del st.session_state["canvas_initial_drawing"]
                import uuid
                st.session_state["canvas_key_id"] = "canvas_" + str(uuid.uuid4())
                st.rerun()
            
            with col_c2:
                bg_img = st.session_state.get("global_canvas_bg", None)
                canvas_key = st.session_state.get("canvas_key_id", "global_drawing_board")
                initial_drawing = st.session_state.get("canvas_initial_drawing", None)
                
                canvas_result = st_canvas(
                    fill_color="rgba(255, 165, 0, 0.3)",
                    stroke_width=stroke_width,
                    stroke_color=stroke_color,
                    background_color=bg_color,
                    background_image=bg_img,
                    height=400,
                    width=600,
                    drawing_mode=drawing_mode,
                    initial_drawing=initial_drawing,
                    key=canvas_key,
                )
                
                # 3D 입체 회전 컨트롤 (프리셋이 3D일 때만 표시)
                if st.session_state.get("active_3d_preset_cat") is not None:
                    st.markdown("---")
                    st.markdown("### 🧭 3D 도표 입체 회전 컨트롤")
                    col_sl1, col_sl2 = st.columns(2)
                    with col_sl1:
                        new_elev = st.slider("상하 회전 (Elevation)", -90, 90, st.session_state.get("preset_elev", 20), key="slider_elev")
                    with col_sl2:
                        new_azim = st.slider("좌우 회전 (Azimuth)", -180, 180, st.session_state.get("preset_azim", 45), key="slider_azim")
                    
                    if new_elev != st.session_state.get("preset_elev") or new_azim != st.session_state.get("preset_azim"):
                        st.session_state["preset_elev"] = new_elev
                        st.session_state["preset_azim"] = new_azim
                        
                        cat = st.session_state["active_3d_preset_cat"]
                        val = st.session_state["active_3d_preset_val"]
                        if cat == "3D 격자":
                            p_stream, _ = draw_unit_cell(val, elev=new_elev, azim=new_azim)
                        elif cat == "분자 오비탈":
                            p_stream, _ = draw_orbital_diagram(val, elev=new_elev, azim=new_azim)
                        else:
                            p_stream = None
                            
                        if p_stream:
                            from PIL import Image
                            st.session_state["global_canvas_bg"] = Image.open(p_stream).copy()
                            # 사용자가 이미 그린 내용 유지
                            if canvas_result and canvas_result.json_data is not None:
                                st.session_state["canvas_initial_drawing"] = canvas_result.json_data
                            import uuid
                            st.session_state["canvas_key_id"] = "canvas_" + str(uuid.uuid4())
                            st.rerun()
                
                st.markdown("---")
                col_t1, col_t2, col_t3, col_t4 = st.columns([2.5, 1, 1.2, 1.5])
                with col_t1:
                    text_to_add = st.text_input("🔤 캔버스 텍스트 도구", placeholder="텍스트 입력 (예: H2O)")
                with col_t2:
                    text_color = st.color_picker("텍스트 색상", stroke_color)
                with col_t3:
                    st.write("") # spacing
                    if st.button("➕ 삽입", use_container_width=True):
                        if text_to_add:
                            import uuid
                            import random
                            current_state = canvas_result.json_data if (canvas_result and canvas_result.json_data is not None) else {"version": "4.4.0", "objects": []}
                            if "objects" not in current_state:
                                current_state["objects"] = []
                                
                            offset_x = random.randint(-20, 20)
                            offset_y = random.randint(-20, 20)
                            
                            new_text_obj = {
                                "type": "text",
                                "version": "4.4.0",
                                "originX": "left",
                                "originY": "top",
                                "left": 250 + offset_x,
                                "top": 180 + offset_y,
                                "fill": text_color,
                                "text": text_to_add,
                                "fontSize": 24,
                                "fontFamily": "sans-serif",
                                "fontWeight": "bold",
                                "angle": 0,
                                "scaleX": 1,
                                "scaleY": 1,
                                "selectable": True,
                                "evented": True,
                                "hasControls": True,
                                "hasBorders": True
                            }
                            current_state["objects"].append(new_text_obj)
                            st.session_state["canvas_initial_drawing"] = current_state
                            st.session_state["canvas_key_id"] = "canvas_" + str(uuid.uuid4())
                            # 텍스트 삽입 후 자동으로 '이동/크기조절' 도구로 변경하여 바로 드래그 가능하게 함
                            st.session_state["canvas_tool_index"] = 4
                            st.rerun()
                with col_t4:
                    st.write("") # spacing
                    if st.button("↩️ 되돌리기", use_container_width=True, help="가장 최근에 추가한 도형, 선, 또는 텍스트 1개를 지웁니다."):
                        import uuid
                        current_state = canvas_result.json_data if (canvas_result and canvas_result.json_data is not None) else None
                        if current_state and "objects" in current_state and len(current_state["objects"]) > 0:
                            # 마지막 객체(텍스트/도형 무관) 삭제
                            current_state["objects"].pop()
                            st.session_state["canvas_initial_drawing"] = current_state
                            st.session_state["canvas_key_id"] = "canvas_" + str(uuid.uuid4())
                            st.rerun()
                
            col_b1, col_b2 = st.columns(2)
            with col_b1:
                save_canvas_btn = st.button("💾 캔버스 화면(배경+그림) 그대로 워드에 추가 (AI 변환 생략)", use_container_width=True)
            with col_b2:
                ai_transform_btn = st.button("🧠 직접 그린 그림 AI 완벽 변환 후 워드 추가", use_container_width=True, type="primary")

        if save_canvas_btn:
            if canvas_result.image_data is not None:
                import numpy as np
                from PIL import Image
                import io
                img_arr = canvas_result.image_data.astype(np.uint8)
                drawn_img = Image.fromarray(img_arr)
                
                bg = st.session_state.get("global_canvas_bg", None)
                if bg:
                    bg = bg.resize((600, 400)).convert("RGBA")
                else:
                    bg = Image.new('RGBA', (600, 400), bg_color)
                
                if drawn_img.mode == 'RGBA':
                    final_img = Image.alpha_composite(bg, drawn_img)
                else:
                    final_img = drawn_img
                
                final_img = final_img.convert('RGB')
                
                img_stream = io.BytesIO()
                final_img.save(img_stream, format="PNG")
                img_stream.seek(0)
                from docx.shared import Inches
                st.session_state.word_doc.add_picture(img_stream, width=Inches(4.0))
                st.success("✅ 캔버스 화면이 워드에 그대로 추가되었습니다!")

        if ai_transform_btn:
            if canvas_result.image_data is not None:
                import numpy as np
                from PIL import Image
                import io
                
                img_arr = canvas_result.image_data.astype(np.uint8)
                drawn_img = Image.fromarray(img_arr)
                background = Image.new('RGB', drawn_img.size, bg_color)
                
                if drawn_img.mode == 'RGBA':
                    background.paste(drawn_img, mask=drawn_img.split()[3])
                else:
                    background = drawn_img.convert('RGB')
                
                if "3D 격자" in draw_cat:
                    target_info = '1. "3D 격자 생성": parameters: {"cell_type": "Face-Centered (FCC)" | "Body-Centered (BCC)" | "Simple Cubic (SC)" | "HCP (Hexagonal)" | "NaCl (Rock Salt)", "lattice_size": 1, "atom_rad": 0.18}'
                elif "분자 오비탈" in draw_cat:
                    target_info = '4. "분자 오비탈 시각화": parameters: {"molecule_type": "에텐 (C2H4)" | "에타인 (C2H2)" | "물 (H2O)"}'
                elif "루이스" in draw_cat:
                    target_info = '3. "루이스 전자점식": parameters: {"molecule": e.g., "O2", "N2", "H2O"}'
                elif "기하 도형" in draw_cat:
                    target_info = '5. "기본 도식 및 기하 도형": parameters: {"shape_type": "정사면체 (Tetrahedral)" | "옥타헤드론 (Octahedral)" | "삼각쌍뿔 (Trigonal Bipyramidal)"}'
                elif "선구조식" in draw_cat:
                    target_info = '6. "2D 분자 구조 / 선구조식": parameters: {"molecule": e.g., "벤젠", "아스피린"}'
                else:
                    target_info = '7. "표 / 기타 그래프": parameters: {"type": "Table", "description": "설명"}'

                prompt = f"""
                You are an expert scientific diagram analyzer. The user has hand-drawn a sketch.
                The user explicitly selected the category: {draw_cat}.
                You MUST interpret the image as the following category and extract the appropriate parameters:
                {target_info}
                
                Return ONLY a valid JSON string (no markdown ticks) with keys:
                {{
                    "category": "<the category selected by the user>",
                    "reasoning": "<short explanation of how you mapped the user's sketch to the parameters>",
                    "parameters": {{ ... }},
                    "crop_boxes": [],
                    "has_complex_annotations": false
                }}
                """
                
                if not st.session_state.get("gemini_api_key"):
                    st.error("Gemini API 키가 필요합니다. 왼쪽 메뉴 하단에 키를 입력해주세요.")
                else:
                    with st.spinner(f"'{draw_cat}' 그림을 분석하여 깔끔한 도표로 변환 중입니다..."):
                        try:
                            import google.generativeai as genai
                            import json
                            genai.configure(api_key=st.session_state.gemini_api_key)
                            
                            res_text = robust_generate_content(prompt, images=[background], use_grounding=False)
                            if not res_text:
                                raise Exception("AI 응답을 받지 못했습니다.")
                                
                            import re
                            match = re.search(r'\{.*\}', res_text, re.DOTALL)
                            if match:
                                res_text = match.group(0)
                                
                            data = json.loads(res_text)
                            st.success(f"✅ AI 인식 완료: {data.get('reasoning', '')}")
                            
                            cat = data.get("category", "")
                            params = data.get("parameters", {})
                            
                            img_stream = None
                            errors = []
                            
                            if "3D 격자" in draw_cat:
                                img_stream, errors = draw_unit_cell(params.get("cell_type", "Face-Centered (FCC)"), params.get("lattice_size", 1), params.get("atom_rad", 0.18))
                            elif "루이스" in draw_cat:
                                img_stream, errors = draw_lewis_structure(params.get("molecule", "O2"))
                            elif "오비탈" in draw_cat:
                                img_stream, errors = draw_orbital_diagram(params.get("molecule_type", "에텐 (C2H4)"))
                            elif "도형" in draw_cat or "도식" in draw_cat:
                                img_stream, errors = draw_schematic(params.get("shape_type", "정사면체 (Tetrahedral)"))
                            elif "2D" in draw_cat or "선구조식" in draw_cat:
                                molecule_name = str(params.get("molecule", ""))
                                supported = ["Butane (뷰테인)", "Hexane (헥세인)", "Cyclohexane (사이클로헥세인)", "Benzene (벤젠)", "Acetone (아세톤)", "Acetic Acid (아세트산)"]
                                matched = next((s for s in supported if molecule_name.lower() in s.lower() or s.lower() in molecule_name.lower()), None)
                                if matched:
                                    img_stream, errors = draw_skeletal_structure(matched)
                                else:
                                    img_stream = io.BytesIO()
                                    background.save(img_stream, format='PNG')
                            else:
                                img_stream = io.BytesIO()
                                background.save(img_stream, format='PNG')
                                
                            if errors:
                                for e in errors: st.error(e)
                                
                            if img_stream:
                                from docx.shared import Inches
                                img_stream.seek(0)
                                st.image(img_stream, caption="AI 변환 결과")
                                img_stream.seek(0)
                                st.session_state.word_doc.add_picture(img_stream, width=Inches(4.0))
                                st.success("변환된 도표가 워드 파일에 추가되었습니다!")
                                
                        except Exception as e:
                            st.error(f"오류가 발생했습니다: {e}")
            else:
                st.warning("먼저 캔버스에 그림을 그려주세요.")

        st.markdown("---")
        st.markdown("### 📥 캔버스 작업물 워드 파일 다운로드")
        st.caption("위 버튼들로 워드에 추가한 도표들을 하나의 워드 파일(.docx)로 다운로드합니다.")
        import io
        final_doc_stream = io.BytesIO()
        st.session_state.word_doc.save(final_doc_stream)
        final_doc_stream.seek(0)
        st.download_button(
            label="📥 지금까지 추가한 그림 모두 워드(.docx)로 다운로드",
            data=final_doc_stream,
            file_name="Canvas_Diagrams.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
            type="primary"
        )

elif menu == "📝 수기 노트 AI 문서화":
    st.title("📝 수기 노트 AI 문서화")
    uploaded_files = st.file_uploader("이미지, PDF 및 HWP 업로드", type=["png", "jpg", "jpeg", "pdf", "hwp"], accept_multiple_files=True)

    col1, col2 = st.columns(2)
    with col1:
        start_conversion = st.button("문서 변환 시작", use_container_width=True)
    with col2:
        start_analysis = st.button("이미지 상세 분석", use_container_width=True)

    if uploaded_files:


        if start_conversion or st.session_state.get("hw_conversion_active", False):
            st.session_state.hw_conversion_active = True
            if not api_key: st.error("API 키 입력 필수!")
            else:
                genai.configure(api_key=api_key)
                
                # 1. 이미지 변환 준비 (PDF 지원)
                hw_imgs = []
                for uploaded_file in uploaded_files:
                    f_bytes = uploaded_file.getvalue()
                    if uploaded_file.name.lower().endswith(".pdf"):
                        import fitz
                        pdf_doc = fitz.open(stream=f_bytes, filetype="pdf")
                        max_pages = min(len(pdf_doc), 1000)
                        for page_idx in range(max_pages):
                            page = pdf_doc.load_page(page_idx)
                            pix = page.get_pixmap(matrix=fitz.Matrix(3.0, 3.0))
                            from PIL import Image
                            import io
                            hw_imgs.append(Image.open(io.BytesIO(pix.tobytes("png"))))
                        pdf_doc.close()
                    elif uploaded_file.name.lower().endswith(".hwp"):
                        st.warning(f"⚠️ **[{uploaded_file.name}] HWP 텍스트 모드로 분석 중... (그림 누락 가능성 있음)**")
                        import tempfile, subprocess, os
                        with tempfile.NamedTemporaryFile(delete=False, suffix=".hwp") as tmp:
                            tmp.write(f_bytes)
                            tmp_path = tmp.name
                        try:
                            result = subprocess.run(["hwp5txt", tmp_path], capture_output=True, text=True)
                            if result.returncode == 0:

                                hw_imgs.append(f"--- [HWP 수기/문서 텍스트 데이터: {uploaded_file.name}] ---\n{result.stdout}")
                            else:
                                import zipfile
                                import xml.etree.ElementTree as ET
                                hwpx_text = ""
                                try:
                                    with zipfile.ZipFile(tmp_path, 'r') as zf:
                                        for item in zf.namelist():
                                            if item.startswith("Contents/section") and item.endswith(".xml"):
                                                root = ET.fromstring(zf.read(item))
                                                for elem in root.iter():
                                                    if elem.tag.endswith('}t') and elem.text:
                                                        hwpx_text += elem.text + " "
                                                    elif elem.tag.endswith('}p'):
                                                        hwpx_text += ""
                                    if hwpx_text.strip():
                                        hw_imgs.append(f"--- [HWPX 수기/문서 텍스트 데이터: {uploaded_file.name}] ---\n{hwpx_text}")
                                        st.toast("HWPX 포맷으로 성공적으로 추출되었습니다.")
                                    else:
                                        st.error(f"❌ {uploaded_file.name} 텍스트 추출 실패. PDF 변환을 권장합니다.")
                                        st.stop()
                                except:
                                    st.error(f"❌ {uploaded_file.name} 추출 불가. 구형/특수 포맷이므로 PDF로 변환 후 업로드해주세요.")
                                    st.stop()
                        except Exception as e:
                            st.error(f"❌ HWP 오류: {e}")
                            st.stop()
                        finally:
                            os.remove(tmp_path)
                    else:
                        from PIL import Image
                        import io
                        hw_imgs.append(Image.open(io.BytesIO(f_bytes)))

                if hw_imgs:
                    import hashlib
                    # 파일 내용 기반 해시 생성
                    hw_file_names = "".join([getattr(f, 'name', 'img') + str(getattr(f, 'size', 0)) for f in uploaded_files])
                    hw_file_hash = hashlib.md5(hw_file_names.encode()).hexdigest()
                    
                    if st.session_state.get("hw_curent_file_hash") != hw_file_hash:
                        st.session_state.hw_analysis_buffer = {}
                        st.session_state.hw_curent_file_hash = hw_file_hash
                        save_state()

                    if "hw_analysis_buffer" not in st.session_state:
                        st.session_state.hw_analysis_buffer = {}
                    
                    total_pages = len(hw_imgs)
                    
                    total_pages = len(hw_imgs)
                    st.success(f"🎯 총 {total_pages}페이지를 감지했습니다. 페이지별 순차 문서화를 시작합니다.")
                    
                    # [페이지별 순차 문서화]
                    for p_idx in range(total_pages):
                        q_key = f"hw_page_{p_idx}"
                        if q_key in st.session_state.hw_analysis_buffer:
                            continue

                        progress_label = f"{p_idx+1}) {total_pages}페이지 중 {p_idx+1}페이지"
                        st.info(f"⏳ [수기 정밀 타격] {progress_label} 문서화 중...")
                        
                        try:
                            # [단일 단계 초정밀 전사 모드: 부모-자식 그룹화]
                            hw_ocr_prompt = fr"""당신은 수기 문서 및 학술 데이터 디지털화 전문가입니다. **인사말이나 "전사를 시작합니다" 등 어떠한 메타 코멘트도 절대 출력하지 마세요.**
오직 원본 내용의 **전사 결과물**만 출력하세요.

**[수행 지침: 소문 반복 템플릿]**
- 각 소문(a, b, c, d, e, f, g...)마다 아래 형식을 반드시 반복하세요:
  1. **(소문 x) 질문 원문 전사**
  2. **원본 시각화 재현**: 도표나 그림은 반드시 일반 텍스트 기반 ASCII 아트나 Markdown 표로만 재현하십시오.
  3. **내용 전사**: (원본 텍스트 그대로되, 아래의 빨간 글씨 에러 방지 규칙을 엄수할 것)

- 대문 하위에 소문들을 그룹으로 묶어 배치하고, 건너뛰는 번호가 없도록 하세요.
- **[웹 렌더링 에러(빨간색 글씨) 방지 절대 규칙]**:
  1. 수식 내부에 한글을 쓰면 에러가 나므로 `\text{{한글}}`이나 `\fbox{{한글}}` 대신 수식 밖으로 빼서 일반 텍스트로 적으세요.
  2. `\begin{{tikzpicture}}`, `\begin{{gathered}}` 등 지원되지 않는 복잡한 LaTeX 환경은 절대 사용 금지입니다. 도형이나 분자 구조는 오직 키보드 특수기호(ASCII)만 사용해서 그리세요.

### 📌 {progress_label} 디지털 문서화 결과 (그룹화 완료)"""
                            
                            res = robust_generate_content(hw_ocr_prompt, images=[hw_imgs[p_idx]], use_grounding=False)
                            
                            if res:
                                import os
                                img_dir = os.path.abspath("hw_cache_images")
                                os.makedirs(img_dir, exist_ok=True)
                                img_path = os.path.join(img_dir, f"{q_key}.png")
                                hw_imgs[p_idx].save(img_path)
                                
                                st.session_state.hw_analysis_buffer[q_key] = res
                                st.session_state.word_doc.add_paragraph(res)
                                save_state()
                                st.toast(f"✅ {progress_label} 문서화 완료")
                            else:
                                st.warning(f"⚠️ {p_idx+1}페이지 분석 중 응답이 없습니다. 잠시 후 다시 시도합니다.")
                                time.sleep(5)
                        except Exception as e:
                            err_msg = str(e)
                            if "QUOTA_EXHAUSTED" in err_msg or "ResourceExhausted" in err_msg or "429" in err_msg:
                                import time, re
                                match = re.search(r'QUOTA_EXHAUSTED:(\d+\.?\d*)', err_msg)
                                wait_time = int(float(match.group(1))) if match else 65
                                msg = st.empty()
                                for i in range(wait_time, 0, -1):
                                    msg.warning(f"🚫 API 한도 초과. 멈춤 방지를 위해 {i}초간 강제 휴식 후 자동 재개합니다...")
                                    time.sleep(1)
                                msg.empty()
                                st.rerun()
                            
                            st.error(f"❌ {p_idx+1}페이지 분석 중 치명적 오류 (스킵 후 다음 진행): {e}")
                            st.session_state.hw_analysis_buffer[q_key] = f"⚠️ {p_idx+1}페이지 오류로 스킵됨: {e}"
                            save_state()
                            continue
                            
                    st.session_state.hw_conversion_active = False
                    st.success("🎉 모든 페이지의 분석이 완료되었습니다!")




        if start_analysis:
            for file_obj in uploaded_files:
                if file_obj.name.lower().endswith('.pdf'):
                    st.warning(f"PDF 파일('{file_obj.name}')은 단일 이미지 상세 분석 대상이 아닙니다. 좌측의 '문서 변환 시작' 버튼을 사용하세요.")
                elif file_obj.name.lower().endswith('.hwp'):
                    st.warning(f"HWP 파일('{file_obj.name}')은 단일 이미지 상세 분석 대상이 아닙니다. 좌측의 '문서 변환 시작' 버튼을 사용하세요.")
                else:
                    handle_image_analysis(file_obj)


    # --- 상시 수기 분석 결과 표시 (새로고침 대응) ---
    if st.session_state.get("hw_analysis_buffer"):
        with st.container(border=True):
            col_t, col_btn = st.columns([0.8, 0.2])
            with col_t:
                st.markdown("### 📋 수기 노트 AI 분석 결과")
            with col_btn:
                if st.button("🗑️ 전체 초기화", use_container_width=True):
                    st.session_state.hw_analysis_buffer = {}
                    save_state()
                    st.rerun()
            # 개별 페이지별 결과 표시 및 관리
            hw_keys = sorted([k for k in st.session_state.hw_analysis_buffer.keys() if k.startswith("hw_page_")])
            for k in hw_keys:
                col_text, col_trash = st.columns([0.9, 0.1])
                with col_text:
                    try:
                        p_num = int(k.split('_')[2]) + 1
                        label = f"📝 수기 전사: {p_num}페이지"
                    except:
                        label = f"📝 수기 분석 ({k})"
                    
                    with st.expander(label, expanded=False):
                        import os
                        img_path = os.path.abspath(f"hw_cache_images/{k}.png")
                        if os.path.exists(img_path):
                            st.image(img_path, use_container_width=True)
                        st.markdown(st.session_state.hw_analysis_buffer[k])
                
                with col_trash:
                    if st.button("🗑️", key=f"del_{k}", help="이 페이지를 분석 결과에서 삭제합니다."):
                        del st.session_state.hw_analysis_buffer[k]
                        save_state()
                        st.rerun()

            if hw_keys:
                combined_res_parts = []
                import os
                for k in hw_keys:
                    img_path = os.path.abspath(f"hw_cache_images/{k}.png")
                    if os.path.exists(img_path):
                        combined_res_parts.append(f"![원본 이미지]({img_path})\n\n{st.session_state.hw_analysis_buffer[k]}")
                    else:
                        combined_res_parts.append(st.session_state.hw_analysis_buffer[k])
                combined_res = "\n---\n\n".join(combined_res_parts)

                # Word 다운로드 섹션 통합
                if st.button("🚀 분석 결과를 Word로 변환하기", key="hw_convert_btn", use_container_width=True):
                    with st.spinner("분석 내용을 워드 파일로 변환 중입니다..."):
                        try:
                            # 1. Pandoc 변환 시도
                            import tempfile, pypandoc, os
                            with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
                                tmp_path = tmp.name
                            pypandoc.convert_text(combined_res, 'docx', format='markdown', outputfile=tmp_path)
                            with open(tmp_path, "rb") as f:
                                st.session_state.hw_docx_bytes = f.read()
                            os.remove(tmp_path)
                        except:
                            # 2. Fallback: python-docx 엔진
                            import io
                            from docx import Document
                            doc_stream = io.BytesIO()
                            doc_fallback = Document()
                            doc_fallback.add_heading("수기 노트 AI 분석 결과", 0)
                            for line in combined_res.split('\n'):
                                if line.startswith("![원본 이미지]("):
                                    img_p = line.split("(")[1].split(")")[0]
                                    if os.path.exists(img_p):
                                        from docx.shared import Inches
                                        doc_fallback.add_picture(img_p, width=Inches(6.0))
                                else:
                                    doc_fallback.add_paragraph(line)
                            doc_fallback.save(doc_stream)
                            st.session_state.hw_docx_bytes = doc_stream.getvalue()
                        st.success("✅ 변환이 완료되었습니다! 아래에서 다운로드하세요.")

                if st.session_state.get("hw_docx_bytes"):
                    render_download_and_open_buttons("수기 분석 결과", st.session_state.hw_docx_bytes, "Handwritten_Analysis.docx", "hw_analysis")

elif menu == "💬 실시간 AI 학술 상담 (ChatGPT 스타일)":
    st.markdown("<h1 class='main-header'>💬 실시간 AI 학술 상담 (ChatGPT 스타일)</h1>", unsafe_allow_html=True)
    st.markdown("전 세계 AI 지능을 결집한 초정밀 학술 고문과 실시간으로 대화하세요. (Gemini 1.5 Pro 기반)")
    
    if "messages" not in st.session_state:
        st.session_state.messages = []

    # 대화 기록 표시
    for message in st.session_state.messages:
        with st.chat_message(message["role"]):
            st.markdown(message["content"])

    # 채팅 입력
    if prompt := st.chat_input("학술적 질문이나 상담하고 싶은 내용을 입력하세요..."):
        st.session_state.messages.append({"role": "user", "content": prompt})
        with st.chat_message("user"):
            st.markdown(prompt)

        with st.chat_message("assistant"):
            with st.spinner("최첨단 AI 군단이 답변을 생성 중입니다... (무료 서버 동원 가능)"):
                # 전수 동원 모드 사용
                response = robust_generate_content(prompt, use_grounding=False)
                
                # 무료 우회/백업망 (API 키가 없거나 사용량 초과 시 자동 가동)
                if not response:
                    try:
                        import urllib.request, urllib.parse, ssl
                        ctx = ssl._create_unverified_context()
                        history_text = "\\n".join([f"{m['role']}: {m['content'][:200]}" for m in st.session_state.messages[-3:]])
                        fallback_prompt = f"당신은 최고 수준의 학술 고문입니다. 한국어로 답변하세요. 이전 대화:\\n{history_text}\\n\\n질문: {prompt}"
                        
                        req = urllib.request.Request(
                            "https://text.pollinations.ai/" + urllib.parse.quote(fallback_prompt),
                            headers={'User-Agent': 'Mozilla/5.0'}
                        )
                        free_res = urllib.request.urlopen(req, context=ctx).read().decode('utf-8')
                        
                        if "⚠️ **IMPORTANT NOTICE** ⚠️" in free_res:
                            free_res = free_res.split("normally.")[-1].strip()
                            
                        if free_res.strip():
                            response = free_res.strip() + "\n\n*(🚀 무료 공용 AI 서버를 통해 생성된 답변입니다)*"
                    except Exception as e:
                        pass

                if response:
                    st.markdown(response)
                    st.session_state.messages.append({"role": "assistant", "content": response})
                    save_state()
                else:
                    st.error("답변 생성에 실패했습니다. (API 제한 및 무료 서버 과부하)")

save_state()
